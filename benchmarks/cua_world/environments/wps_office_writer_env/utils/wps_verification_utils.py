#!/usr/bin/env python3
"""
WPS Office Writer verification utilities for gym-anything tasks.
Provides helper functions to verify WPS Writer document tasks using DOCX/ODT parsing.

WPS Office uses the same file formats as MS Office, so python-docx works for verification.
"""

import logging
import os
import re
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional, Callable

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import python-docx for DOCX parsing
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - DOCX parsing will be limited")
    DOCX_AVAILABLE = False

# Import odfpy for ODT parsing
try:
    from odf import opendocument, text as odf_text, style as odf_style
    ODF_AVAILABLE = True
except ImportError:
    logger.warning("odfpy not available - ODT parsing will be limited")
    ODF_AVAILABLE = False


# ============================================================================
# Core copy/parse functions
# ============================================================================

def copy_and_parse_document(container_path: str, copy_from_env_fn: Callable,
                            file_format: str = 'auto') -> Tuple[bool, Any, str, str]:
    """
    Copy a document from the container and parse it.

    Args:
        container_path: Path to the file inside the container.
        copy_from_env_fn: Framework-provided copy function.
        file_format: 'docx', 'odt', or 'auto' (detect from extension).

    Returns:
        (success, document_object, error_message, temp_dir)
    """
    temp_dir = tempfile.mkdtemp(prefix='wps_verify_')
    ext = Path(container_path).suffix.lower()
    host_file = os.path.join(temp_dir, f"result{ext}")

    try:
        copy_from_env_fn(container_path, host_file)
    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return False, None, f"Failed to copy file: {e}", ''

    if not os.path.exists(host_file) or os.path.getsize(host_file) == 0:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return False, None, f"File not found or empty: {container_path}", ''

    try:
        if file_format == 'auto':
            file_format = 'docx' if ext in ('.docx', '.doc', '.wps') else 'odt'

        if file_format == 'docx' and DOCX_AVAILABLE:
            doc = Document(host_file)
            return True, doc, "", temp_dir
        elif file_format == 'odt' and ODF_AVAILABLE:
            doc = opendocument.load(host_file)
            return True, doc, "", temp_dir
        else:
            return False, None, f"Unsupported format or missing library: {ext}", temp_dir
    except Exception as e:
        return False, None, f"Parse error: {e}", temp_dir


def cleanup_verification_temp(temp_dir: Optional[str]):
    """Clean up temporary verification files."""
    if temp_dir and os.path.exists(temp_dir):
        shutil.rmtree(temp_dir, ignore_errors=True)


# ============================================================================
# Document text extraction
# ============================================================================

def get_document_text(doc) -> str:
    """Extract all text from a DOCX document."""
    return '\n'.join(para.text for para in doc.paragraphs)


def get_paragraph_styles(doc) -> List[Dict[str, Any]]:
    """Get list of paragraphs with their style and run info."""
    results = []
    for para in doc.paragraphs:
        results.append({
            'text': para.text,
            'style_name': para.style.name if para.style else None,
            'alignment': para.alignment,
            'runs': [{
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font else None,
                'font_size': run.font.size.pt if run.font and run.font.size else None,
            } for run in para.runs]
        })
    return results


# ============================================================================
# Heading and style checks
# ============================================================================

def check_heading_styles(doc, expected_headings: Dict[str, str]) -> Tuple[int, int, List[str]]:
    """
    Check that text content has correct heading styles applied.

    Args:
        doc: python-docx Document object
        expected_headings: {'text_fragment': 'Heading 1', ...}

    Returns:
        (matched_count, total_expected, feedback_list)
    """
    matched = 0
    total = len(expected_headings)
    feedback = []

    for text_fragment, expected_style in expected_headings.items():
        found = False
        for para in doc.paragraphs:
            if text_fragment.lower() in para.text.lower():
                actual_style = para.style.name if para.style else 'None'
                if expected_style.lower() in actual_style.lower():
                    matched += 1
                    found = True
                    feedback.append(f"OK: '{text_fragment[:40]}' has style '{actual_style}'")
                else:
                    feedback.append(
                        f"WRONG: '{text_fragment[:40]}' has '{actual_style}' "
                        f"(expected '{expected_style}')"
                    )
                    found = True
                break
        if not found:
            feedback.append(f"MISSING: '{text_fragment[:40]}' not found in document")

    return matched, total, feedback


def count_headings_by_level(doc) -> Dict[str, int]:
    """Count headings at each level in the document."""
    counts = {}
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            level = para.style.name
            counts[level] = counts.get(level, 0) + 1
    return counts


# ============================================================================
# Text formatting checks
# ============================================================================

def _resolve_run_bold(run, para) -> Optional[bool]:
    """Resolve bold formatting considering style inheritance."""
    if run.bold is not None:
        return run.bold
    # Check paragraph style's font bold setting
    if para.style and para.style.font and para.style.font.bold is not None:
        return para.style.font.bold
    # Check if paragraph style name implies bold (Heading, Title styles)
    if para.style and para.style.name:
        style_name = para.style.name.lower()
        if any(kw in style_name for kw in ('heading', 'title')):
            return True
    return None


def _resolve_run_italic(run, para) -> Optional[bool]:
    """Resolve italic formatting considering style inheritance."""
    if run.italic is not None:
        return run.italic
    # Check paragraph style's font italic setting
    if para.style and para.style.font and para.style.font.italic is not None:
        return para.style.font.italic
    # Check if paragraph style name implies italic (Subtitle styles)
    if para.style and para.style.name:
        style_name = para.style.name.lower()
        if 'subtitle' in style_name:
            return True
    return None


def check_text_formatting(doc, text_fragment: str, bold=None, italic=None,
                          underline=None, font_size=None) -> bool:
    """
    Check if text in the document has specific formatting.
    Handles style-inherited formatting and run fragmentation.

    Args:
        doc: python-docx Document object
        text_fragment: Text to search for (case-insensitive)
        bold: Expected bold state (True/False/None to skip)
        italic: Expected italic state (True/False/None to skip)
        underline: Expected underline state (True/False/None to skip)
        font_size: Expected font size in Pt (None to skip)

    Returns:
        True if matching text with all specified formatting is found.
    """
    for para in doc.paragraphs:
        # First try: match within individual runs
        for run in para.runs:
            if text_fragment.lower() in run.text.lower():
                if bold is not None:
                    resolved_bold = _resolve_run_bold(run, para)
                    if resolved_bold != bold:
                        continue
                if italic is not None:
                    resolved_italic = _resolve_run_italic(run, para)
                    if resolved_italic != italic:
                        continue
                if underline is not None and run.underline != underline:
                    continue
                if font_size is not None:
                    if run.font.size is None:
                        continue
                    if abs(run.font.size.pt - font_size) > 1.0:
                        continue
                return True

        # Second try: text may be split across multiple runs
        if text_fragment.lower() in para.text.lower() and para.runs:
            all_match = True
            if bold is not None:
                has_bold = any(
                    _resolve_run_bold(r, para) == bold
                    for r in para.runs if r.text.strip()
                )
                if not has_bold:
                    all_match = False
            if italic is not None:
                has_italic = any(
                    _resolve_run_italic(r, para) == italic
                    for r in para.runs if r.text.strip()
                )
                if not has_italic:
                    all_match = False
            if font_size is not None:
                has_size = any(
                    r.font.size is not None and abs(r.font.size.pt - font_size) <= 1.0
                    for r in para.runs if r.text.strip()
                )
                if not has_size:
                    all_match = False
            if all_match:
                return True
    return False


def check_paragraph_alignment(doc, text_fragment: str, alignment: str) -> bool:
    """
    Check if paragraph containing text has specified alignment.

    Args:
        doc: python-docx Document object
        text_fragment: Text to search for in paragraph
        alignment: 'left', 'center', 'right', or 'justify'

    Returns:
        True if paragraph with matching text has the specified alignment.
    """
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT if DOCX_AVAILABLE else 0,
        'center': WD_ALIGN_PARAGRAPH.CENTER if DOCX_AVAILABLE else 1,
        'right': WD_ALIGN_PARAGRAPH.RIGHT if DOCX_AVAILABLE else 2,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY if DOCX_AVAILABLE else 3,
    }
    target = alignment_map.get(alignment.lower())

    for para in doc.paragraphs:
        if text_fragment.lower() in para.text.lower():
            if para.alignment == target:
                return True
    return False


# ============================================================================
# Indentation checks
# ============================================================================

def check_hanging_indent(para) -> bool:
    """
    Check if a paragraph has hanging indent formatting.

    Args:
        para: python-docx Paragraph object

    Returns:
        True if paragraph has hanging indent.
    """
    try:
        pf = para.paragraph_format
        left_indent = pf.left_indent
        first_line_indent = pf.first_line_indent

        if left_indent is None:
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            ind = para._element.find('.//w:pPr/w:ind', nsmap)
            if ind is not None:
                hanging = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging')
                if hanging and int(hanging) > 0:
                    return True
            return False

        left_inches = left_indent / 914400

        if first_line_indent is None:
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            ind = para._element.find('.//w:pPr/w:ind', nsmap)
            if ind is not None:
                hanging = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging')
                if hanging and int(hanging) > 0:
                    return True
            return False

        first_inches = first_line_indent / 914400

        return (left_inches >= 0.3) and (first_inches < 0)
    except Exception:
        return False


# ============================================================================
# Table detection and verification
# ============================================================================

def count_tables(doc) -> int:
    """Count the number of tables in the document."""
    return len(doc.tables)


def get_table_dimensions(doc, table_index: int = 0) -> Tuple[int, int]:
    """
    Get dimensions of a specific table.

    Returns:
        (rows, columns) or (0, 0) if table doesn't exist.
    """
    if table_index >= len(doc.tables):
        return (0, 0)
    table = doc.tables[table_index]
    rows = len(table.rows)
    cols = len(table.columns) if table.rows else 0
    return (rows, cols)


def get_table_content(doc, table_index: int = 0) -> List[List[str]]:
    """
    Extract content from a table as a 2D list.

    Returns:
        List of rows, each row is a list of cell texts.
    """
    if table_index >= len(doc.tables):
        return []
    table = doc.tables[table_index]
    content = []
    for row in table.rows:
        row_content = [cell.text for cell in row.cells]
        content.append(row_content)
    return content


def check_table_header_formatting(doc, table_index: int = 0) -> Dict[str, Any]:
    """
    Check formatting of the header row (first row) of a table.

    Returns:
        Dict with: has_bold, has_shading, shading_color, cell_count
    """
    result = {
        'has_bold': False,
        'has_shading': False,
        'shading_color': None,
        'cell_count': 0
    }

    if table_index >= len(doc.tables):
        return result

    table = doc.tables[table_index]
    if not table.rows:
        return result

    header_row = table.rows[0]
    result['cell_count'] = len(header_row.cells)

    # Check each cell in header row
    for cell in header_row.cells:
        # Check for bold text in cell
        for para in cell.paragraphs:
            for run in para.runs:
                if run.bold or (run.font and run.font.bold):
                    result['has_bold'] = True
                # Also check if the whole paragraph style implies bold
                if para.style and para.style.name and 'heading' in para.style.name.lower():
                    result['has_bold'] = True

        # Check for cell shading
        try:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shd is not None:
                fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill and fill.upper() not in ('AUTO', 'FFFFFF', ''):
                    result['has_shading'] = True
                    result['shading_color'] = fill
        except Exception:
            pass

    return result


def check_table_cell_alignment(doc, table_index: int = 0) -> Dict[str, Any]:
    """
    Check alignment of cells in a table.

    Returns:
        Dict with alignment info per column.
    """
    result = {
        'columns': [],
        'has_right_aligned_numbers': False
    }

    if table_index >= len(doc.tables):
        return result

    table = doc.tables[table_index]
    if not table.rows or len(table.rows) < 2:
        return result

    # Get column count from first row
    col_count = len(table.rows[0].cells)
    column_alignments = [[] for _ in range(col_count)]

    # Check alignment of data rows (skip header)
    for row_idx, row in enumerate(table.rows[1:], 1):
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= col_count:
                continue
            for para in cell.paragraphs:
                if para.alignment:
                    column_alignments[col_idx].append(para.alignment)

    # Analyze column alignments
    alignment_names = {
        WD_ALIGN_PARAGRAPH.LEFT if DOCX_AVAILABLE else 0: 'left',
        WD_ALIGN_PARAGRAPH.CENTER if DOCX_AVAILABLE else 1: 'center',
        WD_ALIGN_PARAGRAPH.RIGHT if DOCX_AVAILABLE else 2: 'right',
        WD_ALIGN_PARAGRAPH.JUSTIFY if DOCX_AVAILABLE else 3: 'justify',
    }

    for col_idx, alignments in enumerate(column_alignments):
        if not alignments:
            result['columns'].append('unknown')
            continue
        # Get most common alignment
        from collections import Counter
        most_common = Counter(alignments).most_common(1)
        if most_common:
            align_val = most_common[0][0]
            result['columns'].append(alignment_names.get(align_val, 'unknown'))
            if align_val == (WD_ALIGN_PARAGRAPH.RIGHT if DOCX_AVAILABLE else 2):
                result['has_right_aligned_numbers'] = True
        else:
            result['columns'].append('unknown')

    return result


def check_table_alternating_colors(doc, table_index: int = 0) -> Dict[str, Any]:
    """
    Check if table has alternating row colors/shading.

    Returns:
        Dict with: has_alternating, row_shadings (list of colors)
    """
    result = {
        'has_alternating': False,
        'row_shadings': []
    }

    if table_index >= len(doc.tables):
        return result

    table = doc.tables[table_index]
    if len(table.rows) < 2:
        return result

    # Get shading for each row
    for row_idx, row in enumerate(table.rows):
        row_shading = None
        for cell in row.cells:
            try:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill and fill.upper() not in ('AUTO', ''):
                        row_shading = fill.upper()
                        break
            except Exception:
                pass
        result['row_shadings'].append(row_shading)

    # Check for alternating pattern (at least 2 different shadings)
    unique_shadings = set(s for s in result['row_shadings'] if s)
    if len(unique_shadings) >= 2:
        result['has_alternating'] = True

    return result


# ============================================================================
# Citation / bibliography verification
# ============================================================================

def extract_citation_paragraphs(doc, start_after: str = "References") -> List:
    """
    Extract citation paragraphs from the document (after the References heading).

    Returns:
        List of paragraph objects that are citations.
    """
    citations = []
    in_references = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if not in_references:
            if start_after.lower() in text.lower():
                in_references = True
            continue

        if not text or len(text) < 20:
            continue

        if text.startswith('-') or text.startswith('TASK') or text.startswith('MISSING'):
            continue

        citations.append(para)

    return citations


# ============================================================================
# Image/figure verification
# ============================================================================

def count_images(doc) -> int:
    """Count the number of inline images in the document."""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            xml = run._element.xml
            if 'w:drawing' in xml or 'w:pict' in xml:
                count += 1
    return count


# ============================================================================
# VLM cross-validation helper
# ============================================================================

def vlm_verify_screenshot(env_info, traj, prompt: str) -> Optional[Dict]:
    """
    Query the VLM with the final screenshot and a structured prompt.
    Returns parsed JSON from VLM response, or None if VLM is unavailable.

    Args:
        env_info: Environment info dict (contains query_vlm, get_final_screenshot).
        traj: Trajectory dict.
        prompt: Structured prompt asking VLM to evaluate the screenshot.

    Returns:
        Parsed dict from VLM response, or None.
    """
    query_vlm = env_info.get('query_vlm')
    get_final_screenshot = env_info.get('get_final_screenshot')

    if not query_vlm or not get_final_screenshot:
        return None

    final_frame = get_final_screenshot(traj)
    if not final_frame:
        return None

    try:
        result = query_vlm(prompt=prompt, image=final_frame)
        if result and result.get("success"):
            return result.get("parsed", {})
        logger.warning(f"VLM query failed: {result.get('error', 'unknown') if result else 'no result'}")
    except Exception as e:
        logger.warning(f"VLM query exception: {e}")

    return None


# ============================================================================
# Word count and statistics
# ============================================================================

def get_word_count(doc) -> int:
    """Get total word count of the document."""
    text = get_document_text(doc)
    return len(text.split())


def get_paragraph_count(doc) -> int:
    """Get total number of non-empty paragraphs."""
    return sum(1 for para in doc.paragraphs if para.text.strip())


def get_page_count_estimate(doc) -> int:
    """
    Estimate page count based on content.
    Note: This is a rough estimate as actual page count depends on formatting.
    """
    # Rough estimate: ~500 words per page
    words = get_word_count(doc)
    return max(1, words // 500)
