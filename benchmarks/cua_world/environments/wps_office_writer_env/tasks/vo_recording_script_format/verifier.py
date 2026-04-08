#!/usr/bin/env python3
"""
Verifier for vo_recording_script_format task.

SCORING CRITERIA:
1. Landscape Orientation (15 points)
2. Text-to-Table Conversion (25 points)
3. Header Row Created correctly (15 points)
4. Column Typography: Character=Bold, Direction=Italic (25 points)
5. Document Headers/Footers present (20 points)
"""

import json
import os
import tempfile
import logging
from docx import Document
from docx.enum.section import WD_ORIENT

# Import VLM utilities from framework
try:
    from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
except ImportError:
    logger = logging.getLogger(__name__)
    logger.warning("VLM utilities not found, visual verification will be bypassed.")
    def sample_trajectory_frames(*args, **kwargs): return []
    def get_final_screenshot(*args, **kwargs): return None
    def query_vlm(*args, **kwargs): return {"success": False}

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_vo_recording_script(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    
    score = 0
    feedback_parts = []
    
    # 1. Fetch JSON result to check timestamps/creation
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result JSON: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Anti-gaming: Ensure file was made and modified during task
    if not result.get("output_exists") or not result.get("file_modified_during_task"):
        return {
            "passed": False, 
            "score": 0, 
            "feedback": "FAIL: Correctly named output file was not created or modified during the task."
        }
        
    # 2. Fetch the output DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/NOTLD_VO_Script.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # --- CRITERION 1: Landscape Orientation (15 pts) ---
    landscape = False
    if doc.sections:
        sec = doc.sections[0]
        if sec.orientation == WD_ORIENT.LANDSCAPE or sec.page_width > sec.page_height:
            landscape = True
            
    if landscape:
        score += 15
        feedback_parts.append("Orientation: Landscape (15/15)")
    else:
        feedback_parts.append("Orientation: Not Landscape (0/15)")

    # --- CRITERION 2: Table Conversion (25 pts) ---
    has_table = len(doc.tables) > 0
    correct_columns = False
    if has_table:
        table = doc.tables[0]
        if len(table.columns) >= 4:
            correct_columns = True
            score += 25
            feedback_parts.append("Table: 4-column table created (25/25)")
        else:
            score += 10
            feedback_parts.append(f"Table: Created but wrong columns ({len(table.columns)} instead of 4) (10/25)")
    else:
        feedback_parts.append("Table: No table found (0/25)")

    # --- CRITERION 3: Header Row (15 pts) ---
    header_correct = False
    if correct_columns:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        expected = metadata.get("expected_headers", [])
        
        # Check if "character" and "dialogue" are in the header row
        if "character" in headers and "dialogue" in headers:
            score += 15
            header_correct = True
            feedback_parts.append("Header Row: Correct headers found (15/15)")
        else:
            feedback_parts.append(f"Header Row: Incorrect titles {headers} (0/15)")

    # --- CRITERION 4: Column Typography (25 pts) ---
    char_bold = False
    dir_italic = False
    
    if correct_columns and len(table.rows) > 1:
        bold_count = 0
        italic_count = 0
        rows_to_check = min(5, len(table.rows) - 1) # Check up to 5 rows
        
        for i in range(1, rows_to_check + 1):
            char_cell = table.rows[i].cells[1]
            dir_cell = table.rows[i].cells[2]
            
            # Check for bold runs in Character cell
            if any(r.bold for p in char_cell.paragraphs for r in p.runs if r.text.strip()):
                bold_count += 1
                
            # Check for italic runs in Voice Direction cell
            if any(r.italic for p in dir_cell.paragraphs for r in p.runs if r.text.strip()):
                italic_count += 1
                
        if bold_count >= rows_to_check - 1:
            char_bold = True
        if italic_count >= rows_to_check - 1:
            dir_italic = True
            
    if char_bold and dir_italic:
        score += 25
        feedback_parts.append("Typography: Character Bold & Direction Italic applied (25/25)")
    elif char_bold or dir_italic:
        score += 12
        feedback_parts.append(f"Typography: Partial (Bold:{char_bold}, Italic:{dir_italic}) (12/25)")
    else:
        feedback_parts.append("Typography: Missing required styling (0/25)")

    # --- CRITERION 5: Document Headers/Footers (20 pts) ---
    header_found = False
    expected_header = metadata.get("expected_document_header", "Project: NOTLD").lower()
    
    if doc.sections:
        for sec in doc.sections:
            if sec.header:
                for p in sec.header.paragraphs:
                    if expected_header in p.text.lower() or "notld" in p.text.lower():
                        header_found = True
                        break
                        
    if header_found:
        score += 20
        feedback_parts.append("Doc Header: 'NOTLD' header present (20/20)")
    else:
        feedback_parts.append("Doc Header: Missing required header text (0/20)")

    # --- VLM Visual Verification (Safeguard) ---
    frames = sample_trajectory_frames(traj, n=3)
    final = get_final_screenshot(traj)
    if final:
        images_to_check = frames + [final]
        prompt = (
            "You are grading a WPS Office task. Did the agent successfully select the text, "
            "convert it into a 4-column table, and change the page to landscape mode? "
            "Reply strictly with a JSON: {'table_visible': true/false, 'is_landscape': true/false}"
        )
        
        vlm_res = query_vlm(images=images_to_check, prompt=prompt)
        if vlm_res.get("success"):
            parsed = vlm_res.get("parsed", {})
            if not parsed.get("table_visible"):
                feedback_parts.append("VLM Penalty: Table not visually confirmed.")
                score = max(0, score - 20)
                
    # Final pass logic
    passed = score >= 75 and has_table and landscape
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }