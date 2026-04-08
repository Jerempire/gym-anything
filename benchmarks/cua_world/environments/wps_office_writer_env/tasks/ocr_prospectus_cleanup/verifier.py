#!/usr/bin/env python3
import json
import os
import tempfile
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure python-docx is available in the verifier environment
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def verify_ocr_prospectus_cleanup(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Copy the result JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result JSON: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not result.get('document_exists'):
        return {
            "passed": False,
            "score": 0,
            "feedback": "Final document not found at expected path: ~/Documents/project_pegasus_final.docx"
        }

    # Copy the finalized DOCX from the container
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/project_pegasus_final.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse final document: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 0
    feedback_parts = []
    
    # 1. Text Cleanup: merged paragraphs (30 points)
    # The agent should have merged ~15 short lines into 3 large paragraphs.
    # We define a "long paragraph" as one having > 150 characters.
    long_paras = [p for p in doc.paragraphs if p.text and len(p.text.strip()) > 150]
    
    if len(long_paras) >= 3:
        score += 30
        feedback_parts.append("Text cleanup: hard returns removed (30/30)")
    elif len(long_paras) > 0:
        score += 15
        feedback_parts.append(f"Text cleanup: partially merged ({len(long_paras)} long paras) (15/30)")
    else:
        feedback_parts.append("Text cleanup: hard returns NOT removed (0/30)")

    # 2. Paragraph Formatting: Justified + Indent (10 points)
    justified_count = 0
    indented_count = 0
    for p in long_paras:
        # 3 is the integer value for WD_ALIGN_PARAGRAPH.JUSTIFY in python-docx
        if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY or p.alignment == 3:
            justified_count += 1
        
        if p.paragraph_format.first_line_indent and p.paragraph_format.first_line_indent > 0:
            indented_count += 1
            
    format_score = 0
    if len(long_paras) > 0:
        if justified_count == len(long_paras):
            format_score += 5
            feedback_parts.append("Paragraphs justified (5/5)")
        elif justified_count > 0:
            format_score += 2
            feedback_parts.append("Paragraphs partially justified (2/5)")
        else:
            feedback_parts.append("Paragraphs NOT justified (0/5)")
            
        if indented_count == len(long_paras):
            format_score += 5
            feedback_parts.append("First-line indent applied (5/5)")
        elif indented_count > 0:
            format_score += 2
            feedback_parts.append("First-line indent partially applied (2/5)")
        else:
            feedback_parts.append("First-line indent NOT applied (0/5)")
            
    score += format_score

    # 3. Heading Styles (10 points)
    h1_found = False
    h2_found = 0
    for p in doc.paragraphs:
        style_name = p.style.name.lower() if p.style else ""
        text_lower = p.text.lower()
        if "prospectus summary" in text_lower and "heading 1" in style_name:
            h1_found = True
        if "heading 2" in style_name:
            if "the company" in text_lower or "risk factors" in text_lower or "selected financial data" in text_lower:
                h2_found += 1
                
    heading_score = 0
    if h1_found:
        heading_score += 4
    if h2_found >= 3:
        heading_score += 6
    elif h2_found > 0:
        heading_score += 2 * h2_found
        
    score += heading_score
    feedback_parts.append(f"Headings: {heading_score}/10")

    # 4. Table Reconstruction (20 points)
    table_score = 0
    if len(doc.tables) >= 1:
        table = doc.tables[-1] # Usually the last table added
        has_data = False
        text_content = ""
        for row in table.rows:
            for cell in row.cells:
                text_content += cell.text.lower()
        
        if "147,787" in text_content or "gross profit" in text_content:
            has_data = True
            
        if has_data:
            if len(table.columns) == 4 or len(table.columns) == 5:
                table_score = 20
                feedback_parts.append("Table reconstructed correctly (20/20)")
            else:
                table_score = 10
                feedback_parts.append("Table reconstructed but wrong column count (10/20)")
        else:
            table_score = 5
            feedback_parts.append("Table exists but missing expected financial data (5/20)")
    else:
        feedback_parts.append("Table NOT reconstructed (0/20)")
        
    score += table_score

    # 5. Table Alignment (10 points)
    align_score = 0
    if table_score >= 10 and len(doc.tables) >= 1:
        table = doc.tables[-1]
        right_aligned = 0
        total_data_cells = 0
        
        # Check rows starting from index 1 (skip header)
        for row in table.rows[1:]:
            for i, cell in enumerate(row.cells):
                if i > 0 and cell.text.strip():  # Skip first column (labels)
                    total_data_cells += 1
                    is_right = False
                    for cp in cell.paragraphs:
                        if cp.alignment == WD_ALIGN_PARAGRAPH.RIGHT or cp.alignment == 2:
                            is_right = True
                    if is_right:
                        right_aligned += 1
                        
        if total_data_cells > 0:
            ratio = right_aligned / total_data_cells
            if ratio > 0.8:
                align_score = 10
                feedback_parts.append("Table data right-aligned (10/10)")
            elif ratio > 0.4:
                align_score = 5
                feedback_parts.append("Table data partially right-aligned (5/10)")
            else:
                feedback_parts.append("Table data NOT right-aligned (0/10)")
                
    score += align_score

    # 6. Document Header (10 points)
    header_found = False
    header_right_aligned = False
    
    for section in doc.sections:
        for hp in section.header.paragraphs:
            if "project pegasus" in hp.text.lower():
                header_found = True
                if hp.alignment == WD_ALIGN_PARAGRAPH.RIGHT or hp.alignment == 2:
                    header_right_aligned = True
                    
    header_score = 0
    if header_found:
        if header_right_aligned:
            header_score = 10
            feedback_parts.append("Header present and right-aligned (10/10)")
        else:
            header_score = 5
            feedback_parts.append("Header present but not right-aligned (5/10)")
    else:
        feedback_parts.append("Header NOT found (0/10)")
        
    score += header_score

    # 7. VLM Watermark Verification (10 points)
    watermark_xml_found = False
    try:
        # Fallback check inside the document XML
        for section in doc.sections:
            for hp in section.header.paragraphs:
                if 'CONFIDENTIAL' in hp._p.xml:
                    watermark_xml_found = True
    except:
        pass

    vlm_score = 0
    vlm_used = False
    try:
        from gym_anything.vlm import query_vlm, sample_trajectory_frames, get_final_screenshot
        
        frames = sample_trajectory_frames(traj, n=3)
        final = get_final_screenshot(traj)
        images = frames + [final] if final else frames
        
        if images:
            prompt = (
                "You are an automated grader evaluating a task performed in WPS Writer.\n"
                "Look at these screenshots. Does the document contain a 'CONFIDENTIAL' watermark?\n"
                "A watermark typically appears as large, light gray, diagonal text behind the main document text.\n"
                "Reply in JSON format: {\"watermark_present\": true/false}"
            )
            vlm_res = query_vlm(images=images, prompt=prompt)
            if vlm_res and vlm_res.get("parsed", {}).get("watermark_present", False):
                vlm_score = 10
                feedback_parts.append("Watermark visually confirmed (10/10)")
            else:
                feedback_parts.append("Watermark NOT confirmed by VLM (0/10)")
            vlm_used = True
    except ImportError:
        pass
        
    if not vlm_used:
        if watermark_xml_found:
            vlm_score = 10
            feedback_parts.append("Watermark found in document XML (10/10)")
        else:
            feedback_parts.append("Watermark NOT found in XML, VLM unavailable (0/10)")

    score += vlm_score

    # Key criteria: Must clean text AND create table to pass
    key_criteria_met = (len(long_paras) >= 1) and (len(doc.tables) >= 1)
    passed = (score >= 70) and key_criteria_met

    if not key_criteria_met:
        feedback_parts.append("FAILED: Key criteria not met (Must complete Text Cleanup and Table Reconstruction)")

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }