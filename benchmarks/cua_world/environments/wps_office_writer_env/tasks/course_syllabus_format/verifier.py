#!/usr/bin/env python3
"""
Verifier for course_syllabus_format task.
Evaluates document styling, table creation, font enforcement, and content addition.
"""

import json
import os
import tempfile
import logging

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Fallback VLM utilities if available
try:
    from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
    VLM_AVAILABLE = True
except ImportError:
    VLM_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_course_syllabus_format(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Framework error: copy_from_env unavailable"}

    score = 0
    max_score = 100
    feedback_parts = []

    # 1. Fetch metadata JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_meta = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    output_exists = result_meta.get('output_exists', False)
    file_created = result_meta.get('file_created_during_task', False)
    modified = result_meta.get('modified_from_draft', False)

    if not output_exists:
        return {
            "passed": False,
            "score": 0,
            "feedback": "Final syllabus file was not saved to /home/ga/Documents/ENVS4350_Syllabus_Final.docx"
        }

    # Base points for creating the file during the task
    if file_created and modified:
        score += 10
        feedback_parts.append("File created and modified successfully (+10)")
    else:
        feedback_parts.append("Warning: File does not appear to be newly modified")

    # 2. Fetch the actual DOCX for content inspection
    if not DOCX_AVAILABLE:
        # If python-docx isn't on the host, we have to fail or skip programmatic
        logger.warning("python-docx not available on host. Programmatic checks skipped.")
        feedback_parts.append("Cannot parse DOCX (missing python-docx).")
        return {"passed": False, "score": score, "feedback": " | ".join(feedback_parts)}

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/ENVS4350_Syllabus_Final_HostCheck.docx", temp_docx.name)
        doc = docx.Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to read docx: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # Helper to get full text
    full_text_lower = "\n".join([p.text.lower() for p in doc.paragraphs])
    for t in doc.tables:
        for r in t.rows:
            full_text_lower += " " + " ".join([c.text.lower() for c in r.cells])

    # 3. Check Headings (H1 and H2)
    h1_found = False
    h2_sections_found = 0
    expected_h2s = task_info.get('metadata', {}).get('expected_h2_sections', [
        "Course Information", "Course Description", "Learning Objectives",
        "Course Schedule", "Grading Policy", "Required Materials", "Course Policies"
    ])

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip().lower()
        if not text:
            continue

        if "envs 4350" in text and "heading 1" in style_name.lower():
            h1_found = True
        
        if "heading 2" in style_name.lower():
            for expected in expected_h2s:
                if expected.lower() in text:
                    h2_sections_found += 1

    if h1_found:
        score += 5
        feedback_parts.append("Title styled with Heading 1 (+5)")
    else:
        feedback_parts.append("Title missing Heading 1")

    # Proportional score for H2 (max 10)
    h2_score = int((min(h2_sections_found, len(expected_h2s)) / len(expected_h2s)) * 10)
    score += h2_score
    feedback_parts.append(f"H2 Sections applied: {h2_sections_found}/{len(expected_h2s)} (+{h2_score})")

    # 4. Check Tables (Schedule and Grading)
    schedule_table_ok = False
    grading_table_ok = False

    for table in doc.tables:
        # Check rows/cols and content to identify
        num_rows = len(table.rows)
        num_cols = len(table.columns) if table.rows else 0
        
        table_text = " ".join([c.text.lower() for r in table.rows for c in r.cells])
        
        # Identify schedule table
        if "week 1" in table_text and "ipcc" in table_text:
            if num_rows >= 14 and num_cols >= 3:
                schedule_table_ok = True
                
        # Identify grading table
        if "participation" in table_text and "midterm exam" in table_text:
            if num_rows >= 5 and num_cols >= 2:
                grading_table_ok = True

    if schedule_table_ok:
        score += 15
        feedback_parts.append("Schedule table successfully created (+15)")
    else:
        feedback_parts.append("Schedule table missing/incorrect")

    if grading_table_ok:
        score += 15
        feedback_parts.append("Grading table successfully created (+15)")
    else:
        feedback_parts.append("Grading table missing/incorrect")

    # 5. Body Text Font Check (Times New Roman 12pt)
    # Sample a few non-heading paragraphs that have decent length
    body_paras = [p for p in doc.paragraphs if len(p.text) > 50 and "heading" not in (p.style.name.lower() if p.style else "")]
    tnr_count = 0
    size_12_count = 0
    
    for p in body_paras[:10]:
        # Check paragraph style font or run font
        font_name = None
        font_size = None
        for run in p.runs:
            if run.text.strip():
                if run.font.name: font_name = run.font.name.lower()
                if run.font.size: font_size = run.font.size.pt
                break
        if not font_name and p.style and p.style.font and p.style.font.name:
            font_name = p.style.font.name.lower()
        if not font_size and p.style and p.style.font and p.style.font.size:
            font_size = p.style.font.size.pt
            
        if font_name and 'times' in font_name:
            tnr_count += 1
        if font_size and abs(font_size - 12.0) < 0.5:
            size_12_count += 1

    if body_paras:
        if tnr_count >= len(body_paras[:10]) * 0.7:
            score += 5
            feedback_parts.append("Body font is Times New Roman (+5)")
        if size_12_count >= len(body_paras[:10]) * 0.7:
            score += 5
            feedback_parts.append("Body font is 12pt (+5)")

    # 6. Policy Additions & H3 Checks
    ai_found = "academic integrity" in full_text_lower or "plagiarism" in full_text_lower
    da_found = "disability" in full_text_lower or "accommodation" in full_text_lower
    h3_count = sum(1 for p in doc.paragraphs if p.style and "heading 3" in p.style.name.lower())

    if ai_found and da_found:
        score += 10
        feedback_parts.append("Both required policies added (+10)")
    elif ai_found or da_found:
        score += 5
        feedback_parts.append("One required policy added (+5)")
    else:
        feedback_parts.append("Missing required policy additions")

    if h3_count >= 2:
        score += 5
        feedback_parts.append("H3 Subsections applied correctly (+5)")

    # 7. Content Preservation Check
    preservation_keywords = ["ipcc ar6", "friedlingstein", "sarah mitchell"]
    preserved = sum(1 for kw in preservation_keywords if kw in full_text_lower)
    if preserved == 3:
        score += 10
        feedback_parts.append("Original content properly preserved (+10)")
    else:
        feedback_parts.append("Warning: Original content may have been deleted!")

    # 8. Trajectory VLM Check (Anti-gaming)
    if VLM_AVAILABLE and traj:
        try:
            frames = sample_trajectory_frames(traj, n=3)
            final_frame = get_final_screenshot(traj)
            if frames and final_frame:
                prompt = (
                    "You are verifying if a computer agent successfully formatted a syllabus document in WPS Writer. "
                    "Look at these screenshots taken during the task. "
                    "Did the agent actively interact with the WPS Writer interface? "
                    "For example, selecting text, applying styles/fonts from the ribbon, inserting tables, or typing new text. "
                    "Respond with a JSON containing a boolean field 'gui_interaction_observed'."
                )
                vlm_resp = query_vlm(images=frames + [final_frame], prompt=prompt)
                if vlm_resp.get('success') and vlm_resp.get('parsed', {}).get('gui_interaction_observed', False):
                    score += 10
                    feedback_parts.append("VLM confirmed GUI interaction (+10)")
                else:
                    feedback_parts.append("VLM did NOT observe meaningful GUI interaction (scripting suspected)")
        except Exception as e:
            logger.warning(f"VLM verification failed: {e}")
            # Do not heavily penalize if VLM simply errors out, but log it
            feedback_parts.append("VLM check skipped/failed.")

    # Determine Pass/Fail
    # To pass: must have 60 points, the output must exist, and core tables must be mostly there
    passed = (score >= 60) and output_exists and (schedule_table_ok or grading_table_ok)

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }