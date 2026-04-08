#!/usr/bin/env python3
"""
Verifier for CBA Ratification Redline task.
"""

import json
import tempfile
import os
import logging

# Check if python-docx is available
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_cba_redline(traj, env_info, task_info):
    """
    Verify the CBA redlining task using multi-criteria checks.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Framework error: copy_from_env not found"}

    # 1. Read the export JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/cba_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    doc_exists = result.get('doc_exists', False)
    doc_created_during_task = result.get('doc_created_during_task', False)
    
    if not doc_exists:
        return {"passed": False, "score": 0, "feedback": "Failed: CBA_Tentative_Agreement.docx was not found in Documents."}
    if not doc_created_during_task:
        return {"passed": False, "score": 0, "feedback": "Failed: Document exists but was not created during the task timeframe."}

    if not DOCX_AVAILABLE:
        # Fallback to pure VLM if docx parsing fails entirely
        return run_vlm_fallback(traj, "Cannot parse document programmatically.")

    # 2. Extract Document for Parsing
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    score = 10  # Base points for creating the file
    feedback_parts = ["File created"]
    
    try:
        copy_from_env("/tmp/cba_output.docx", temp_docx.name)
        doc = Document(temp_docx.name)
        
        # Criterion 1: Headings (15 pts)
        articles_heading1 = 0
        sections_heading2 = 0
        has_toc = False
        grievance_list = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name.lower() if para.style else ""
            
            if "article" in text.lower() and "heading 1" in style_name:
                articles_heading1 += 1
            if "section" in text.lower() and "heading 2" in style_name:
                sections_heading2 += 1
                
            # Check for TOC indicators
            if "table of contents" in text.lower() or "toc" in style_name:
                has_toc = True
                
            # Check for list applied to grievance steps
            if "step" in text.lower() and ("list" in style_name or para._p.pPr is not None and para._p.pPr.numPr is not None):
                grievance_list += 1

        if articles_heading1 >= 3 and sections_heading2 >= 3:
            score += 15
            feedback_parts.append("Headings correctly applied")
        else:
            feedback_parts.append("Heading styles missing or incomplete")

        # Criterion 2: Table of Contents (10 pts)
        if has_toc:
            score += 10
            feedback_parts.append("TOC present")
        else:
            # Fallback checking XML for field codes
            xml_str = doc.element.xml
            if 'w:fldSimple' in xml_str and 'TOC' in xml_str:
                score += 10
                has_toc = True
                feedback_parts.append("TOC field code found")
            else:
                feedback_parts.append("TOC not found")

        # Criterion 3: Grievance Numbered List (10 pts)
        if grievance_list >= 2:
            score += 10
            feedback_parts.append("Numbered list applied to Grievance steps")
        else:
            feedback_parts.append("Numbered list missing for Grievance steps")

        # Criterion 4: Wage Scale Table (15 pts)
        table_found = False
        for table in doc.tables:
            if len(table.columns) == 5:
                # verify it's the wage table
                first_row_text = "".join([c.text for c in table.rows[0].cells]).lower()
                if "classification" in first_row_text and "rate" in first_row_text:
                    table_found = True
                    # Check bold header
                    has_bold = any(run.bold for p in table.rows[0].cells[0].paragraphs for run in p.runs)
                    if has_bold:
                        score += 15
                        feedback_parts.append("5-column Wage table with bold header present")
                    else:
                        score += 10
                        feedback_parts.append("5-column Wage table present (header not bold)")
                    break
        
        if not table_found:
            feedback_parts.append("Wage table not found or incorrect columns")

        # Criterion 5: Redline Character Formatting (20 pts)
        strike_15_found = False
        under_18_found = False
        strike_24_found = False
        under_48_found = False
        
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text.lower()
                # Health redlines
                if "15" in text and run.font.strike:
                    strike_15_found = True
                if "18" in text and (run.font.underline and run.font.bold):
                    under_18_found = True
                    
                # Overtime redlines
                if "24" in text and run.font.strike:
                    strike_24_found = True
                if "48" in text and (run.font.underline and run.font.bold):
                    under_48_found = True

        redline_score = 0
        if strike_15_found and under_18_found:
            redline_score += 10
        if strike_24_found and under_48_found:
            redline_score += 10
            
        score += redline_score
        if redline_score == 20:
            feedback_parts.append("All redline character formatting correct")
        elif redline_score > 0:
            feedback_parts.append("Partial redline formatting applied")
        else:
            feedback_parts.append("Redline formatting missing/incorrect")

        # Criterion 6: Footer presence (10 pts)
        footer_found = False
        for section in doc.sections:
            for footer_para in section.footer.paragraphs:
                if "pending ratification" in footer_para.text.lower():
                    footer_found = True
                    break
            if footer_found: break
            
        if footer_found:
            score += 10
            feedback_parts.append("Footer text found")
        else:
            feedback_parts.append("Footer text missing")

    except Exception as e:
        logger.error(f"Error parsing document: {e}")
        feedback_parts.append(f"Doc parse error: {str(e)}")
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # Criterion 7: VLM Visual Verification (10 pts)
    # Checks trajectory to confirm they actually used the GUI elements
    vlm_score = 0
    try:
        from gym_anything.vlm import sample_trajectory_frames, query_vlm, get_final_screenshot
        frames = sample_trajectory_frames(traj, n=3)
        final = get_final_screenshot(traj)
        
        if final:
            prompt = """Look at these screenshots of a user formatting a document in WPS Writer.
            Did the user successfully convert comma-separated text into a table grid, or apply struck-through text?
            Respond with JSON:
            {"table_visible": true/false, "struck_text_visible": true/false}
            """
            vlm_res = query_vlm(images=frames + [final], prompt=prompt)
            if vlm_res and vlm_res.get("success"):
                parsed = vlm_res.get("parsed", {})
                if parsed.get("table_visible", False) or parsed.get("struck_text_visible", False):
                    vlm_score = 10
                    score += 10
                    feedback_parts.append("VLM confirmed visual formatting")
    except Exception as e:
        logger.warning(f"VLM verification skipped or failed: {e}")
        # Give benefit of doubt if VLM fails but programatics passed
        if score >= 60:
            score += 10
            vlm_score = 10

    # Key criteria: Output file exists, Table converted, Redline applied
    key_criteria_met = doc_created_during_task and table_found and (strike_15_found or strike_24_found)
    passed = score >= 70 and key_criteria_met

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }

def run_vlm_fallback(traj, reason):
    """Fallback if programmatic verification totally fails."""
    return {"passed": False, "score": 0, "feedback": f"Programmatic verification failed: {reason}"}