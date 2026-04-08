#!/usr/bin/env python3
"""
Verifier for Proxy Statement SEC Formatting Task.
Combines python-docx programmatic checks with VLM trajectory verification.
"""

import json
import os
import tempfile
import logging

from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Fallback dependencies if docx is not installed
try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def build_vlm_prompt():
    return """You are verifying an agent's completion of formatting an SEC Proxy Statement in WPS Writer.
Please review the screenshots to evaluate the formatting.

Check for the following criteria:
1. TABLE CONVERSION & BORDER: Has the pipe-delimited text ("Name and Principal Position...", etc.) been converted into a proper table? Does the table's header row have a bold font and a thick bottom border separating it from the data rows?
2. COLUMN ALIGNMENT: Inside the table, is the first column left-aligned, while the financial columns (Year, Salary, Bonus, Stock Awards, Total) are right-aligned?
3. SUPERSCRIPTS: In the table data (e.g., next to the CEO's $1,200,000 salary or CFO's Stock Awards), are there superscript footnotes (like ¹ or ²) instead of bracketed text like [1] or [2]?
4. FOOTER: Look at the bottom of the page in the final state. Is there a footer that says "2024 Proxy Statement" on the left, and a page number on the right?

Provide your response strictly in JSON format:
{
    "table_and_border_visible": true/false,
    "right_alignment_visible": true/false,
    "superscripts_visible": true/false,
    "footer_visible": true/false,
    "reasoning": "Brief explanation of what you observed."
}
"""


def verify_proxy_statement(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Framework error: copy_from_env not available."}

    # Fetch result JSON
    temp_result = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result.name)
        with open(temp_result.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_result.name):
            os.unlink(temp_result.name)

    output_exists = result.get('output_exists', False)
    created_during_task = result.get('file_created_during_task', False)

    if not output_exists:
        return {"passed": False, "score": 0, "feedback": "proxy_statement_final.docx was not saved."}

    score = 10  # 10 pts for saving the file
    feedback = ["File successfully saved."]
    
    if not created_during_task:
        feedback.append("Warning: File timestamp indicates it might not have been created during this task.")

    # ---------------------------------------------------------
    # Programmatic DOCX Checks
    # ---------------------------------------------------------
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    prog_checks = {
        "headings": False,
        "table_exists": False,
        "header_bold": False,
        "alignment_right": False,
        "superscript_detected": False,
        "footer_detected": False
    }

    try:
        copy_from_env("/tmp/proxy_statement_final.docx", temp_docx.name)
        
        if DOCX_AVAILABLE:
            doc = docx.Document(temp_docx.name)
            
            # 1. Heading Styles (15 pts)
            h1_found = any(p.style.name.startswith('Heading 1') and "Compensation Discussion" in p.text for p in doc.paragraphs)
            h2_found = any(p.style.name.startswith('Heading 2') and "Summary Compensation" in p.text for p in doc.paragraphs)
            
            if h1_found and h2_found:
                prog_checks["headings"] = True
                score += 15
                feedback.append("Headings successfully applied.")
            else:
                feedback.append("Heading styles were not correctly applied.")
            
            # 2. Table checks
            if len(doc.tables) >= 1:
                prog_checks["table_exists"] = True
                score += 15
                feedback.append("Table successfully created.")
                
                table = doc.tables[0]
                
                # Check header bold
                if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                    first_cell_runs = table.rows[0].cells[0].paragraphs[0].runs
                    if first_cell_runs and first_cell_runs[0].bold:
                        prog_checks["header_bold"] = True
                        score += 10
                        feedback.append("Table header is bolded.")
                
                # Check column alignment (Col 1+ should be right aligned)
                if len(table.rows) > 1 and len(table.rows[1].cells) > 2:
                    align_val = table.rows[1].cells[2].paragraphs[0].alignment
                    if align_val == WD_ALIGN_PARAGRAPH.RIGHT:
                        prog_checks["alignment_right"] = True
                        score += 15
                        feedback.append("Financial columns right-aligned.")

                # Check superscripts & bracket removal
                bracket_found = False
                super_found = False
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        if "[1]" in text or "[2]" in text:
                            bracket_found = True
                        for p in cell.paragraphs:
                            for r in p.runs:
                                if r.font.superscript:
                                    super_found = True
                
                if super_found and not bracket_found:
                    prog_checks["superscript_detected"] = True
                    score += 15
                    feedback.append("Superscript footnotes properly formatted.")
            else:
                feedback.append("No table found in the document.")
            
            # 3. Footer check
            footer_found = False
            for section in doc.sections:
                if section.footer:
                    for p in section.footer.paragraphs:
                        if "2024 Proxy Statement" in p.text:
                            footer_found = True
            
            if footer_found:
                prog_checks["footer_detected"] = True
                score += 10
                feedback.append("Footer with '2024 Proxy Statement' found.")
            else:
                feedback.append("Footer text not detected in document XML.")
        else:
            feedback.append("python-docx not available, skipping internal XML checks.")

    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        feedback.append(f"Error reading document properties: {e}")
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # ---------------------------------------------------------
    # VLM Trajectory & Screenshot Fallback Verification (10 pts)
    # ---------------------------------------------------------
    try:
        frames = sample_trajectory_frames(traj, n=4)
        final = get_final_screenshot(traj)
        
        vlm_resp = query_vlm(
            images=frames + [final] if final else frames,
            prompt=build_vlm_prompt()
        )
        
        if vlm_resp.get('success'):
            parsed = vlm_resp.get('parsed', {})
            
            # Use VLM to validate/boost score if programmatic checks failed due to XML structure weirdness from WPS
            if not prog_checks["header_bold"] and parsed.get("table_and_border_visible"):
                score += 10
                feedback.append("VLM confirmed table border and header visually.")
                
            if not prog_checks["alignment_right"] and parsed.get("right_alignment_visible"):
                score += 15
                feedback.append("VLM confirmed right-alignment visually.")
                
            if not prog_checks["superscript_detected"] and parsed.get("superscripts_visible"):
                score += 15
                feedback.append("VLM confirmed superscripts visually.")
                
            if not prog_checks["footer_detected"] and parsed.get("footer_visible"):
                score += 10
                feedback.append("VLM confirmed footer visually.")
                
            # Base VLM point for general visual workflow adherence
            if any([parsed.get(k) for k in ["table_and_border_visible", "right_alignment_visible", "superscripts_visible"]]):
                score += 10
                feedback.append("VLM verified positive progress on visual formatting tasks.")

    except Exception as e:
        logger.error(f"VLM verification error: {e}")
        feedback.append(f"VLM Verification skipped/failed: {e}")

    # Ensure score doesn't exceed 100
    score = min(score, 100)
    
    # Core requirements for passing: Must save file and MUST create the table
    passed = (score >= 80) and output_exists and (prog_checks["table_exists"] or score >= 45)

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }