#!/usr/bin/env python3
"""
Verifier for employee_handbook_format task.
Evaluates formatting, find/replace, and layout of a DOCX file.
"""

import json
import os
import tempfile
import logging
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_employee_handbook_format(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    
    feedback_parts = []
    score = 0
    max_score = 100
    
    # 1. Check basic export results (File exists & created during task)
    result_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", result_json.name)
        with open(result_json.name, 'r') as f:
            export_data = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(result_json.name):
            os.unlink(result_json.name)

    if not export_data.get('file_exists'):
        return {
            "passed": False,
            "score": 0,
            "feedback": "Final document not found at ~/Documents/results/employee_handbook_final.docx."
        }
    
    if export_data.get('file_created_during_task'):
        score += 10
        feedback_parts.append("File correctly saved during task.")
    else:
        feedback_parts.append("File was not saved after task started.")

    # 2. Parse the document
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env('/tmp/employee_handbook_final.docx', temp_doc.name)
        doc = Document(temp_doc.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse docx: {e}"}
    finally:
        if os.path.exists(temp_doc.name):
            os.unlink(temp_doc.name)

    # Compile full text
    paragraphs = doc.paragraphs
    full_text_lower = '\n'.join([p.text for p in paragraphs]).lower()
    
    # 3. Find and Replace (20 pts)
    placeholder_count = full_text_lower.count("[company name]")
    company_count = full_text_lower.count("medtech solutions")
    
    if placeholder_count == 0 and company_count >= 10:
        score += 20
        feedback_parts.append("Find and Replace successful.")
    elif placeholder_count > 0 and company_count > 0:
        score += 10
        feedback_parts.append(f"Partial Find and Replace: {placeholder_count} placeholders remain.")
    else:
        feedback_parts.append("Find and Replace failed.")

    # 4. Title Style (10 pts)
    title_ok = False
    if len(paragraphs) > 0:
        p0 = paragraphs[0]
        if "2026 employee handbook" in p0.text.lower() and "title" in p0.style.name.lower():
            title_ok = True
    if title_ok:
        score += 10
        feedback_parts.append("Title style correctly applied.")
    else:
        feedback_parts.append("Title style missing or incorrect on the first line.")

    # 5. Heading 1 (15 pts) & Heading 2 (15 pts)
    h1_targets = [s.lower() for s in metadata.get("h1_sections", [])]
    h2_targets = [s.lower() for s in metadata.get("h2_sections", [])]
    
    h1_found = 0
    h2_found = 0
    
    for p in paragraphs:
        ptxt = p.text.strip().lower()
        if p.style and "heading 1" in p.style.name.lower():
            if any(h1 in ptxt for h1 in h1_targets):
                h1_found += 1
        elif p.style and "heading 2" in p.style.name.lower():
            if any(h2 in ptxt for h2 in h2_targets):
                h2_found += 1

    h1_score = int((min(h1_found, 5) / 5.0) * 15)
    score += h1_score
    feedback_parts.append(f"Heading 1s formatted: {min(h1_found, 5)}/5.")
    
    h2_score = int((min(h2_found, 7) / 7.0) * 15)
    score += h2_score
    feedback_parts.append(f"Heading 2s formatted: {min(h2_found, 7)}/7.")

    # 6. Pagination (10 pts)
    # Check for <w:br w:type="page"/> in document XML or page_break_before in H1 style
    xml_str = doc._element.xml
    page_break_count = xml_str.count('w:type="page"')
    
    # We expect at least 2 page breaks (after title, after TOC) 
    # and ideally 5 more if they used manual page breaks for chapters, 
    # but they might have just used paragraph formatting `page_break_before`
    has_page_break_before = False
    for style in doc.styles:
        if "heading 1" in style.name.lower():
            if style.paragraph_format.page_break_before:
                has_page_break_before = True

    if page_break_count >= 2 or has_page_break_before:
        score += 10
        feedback_parts.append("Pagination requirements met.")
    else:
        feedback_parts.append("Page breaks missing.")

    # 7. Headers (10 pts)
    header_found = False
    for section in doc.sections:
        for p in section.header.paragraphs:
            if "medtech solutions employee handbook" in p.text.lower():
                header_found = True
                break
    if header_found:
        score += 10
        feedback_parts.append("Document header verified.")
    else:
        feedback_parts.append("Document header missing or incorrect.")

    # 8. Table of Contents (10 pts)
    # Check for standard MS/WPS TOC fields
    has_toc = False
    if 'TOC' in xml_str or 'w:sdt' in xml_str:
        has_toc = True
    
    if has_toc:
        score += 10
        feedback_parts.append("Table of Contents found.")
    else:
        feedback_parts.append("Table of Contents not found.")

    # Evaluate final VLM verification to add robustness
    try:
        from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
        
        frames = sample_trajectory_frames(traj, n=3)
        final = get_final_screenshot(traj)
        
        vlm_prompt = (
            "You are reviewing an agent's work formatting an Employee Handbook in a Word Processor. "
            "Examine these screenshots. "
            "Do you see evidence that the agent successfully generated an automated Table of Contents (a block with dot leaders and page numbers) "
            "or successfully added a document header with the text 'MedTech Solutions Employee Handbook'? "
            "Reply strictly in JSON: {\"toc_or_header_visible\": true/false}"
        )
        vlm_res = query_vlm(images=frames + [final], prompt=vlm_prompt)
        
        if vlm_res.get("parsed", {}).get("toc_or_header_visible", False):
            feedback_parts.append("VLM visual verification confirmed UI state.")
        else:
            feedback_parts.append("VLM could not confirm formatting visually.")
    except Exception as e:
        logger.warning(f"VLM verification failed or unavailable: {e}")

    # Calculate success
    passed = score >= 75 and placeholder_count == 0 and h1_found >= 4
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }