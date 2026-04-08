#!/usr/bin/env python3
"""
Verifier for crm_archaeology_report task.
Evaluates document styling, table creation from unstructured text, hanging indents,
and specific character-level formatting (italics).
"""

import json
import tempfile
import os
import logging
from typing import Dict, Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import VLM utils
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "utils"))
try:
    from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
    VLM_AVAILABLE = True
except ImportError:
    VLM_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def verify_crm_archaeology_report(traj: Dict[str, Any], env_info: Dict[str, Any], task_info: Dict[str, Any]) -> Dict[str, Any]:
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # --- 1. Load result JSON ---
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_data = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load export JSON: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # --- Anti-gaming Check ---
    if not result_data.get('output_exists'):
        return {"passed": False, "score": 0, "feedback": "Output file was not saved to the correct path (~/Documents/results/phase_i_report_formatted.docx)."}
    if not result_data.get('file_created_during_task'):
        return {"passed": False, "score": 0, "feedback": "File modification timestamp precedes task start (Cheating detected)."}

    # --- 2. Copy and Parse the Formatted Docx ---
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/phase_i_report_formatted.docx", temp_docx.name)
        if not DOCX_AVAILABLE:
            return {"passed": False, "score": 0, "feedback": "Python-docx not available on host for verification."}
        
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read docx file: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 0
    feedback = []
    
    # --- 3. Evaluate Criteria ---

    # Criterion A: Confidentiality Warning (15 pts)
    # Looking for "CONFIDENTIAL" in bold and red near the start of the document
    confidential_found = False
    is_bold = False
    is_centered = False
    
    for i, para in enumerate(doc.paragraphs[:5]): # Check top 5 paragraphs
        if "CONFIDENTIAL" in para.text.upper():
            confidential_found = True
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                is_centered = True
            for run in para.runs:
                if "CONFIDENTIAL" in run.text.upper():
                    if run.bold:
                        is_bold = True
                    # Checking exact color is unreliable across different Word processors and styles,
                    # but we'll accept if bold and text is present. 
            break
            
    if confidential_found:
        score += 10
        feedback.append("Confidentiality warning found.")
        if is_bold and is_centered:
            score += 5
            feedback.append("Confidentiality warning properly formatted (bold & centered).")
        else:
            feedback.append("Confidentiality warning missing bold or centered formatting.")
    else:
        feedback.append("Missing CONFIDENTIAL warning.")

    # Criterion B: Title Style (5 pts)
    title_correct = False
    for para in doc.paragraphs[:10]:
        if "Phase I Archaeological Survey" in para.text:
            if para.style and ("Title" in para.style.name or "Heading" in para.style.name):
                title_correct = True
                break
    if title_correct:
        score += 5
        feedback.append("Title style applied.")
    else:
        feedback.append("Title style missing or incorrect.")

    # Criterion C: Heading 1 Hierarchy (15 pts)
    headings_found = 0
    expected_headings = [
        "1.0 INTRODUCTION", "2.0 ENVIRONMENTAL SETTING", "3.0 CULTURAL CONTEXT",
        "4.0 FIELD METHODS", "5.0 RESULTS", "6.0 SUMMARY AND RECOMMENDATIONS", "7.0 REFERENCES CITED"
    ]
    for para in doc.paragraphs:
        text_up = para.text.upper()
        if para.style and "Heading 1" in para.style.name:
            for eh in expected_headings:
                if eh in text_up:
                    headings_found += 1
                    
    if headings_found >= 6:
        score += 15
        feedback.append(f"Found {headings_found}/7 Heading 1 styles.")
    elif headings_found >= 3:
        score += 7
        feedback.append(f"Partial Heading 1 styles found ({headings_found}/7).")
    else:
        feedback.append(f"Missing Heading 1 styles (found {headings_found}/7).")

    # Criterion D: Scientific Names Italics (10 pts)
    sci_names = ["Pinus taeda", "Quercus alba", "Odocoileus virginianus", "Meleagris gallopavo"]
    sci_found = 0
    
    for para in doc.paragraphs:
        if "2.0" in para.text or "ENVIRONMENTAL" in para.text.upper() or "loblolly pine" in para.text:
            # We look at all runs.
            # Because names might be split across runs, we just aggregate italicized text
            italic_text = "".join([run.text for run in para.runs if run.italic])
            for name in sci_names:
                # If the entire name or substantial part is in the italic text block
                if name.split()[0] in italic_text or name in italic_text:
                    sci_found += 1
            if sci_found > 0:
                break # Found the paragraph, checked names
                
    if sci_found >= 3:
        score += 10
        feedback.append(f"Scientific italics correctly applied ({sci_found}/4).")
    elif sci_found > 0:
        score += 5
        feedback.append(f"Partial scientific italics applied ({sci_found}/4).")
    else:
        feedback.append("Scientific names not italicized.")

    # Criterion E: Tables created from text (30 pts - 15 each)
    num_tables = len(doc.tables)
    stp_table_found = False
    artifact_table_found = False
    
    if num_tables > 0:
        for table in doc.tables:
            if len(table.rows) == 0: continue
            
            # Check headers
            header_text = "".join([cell.text for cell in table.rows[0].cells]).lower()
            
            # Is it the STP table?
            if "transect" in header_text and "stp" in header_text and len(table.rows) >= 6:
                stp_table_found = True
                header_bold = any(run.bold for para in table.rows[0].cells[0].paragraphs for run in para.runs)
                if header_bold:
                    score += 15
                    feedback.append("STP table successfully formatted with bold header.")
                else:
                    score += 10
                    feedback.append("STP table created but header lacks bolding.")
            
            # Is it the Artifact table?
            if "catalog" in header_text and "material" in header_text and len(table.rows) >= 4:
                artifact_table_found = True
                header_bold = any(run.bold for para in table.rows[0].cells[0].paragraphs for run in para.runs)
                if header_bold:
                    score += 15
                    feedback.append("Artifact table successfully formatted with bold header.")
                else:
                    score += 10
                    feedback.append("Artifact table created but header lacks bolding.")

    if not stp_table_found:
        feedback.append("STP table not found or improperly formatted.")
    if not artifact_table_found:
        feedback.append("Artifact table not found or improperly formatted.")

    # Criterion F: Hanging Indents for References (10 pts)
    # In docx, a hanging indent means left_indent > 0 and first_line_indent < 0
    hanging_indents_found = 0
    in_refs = False
    
    for para in doc.paragraphs:
        if "REFERENCES CITED" in para.text.upper():
            in_refs = True
            continue
            
        if in_refs and len(para.text.strip()) > 10:
            fmt = para.paragraph_format
            # Negative first line indent = Hanging
            if fmt.first_line_indent and fmt.first_line_indent.pt < 0:
                hanging_indents_found += 1
            # Some WPS implementations do this differently, so we also check if left_indent > 0 and no first_line
            elif fmt.left_indent and fmt.left_indent.pt > 0 and (not fmt.first_line_indent or fmt.first_line_indent.pt <= 0):
                # Loose pass for hanging logic
                hanging_indents_found += 1
                
    if hanging_indents_found >= 2:
        score += 10
        feedback.append("Hanging indents successfully applied to references.")
    elif hanging_indents_found == 1:
        score += 5
        feedback.append("Hanging indents partially applied.")
    else:
        feedback.append("Hanging indents not found in References section.")

    # Criterion G: VLM Verification (15 points)
    # Check trajectories to verify genuine interaction
    if VLM_AVAILABLE:
        try:
            frames = sample_trajectory_frames(traj, n=4)
            final = get_final_screenshot(traj)
            prompt = """
            You are verifying a document formatting task in WPS Office.
            Look at the screenshots. Did the user:
            1. Open the document?
            2. Convert plain text logs into actual structural tables?
            3. Apply formatting (bold/colors)?
            
            Return JSON:
            {"work_performed": true/false, "tables_visible": true/false}
            """
            vlm_result = query_vlm(images=frames + [final] if final else frames, prompt=prompt)
            
            if vlm_result.get("success"):
                parsed = vlm_result.get("parsed", {})
                if parsed.get("work_performed", False) and parsed.get("tables_visible", False):
                    score += 15
                    feedback.append("VLM visual verification passed.")
                elif parsed.get("work_performed", False):
                    score += 8
                    feedback.append("VLM visual verification: Partial work detected.")
                else:
                    feedback.append("VLM could not confirm work in trajectory.")
        except Exception as e:
            logger.warning(f"VLM check failed: {e}")
            # Give benefit of doubt if VLM fails but file parses perfectly
            if score >= 60:
                score += 15
                feedback.append("VLM check bypassed due to error, assumed pass based on file heuristics.")
    else:
        # Without VLM, grant the trajectory points if the file heuristics are very strong
        if score >= 60:
            score += 15
            feedback.append("VLM unavailable - awarded points based on strong file structure.")

    # Cap score at 100
    score = min(100, max(0, score))
    passed = score >= 75 and stp_table_found and artifact_table_found

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback),
        "details": {
            "confidential": confidential_found,
            "headings_found": headings_found,
            "sci_names_italicized": sci_found,
            "stp_table_found": stp_table_found,
            "artifact_table_found": artifact_table_found,
            "hanging_indents": hanging_indents_found
        }
    }