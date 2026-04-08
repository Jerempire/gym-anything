#!/usr/bin/env python3
"""Verifier for municipal_bond_official_statement task."""

import sys
import os
import re
import logging

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from wps_verification_utils import (
    copy_and_parse_document,
    cleanup_verification_temp,
    get_document_text,
    check_heading_styles,
    count_tables,
    get_table_content,
    get_table_dimensions,
    check_table_header_formatting,
    count_headings_by_level,
    vlm_verify_screenshot,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def verify_municipal_bond_official_statement(traj, env_info, task_info):
    """
    Verify that the draft Official Statement errors were corrected.

    PREREQUISITE: Core bond document content must be preserved.

    SCORING CRITERIA (10 criteria):
    1. Maturity schedule total corrected: $43,500,000 -> $45,000,000
    2. S&P credit rating corrected: AA- -> AA
    3. Risk factors / Bondholders' risks section added
    4. Continuing disclosure italic formatting removed
    5. Heading hierarchy fixed (Maturity Schedule and Financial Info were wrong levels)
    6. Credit Ratings section has heading style (was Normal+bold)
    7. Underwriting heading style fixed (was Normal+bold+centered)
    8. Revenue data properly formatted (not right-aligned)
    9. Table header formatting on maturity and/or debt service tables
    10. VLM visual verification
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    container_path = "/home/ga/Documents/os_draft_greenfield.docx"
    success, doc, error, temp_dir = copy_and_parse_document(
        container_path, copy_from_env, file_format='docx'
    )

    if not success:
        return {"passed": False, "score": 0, "feedback": error}

    try:
        feedback_parts = []
        full_text = get_document_text(doc).lower()

        # PREREQUISITE: Core content must be preserved
        key_phrases = [
            "city of greenfield",
            "general obligation bonds",
            "series 2024a",
            "45,000,000",
            "chapman & cutler",
        ]
        preserved = sum(1 for p in key_phrases if p.lower() in full_text)
        if preserved < 3:
            return {
                "passed": False,
                "score": 0,
                "feedback": f"PREREQUISITE FAILED: Content corrupted ({preserved}/{len(key_phrases)} key phrases)",
            }
        feedback_parts.append(f"Prerequisite: content preserved ({preserved}/{len(key_phrases)})")

        criteria_passed = 0
        total_criteria = 10

        # Criterion 1: Maturity schedule total corrected ($43,500,000 -> $45,000,000)
        num_tables = count_tables(doc)
        maturity_total_fixed = False
        for t_idx in range(num_tables):
            content = get_table_content(doc, t_idx)
            for row in content:
                row_text = ' '.join(row).lower()
                if 'total' in row_text:
                    # Check for correct total
                    if '45,000,000' in ' '.join(row) or '45000000' in ' '.join(row):
                        maturity_total_fixed = True
                    # Make sure wrong total is gone
                    if '43,500,000' in ' '.join(row) or '43500000' in ' '.join(row):
                        maturity_total_fixed = False
        if maturity_total_fixed:
            criteria_passed += 1
            feedback_parts.append("Maturity total: corrected to $45,000,000")
        else:
            feedback_parts.append("Maturity total: NOT corrected (still $43,500,000 or missing)")

        # Criterion 2: S&P rating corrected (AA- -> AA)
        has_wrong_rating = False
        has_correct_rating = False
        for para in doc.paragraphs:
            text_l = para.text.lower()
            if "s&p" in text_l or "standard" in text_l:
                if "aa-" in text_l:
                    has_wrong_rating = True
                # Check for corrected 'AA' near S&P context (ignore Moody's 'Aa2')
                sp_match = re.search(r"(?:s&p|standard).{0,60}", text_l)
                if sp_match:
                    sp_context = sp_match.group()
                    if ("'aa'" in sp_context or '"aa"' in sp_context) and 'aa-' not in sp_context:
                        has_correct_rating = True
        # Fallback: check full text for aa- near s&p
        if not has_wrong_rating and not has_correct_rating:
            sp_sections = re.findall(r"(?:s&p|standard.+poor).{0,100}", full_text)
            for section in sp_sections:
                if 'aa-' in section:
                    has_wrong_rating = True
                if ("'aa'" in section or '"aa"' in section) and 'aa-' not in section:
                    has_correct_rating = True
        if has_correct_rating and not has_wrong_rating:
            criteria_passed += 1
            feedback_parts.append("S&P rating: corrected to AA")
        elif not has_wrong_rating:
            criteria_passed += 1
            feedback_parts.append("S&P rating: AA- removed (assuming corrected)")
        else:
            feedback_parts.append("S&P rating: NOT corrected (still AA-)")

        # Criterion 3: Risk factors section added
        risk_found = False
        for para in doc.paragraphs:
            text_l = para.text.strip().lower()
            if ('risk factor' in text_l or "bondholders' risk" in text_l or
                'risk' in text_l and ('factor' in text_l or 'bondholder' in text_l)):
                if para.style and 'heading' in para.style.name.lower():
                    risk_found = True
                    break
                elif any(r.bold for r in para.runs if r.text.strip()):
                    risk_found = True
                    break
        if risk_found:
            criteria_passed += 1
            feedback_parts.append("Risk factors: section added")
        else:
            feedback_parts.append("Risk factors: section NOT found (MSRB requirement)")

        # Criterion 4: Continuing disclosure italic removed
        cd_italic_fixed = True
        for para in doc.paragraphs:
            text_l = para.text.lower()
            if 'continuing disclosure' in text_l or 'rule 15c2-12' in text_l:
                for run in para.runs:
                    if run.text.strip() and run.italic:
                        cd_italic_fixed = False
                        break
                if not cd_italic_fixed:
                    break
        if cd_italic_fixed:
            criteria_passed += 1
            feedback_parts.append("Continuing disclosure: italic formatting removed")
        else:
            feedback_parts.append("Continuing disclosure: still has italic formatting")

        # Criterion 5: Heading hierarchy fixed
        heading_fixes = 0
        for para in doc.paragraphs:
            text_l = para.text.strip().lower()
            if 'maturity schedule' in text_l and para.style:
                if 'heading 1' in para.style.name.lower():
                    heading_fixes += 1
            if 'financial information' in text_l and para.style:
                if 'heading 1' in para.style.name.lower():
                    heading_fixes += 1
        if heading_fixes >= 1:
            criteria_passed += 1
            feedback_parts.append(f"Heading hierarchy: {heading_fixes}/2 sections fixed to Heading 1")
        else:
            feedback_parts.append("Heading hierarchy: NOT fixed")

        # Criterion 6: Credit Ratings has heading style
        cr_fixed = False
        for para in doc.paragraphs:
            text_l = para.text.strip().lower()
            if 'credit rating' in text_l:
                if para.style and 'heading' in para.style.name.lower():
                    cr_fixed = True
                break
        if cr_fixed:
            criteria_passed += 1
            feedback_parts.append("Credit Ratings: heading style applied")
        else:
            feedback_parts.append("Credit Ratings: still Normal+bold (not heading style)")

        # Criterion 7: Underwriting heading style fixed
        uw_fixed = False
        for para in doc.paragraphs:
            text_l = para.text.strip().lower()
            if text_l == 'underwriting' or 'underwriting' == text_l.strip():
                if para.style and 'heading' in para.style.name.lower():
                    uw_fixed = True
                break
        if uw_fixed:
            criteria_passed += 1
            feedback_parts.append("Underwriting: heading style applied")
        else:
            feedback_parts.append("Underwriting: still Normal+bold+centered")

        # Criterion 8: Revenue data alignment fixed
        revenue_aligned = True
        for para in doc.paragraphs:
            text_l = para.text.lower()
            if 'property tax' in text_l and 'sales tax' in text_l:
                from docx.enum.text import WD_ALIGN_PARAGRAPH as WD
                if para.alignment == WD.RIGHT:
                    revenue_aligned = False
                break
        if revenue_aligned:
            criteria_passed += 1
            feedback_parts.append("Revenue data: properly aligned")
        else:
            feedback_parts.append("Revenue data: still right-aligned")

        # Criterion 9: Table header formatting
        if num_tables >= 1:
            any_formatted = False
            for t_idx in range(min(num_tables, 5)):
                header_fmt = check_table_header_formatting(doc, t_idx)
                if header_fmt['has_bold']:
                    any_formatted = True
                    break
            if any_formatted:
                criteria_passed += 1
                feedback_parts.append("Table headers: formatted")
            else:
                feedback_parts.append("Table headers: no formatting")
        else:
            feedback_parts.append("Table headers: no tables found")

        # Criterion 10: VLM verification
        vlm_result = vlm_verify_screenshot(env_info, traj, """
Analyze this WPS Writer screenshot of a municipal bond Official Statement. Answer in JSON:
{
    "has_clear_sections": true/false,
    "has_tables": true/false,
    "appears_professionally_formatted": true/false,
    "has_structured_headings": true/false,
    "looks_like_official_document": true/false
}
Does the document show:
1. Clear section headings with consistent formatting?
2. Financial tables with proper formatting?
3. Professional municipal finance document formatting?
4. Structured heading hierarchy?
5. Formal Official Statement appearance?
""")
        if vlm_result is not None:
            has_sections = vlm_result.get("has_clear_sections", False)
            looks_official = vlm_result.get("looks_like_official_document", False)
            if has_sections or looks_official:
                criteria_passed += 1
                feedback_parts.append("VLM: professional formatting confirmed")
            else:
                feedback_parts.append("VLM: formatting not confirmed")
        else:
            total_criteria -= 1
            feedback_parts.append("VLM: unavailable (skipped)")

        score = int((criteria_passed / total_criteria) * 100)
        passed = score >= 55

        return {
            "passed": passed,
            "score": score,
            "feedback": " | ".join(feedback_parts),
        }

    except Exception as e:
        logger.error(f"Verification error: {e}", exc_info=True)
        return {"passed": False, "score": 0, "feedback": f"Verification error: {str(e)}"}
    finally:
        cleanup_verification_temp(temp_dir)
