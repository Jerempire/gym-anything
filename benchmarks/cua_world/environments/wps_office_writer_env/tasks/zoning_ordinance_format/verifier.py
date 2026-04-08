#!/usr/bin/env python3
"""Verifier for zoning_ordinance_format task."""

import os
import json
import tempfile
import logging

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import VLM utils
from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_zoning_ordinance_format(traj, env_info, task_info):
    """
    Verify the zoning ordinance draft formatting.
    
    Criteria:
    1. File Saved & Modified during task (10 pts)
    2. Legal Page Size: 8.5" x 14" (20 pts)
    3. Heading Hierarchy: Article X (H1), Section 1-6 (H2) (20 pts)
    4. Permitted Uses Table: 4 cols, expected headers, 4 rows of data (25 pts)
    5. Draft Header: "DRAFT ORDINANCE - NOT ADOPTED" right aligned (15 pts)
    6. Bold Definitions: 4 specific terms bolded (10 pts)
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # 1. Check basic file JSON metrics
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    result_data = {}
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_data = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result JSON: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)
            
    if not result_data.get("document_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output document was not saved to expected path."}
        
    score = 0
    feedback = []
    
    if result_data.get("file_created_during_task", False):
        score += 10
        feedback.append("File properly saved during task.")
    else:
        feedback.append("File saved, but timestamps suggest it was not modified during task.")

    # 2. Extract Document for Parsing
    if not DOCX_AVAILABLE:
        return {"passed": False, "score": score, "feedback": "python-docx not available for deep verification."}

    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/tod_ordinance_draft_result.docx", temp_doc.name)
        doc = Document(temp_doc.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse docx: {e}"}
    finally:
        if os.path.exists(temp_doc.name):
            os.unlink(temp_doc.name)

    # 3. Verify Legal Page Size (8.5 x 14 inches)
    section = doc.sections[0]
    height_inches = section.page_height.inches
    width_inches = section.page_width.inches
    
    if abs(height_inches - 14.0) < 0.2 and abs(width_inches - 8.5) < 0.2:
        score += 20
        feedback.append("Page Size correctly set to Legal (8.5x14).")
    else:
        feedback.append(f"Page Size incorrect. Expected ~8.5x14, got {width_inches:.2f}x{height_inches:.2f}")

    # 4. Verify Heading Hierarchy
    h1_correct = False
    h2_count = 0
    
    for para in doc.paragraphs:
        text = para.text.lower()
        if "article x: transit-oriented development" in text:
            if para.style and 'heading 1' in para.style.name.lower():
                h1_correct = True
        elif "section " in text and ("purpose" in text or "applicability" in text or "definitions" in text or "permitted uses" in text or "development standards" in text or "parking" in text):
            if para.style and 'heading 2' in para.style.name.lower():
                h2_count += 1
                
    h_score = 0
    if h1_correct:
        h_score += 10
    if h2_count >= 5: # Allow slightly fuzzy match in case one was missed
        h_score += 10
    score += h_score
    feedback.append(f"Headings formatting score: {h_score}/20 (H1: {h1_correct}, H2s: {h2_count}/6).")

    # 5. Verify Permitted Uses Table
    table_score = 0
    if len(doc.tables) >= 1:
        table = doc.tables[0]
        # Check Columns
        if len(table.columns) == 4:
            table_score += 10
            # Check Headers
            headers = [cell.text.lower().strip() for cell in table.rows[0].cells]
            if "use category" in headers[0] and "specific use" in headers[1]:
                table_score += 5
            # Check Data Rows
            if len(table.rows) >= 5: # 1 header + 4 data rows
                table_score += 10
        else:
            feedback.append(f"Table found but incorrect column count ({len(table.columns)} instead of 4).")
    else:
        feedback.append("No table found in the document.")
        
    score += table_score
    feedback.append(f"Table formatting score: {table_score}/25.")

    # 6. Verify Draft Header
    header_found = False
    for sect in doc.sections:
        for para in sect.header.paragraphs:
            if "DRAFT ORDINANCE - NOT ADOPTED" in para.text:
                header_found = True
                break
        if header_found:
            break
            
    if header_found:
        score += 15
        feedback.append("Draft Header text found in document header.")
    else:
        feedback.append("Draft Header text missing from document header.")

    # 7. Verify Bold Definitions
    terms = ["Transit Station Area", "Active Floor Area", "Pedestrian-Oriented Facade", "Shared Parking"]
    bold_terms_count = 0
    
    for para in doc.paragraphs:
        for run in para.runs:
            for term in terms:
                if term.lower() in run.text.lower() and run.bold:
                    bold_terms_count += 1
                    terms.remove(term) # Prevent double counting

    bold_score = min(10, int((bold_terms_count / 4.0) * 10))
    score += bold_score
    feedback.append(f"Bold definitions score: {bold_score}/10 ({bold_terms_count}/4 found).")

    # VLM Trajectory Verification to prevent "pure script" gaming
    try:
        frames = sample_trajectory_frames(traj, n=4)
        final_img = get_final_screenshot(traj)
        if final_img:
            frames.append(final_img)
            
        vlm_prompt = """Review these screenshots of a user formatting a document in WPS Office Writer.
        1. Is the WPS Writer application visibly open and being used?
        2. Can you see signs of document formatting occurring (e.g. creating tables, adjusting styles, editing headers)?
        
        Respond in JSON:
        {
            "wps_used": true/false,
            "formatting_visible": true/false
        }"""
        
        vlm_res = query_vlm(images=frames, prompt=vlm_prompt)
        if vlm_res and vlm_res.get("parsed", {}).get("wps_used", False):
            feedback.append("VLM confirms WPS Writer was actively used.")
        else:
            # Penalize heavily if VLM detects cheating (e.g. Python script generation from shell)
            feedback.append("VLM WARNING: Could not confirm visual usage of WPS Writer in trajectory. Possible gaming detected.")
            score = int(score * 0.5) 
    except Exception as e:
        logger.warning(f"VLM verification failed/skipped: {e}")

    # Determine final outcome
    key_criteria_met = result_data.get("file_created_during_task", False) and (h1_correct or table_score >= 10)
    passed = score >= 75 and key_criteria_met

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback),
        "details": {
            "score": score,
            "h1_correct": h1_correct,
            "h2_count": h2_count,
            "table_score": table_score,
            "header_found": header_found,
            "bold_terms_count": bold_terms_count
        }
    }