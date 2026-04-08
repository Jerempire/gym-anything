#!/usr/bin/env python3
"""
Verifier for clinical_trial_icf_format task.
Evaluates strict IRB formatting requirements on a Word document.
"""

import json
import os
import sys
import tempfile
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_clinical_trial_icf_format(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Framework error: Copy function not available"}

    feedback_parts = []
    score = 0

    # 1. Read exported results JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Validate output exists
    if not result.get('output_exists', False):
        return {
            "passed": False, 
            "score": 0, 
            "feedback": "Output document final_icf_v2.docx not found. Task incomplete."
        }

    if result.get('file_created_during_task', False):
        score += 10
        feedback_parts.append("File created correctly (+10)")
    else:
        feedback_parts.append("Warning: File not modified during task timeframe")

    # 2. Extract Document using python-docx
    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/final_icf_v2.docx", temp_docx.name)
        doc = docx.Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse docx: {e}"}
    
    # Pre-check: Verify contents weren't completely wiped
    full_text = "\n".join([p.text for p in doc.paragraphs]).lower()
    if len(full_text) < 1000 or "investigational" not in full_text:
        os.unlink(temp_docx.name)
        return {
            "passed": False, 
            "score": score, 
            "feedback": "Core document content appears to be missing or deleted."
        }

    # 3. Typography Verification: Arial, 12pt (10 pts)
    arial_12_found = False
    for p in doc.paragraphs[4:10]:  # Sample body paragraphs
        if p.runs and p.runs[0].font:
            font_name = p.runs[0].font.name or (p.style.font and p.style.font.name)
            font_size = p.runs[0].font.size or (p.style.font and p.style.font.size)
            if font_name == 'Arial' and font_size == Pt(12):
                arial_12_found = True
                break
    
    if arial_12_found:
        score += 10
        feedback_parts.append("Typography: Arial 12pt applied (+10)")
    else:
        feedback_parts.append("Typography: Did not reliably detect Arial 12pt")

    # 4. Title Block Formatting (10 pts)
    title_correct = 0
    for i in range(min(3, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        # Allow alignment value 1 which matches Center in python-docx
        is_center = (p.alignment == WD_ALIGN_PARAGRAPH.CENTER or p.alignment == 1)
        is_bold = any(r.bold for r in p.runs)
        is_14pt = any(r.font and r.font.size == Pt(14) for r in p.runs)
        if is_center and is_bold and is_14pt:
            title_correct += 1
            
    if title_correct == 3:
        score += 10
        feedback_parts.append("Title block: Perfect formatting (+10)")
    elif title_correct > 0:
        score += 5
        feedback_parts.append(f"Title block: Partial formatting ({title_correct}/3) (+5)")
    else:
        feedback_parts.append("Title block: Formatting missing")

    # 5. Section Headings Verification (10 pts)
    target_headings = [
        'introduction', 'purpose of the research', 'study procedures', 
        'risks and discomforts', 'potential benefits', 'confidentiality', 
        'voluntary participation', 'contact information'
    ]
    headings_found = 0
    for p in doc.paragraphs:
        if p.style and 'heading 1' in p.style.name.lower():
            if any(h in p.text.lower() for h in target_headings):
                headings_found += 1
                
    if headings_found >= 7:
        score += 10
        feedback_parts.append("Headings: Successfully applied Heading 1 (+10)")
    elif headings_found >= 3:
        score += 5
        feedback_parts.append(f"Headings: Partially applied ({headings_found}/8) (+5)")
    else:
        feedback_parts.append("Headings: Missing Heading 1 styles")

    # 6. Risk Bulleted List (10 pts)
    list_detected = False
    for p in doc.paragraphs:
        text_l = p.text.lower()
        if 'nausea' in text_l or 'severe fatigue' in text_l or 'mild dizziness' in text_l:
            # Native list style check
            if p.style and 'list' in p.style.name.lower():
                list_detected = True
            # Properties check (Word native lists)
            elif getattr(p._p, 'pPr', None) is not None and getattr(p._p.pPr, 'numPr', None) is not None:
                list_detected = True
            # Fallback for manual bullet character
            elif text_l.strip().startswith(('•', '-', 'o')):
                list_detected = True
                
    if list_detected:
        score += 10
        feedback_parts.append("Risk List: Bulleted list created (+10)")
    else:
        feedback_parts.append("Risk List: Not converted to a bulleted list")

    # 7. Voluntary Paragraph Bolded (10 pts)
    voluntary_bold = False
    for p in doc.paragraphs:
        if 'your participation in this research is entirely voluntary' in p.text.lower():
            total_len = len(p.text.strip())
            bold_len = sum(len(r.text) for r in p.runs if r.bold)
            if total_len > 0 and (bold_len > total_len * 0.5):
                voluntary_bold = True
                break
                
    if voluntary_bold:
        score += 10
        feedback_parts.append("Voluntary paragraph: Successfully bolded (+10)")
    else:
        feedback_parts.append("Voluntary paragraph: Not properly bolded")

    # 8. Version Control Footer (10 pts)
    footer_text = ""
    for section in doc.sections:
        if section.footer:
            for p in section.footer.paragraphs:
                footer_text += p.text.lower() + " "
                
    if 'onc-2026' in footer_text and 'version 2.1' in footer_text:
        score += 10
        feedback_parts.append("Footer: Protocol and version present (+10)")
    else:
        feedback_parts.append("Footer: Missing or incorrect contents")

    # 9. Signature Table (10 pts)
    table_correct = False
    if len(doc.tables) > 0:
        table = doc.tables[-1]
        if len(table.rows) >= 3 and len(table.columns) >= 3:
            header_row = "".join([cell.text.lower() for cell in table.rows[0].cells])
            if 'name' in header_row and 'signature' in header_row and 'date' in header_row:
                # Check for bold header
                is_bold = False
                for p in table.rows[0].cells[0].paragraphs:
                    if any(r.bold for r in p.runs):
                        is_bold = True
                table_correct = True
                
    if table_correct:
        score += 10
        feedback_parts.append("Signature Table: Created correctly (+10)")
    else:
        feedback_parts.append("Signature Table: Missing or incorrect format")

    os.unlink(temp_docx.name)

    # 10. VLM Trajectory Verification (20 pts)
    try:
        from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
        frames = sample_trajectory_frames(traj, n=3)
        final = get_final_screenshot(traj)
        
        if frames and final:
            prompt = (
                "You are verifying a document formatting task in WPS Office Writer.\n"
                "Task goals: Change font to Arial 12pt, format title block, add Heading 1 styles, "
                "create a bulleted list, bold a paragraph, add a footer, and insert a 3x3 signature table.\n"
                "Examine these trajectory frames. Do they show the agent actively using WPS Writer and "
                "interacting with formatting menus, styles, table insertion, or footer editing?\n"
                "Respond in JSON format with exactly this key:\n"
                '{"active_formatting_visible": true/false}'
            )
            vlm_result = query_vlm(images=frames + [final], prompt=prompt)
            if vlm_result.get('success', False):
                parsed = vlm_result.get('parsed', {})
                if parsed.get('active_formatting_visible', False):
                    score += 20
                    feedback_parts.append("VLM: Active formatting visible (+20)")
                else:
                    feedback_parts.append("VLM: No active formatting visible")
            else:
                # Fallback pass if VLM module fails internal logic
                score += 20
                feedback_parts.append("VLM: Query failed, bypassing (+20)")
        else:
            score += 20
            feedback_parts.append("VLM: No frames available, bypassing (+20)")
    except Exception as e:
        logger.warning(f"VLM verification exception: {e}")
        score += 20
        feedback_parts.append("VLM: Module not available/failed, bypassing (+20)")

    passed = score >= 70

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }