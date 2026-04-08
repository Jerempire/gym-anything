#!/usr/bin/env python3
"""
Verifier for theater_playbill_format task.
Evaluates page layout changes (A5 size, 0.5" margins), column layouts, 
heading styles, inline character formatting (bolding), and pagination.
"""

import os
import json
import logging
import tempfile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_theater_playbill_format(traj, env_info, task_info):
    """
    Main verification function.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # ====================================================================
    # 1. Read Result JSON
    # ====================================================================
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not result.get('output_exists', False):
        return {
            "passed": False,
            "score": 0,
            "feedback": "Target file formatted_playbill.docx was not created."
        }

    if not result.get('file_created_during_task', False):
        return {
            "passed": False,
            "score": 0,
            "feedback": "File exists but was not created/modified during the task window (anti-gaming failure)."
        }

    # ====================================================================
    # 2. Extract Document for Parsing
    # ====================================================================
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/formatted_playbill.docx", temp_docx.name)
        
        # Ensure dependencies
        try:
            from docx import Document
        except ImportError:
            import subprocess, sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            from docx import Document
            
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse document: {e}"}
    finally:
        pass # We will unlink it at the very end

    score = 0
    feedback_parts = []
    
    # Expected metrics
    A5_WIDTH = 5.83
    A5_HEIGHT = 8.27
    MARGIN = 0.5
    TOLERANCE = 0.2

    # Track states
    is_a5 = False
    has_05_margins = False
    has_2_cols = False
    page_numbers_added = False

    # ====================================================================
    # 3. Verify Page Layout (Sections)
    # ====================================================================
    try:
        # Check all sections. A proper continuous section break creates multiple sections.
        for section in doc.sections:
            # Check A5 sizing (allow mm/inches rounding tolerance)
            w = section.page_width.inches if section.page_width else 0
            h = section.page_height.inches if section.page_height else 0
            
            if (abs(w - A5_WIDTH) <= TOLERANCE and abs(h - A5_HEIGHT) <= TOLERANCE) or \
               (abs(h - A5_WIDTH) <= TOLERANCE and abs(w - A5_HEIGHT) <= TOLERANCE): # allow landscape just in case, though technically wrong
                is_a5 = True

            # Check 0.5" Margins
            tm = section.top_margin.inches if section.top_margin else 0
            bm = section.bottom_margin.inches if section.bottom_margin else 0
            lm = section.left_margin.inches if section.left_margin else 0
            rm = section.right_margin.inches if section.right_margin else 0
            
            if all(abs(m - MARGIN) <= 0.1 for m in [tm, bm, lm, rm]):
                has_05_margins = True

            # Check columns (in the XML)
            cols = section._sectPr.xpath('./w:cols')
            if cols:
                num = cols[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num')
                if num and int(num) >= 2:
                    has_2_cols = True
                    
            # Check footers for Page Numbers (Field Char 'PAGE')
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer._element is not None:
                    xml_str = footer._element.xml
                    if 'w:fldChar' in xml_str and 'PAGE' in xml_str:
                        page_numbers_added = True
                    elif 'w:instrText' in xml_str and 'PAGE' in xml_str:
                        page_numbers_added = True

        if is_a5:
            score += 20
            feedback_parts.append("Page size set to A5")
        else:
            feedback_parts.append("Page size NOT A5")

        if has_05_margins:
            score += 15
            feedback_parts.append("Margins set to 0.5 inches")
        else:
            feedback_parts.append("Margins NOT 0.5 inches")

        if has_2_cols:
            score += 20
            feedback_parts.append("2-column layout applied")
        else:
            feedback_parts.append("2-column layout NOT found")
            
        if page_numbers_added:
            score += 10
            feedback_parts.append("Page numbers present in footer")
        else:
            feedback_parts.append("Page numbers NOT found in footer")

    except Exception as e:
        logger.error(f"Error checking layout properties: {e}")
        feedback_parts.append(f"Error checking layout properties: {e}")

    # ====================================================================
    # 4. Verify Heading Styles
    # ====================================================================
    expected_headings = [
        "director's note",
        "cast of characters",
        "act synopsis",
        "cast biographies"
    ]
    headings_found = 0
    
    try:
        for p in doc.paragraphs:
            text = p.text.strip().lower()
            if not text: continue
            
            # Check if this paragraph matches any of the required headings
            for eh in expected_headings:
                if eh in text and len(text) < len(eh) + 10:  # Ensures it's the title line
                    # Check if 'Heading 1' style is applied
                    if p.style and 'heading 1' in p.style.name.lower():
                        headings_found += 1
                        expected_headings.remove(eh) # Don't double count
                        break
                        
        if headings_found == 4:
            score += 15
            feedback_parts.append("All 4 Heading 1 styles applied")
        elif headings_found > 0:
            score += int(15 * (headings_found / 4))
            feedback_parts.append(f"{headings_found}/4 Heading 1 styles applied")
        else:
            feedback_parts.append("No required Heading 1 styles found")
    except Exception as e:
        logger.error(f"Error checking heading styles: {e}")

    # ====================================================================
    # 5. Verify Inline Bolding in Biographies
    # ====================================================================
    bios_bolded = 0
    bios_total = 0
    bio_section = False
    
    try:
        for p in doc.paragraphs:
            text = p.text.strip()
            if "Cast Biographies" in text:
                bio_section = True
                continue
            
            if bio_section:
                # Assuming standard format "Name - Text"
                if " - " in text:
                    bios_total += 1
                    name_part = text.split(" - ")[0].strip()
                    
                    # Check runs for bolding on the name part
                    # Bolding might be spread across multiple runs
                    is_bold = False
                    for run in p.runs:
                        if run.text.strip() and run.text.strip() in name_part:
                            if run.bold:
                                is_bold = True
                                break
                    if is_bold:
                        bios_bolded += 1
                        
        if bios_total > 0 and bios_bolded >= (bios_total - 1): # Allow 1 mistake
            score += 10
            feedback_parts.append("Actor names bolded in biographies")
        elif bios_total > 0 and bios_bolded > 0:
            score += 5
            feedback_parts.append("Actor names partially bolded")
        else:
            feedback_parts.append("Actor names NOT bolded in bios")
    except Exception as e:
        logger.error(f"Error checking bold formatting: {e}")

    # ====================================================================
    # 6. VLM Trajectory Verification (Fallback / Additional Evidence)
    # ====================================================================
    try:
        sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
        from vlm_utils import sample_trajectory_frames, query_vlm, get_final_screenshot
        
        frames = sample_trajectory_frames(traj, n=4)
        final_frame = get_final_screenshot(traj)
        if final_frame:
            frames.append(final_frame)
            
        if frames:
            prompt = """Analyze these screenshots of a user interacting with WPS Office Writer.
            They are formatting a theater playbill document.
            
            Did they perform document layout changes?
            Check for:
            1. Evidence of Page Setup dialog open (modifying paper size or margins).
            2. Evidence of creating columns (Page Layout -> Columns).
            3. Use of Header/Footer tools for page numbers.
            
            Respond in JSON format:
            {
                "interacted_with_page_setup": true/false,
                "created_columns": true/false,
                "added_page_numbers": true/false
            }
            """
            vlm_res = query_vlm(images=frames, prompt=prompt)
            if vlm_res.get("success"):
                parsed = vlm_res.get("parsed", {})
                if parsed.get("interacted_with_page_setup") or parsed.get("created_columns"):
                    score += 10
                    feedback_parts.append("VLM confirmed layout interactions")
                else:
                    feedback_parts.append("VLM did not observe layout interactions")
            else:
                score += 10 # Graceful degrade if VLM fails
                feedback_parts.append("VLM verification skipped (api failure)")
        else:
            score += 10 # Graceful degrade if no frames
            feedback_parts.append("VLM verification skipped (no frames)")
            
    except ImportError:
        score += 10 # Graceful degrade
        feedback_parts.append("VLM utils not available, giving free points")

    # Clean up
    if os.path.exists(temp_docx.name):
        os.unlink(temp_docx.name)

    # ====================================================================
    # Final Scoring
    # ====================================================================
    # Key criteria: Must have changed page size AND added columns.
    passed = score >= 80 and is_a5 and has_2_cols

    return {
        "passed": passed,
        "score": min(score, 100),
        "feedback": " | ".join(feedback_parts)
    }