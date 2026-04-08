#!/usr/bin/env python3
"""
Verifier for format_tech_rider task.
Evaluates WPS Office Writer document formatting: tables, margins, text styles, and bullet points.
"""

import json
import os
import tempfile
import logging
from typing import Dict, Any

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_format_tech_rider(traj, env_info, task_info) -> Dict[str, Any]:
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    feedback_parts = []
    score = 0
    max_score = 100

    # Read result.json
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Criterion 1: File Saved (10 pts)
    if not result.get('output_exists'):
        return {"passed": False, "score": 0, "feedback": "formatted_tech_rider.docx was not saved to the results directory."}
    if not result.get('file_created_during_task'):
        feedback_parts.append("Warning: File timestamp indicates it might not be new.")
    
    score += 10
    feedback_parts.append("File Saved (10/10)")

    # Read the docx
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/formatted_tech_rider.docx", temp_docx.name)
        try:
            from docx import Document
            from docx.shared import RGBColor
            doc = Document(temp_docx.name)
            docx_available = True
        except ImportError:
            docx_available = False
            feedback_parts.append("python-docx not available for deep verification.")
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to load document: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    if not docx_available:
        return {"passed": False, "score": score, "feedback": " | ".join(feedback_parts)}

    # Criterion 2: Page Margins (15 pts) - check if <= 0.6 inches (Narrow is 0.5)
    margin_ok = False
    if len(doc.sections) > 0:
        sec = doc.sections[0]
        # Docx represents narrow margin as ~0.5 inches
        left = sec.left_margin.inches
        right = sec.right_margin.inches
        top = sec.top_margin.inches
        bottom = sec.bottom_margin.inches
        
        if left <= 0.6 and right <= 0.6 and top <= 0.6 and bottom <= 0.6:
            margin_ok = True
            
    if margin_ok:
        score += 15
        feedback_parts.append("Margins Narrow (15/15)")
    else:
        feedback_parts.append("Margins not set to Narrow")

    # Criterion 3: Headings Applied (15 pts)
    expected_headings = [
        "1. AUDIO SPECIFICATIONS",
        "2. LIGHTING & VIDEO",
        "3. BACKLINE REQUIREMENTS",
        "4. INPUT LIST",
        "5. HOSPITALITY"
    ]
    headings_found = 0
    for para in doc.paragraphs:
        text = para.text.strip().upper()
        if para.style and para.style.name.startswith('Heading'):
            for eh in expected_headings:
                if eh in text:
                    headings_found += 1
                    
    # Cap at 5 (in case of duplicates)
    headings_found = min(headings_found, 5)
    heading_pts = int((headings_found / 5.0) * 15)
    score += heading_pts
    if heading_pts == 15:
        feedback_parts.append("Headings applied (15/15)")
    else:
        feedback_parts.append(f"Headings applied ({headings_found}/5)")

    # Criterion 4: Table Creation (20 pts)
    # Target table should have 4 columns and >= 24 rows
    table_created = False
    target_table = None
    for tbl in doc.tables:
        cols = len(tbl.columns)
        rows = len(tbl.rows)
        # Check if it has 4 columns and at least 24 rows (the input list)
        if cols == 4 and rows >= 24:
            table_created = True
            target_table = tbl
            break
            
    if table_created:
        score += 20
        feedback_parts.append("Table created (20/20)")
    else:
        feedback_parts.append("Table not created properly (requires 4 cols and 24+ rows)")

    # Criterion 5: Table Header (15 pts)
    header_ok = False
    if target_table:
        first_row = target_table.rows[0]
        row_text = "".join([cell.text.strip().upper() for cell in first_row.cells])
        if "CH" in row_text and "SOURCE" in row_text and "MIC/DI" in row_text and "NOTES" in row_text:
            # Check bold in the first row's cells
            bold_count = 0
            for cell in first_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold:
                            bold_count += 1
            if bold_count > 0:
                header_ok = True
                
    if header_ok:
        score += 15
        feedback_parts.append("Table Header added and Bold (15/15)")
    else:
        if target_table:
            feedback_parts.append("Table header missing or not bold")

    # Criterion 6: Warning Highlights - CRITICAL: Bold and Red (15 pts)
    warnings_found = 0
    warnings_correct = 0
    for para in doc.paragraphs:
        if "CRITICAL:" in para.text:
            warnings_found += 1
            # We look for the specific run containing "CRITICAL:"
            for run in para.runs:
                if "CRITICAL:" in run.text:
                    is_bold = run.bold
                    is_red = False
                    if run.font.color and run.font.color.rgb:
                        # R is high, G and B are low for red
                        color_hex = str(run.font.color.rgb)
                        if color_hex.startswith('FF') and len(color_hex) == 6:
                            is_red = True
                            
                    if is_bold and is_red:
                        warnings_correct += 1
                        break
                        
    warning_pts = int((min(warnings_correct, 3) / 3.0) * 15)
    score += warning_pts
    if warning_pts == 15:
        feedback_parts.append("Warnings highlighted properly (15/15)")
    else:
        feedback_parts.append(f"Warnings highlighted ({warnings_correct}/3)")

    # Criterion 7: Bulleted List (10 pts)
    # Check if lines starting with "*" in the raw doc have been converted to list elements
    list_items_found = 0
    for para in doc.paragraphs:
        text = para.text.lower()
        if "spring water" in text or "deli tray" in text or "craft beers" in text or "stage towels" in text:
            # Check for xml element w:numPr which indicates a list
            if para._element.xpath('./w:pPr/w:numPr') or (para.style and 'List' in para.style.name):
                list_items_found += 1
                
    list_pts = int((min(list_items_found, 6) / 6.0) * 10)
    score += list_pts
    if list_pts == 10:
        feedback_parts.append("Hospitality list formatted (10/10)")
    else:
        feedback_parts.append(f"Hospitality list formatted ({list_items_found}/6)")

    # VLM Trajectory Verification check (Anti-gaming check)
    # Not strictly adding points, but we verify workflow isn't completely faked
    try:
        from vlm_utils import sample_trajectory_frames, get_final_screenshot, query_vlm
        
        frames = sample_trajectory_frames(traj, n=3)
        final = get_final_screenshot(traj)
        
        vlm_prompt = """
        Review these screenshots from a session using a Word Processor (WPS Writer).
        Did the user actively modify the document? Specifically, did they attempt to:
        1. Insert/format a table
        2. Change text colors/styles (like making "CRITICAL:" red)
        3. Create lists or change margins
        
        Respond with {"work_observed": true} if there's clear visual evidence of work being performed 
        (menus open, text highlighted, table tools visible, document changes from start to end).
        Otherwise respond with {"work_observed": false}.
        """
        
        vlm_result = query_vlm(images=frames + [final] if final else frames, prompt=vlm_prompt)
        vlm_parsed = vlm_result.get("parsed", {})
        
        if not vlm_parsed.get("work_observed", False) and score > 50:
            feedback_parts.append("VLM Warning: Document parsed successfully but trajectory lacks visual evidence of work.")
            # We don't necessarily fail them, but we flag it.
    except Exception as e:
        logger.warning(f"VLM check failed: {e}")

    # Pass Condition
    # Threshold 80 points, and table MUST be created
    key_criteria = table_created
    passed = (score >= 80) and key_criteria

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }