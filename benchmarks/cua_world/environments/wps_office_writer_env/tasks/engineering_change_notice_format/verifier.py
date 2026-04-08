#!/usr/bin/env python3
"""
Verifier for engineering_change_notice_format task.

Verifies that the ECN draft was properly formatted into tables with correct
heading hierarchy and document control header. Uses python-docx for programmatic
validation and VLM for final visual structure validation using trajectory screenshots.
"""

import os
import json
import tempfile
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_engineering_change_notice_format(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    feedback_parts = []
    score = 0
    max_score = 100
    
    # 1. Read task result JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/ecn_task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_info = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    doc_exists = result_info.get('document_exists', False)
    task_start = int(result_info.get('task_start', 0))
    doc_mtime = int(result_info.get('document_mtime', 0))

    if not doc_exists:
        return {"passed": False, "score": 0, "feedback": "Final document ECN-2024-0847_final.docx not found."}
    
    if doc_mtime < task_start:
        return {"passed": False, "score": 0, "feedback": "Document appears to be older than task start time (Anti-gaming check failed)."}
        
    score += 10
    feedback_parts.append("File created correctly")

    # 2. Extract and Parse the DOCX File
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/ECN-2024-0847_final.docx", temp_docx.name)
        
        # We use python-docx to inspect the internal structure
        import docx
        doc = docx.Document(temp_docx.name)
        
        # Check Document Control Header Text
        full_text = "\n".join([p.text.lower() for p in doc.paragraphs])
        header_requirements = [
            "precision surgical devices", 
            "engineering change notice", 
            "ecn-2024-0847", 
            "class ii"
        ]
        headers_found = sum([1 for req in header_requirements if req in full_text])
        if headers_found == 4:
            score += 10
            feedback_parts.append("Document control header present")
        elif headers_found > 0:
            score += 5
            feedback_parts.append(f"Partial document control header ({headers_found}/4)")
        else:
            feedback_parts.append("Missing document control header")

        # Check Heading Hierarchy
        expected_headings = [
            "description of change", "reason for change", "affected documents", 
            "impact assessment", "risk assessment", "implementation plan", "approval signatures"
        ]
        
        h1_texts = []
        for para in doc.paragraphs:
            if para.style and 'Heading 1' in para.style.name:
                h1_texts.append(para.text.lower().strip())
                
        matched_headings = sum([1 for eh in expected_headings if any(eh in h for h in h1_texts)])
        if matched_headings == 7:
            score += 20
            feedback_parts.append("All 7 Heading 1 sections found")
        else:
            score += (matched_headings * 2)
            feedback_parts.append(f"Found {matched_headings}/7 expected Headings")

        # Check Tables
        tables = doc.tables
        if len(tables) >= 5:
            score += 20
            feedback_parts.append(f"Found {len(tables)} tables (expected 5)")
        else:
            score += (len(tables) * 4)
            feedback_parts.append(f"Found {len(tables)} tables (expected 5)")

        # Validate Table Structures (Columns & Rows logic)
        table_validations = 0
        has_bold_headers = False
        
        for table in tables:
            rows = len(table.rows)
            cols = len(table.columns) if rows > 0 else 0
            
            # Identify bold headers in first row
            if rows > 0:
                first_row_bold = any([
                    run.bold 
                    for cell in table.rows[0].cells 
                    for para in cell.paragraphs 
                    for run in para.runs
                ])
                if first_row_bold:
                    has_bold_headers = True

            # Roughly classify and check dimensions
            if cols == 5 and rows >= 6: # Affected documents or Implementation
                table_validations += 1
            elif cols == 3 and rows >= 4: # Impact assessment
                table_validations += 1
            elif cols == 4 and rows >= 3: # Risk assessment or Approvals
                table_validations += 1
                
        # Cap at 5 valid tables
        valid_tables = min(table_validations, 5)
        score += (valid_tables * 3)
        feedback_parts.append(f"Structurally valid tables: {valid_tables}/5")
        
        if has_bold_headers:
            score += 10
            feedback_parts.append("Table headers are bolded")
        else:
            feedback_parts.append("Table headers lack bold formatting")

    except ImportError:
        logger.warning("python-docx not installed on host. Falling back to VLM evaluation.")
        feedback_parts.append("Programmatic inspection skipped (missing python-docx)")
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        feedback_parts.append(f"Failed to parse document: {e}")
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # 3. VLM Verification (Trajectory & Final)
    try:
        from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
        
        frames = sample_trajectory_frames(traj, n=3)
        final_frame = get_final_screenshot(traj)
        
        prompt = (
            "You are grading a Word Processing task. The user was supposed to format an "
            "Engineering Change Notice. Please look at these trajectory and final screenshots.\n"
            "Respond in JSON with the following keys:\n"
            " - 'tables_created': true/false (Are there multiple distinct formatted tables visible?)\n"
            " - 'headings_formatted': true/false (Are section headings visibly styled larger/bolder?)\n"
            " - 'professional_layout': true/false (Does it look like a structured formal report instead of a wall of text?)\n"
        )
        
        vlm_resp = query_vlm(images=frames + [final_frame], prompt=prompt)
        
        if vlm_resp and vlm_resp.get("success"):
            parsed = vlm_resp.get("parsed", {})
            if parsed.get("tables_created"):
                score += 5
            if parsed.get("headings_formatted"):
                score += 5
            if parsed.get("professional_layout"):
                score += 5
            feedback_parts.append("VLM confirmed visual structure")
        else:
            feedback_parts.append("VLM visual verification failed/skipped")
    except Exception as e:
        logger.warning(f"VLM verification error: {e}")
        feedback_parts.append("VLM check bypassed")

    # Score clamping and pass calculation
    score = min(score, max_score)
    passed = score >= 70

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }