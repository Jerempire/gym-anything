#!/usr/bin/env python3
"""
Verifier for ir_press_release_format task.
"""

import os
import json
import logging
import tempfile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import python-docx dynamically
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

def verify_ir_press_release(traj, env_info, task_info):
    """
    Verifies the formatted PR Newswire document.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Framework error: copy_from_env missing"}

    # 1. Read the export JSON result
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read export data: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    output_exists = result.get('output_exists', False)
    file_modified = result.get('file_modified_during_task', False)

    if not output_exists:
        return {
            "passed": False,
            "score": 0,
            "feedback": "Earnings_Release_Formatted.docx was not saved."
        }

    # 2. Extract Document
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    success_doc = False
    doc = None
    try:
        copy_from_env("/tmp/Earnings_Release_Formatted.docx", temp_docx.name)
        if DOCX_AVAILABLE:
            try:
                doc = Document(temp_docx.name)
                success_doc = True
            except Exception as e:
                logger.error(f"Failed to parse docx: {e}")
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    if not success_doc or doc is None:
        return {"passed": False, "score": 10, "feedback": "File exists but could not be parsed as DOCX."}

    # Track scoring
    score = 0
    feedback = []

    # Criterion 1: Anti-gaming and file creation (10 pts)
    if file_modified:
        score += 10
        feedback.append("File properly saved during task.")
    else:
        feedback.append("File was not modified during the task window.")

    # Parse full text to facilitate checks
    paragraphs = doc.paragraphs
    full_text = "\n".join([p.text for p in paragraphs]).strip()

    # Criterion 2: PR Headers and Dateline (10 pts)
    has_for_immediate_release = False
    has_dateline = False
    for p in paragraphs[:5]:  # Look in the first few paragraphs
        text_upper = p.text.strip().upper()
        if "FOR IMMEDIATE RELEASE" in text_upper:
            has_for_immediate_release = True
            # Extra points if bolded (but don't strictly fail if run structure splits it)
            if any(run.bold for run in p.runs if "IMMEDIATE" in run.text.upper()):
                score += 2
        
        if "NEW YORK, NY" in p.text.upper():
            has_dateline = True
            if any(run.bold for run in p.runs if "NEW YORK" in run.text.upper()):
                score += 3
    
    if has_for_immediate_release and has_dateline:
        score += 5
        feedback.append("PR Header and Dateline correctly inserted.")
    else:
        feedback.append(f"Missing header/dateline (Header: {has_for_immediate_release}, Dateline: {has_dateline}).")

    # Criterion 3: Headline Formatting (10 pts)
    headline_found = False
    for p in paragraphs:
        if "Meridian Corp Announces" in p.text:
            headline_found = True
            is_centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
            
            # Check font size and bold in runs
            has_16pt = False
            is_bold = False
            for run in p.runs:
                if run.bold:
                    is_bold = True
                if run.font.size and run.font.size.pt == 16.0:
                    has_16pt = True

            if is_centered: score += 4
            if is_bold: score += 3
            if has_16pt: score += 3
            
            if is_centered and is_bold and has_16pt:
                feedback.append("Headline perfectly formatted (Center, Bold, 16pt).")
            else:
                feedback.append(f"Headline formatting partial (Center:{is_centered}, Bold:{is_bold}, 16pt:{has_16pt}).")
            break
            
    if not headline_found:
        feedback.append("Headline text missing.")

    # Criterion 4: Financial Table Created (30 pts)
    has_table = len(doc.tables) >= 1
    table_correct = False
    if has_table:
        score += 10
        feedback.append("Table created.")
        t = doc.tables[0]
        
        # Check rows/cols
        if len(t.columns) == 3 and len(t.rows) >= 4:
            score += 10
            
            # Look for specific financial metrics in the table
            table_text = "\n".join(["\t".join([cell.text for cell in row.cells]) for row in t.rows]).lower()
            expected = ["revenue", "operating income", "net income", "diluted eps"]
            found_count = sum(1 for metric in expected if metric in table_text)
            
            score += min(10, found_count * 3) # Up to 10 points
            if found_count >= 3:
                table_correct = True
                feedback.append(f"Extracted {found_count}/4 financial metrics into table.")
                
            # Check if prose was deleted
            if "For the fourth quarter, Revenue was" not in full_text:
                feedback.append("Original financial prose successfully removed.")
            else:
                score -= 5
                feedback.append("Original prose NOT removed (penalty).")
        else:
            feedback.append(f"Table dimensions incorrect ({len(t.columns)} cols, {len(t.rows)} rows).")
    else:
        feedback.append("Financial data table was NOT created.")

    # Criterion 5: Boilerplate Heading 2 (10 pts)
    heading_ok = False
    for p in paragraphs:
        if "About Meridian Corp" in p.text:
            if p.style and "Heading 2" in p.style.name:
                heading_ok = True
                break
    if heading_ok:
        score += 10
        feedback.append("'About Meridian Corp' properly formatted as Heading 2.")
    else:
        feedback.append("'About Meridian Corp' missing Heading 2 style.")

    # Criterion 6: Safe Harbor Formatting (Italic, 9pt) (10 pts)
    safe_harbor_ok = False
    for p in paragraphs:
        if "Safe Harbor" in p.text or "forward-looking statements" in p.text.lower():
            # Need to check runs for Italic and 9pt
            italic_runs = 0
            size9_runs = 0
            total_runs = len([r for r in p.runs if r.text.strip()])
            
            for run in p.runs:
                if not run.text.strip(): continue
                if run.italic: italic_runs += 1
                if run.font.size and abs(run.font.size.pt - 9.0) < 0.5:
                    size9_runs += 1
            
            if total_runs > 0:
                if italic_runs >= (total_runs / 2): score += 5
                if size9_runs >= (total_runs / 2): score += 5
                if italic_runs > 0 and size9_runs > 0:
                    safe_harbor_ok = True
                    feedback.append("Safe Harbor text formatting applied (Italics + 9pt).")
            break
            
    if not safe_harbor_ok:
        feedback.append("Safe Harbor text missing correct Italics or 9pt size.")

    # Criterion 7: End Mark (10 pts)
    end_mark_found = False
    for p in reversed(paragraphs):
        if "###" in p.text:
            end_mark_found = True
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                score += 10
                feedback.append("End mark '###' centered at bottom.")
            else:
                score += 5
                feedback.append("End mark '###' found but not centered.")
            break
            
    if not end_mark_found:
        feedback.append("End mark '###' is missing.")

    # Criterion 8: VLM Verification using Trajectory (10 pts)
    # Use trajectory frames to verify the agent actually operated the UI.
    from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
    try:
        frames = sample_trajectory_frames(traj, n=3)
        final_frame = get_final_screenshot(traj)
        if frames and final_frame:
            all_frames = frames + [final_frame]
            vlm_prompt = (
                "You are auditing an agent operating WPS Office Writer. "
                "Look at these screenshots representing the agent's workflow.\n"
                "1. Did the agent actually use the WPS Writer application?\n"
                "2. Do you see a table with financial data inserted into the document?\n"
                "Respond in JSON: {\"used_wps\": true/false, \"table_visible\": true/false}"
            )
            vlm_res = query_vlm(images=all_frames, prompt=vlm_prompt)
            if vlm_res and vlm_res.get("success"):
                parsed = vlm_res.get("parsed", {})
                if parsed.get("used_wps", False):
                    score += 5
                if parsed.get("table_visible", False):
                    score += 5
                feedback.append(f"VLM verification completed. Used WPS: {parsed.get('used_wps')}, Table visible: {parsed.get('table_visible')}.")
            else:
                feedback.append("VLM query failed or format was incorrect.")
        else:
            feedback.append("Could not retrieve trajectory frames for VLM check.")
    except Exception as e:
        logger.warning(f"VLM verification skipped or failed: {e}")
        feedback.append("VLM verification skipped/failed.")

    # Determine Pass/Fail (Threshold 70 and table must be created)
    passed = score >= 70 and table_correct

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }