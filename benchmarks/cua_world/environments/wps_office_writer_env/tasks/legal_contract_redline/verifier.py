#!/usr/bin/env python3
"""Verifier for legal_contract_redline task."""

import sys
import os
import logging

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from wps_verification_utils import (
    copy_and_parse_document,
    cleanup_verification_temp,
    get_document_text,
    check_text_formatting,
    check_paragraph_alignment,
    check_heading_styles,
    count_tables,
    check_table_header_formatting,
    vlm_verify_screenshot,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def verify_legal_contract_redline(traj, env_info, task_info):
    """
    Verify that the legal contract was properly cleaned up.

    PREREQUISITE: Core contract content must be preserved.

    SCORING CRITERIA:
    1. Heading hierarchy: Major sections (1-8) have Heading 1 style
    2. Sub-heading consistency: Section 3 and 6 headers have heading styles (not Normal)
    3. Party name consistency: Informal abbreviations reduced
    4. Definitions completeness: Missing terms added
    5. Table header formatting: Bold + shading on deliverables table
    6. Body alignment: Centered paragraphs fixed to left
    7. Formatting cleanup: Spurious italic/bold removed from 6.2, 8.1, 8.2
    8. Signature block: Present at end of document
    9. VLM visual verification
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    container_path = "/home/ga/Documents/vendor_agreement_draft.docx"
    success, doc, error, temp_dir = copy_and_parse_document(
        container_path, copy_from_env, file_format='docx'
    )

    if not success:
        return {"passed": False, "score": 0, "feedback": error}

    try:
        feedback_parts = []
        full_text = get_document_text(doc).lower()

        # PREREQUISITE: Content must be preserved
        key_phrases = [
            "vendor services agreement",
            "cloudFirst industries",
            "meridian technology solutions",
            "scope of services",
            "limitation of liability",
            "general provisions",
        ]
        preserved = sum(1 for p in key_phrases if p.lower() in full_text)
        if preserved < 4:
            return {
                "passed": False,
                "score": 0,
                "feedback": f"PREREQUISITE FAILED: Content corrupted ({preserved}/{len(key_phrases)} key phrases)",
            }
        feedback_parts.append(f"Prerequisite: content preserved ({preserved}/{len(key_phrases)})")

        criteria_passed = 0
        total_criteria = 9

        # Criterion 1: Heading hierarchy - major sections should have Heading 1
        heading_checks = {
            "definitions": "Heading 1",
            "scope of services": "Heading 1",
            "compensation": "Heading 1",
            "term and termination": "Heading 1",
            "confidentiality": "Heading 1",
            "limitation of liability": "Heading 1",
            "general provisions": "Heading 1",
        }
        matched, total, heading_feedback = check_heading_styles(doc, heading_checks)
        # Sections 1 and 4 were Heading 2 (wrong), sections 3 and 6 were Normal (wrong)
        # If agent fixed at least 5 out of 7 to Heading 1, that's good
        if matched >= 5:
            criteria_passed += 1
            feedback_parts.append(f"Heading hierarchy: {matched}/{total} correct")
        else:
            feedback_parts.append(f"Heading hierarchy: only {matched}/{total} correct (need 5+)")

        # Criterion 2: Section 3 and 6 specifically fixed (were Normal+bold, not headings)
        section3_fixed = False
        section6_fixed = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            if 'deliverables' in text_lower and para.style and 'heading' in para.style.name.lower():
                section3_fixed = True
            if 'confidentiality' in text_lower and '6' in para.text and para.style and 'heading' in para.style.name.lower():
                section6_fixed = True
        if section3_fixed and section6_fixed:
            criteria_passed += 1
            feedback_parts.append("Non-heading sections fixed: both 3 and 6")
        elif section3_fixed or section6_fixed:
            feedback_parts.append("Non-heading sections: only one of 3/6 fixed")
        else:
            feedback_parts.append("Non-heading sections: neither 3 nor 6 fixed to heading style")

        # Criterion 3: Party name consistency
        # Count informal/inconsistent references
        informal_count = 0
        if 'meridian tech' in full_text:
            # Count occurrences but exclude "meridian technology"
            import re
            matches = re.findall(r'meridian tech(?!nology)', full_text)
            informal_count += len(matches)
        # Count all-caps MERIDIAN used alone (not in a heading or legal caps section)
        meridian_caps = full_text.count('meridian shall') + full_text.count('meridian ')
        # Check if the text has been cleaned up - fewer informal references is better
        # Original has 3+ informal refs; if reduced to <=1, that's good
        original_informal = 4  # Meridian Tech (x2) + MERIDIAN (x2) in original
        if informal_count <= 1:
            criteria_passed += 1
            feedback_parts.append(f"Party name consistency: cleaned ({informal_count} informal refs remaining)")
        else:
            feedback_parts.append(f"Party name consistency: {informal_count} informal references still present")

        # Criterion 4: Definitions completeness
        # Check if "Deliverables", "Service Level", and "Term" are defined near "means"
        defs_found = 0
        for term in ["deliverables", "service level", "term"]:
            # Look for pattern: "Term" means... in the definitions section
            pattern = f'"{term}' if term != "term" else '"term"'
            # Simple check: the term appears near "means" in the document
            for i, para in enumerate(doc.paragraphs):
                text_l = para.text.lower()
                if pattern in text_l and 'means' in text_l:
                    defs_found += 1
                    break
        if defs_found >= 2:
            criteria_passed += 1
            feedback_parts.append(f"Definitions: {defs_found}/3 missing terms added")
        else:
            feedback_parts.append(f"Definitions: only {defs_found}/3 missing terms defined")

        # Criterion 5: Table header formatting (bold + shading)
        if count_tables(doc) >= 1:
            header_fmt = check_table_header_formatting(doc, 0)
            if header_fmt['has_bold'] and header_fmt['has_shading']:
                criteria_passed += 1
                feedback_parts.append(f"Table header: bold with shading ({header_fmt['shading_color']})")
            elif header_fmt['has_bold'] or header_fmt['has_shading']:
                feedback_parts.append("Table header: partial formatting (need both bold AND shading)")
            else:
                feedback_parts.append("Table header: no formatting applied")
        else:
            feedback_parts.append("Table: not found in document")

        # Criterion 6: Body alignment - check that formerly centered paragraphs are fixed
        # Section 4.3 was centered; section 6 heading was centered
        centered_issues = 0
        for para in doc.paragraphs:
            text_l = para.text.strip().lower()
            # Check if body paragraphs (not title) are improperly centered
            if len(text_l) > 50 and para.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH as WD
                if para.alignment == WD.CENTER:
                    # Skip the title line
                    if 'vendor services agreement' not in text_l:
                        centered_issues += 1
        if centered_issues == 0:
            criteria_passed += 1
            feedback_parts.append("Body alignment: all correct (no improper centering)")
        else:
            feedback_parts.append(f"Body alignment: {centered_issues} paragraphs still improperly centered")

        # Criterion 7: Formatting cleanup - spurious italic/bold removed
        # Check 6.2 (was all italic), 8.1 (was all italic), 8.2 (was bold+italic)
        formatting_fixed = 0
        checks = [
            ("obligations of confidentiality", False, None),   # 6.2 should NOT be italic
            ("governing law", False, None),                     # 8.1 should NOT be italic
            ("dispute resolution", None, False),                # 8.2 should NOT be bold+italic
        ]
        for text_frag, expected_italic, expected_bold in checks:
            for para in doc.paragraphs:
                if text_frag in para.text.lower():
                    has_issue = False
                    for run in para.runs:
                        if run.text.strip():
                            if expected_italic is False and run.italic:
                                has_issue = True
                            if expected_bold is False and run.bold and run.italic:
                                has_issue = True
                    if not has_issue:
                        formatting_fixed += 1
                    break

        if formatting_fixed >= 2:
            criteria_passed += 1
            feedback_parts.append(f"Formatting cleanup: {formatting_fixed}/3 spurious formats fixed")
        else:
            feedback_parts.append(f"Formatting cleanup: only {formatting_fixed}/3 fixed")

        # Criterion 8: Signature block present
        sig_indicators = ['by:', 'name:', 'title:', 'date:', 'signature', 'authorized']
        sig_found = sum(1 for ind in sig_indicators if ind in full_text[-2000:])
        # Also check for both party names near end
        has_client_sig = 'cloudFirst' in full_text[-2000:].lower() or 'client' in full_text[-2000:]
        has_vendor_sig = 'meridian' in full_text[-2000:].lower() or 'vendor' in full_text[-2000:]

        if sig_found >= 3 and has_client_sig and has_vendor_sig:
            criteria_passed += 1
            feedback_parts.append(f"Signature block: present ({sig_found} indicators, both parties)")
        elif sig_found >= 2:
            feedback_parts.append(f"Signature block: partial ({sig_found} indicators)")
        else:
            feedback_parts.append("Signature block: NOT found")

        # Criterion 9: VLM verification
        vlm_result = vlm_verify_screenshot(env_info, traj, """
Analyze this WPS Writer screenshot of a legal contract. Answer in JSON:
{
    "has_headings": true/false,
    "has_table": true/false,
    "has_signature_block": true/false,
    "appears_professionally_formatted": true/false,
    "has_consistent_structure": true/false
}
Does the document show:
1. Clear section headings with consistent formatting?
2. A formatted table (with header row visually distinct)?
3. A signature block at the bottom?
4. Professional legal document formatting overall?
5. Consistent visual structure throughout?
""")
        if vlm_result is not None:
            has_headings = vlm_result.get("has_headings", False)
            has_structure = vlm_result.get("has_consistent_structure", False)
            appears_pro = vlm_result.get("appears_professionally_formatted", False)

            if (has_headings or has_structure) and appears_pro:
                criteria_passed += 1
                feedback_parts.append("VLM: professional formatting confirmed")
            else:
                feedback_parts.append("VLM: professional formatting not confirmed")
        else:
            total_criteria -= 1
            feedback_parts.append("VLM: unavailable (skipped)")

        score = int((criteria_passed / total_criteria) * 100)
        passed = score >= 55

        return {
            "passed": passed,
            "score": score,
            "feedback": " | ".join(feedback_parts),
        }

    except Exception as e:
        logger.error(f"Verification error: {e}", exc_info=True)
        return {"passed": False, "score": 0, "feedback": f"Verification error: {str(e)}"}
    finally:
        cleanup_verification_temp(temp_dir)
