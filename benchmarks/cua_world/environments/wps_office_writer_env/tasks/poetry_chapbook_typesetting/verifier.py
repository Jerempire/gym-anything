#!/usr/bin/env python3
"""
Verifier for poetry_chapbook_typesetting task.

Verifies:
1. File exists and was created during the task
2. Page Dimensions (5.5" x 8.5")
3. Margins (L: 1.0", R: 0.5", T/B: 0.75")
4. Styles (Title, Heading 1)
5. TOC Generation
6. Pagination
7. VLM verification of Trajectory
"""

import json
import os
import tempfile
import logging

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import VLM utils
try:
    from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
    VLM_AVAILABLE = True
except ImportError:
    VLM_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Helper to check EMU closeness (Word uses EMUs: 1 Inch = 914400 EMUs)
def is_close(expected_inches, actual_emus, tolerance_inches=0.1):
    if actual_emus is None:
        return False
    actual_inches = actual_emus / 914400.0
    return abs(expected_inches - actual_inches) <= tolerance_inches

def verify_poetry_chapbook_typesetting(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}
        
    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "python-docx not installed on host."}

    metadata = task_info.get('metadata', {})
    poem_titles = metadata.get('poem_titles', [])

    # Retrieve the result json
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    result = {}
    try:
        copy_from_env("/tmp/poetry_chapbook_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    doc_exists = result.get('document_exists', False)
    task_start = int(result.get('task_start', 0))
    doc_mtime = int(result.get('document_mtime', 0))
    
    if not doc_exists:
        return {"passed": False, "score": 0, "feedback": "Print-ready document was not saved to the specified path."}
    
    if doc_mtime < task_start:
        return {"passed": False, "score": 0, "feedback": "Document was not modified during the task execution (Anti-gaming check failed)."}

    # Copy the actual DOCX for parsing
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/tmp/yeats_chapbook_print.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 10  # 10 pts for saving the file successfully and passing anti-gaming
    feedback_parts = ["File saved and modified"]
    
    # 1. Page Dimensions & Margins
    dimensions_correct = False
    margins_correct = False
    if len(doc.sections) > 0:
        sec = doc.sections[0]
        # Dimensions: 5.5 x 8.5
        w_ok = is_close(5.5, sec.page_width)
        h_ok = is_close(8.5, sec.page_height)
        if w_ok and h_ok:
            score += 20
            dimensions_correct = True
            feedback_parts.append("Page dimensions set to 5.5x8.5")
        else:
            feedback_parts.append(f"Dimensions incorrect (Expected 5.5x8.5, got approx {sec.page_width/914400.0:.2f}x{sec.page_height/914400.0:.2f})")
            
        # Margins: L:1.0, R:0.5, T/B:0.75
        l_ok = is_close(1.0, sec.left_margin)
        r_ok = is_close(0.5, sec.right_margin)
        t_ok = is_close(0.75, sec.top_margin)
        b_ok = is_close(0.75, sec.bottom_margin)
        
        if l_ok and r_ok and t_ok and b_ok:
            score += 20
            margins_correct = True
            feedback_parts.append("Margins correct (Asymmetrical binding setup)")
        else:
            feedback_parts.append("Margins incorrect")
            if (l_ok or r_ok) and (t_ok or b_ok):
                score += 10 # Partial margin credit

    # 2. Styles Check
    title_styled = False
    headings_styled = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name.lower() if para.style else ""
        
        if "the wild swans at coole" in text.lower() and "title" in style_name:
            title_styled = True
            
        for pt in poem_titles:
            if pt.lower() == text.lower() and "heading 1" in style_name:
                headings_styled += 1

    if title_styled:
        score += 10
        feedback_parts.append("Title style applied")
        
    if headings_styled >= 3: # Allowance for a few missed
        score += 15
        feedback_parts.append(f"Heading 1 applied to poem titles ({headings_styled}/{len(poem_titles)})")
    elif headings_styled > 0:
        score += 5
        feedback_parts.append(f"Partial Heading 1 application ({headings_styled}/{len(poem_titles)})")
        
    # 3. TOC Check
    has_toc = False
    xml_str = str(doc._element.xml)
    if 'w:instrText' in xml_str and 'TOC' in xml_str:
        has_toc = True
    else:
        # Fallback: check for standard TOC styles if fields aren't updated
        for para in doc.paragraphs:
            if para.style and "toc" in para.style.name.lower():
                has_toc = True
                break

    if has_toc:
        score += 15
        feedback_parts.append("TOC inserted")
    else:
        if "[insert table of contents here]" not in xml_str.lower():
            feedback_parts.append("TOC placeholder removed but no active TOC field found")
        else:
            feedback_parts.append("TOC not generated")

    # 4. Pagination / Footer check
    has_pagination = False
    for sec in doc.sections:
        if sec.footer:
            try:
                footer_xml = str(sec.footer._element.xml)
                if 'w:instrText' in footer_xml and ('PAGE' in footer_xml or 'NUMPAGES' in footer_xml):
                    has_pagination = True
                    break
            except:
                pass

    if has_pagination:
        score += 10
        feedback_parts.append("Page numbers present in footer")

    # 5. VLM Trajectory Verification
    # Ensure they actually worked in WPS Writer, and didn't just python-script the solution from a terminal
    vlm_passed = False
    if VLM_AVAILABLE and traj:
        try:
            frames = sample_trajectory_frames(traj, n=4)
            final = get_final_screenshot(traj)
            images = frames + [final] if final else frames
            
            prompt = """
            Review these trajectory frames of an agent interacting with a word processor.
            1. Is the agent actively using WPS Writer?
            2. Is there evidence of the agent modifying page layouts (margins/size), opening style menus, or inserting a Table of Contents?
            3. Did the agent perform genuine desktop publishing work through the GUI?
            
            Respond in JSON:
            {"gui_interaction": true/false, "layout_work_visible": true/false}
            """
            
            vlm_res = query_vlm(images=images, prompt=prompt)
            if vlm_res.get("success"):
                parsed = vlm_res.get("parsed", {})
                if parsed.get("gui_interaction") and parsed.get("layout_work_visible"):
                    vlm_passed = True
        except Exception as e:
            logger.error(f"VLM verification failed: {e}")

    if not vlm_passed:
        # We don't deduct points, but we note it.
        feedback_parts.append("VLM visual confirmation inconclusive/failed")

    # Final Evaluation
    # Must have correct page dimensions and a TOC to pass as "Typesetting"
    key_criteria_met = dimensions_correct and has_toc
    passed = (score >= 75) and key_criteria_met

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }