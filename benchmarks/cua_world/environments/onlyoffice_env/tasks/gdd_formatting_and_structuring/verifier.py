#!/usr/bin/env python3
"""
Verifier for GDD Formatting and Structuring Task.

Parses the exported `Eldoria_GDD.docx` file's DOM to verify:
1. File creation (15 pts)
2. Title alignment and styling (10 pts)
3. Heading 1 styles (20 pts)
4. Heading 2 styles (15 pts)
5. Table conversion (20 pts)
6. Table Header bolding (10 pts)
7. Italics on keybinds (10 pts)

Uses VLM trajectory analysis as a supplementary check.
"""

import json
import os
import tempfile
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try importing docx, required for parsing ONLYOFFICE DOM
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.error("python-docx is not installed.")

def verify_gdd_formatting(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Load result metadata
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    score = 0
    feedback_parts = []

    # 1. Check file existence & integrity
    if not result.get('file_exists'):
        return {"passed": False, "score": 0, "feedback": "Eldoria_GDD.docx was not found in the target directory."}
    if not result.get('file_created_during_task'):
        return {"passed": False, "score": 0, "feedback": "File exists but was not created/modified during the task window."}

    score += 15
    feedback_parts.append("File created successfully")

    # Load the DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/TextDocuments/Eldoria_GDD.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse DOCX file: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # Variables for DOM parsing
    title_aligned = False
    h1_count = 0
    h2_count = 0
    table_exists = len(doc.tables) > 0
    header_bold = False
    
    # 2-4. Parse Paragraphs for Headings, Alignments, and Italics
    italicized_text = ""
    for p in doc.paragraphs:
        # Check Title Alignment (1 = CENTER in WD_ALIGN_PARAGRAPH)
        if 'Eldoria: Awakening' in p.text:
            if p.alignment == 1 or (p.style and 'Title' in p.style.name):
                title_aligned = True
                
        # Count Headings
        if p.style and p.style.name:
            if p.style.name.startswith('Heading 1'):
                h1_count += 1
            elif p.style.name.startswith('Heading 2'):
                h2_count += 1
                
        # Collect Italicized text for later evaluation
        for run in p.runs:
            if run.italic:
                italicized_text += run.text

    # Scoring structural formatting
    if title_aligned:
        score += 10
        feedback_parts.append("Title aligned")
    else:
        feedback_parts.append("Title alignment missing")

    if h1_count >= 5:
        score += 20
        feedback_parts.append("Heading 1s applied")
    else:
        feedback_parts.append(f"Found {h1_count}/5 Heading 1s")

    if h2_count >= 4:
        score += 15
        feedback_parts.append("Heading 2s applied")
    else:
        feedback_parts.append(f"Found {h2_count}/4 Heading 2s")

    # 5. Check Table Creation
    if table_exists:
        # Verify it's the correct table by checking contents of a cell
        table = doc.tables[0]
        try:
            cell_text = table.rows[1].cells[0].text
            if 'Spellblade' in cell_text or 'Class' in table.rows[0].cells[0].text:
                score += 20
                feedback_parts.append("Data table created")
                
                # 6. Check Table Header Bolding
                header_bold = True
                for cell in table.rows[0].cells:
                    # Text inside cells can be split into multiple runs
                    cell_is_bold = any(run.bold for p in cell.paragraphs for run in p.runs)
                    if not cell_is_bold and cell.text.strip():
                        header_bold = False
                
                if header_bold:
                    score += 10
                    feedback_parts.append("Table header bolded")
                else:
                    feedback_parts.append("Table header not fully bolded")
            else:
                feedback_parts.append("Table created, but missing expected character data")
        except IndexError:
            feedback_parts.append("Table created but has improper dimensions")
    else:
        feedback_parts.append("Data table was not created")

    # 7. Check Italics for keybinds
    keybinds = task_info.get('metadata', {}).get('keybinds', ["Spacebar", "Left Shift", "F", "Left Mouse Button", "Right Mouse Button", "Q"])
    italics_found = 0
    for kb in keybinds:
        if kb in italicized_text:
            italics_found += 1
            
    if italics_found == len(keybinds):
        score += 10
        feedback_parts.append("All keybinds italicized")
    elif italics_found > 0:
        score += 5
        feedback_parts.append(f"Partial keybinds italicized ({italics_found}/{len(keybinds)})")
    else:
        feedback_parts.append("Keybinds not italicized")

    # Final Evaluation (Pass threshold is 65 points)
    passed = score >= 65

    # Optional VLM Trajectory Check (to detect UI interaction vs macro/script cheating)
    # We load it but don't strictly require it to fail the task if DOM parsing is perfect, 
    # but we can log it for trajectory auditing.
    try:
        from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
        frames = sample_trajectory_frames(traj, n=3)
        if frames:
            vlm_prompt = "Looking at these frames, did the user actively interact with the top ribbon menu (Styles, Tables, Formatting) in the word processor? Answer 'yes' or 'no'."
            vlm_res = query_vlm(images=frames, prompt=vlm_prompt)
            if vlm_res and 'yes' in vlm_res.get('response', '').lower():
                feedback_parts.append("VLM confirmed UI interaction")
    except Exception as e:
        logger.warning(f"VLM trajectory check skipped/failed: {e}")

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }