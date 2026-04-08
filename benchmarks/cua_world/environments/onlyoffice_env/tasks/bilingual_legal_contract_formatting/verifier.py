#!/usr/bin/env python3
"""
Verifier for Bilingual Legal Contract Formatting task.

Scores (100 pts total, Pass Threshold: 70):
1. Document Created/Saved: 10 pts
2. Font Family Standardized (Times New Roman): 15 pts
3. Title Formatting (Center, Bold, 14pt): 15 pts
4. Section Headers (Bold, 12pt): 15 pts
5. Defined Terms (Bold, Italic): 15 pts
6. Signature Alignment (Right-aligned): 15 pts
7. VLM Trajectory (UI Formatting Activity): 15 pts
"""

import sys
import os
import json
import tempfile
import logging
import re
import subprocess

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure python-docx is available
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
    import docx


def extract_text_with_styles(doc):
    """
    Extracts text down to the character level to immune parsing against
    arbitrary python-docx run splits.
    """
    all_chars = []
    
    # Fallback default if not explicitly set
    default_font = 'Arial'
    if 'Normal' in doc.styles and doc.styles['Normal'].font.name:
        default_font = doc.styles['Normal'].font.name
        
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paragraphs.append(p)
                    
    for p in paragraphs:
        align = p.alignment
        p_text = p.text
        if not p_text.strip():
            continue
            
        char_styles = []
        for run in p.runs:
            font_name = run.font.name if run.font.name else default_font
            size = run.font.size.pt if run.font.size else None
            bold = run.bold if run.bold is not None else False
            italic = run.italic if run.italic is not None else False
            
            for c in run.text:
                char_styles.append({
                    'char': c,
                    'font': font_name,
                    'size': size,
                    'bold': bold,
                    'italic': italic
                })
        
        # Verify length matches up safely
        if len(char_styles) == len(p_text):
            all_chars.append({
                'text': p_text,
                'align': align,
                'char_styles': char_styles
            })
            
    return all_chars


def verify_bilingual_contract(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    query_vlm = env_info.get('query_vlm')
    
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    score = 0
    feedback_parts = []
    
    # 1. Read result JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    output_exists = result.get('output_exists', False)
    file_modified = result.get('file_modified', False)
    exact_name = result.get('exact_name', False)
    file_path = result.get('file_path', '')

    if not output_exists:
        return {
            "passed": False, 
            "score": 0, 
            "feedback": "Document not saved. File is missing."
        }

    # CRITERION 1: File Existence / Name
    if exact_name and file_modified:
        score += 10
        feedback_parts.append("Final document saved correctly")
    elif file_modified:
        score += 5
        feedback_parts.append("Draft document overwritten instead of 'Save As'")
    else:
        feedback_parts.append("Document not modified during task (anti-gaming)")

    # 2. Parse DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(file_path, temp_docx.name)
        doc = docx.Document(temp_docx.name)
        all_chars = extract_text_with_styles(doc)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse docx: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # CRITERION 2: Base Font (Times New Roman) - 15 pts
    total_chars = sum(len(p['char_styles']) for p in all_chars)
    tnr_chars = sum(1 for p in all_chars for c in p['char_styles'] if c['font'] == 'Times New Roman')
    if total_chars > 0 and (tnr_chars / total_chars) > 0.8:
        score += 15
        feedback_parts.append("Font standardized to Times New Roman")
    else:
        feedback_parts.append("Font standardization incomplete")

    # CRITERION 3: Title Formatting - 15 pts
    titles_found = 0
    titles_correct = 0
    for p in all_chars:
        if "MUTUAL NON-DISCLOSURE AGREEMENT" in p['text'] or "ACUERDO DE CONFIDENCIALIDAD MUTUA" in p['text']:
            titles_found += 1
            if p['align'] in (1, 1.0, 'CENTER', 'center'):
                chars = p['char_styles']
                bold_pct = sum(1 for c in chars if c['bold']) / len(chars)
                size_pct = sum(1 for c in chars if c['size'] == 14.0) / len(chars)
                if bold_pct > 0.8 and size_pct > 0.8:
                    titles_correct += 1
    
    if titles_found > 0 and titles_correct == titles_found:
        score += 15
        feedback_parts.append("Title formatting correct (Center, Bold, 14pt)")
    elif titles_correct > 0:
        score += 7
        feedback_parts.append("Title formatting partially correct")

    # CRITERION 4: Section Headers - 15 pts
    headers_found = 0
    headers_correct = 0
    for p in all_chars:
        text = p['text'].strip()
        if text.startswith("1. ") or text.startswith("2. "):
            # Split out just the header line if it shares a paragraph block
            lines = text.split('\n')
            header_text = lines[0]
            match_len = len(header_text)
            headers_found += 1
            
            header_chars = p['char_styles'][:match_len]
            bold_pct = sum(1 for c in header_chars if c['bold']) / match_len
            size_pct = sum(1 for c in header_chars if c['size'] == 12.0) / match_len
            
            if bold_pct > 0.8 and size_pct > 0.8:
                headers_correct += 1

    if headers_found > 0 and headers_correct >= 2:
        score += 15
        feedback_parts.append("Section headers formatting correct (Bold, 12pt)")
    elif headers_correct > 0:
        score += 7
        feedback_parts.append("Section headers partially formatted")

    # CRITERION 5: Defined Terms - 15 pts
    terms_correct = 0
    terms_checked = 0
    for p in all_chars:
        text = p['text']
        for term in ["Confidential Information", "Información Confidencial"]:
            for match in re.finditer(re.escape(term), text):
                # Skip header occurrences
                if match.start() < 5 and text.startswith("1. "):
                    continue
                    
                terms_checked += 1
                term_chars = p['char_styles'][match.start():match.end()]
                bold_pct = sum(1 for c in term_chars if c['bold']) / len(term_chars)
                italic_pct = sum(1 for c in term_chars if c['italic']) / len(term_chars)
                
                if bold_pct > 0.8 and italic_pct > 0.8:
                    terms_correct += 1

    if terms_checked > 0 and terms_correct >= 2:
        score += 15
        feedback_parts.append("Defined terms formatted correctly (Bold, Italic)")
    elif terms_correct > 0:
        score += 7
        feedback_parts.append("Defined terms partially formatted")

    # CRITERION 6: Signature Alignment - 15 pts
    sigs_found = 0
    sigs_correct = 0
    for p in all_chars:
        text = p['text'].strip()
        if text.startswith("Date:") or text.startswith("Fecha:"):
            sigs_found += 1
            if p['align'] in (2, 2.0, 'RIGHT', 'right'):
                sigs_correct += 1

    if sigs_found > 0 and sigs_correct == sigs_found:
        score += 15
        feedback_parts.append("Signature alignment correct (Right)")
    elif sigs_correct > 0:
        score += 7
        feedback_parts.append("Signature alignment partially correct")

    # CRITERION 7: VLM Trajectory Check (Anti-gaming) - 15 pts
    if query_vlm:
        from gym_anything.vlm import sample_trajectory_frames
        frames = sample_trajectory_frames(traj, n=4)
        if frames:
            prompt = """Analyze these trajectory frames of an agent formatting a bilingual legal contract in ONLYOFFICE Document Editor.
            
            1. Is the agent actively using the ONLYOFFICE interface?
            2. Do you see the agent selecting text and using the formatting toolbar (e.g., changing fonts to Times New Roman, clicking Bold/Italic, or changing alignment)?
            
            Respond with JSON:
            {
                "used_onlyoffice": true/false,
                "applied_formatting": true/false,
                "confidence": "high"
            }"""
            
            vlm_result = query_vlm(images=frames, prompt=prompt)
            if vlm_result and vlm_result.get('success'):
                parsed = vlm_result.get('parsed', {})
                if parsed.get('used_onlyoffice') and parsed.get('applied_formatting'):
                    score += 15
                    feedback_parts.append("VLM confirmed ONLYOFFICE formatting activity")
                else:
                    feedback_parts.append("VLM did not confirm UI formatting activity")

    passed = score >= 70
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }