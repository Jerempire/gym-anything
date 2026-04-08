#!/usr/bin/env python3
"""
Verifier for ISO 9001 SOP Document Formatting task.
"""

import json
import tempfile
import os
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_iso9001_sop_formatting(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available. Framework error."}
        
    # Copy evaluation metadata
    temp_result = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result.name)
        with open(temp_result.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_result.name):
            os.unlink(temp_result.name)
            
    output_exists = result.get('output_exists', False)
    file_created_during_task = result.get('file_created_during_task', False)
    
    if not output_exists:
        return {"passed": False, "score": 0, "feedback": "Final document 'QA_SOP_08_Final.docx' was not found."}
        
    if not file_created_during_task:
        return {"passed": False, "score": 0, "feedback": "Anti-gaming failure: Final document was not created/modified during the task timeframe."}

    # Extract the resultant DOCX file for deep inspection
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/TextDocuments/QA_SOP_08_Final.docx", temp_docx.name)
        import docx
        try:
            doc = docx.Document(temp_docx.name)
        except Exception as e:
            return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX file: {e}"}
            
        score = 10  # Base points for successfully creating the file
        feedback_parts = ["File created"]
        
        # 1. Check Heading Styles (20 pts)
        h1_found = False
        h2_count = 0
        h2_targets = ["Purpose", "Scope", "Procedure", "Revision"]
        
        for p in doc.paragraphs:
            if "Heading 1" in p.style.name and "SOP-08" in p.text:
                h1_found = True
            if "Heading 2" in p.style.name:
                if any(t in p.text for t in h2_targets):
                    h2_count += 1
                    
        if h1_found and h2_count >= 4:
            score += 20
            feedback_parts.append("Heading styles perfectly mapped")
        elif h1_found or h2_count > 0:
            score += 10
            feedback_parts.append("Heading styles partially mapped")
        else:
            feedback_parts.append("Missing required heading styles")
            
        # 2. Check Body Font: Arial 11pt (15 pts)
        font_correct = False
        for p in doc.paragraphs:
            # Look at specific content paragraphs to assess font logic
            if "purpose of this procedure" in p.text.lower() or "medical devices manufactured" in p.text.lower():
                # Direct formatting on runs takes precedence
                for run in p.runs:
                    if run.text.strip():
                        font_name = run.font.name or p.style.font.name
                        font_size = run.font.size or p.style.font.size
                        if font_name == 'Arial' and font_size and font_size.pt == 11.0:
                            font_correct = True
                            break
                # Fallback checking paragraph-level style attributes if no run overrides are found
                if not font_correct:
                    font_name = p.style.font.name
                    font_size = p.style.font.size
                    if font_name == 'Arial' and font_size and font_size.pt == 11.0:
                        font_correct = True
        
        if font_correct:
            score += 15
            feedback_parts.append("Body font confirmed as Arial 11pt")
        else:
            feedback_parts.append("Body font incorrect/missing")
            
        # 3. Check Warning Formatting: Bold & Colored (20 pts)
        warning_bold = False
        warning_colored = False
        for p in doc.paragraphs:
            if "WARNING:" in p.text:
                for run in p.runs:
                    if "WARNING" in run.text or run.bold:
                        if run.bold:
                            warning_bold = True
                        if run.font.color and (run.font.color.rgb or run.font.color.type is not None):
                            # Accepts any directly-applied color override as indicating intent
                            warning_colored = True
                break
                
        if warning_bold and warning_colored:
            score += 20
            feedback_parts.append("Warning formatted (bold and colored)")
        elif warning_bold or warning_colored:
            score += 10
            feedback_parts.append("Warning partially formatted")
        else:
            feedback_parts.append("Warning formatting missing")
            
        # 4. Check Table Conversion (20 pts)
        table_found = False
        if len(doc.tables) >= 1:
            for tbl in doc.tables:
                if len(tbl.columns) == 4 and len(tbl.rows) >= 4:
                    text_content = ""
                    for row in tbl.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                    if "Rev" in text_content and "Author" in text_content:
                        table_found = True
                        break
                        
        if table_found:
            score += 20
            feedback_parts.append("Revision history table converted correctly")
        else:
            feedback_parts.append("Table conversion missing or column/row counts incorrect")
            
        # 5. Check Document Header (15 pts)
        header_found = False
        for section in doc.sections:
            if section.header is not None:
                for p in section.header.paragraphs:
                    if "ISO 9001:2015" in p.text:
                        header_found = True
                        break
            if header_found:
                break
                
        if header_found:
            score += 15
            feedback_parts.append("Document header added correctly")
        else:
            feedback_parts.append("Document header missing")
            
        # Final evaluation
        passed = score >= 70 and table_found and h1_found
        
        return {
            "passed": passed,
            "score": score,
            "feedback": " | ".join(feedback_parts)
        }
        
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)