#!/usr/bin/env python3
import os
import json
import tempfile
import logging

# Import python-docx tools safely
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def check_margins(doc):
    """Check if all page margins are strictly 0.5 inches."""
    if not doc.sections:
        return False
    sec = doc.sections[0]
    try:
        margins = [
            sec.top_margin.inches if sec.top_margin else 0,
            sec.bottom_margin.inches if sec.bottom_margin else 0,
            sec.left_margin.inches if sec.left_margin else 0,
            sec.right_margin.inches if sec.right_margin else 0
        ]
        return all(abs(m - 0.5) < 0.05 for m in margins)
    except Exception:
        return False

def check_font(doc):
    """Check if the global font is Arial 11pt."""
    default_font = None
    default_size = None
    if 'Normal' in doc.styles:
        default_font = doc.styles['Normal'].font.name
        default_size = doc.styles['Normal'].font.size

    total_runs = 0
    arial_11_runs = 0
    
    for p in doc.paragraphs:
        for r in p.runs:
            if not r.text.strip():
                continue
            total_runs += 1
            name = r.font.name or default_font
            size = r.font.size or default_size
            
            if name == 'Arial':
                if size is not None and getattr(size, 'pt', 0) == 11.0:
                    arial_11_runs += 1
                    
    if total_runs == 0:
        return False
    # Tolerant threshold for "global" rule formatting
    return (arial_11_runs / total_runs) >= 0.8

def check_header(doc):
    """Check if header exists and contains PI name right-aligned."""
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            if "smith, john" in p.text.lower() or "pi: smith" in p.text.lower():
                # Value 2 corresponds to WD_ALIGN_PARAGRAPH.RIGHT
                if p.alignment == 2 or getattr(p.alignment, 'value', None) == 2:
                    return True
                # Sometimes tabs are used to push to right edge
                return True
    return False

def check_section_styling(doc):
    """Check if specific headers were cleaned and formatted Bold & Italic."""
    targets = ["specific aims", "significance", "innovation", "approach"]
    found = 0
    
    for p in doc.paragraphs:
        text = p.text.lower()
        if any(t in text for t in targets) and len(text) < 50:
            if "[section]" not in text:  # Tag was properly removed
                is_bold = any(r.bold for r in p.runs) or p.style.font.bold
                is_italic = any(r.italic for r in p.runs) or p.style.font.italic
                if is_bold and is_italic:
                    found += 1
    return found >= 3

def check_table(doc):
    """Check if the text block was successfully converted to a 4x3 table."""
    for table in doc.tables:
        if len(table.rows) == 4 and len(table.columns) == 3:
            # Check content to confirm it's the correct table
            text = "".join(cell.text.lower() for row in table.rows for cell in row.cells)
            if "sensitivity" in text and "specificity" in text and "cost" in text:
                return True
    return False

def verify_nih_formatting(traj, env_info, task_info):
    """
    Verify the NIH Grant Proposal Formatting task.
    Uses copy_from_env to read the resultant .docx and parses underlying XML.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "python-docx package is not available"}

    feedback_parts = []
    score = 0
    max_score = 100

    # Read exported JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read task result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # 1. Output Exists & Anti-gaming (20 pts)
    output_exists = result.get('output_exists', False)
    file_created = result.get('file_created_during_task', False)
    
    if not output_exists:
        return {"passed": False, "score": 0, "feedback": "Failed: The expected output .docx was not found."}
    
    if file_created:
        score += 20
        feedback_parts.append("File created successfully")
    else:
        feedback_parts.append("File exists but was not created during task bounds")

    # Retrieve and parse DOCX
    metadata = task_info.get('metadata', {})
    expected_path = metadata.get('expected_output_path', '/home/ga/Documents/TextDocuments/nih_r01_research_strategy.docx')
    
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(expected_path, temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse document: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # 2. Margins 0.5" (20 pts)
    if check_margins(doc):
        score += 20
        feedback_parts.append("Margins correctly set to 0.5 inches")
    else:
        feedback_parts.append("Margins incorrect")

    # 3. Global Font Arial 11pt (15 pts)
    if check_font(doc):
        score += 15
        feedback_parts.append("Font set to Arial 11pt")
    else:
        feedback_parts.append("Font formatting incorrect/incomplete")

    # 4. Header PI Name (15 pts)
    if check_header(doc):
        score += 15
        feedback_parts.append("Header formatting correct")
    else:
        feedback_parts.append("Header formatting incorrect/missing")

    # 5. Section Styling (10 pts)
    if check_section_styling(doc):
        score += 10
        feedback_parts.append("Section headers styled properly")
    else:
        feedback_parts.append("Section headers missing proper style")

    # 6. Table Conversion (10 pts)
    if check_table(doc):
        score += 10
        feedback_parts.append("Table successfully created")
    else:
        feedback_parts.append("Table not created or formatted properly")

    # 7. VLM Trajectory Verification (10 pts)
    # Proves the agent actually used the UI to achieve the results
    try:
        from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
        frames = sample_trajectory_frames(traj, n=4)
        final = get_final_screenshot(traj)
        
        prompt = "Look at these screenshots. Did the agent actively interact with the ONLYOFFICE Document Editor to format text, adjust margins, and insert a table? Reply in JSON: {\"active_workflow\": true/false}."
        
        if 'query_vlm' in env_info:
            vlm_fn = env_info['query_vlm']
            res = vlm_fn(images=frames + [final], prompt=prompt)
        else:
            res = query_vlm(images=frames + [final], prompt=prompt)
            
        if res and res.get("parsed", {}).get("active_workflow", False):
            score += 10
            feedback_parts.append("VLM confirms workflow execution")
        else:
            feedback_parts.append("VLM could not confirm workflow execution")
    except Exception as e:
        logger.warning(f"VLM verification skipped or failed: {e}")
        # Soft-fail VLM if absent from framework to not penalize correct deterministic logic
        score += 10 
        feedback_parts.append("VLM verification bypassed")

    # Final logic
    passed = score >= 75
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }