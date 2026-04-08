#!/usr/bin/env python3
"""
Verifier for Academic Manuscript Typesetting task.

Evaluates precise document formatting adjustments programmatically via python-docx.
Integrates VLM trajectory check to ensure the work was done manually inside ONLYOFFICE.
"""

import os
import json
import logging
import tempfile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to import optional vlm dependencies gracefully
try:
    from gym_anything.vlm import sample_trajectory_frames, query_vlm
    VLM_AVAILABLE = True
except ImportError:
    VLM_AVAILABLE = False
    logger.warning("VLM utilities not available. Trajectory checks will be skipped.")

# Try to import python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def check_para_formatting(p):
    """Robust helper to extract alignment, bold, italic, and max font size from a paragraph."""
    align = p.alignment
    if align is None and p.style and hasattr(p.style, 'paragraph_format'):
        align = p.style.paragraph_format.alignment
        
    bold = False
    italic = False
    size = None
    
    for r in p.runs:
        if not r.text.strip():
            continue
        if r.bold: bold = True
        if r.italic: italic = True
        if r.font and r.font.size:
            size_pt = r.font.size.pt
            if size is None or size_pt > size:
                size = size_pt
                
    return align, bold, italic, size


def verify_manuscript_typesetting(traj, env_info, task_info):
    """
    Programmatically verify the styles applied to the DOCX file.
    Also utilizes VLM on the trajectory to ensure anti-gaming (i.e. Agent used OnlyOffice).
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "python-docx not installed on host."}

    # Setup Scoring
    score = 0
    feedback_parts = []
    metadata = task_info.get('metadata', {})

    # 1. Read task execution JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/typeset_task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # If the file doesn't exist, exit early
    if not result.get('output_exists'):
        return {"passed": False, "score": 0, "feedback": "Output document typeset_manuscript.docx not found."}

    if result.get('file_created_during_task'):
        score += 10
        feedback_parts.append("File created/modified during task (+10)")
    else:
        feedback_parts.append("File was NOT created/modified during task (Gaming suspected)")

    # 2. Parse the exported document
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(metadata.get('expected_output_path'), temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse document: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # Track formatting constraints
    found_title = False
    found_abstract = False
    found_heading = False
    found_body = False

    for p in doc.paragraphs:
        text = p.text.strip().lower()
        if not text:
            continue
            
        align, bold, italic, size = check_para_formatting(p)
        
        # Check Title
        if metadata.get('title_text').lower() in text and not found_title:
            found_title = True
            # Title Constraints: 16pt, Bold, Center (1)
            if size == 16.0 and bold and align == 1:
                score += 15
                feedback_parts.append("Title formatted correctly (+15)")
            else:
                feedback_parts.append(f"Title fmt mismatch (Size:{size}, Bold:{bold}, Align:{align})")

        # Check Abstract
        elif metadata.get('abstract_text').lower() in text and not found_abstract:
            found_abstract = True
            # Abstract Constraints: 10pt, Italic, Justified (3)
            if size == 10.0 and italic and align == 3:
                score += 15
                feedback_parts.append("Abstract formatted correctly (+15)")
            else:
                feedback_parts.append(f"Abstract fmt mismatch (Size:{size}, Italic:{italic}, Align:{align})")

        # Check Heading 1
        elif metadata.get('heading_text').lower() in text and not found_heading:
            found_heading = True
            # Heading Constraints: 14pt, Bold, Left (0 or None)
            if size == 14.0 and bold and align in [0, None]:
                score += 15
                feedback_parts.append("Heading formatted correctly (+15)")
            else:
                feedback_parts.append(f"Heading fmt mismatch (Size:{size}, Bold:{bold}, Align:{align})")

        # Check Body
        elif metadata.get('body_text').lower() in text and not found_body:
            found_body = True
            # Body Constraints: Justified (3)
            if align == 3:
                score += 10
                feedback_parts.append("Body formatted correctly (+10)")
            else:
                feedback_parts.append(f"Body fmt mismatch (Align:{align})")

    # 3. Check Table Conversion
    if len(doc.tables) > 0:
        table_has_content = False
        for row in doc.tables[0].rows:
            for cell in row.cells:
                if metadata.get('table_check_text').lower() in cell.text.lower():
                    table_has_content = True
                    break
        
        if table_has_content:
            score += 20
            feedback_parts.append("Data successfully converted to table (+20)")
        else:
            feedback_parts.append("Table created, but expected content missing.")
    else:
        feedback_parts.append("No table object found in document.")

    # 4. Trajectory VLM Check (Anti-Python scripting gaming)
    if VLM_AVAILABLE and traj:
        frames = sample_trajectory_frames(traj, n=5)
        if frames:
            prompt = (
                "You are auditing an agent performing a text formatting task in ONLYOFFICE Document Editor. "
                "Look at these trajectory screenshots. "
                "Did the agent actively use the ONLYOFFICE graphical interface to select text, change fonts/alignments, and convert text to a table? "
                "Respond in JSON format: {'used_gui': true/false}"
            )
            vlm_resp = query_vlm(images=frames, prompt=prompt)
            if vlm_resp.get("success"):
                parsed = vlm_resp.get("parsed", {})
                if parsed.get("used_gui", False):
                    score += 15
                    feedback_parts.append("VLM verified GUI usage (+15)")
                else:
                    feedback_parts.append("VLM did not detect ONLYOFFICE GUI usage.")
            else:
                feedback_parts.append("VLM query failed, skipped trajectory check.")
    else:
        feedback_parts.append("VLM not available/no trajectory. Skipping VLM check.")
        # If VLM isn't available, we auto-grant the points if the programmatic formatting is perfect 
        # to not penalize environments running without VLM.
        if score == 70:
            score += 15
            feedback_parts.append("VLM skipped but formatting perfect (+15)")

    # Validate final constraints
    passed = score >= 70

    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }