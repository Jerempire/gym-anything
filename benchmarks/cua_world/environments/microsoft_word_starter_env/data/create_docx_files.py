#!/usr/bin/env python3
"""Create .docx data files for microsoft_word_starter_env tasks.

All content is real, publicly available text from US government sources
(public domain — US government works are not subject to copyright),
or modeled on standard industry templates in the public domain.
"""
import io
import os
import re
import zipfile

import struct
import zlib

from docx import Document
from docx.shared import Inches, Pt

DATA_DIR = os.path.dirname(os.path.abspath(__file__))


def create_census_press_release():
    """US Census Bureau press release - all plain text, no formatting applied.

    Source: Adapted from real US Census Bureau press release CB24-SFS.17
    "Income and Poverty in the United States: 2023" (September 2024).
    US government works are public domain.
    """
    doc = Document()

    # Remove all default styles - make everything plain Normal text
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = [
        "U.S. Census Bureau News",
        "",
        "Income and Poverty in the United States: 2023",
        "",
        "Release Number: CB24-SFS.17",
        "September 10, 2024",
        "",
        "Overview",
        "",
        "The U.S. Census Bureau announced today that real median household income "
        "was $80,610 in 2023, not statistically different from the 2022 estimate "
        "of $80,440. This is the second year since the redesign of income questions "
        "used in the Current Population Survey Annual Social and Economic Supplement. "
        "The official poverty rate in 2023 was 11.1 percent, with 36.8 million people "
        "in poverty, neither estimate being statistically different from 2022.",
        "",
        "Income",
        "",
        "Real median household income was $80,610 in 2023, not statistically different "
        "from the 2022 estimate. Median earnings of all workers aged 15 and older who "
        "worked full-time, year-round decreased 1.4 percent between 2022 and 2023 from "
        "$60,580 to $59,740. The Gini index of income inequality was 0.485 in 2023, not "
        "statistically different from 2022. Income inequality declined from 2021 to 2022 "
        "following the expiration of pandemic-era economic stimulus payments.",
        "",
        "Median household income varied across demographic groups. Asian households had "
        "the highest median income at $112,800, followed by non-Hispanic White households "
        "at $86,370. Hispanic households had a median income of $62,590 and Black households "
        "had a median income of $56,490.",
        "",
        "Poverty",
        "",
        "The official poverty rate in 2023 was 11.1 percent, with 36.8 million people in "
        "poverty. Neither the poverty rate nor the number of people in poverty was "
        "statistically different from 2022. The poverty rate for children under 18 was "
        "13.7 percent in 2023, not statistically different from the 2022 estimate of "
        "13.4 percent. The poverty rate for people aged 65 and older was 10.3 percent, "
        "not statistically different from 2022.",
        "",
        "Supplemental Poverty Measure",
        "",
        "Under the Supplemental Poverty Measure (SPM), which accounts for government "
        "benefits and necessary expenses not reflected in the official measure, the overall "
        "SPM rate was 12.9 percent in 2023, an increase from 12.4 percent in 2022. The SPM "
        "child poverty rate was 8.9 percent in 2023, not statistically different from 2022. "
        "Social Security continued to be the most important anti-poverty program, keeping "
        "approximately 28.7 million people out of poverty in 2023.",
        "",
        "Health Insurance Coverage",
        "",
        "The uninsured rate in 2023 was 7.9 percent, or 26.4 million people, not "
        "statistically different from 2022. The percentage of people covered by "
        "private health insurance was 65.4 percent, and public coverage was 37.0 percent "
        "in 2023. Among children under 19, 5.0 percent were uninsured in 2023. Medicaid "
        "and the Children's Health Insurance Program covered 30.8 million children.",
        "",
        "Contact Information",
        "",
        "Public Information Office",
        "U.S. Census Bureau",
        "301-763-3030",
        "pio@census.gov",
    ]

    for line in lines:
        doc.add_paragraph(line)

    path = os.path.join(DATA_DIR, "census_press_release.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_meeting_notes_raw():
    """Unformatted meeting notes with quarterly sales data.

    Content based on publicly available corporate meeting minute templates
    and fictional quarterly revenue figures typical of a mid-size retail company.
    The quarterly sales figures are embedded as plain text for the table-creation task.
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = [
        "Regional Sales Team Meeting Notes",
        "",
        "Date: November 15, 2024",
        "Location: Conference Room B, Building 3",
        "Attendees: Sarah Mitchell (VP Sales), James Chen (Regional Manager - East), "
        "Maria Garcia (Regional Manager - West), David Kim (Regional Manager - Central), "
        "Lisa Thompson (Finance Director)",
        "",
        "Meeting called to order at 10:00 AM by Sarah Mitchell.",
        "",
        "Agenda Item 1: Review of 2024 Quarterly Revenue Performance",
        "",
        "Lisa Thompson presented the quarterly revenue figures for fiscal year 2024. "
        "The team reviewed each quarter's performance against the annual targets "
        "established in January.",
        "",
        "Q1: $45,200",
        "Q2: $52,100",
        "Q3: $48,900",
        "Q4: $61,300",
        "",
        "Total annual revenue reached $207,500, exceeding the projected target of "
        "$195,000 by 6.4 percent. The strong Q4 performance was attributed to the "
        "holiday season promotional campaigns and the successful launch of three "
        "new product lines in October.",
        "",
        "Agenda Item 2: Regional Performance Analysis",
        "",
        "James Chen reported that the Eastern region contributed 38 percent of total "
        "revenue, driven by strong performance in the New York and Boston markets. "
        "The Philadelphia office exceeded its individual target by 12 percent.",
        "",
        "Maria Garcia noted that the Western region accounted for 35 percent of "
        "total revenue. The Los Angeles market showed significant growth in Q3 and Q4, "
        "while the Seattle office faced challenges due to increased competition from "
        "online retailers.",
        "",
        "David Kim reported that the Central region contributed 27 percent of total "
        "revenue. The Chicago office performed above expectations, while the Dallas "
        "and Denver offices showed steady growth throughout the year.",
        "",
        "Agenda Item 3: Action Items for Q1 2025",
        "",
        "The following action items were agreed upon:",
        "Sarah Mitchell to present annual results to the Board of Directors by December 6.",
        "Each Regional Manager to submit Q1 2025 targets by November 29.",
        "Lisa Thompson to distribute the final budget allocation by December 13.",
        "All teams to complete annual performance reviews by December 20.",
        "",
        "Meeting adjourned at 11:45 AM.",
        "",
        "Minutes prepared by Lisa Thompson, Finance Director.",
    ]

    for line in lines:
        doc.add_paragraph(line)

    path = os.path.join(DATA_DIR, "meeting_notes_raw.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_company_memo_draft():
    """Draft company memo - used for business letter formatting task.

    This provides a blank starting point. The task asks the agent to create
    a business letter from scratch, so this file is intentionally minimal.
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Just a blank document with a single empty paragraph
    # The create_business_letter task starts from blank
    doc.add_paragraph("")

    path = os.path.join(DATA_DIR, "company_memo_draft.docx")
    doc.save(path)
    print(f"Created: {path}")


# ---------------------------------------------------------------------------
# Helper: inject tracked changes into patent license document XML
# ---------------------------------------------------------------------------

def _inject_patent_track_changes(xml_str):
    """Replace marker paragraphs with OOXML tracked-change paragraphs."""

    # Track Change 1 – Grant of License: non-exclusive → non-exclusive and sublicensable
    grant_xml = (
        '<w:p>'
        '<w:r><w:t xml:space="preserve">2.1 Subject to the terms and conditions of this '
        'Agreement, Licensor hereby grants to Licensee a worldwide, royalty-bearing, '
        '</w:t></w:r>'
        '<w:del w:id="1" w:author="J. Morrison" w:date="2024-11-15T10:00:00Z">'
        '<w:r><w:delText>non-exclusive</w:delText></w:r>'
        '</w:del>'
        '<w:ins w:id="2" w:author="R. Chen" w:date="2024-11-16T09:00:00Z">'
        '<w:r><w:t>non-exclusive and sublicensable</w:t></w:r>'
        '</w:ins>'
        '<w:r><w:t xml:space="preserve"> license under the Patent Rights to make, have made, '
        'use, sell, offer for sale, import, and otherwise commercialize Licensed Products '
        'throughout the Territory during the Term of this Agreement.</w:t></w:r>'
        '</w:p>'
    )
    xml_str = re.sub(
        r'<w:p\b[^>]*>(?:(?!</w:p>).)*?GRANT_CLAUSE_PLACEHOLDER(?:(?!</w:p>).)*?</w:p>',
        grant_xml,
        xml_str,
        flags=re.DOTALL,
    )

    # Track Change 2 – Termination: sixty (60) days → thirty (30) days
    termination_xml = (
        '<w:p>'
        '<w:r><w:t xml:space="preserve">5.2 Either Party may terminate this Agreement upon '
        'written notice if the other Party materially breaches this Agreement and fails to '
        'cure such breach within </w:t></w:r>'
        '<w:del w:id="3" w:author="J. Morrison" w:date="2024-11-15T10:00:00Z">'
        '<w:r><w:delText>sixty (60) days</w:delText></w:r>'
        '</w:del>'
        '<w:ins w:id="4" w:author="R. Chen" w:date="2024-11-16T09:00:00Z">'
        '<w:r><w:t>thirty (30) days</w:t></w:r>'
        '</w:ins>'
        '<w:r><w:t xml:space="preserve"> after receipt of written notice specifying in '
        'reasonable detail the nature of such breach; provided that no cure period shall '
        'apply to payment obligations, which must be cured within fifteen (15) days.</w:t></w:r>'
        '</w:p>'
    )
    xml_str = re.sub(
        r'<w:p\b[^>]*>(?:(?!</w:p>).)*?TERMINATION_CLAUSE_PLACEHOLDER(?:(?!</w:p>).)*?</w:p>',
        termination_xml,
        xml_str,
        flags=re.DOTALL,
    )

    return xml_str


def create_patent_license_draft():
    """Patent License Agreement with tracked changes injected via XML manipulation.

    Source: Modeled after the USPTO Model Patent License Agreement structure
    (public domain). Track changes simulate a legal review: J. Morrison
    (associate) proposed changes; R. Chen (senior partner) counter-proposed.

    The agent task: accept all tracked changes, insert two footnotes citing
    relevant patent-law authority, and save as patent_license_final.docx.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    paragraphs = [
        ("h0", "PATENT LICENSE AGREEMENT"),
        ("n", ""),
        ("n", "This Patent License Agreement (\"Agreement\") is entered into as of "
              "November 1, 2024 (\"Effective Date\") by and between:"),
        ("n", ""),
        ("n", "BioSynthetics Corporation, a Delaware corporation, having its principal "
              "place of business at 2400 Innovation Drive, Research Triangle Park, NC 27709 "
              "(\"Licensor\"); and"),
        ("n", ""),
        ("n", "MedTech Innovations LLC, a California limited liability company, having its "
              "principal place of business at 500 Biotech Boulevard, San Diego, CA 92121 "
              "(\"Licensee\")."),
        ("n", ""),
        ("h1", "SECTION 1. DEFINITIONS"),
        ("n", ""),
        ("n", "1.1 \"Affiliate\" means any corporation or other entity that directly or "
              "indirectly controls, is controlled by, or is under common control with a Party. "
              "\"Control\" means ownership of more than fifty percent (50%) of the outstanding "
              "voting securities of a corporation, or equivalent decision-making authority."),
        ("n", ""),
        ("n", "1.2 \"Licensed Products\" means any product, composition, device, or method "
              "whose development, manufacture, use, sale, or importation would, absent the "
              "license granted herein, infringe one or more Valid Claims of the Patent Rights."),
        ("n", ""),
        ("n", "1.3 \"Net Sales\" means the gross invoice price billed by Licensee or its "
              "Affiliates for Licensed Products sold, less: (i) trade, cash, and quantity "
              "discounts; (ii) credits or allowances for returned goods; (iii) freight, "
              "insurance, and other transportation expenses; and (iv) sales and use taxes."),
        ("n", ""),
        ("n", "1.4 \"Patent Rights\" means United States Patent No. 11,584,921 entitled "
              "\"Compositions and Methods for CRISPR-Mediated Genomic Modification,\" and any "
              "continuations, divisionals, reissues, reexaminations, or foreign counterparts."),
        ("n", ""),
        ("n", "1.5 \"Territory\" means worldwide."),
        ("n", ""),
        ("n", "1.6 \"Valid Claim\" means a claim of an issued, unexpired patent within the "
              "Patent Rights that has not been held invalid or unenforceable by a court or "
              "other governmental agency of competent jurisdiction."),
        ("n", ""),
        ("h1", "SECTION 2. GRANT OF LICENSE"),
        ("n", ""),
        ("GRANT_CLAUSE_PLACEHOLDER", "GRANT_CLAUSE_PLACEHOLDER"),
        ("n", ""),
        ("n", "2.2 Licensor reserves the right to practice the Patent Rights for research, "
              "educational, and other non-commercial purposes, and to grant licenses to "
              "non-profit research institutions for non-commercial research purposes."),
        ("n", ""),
        ("n", "2.3 Licensee shall have the right to grant sublicenses to Affiliates only, "
              "provided that: (a) Licensee provides written notice to Licensor within thirty "
              "(30) days of granting any such sublicense; (b) each sublicensee agrees to be "
              "bound by terms at least as protective of Licensor as those in this Agreement; "
              "and (c) Licensee remains liable for all obligations of its sublicensees."),
        ("n", ""),
        ("h1", "SECTION 3. ROYALTIES AND PAYMENTS"),
        ("n", ""),
        ("n", "3.1 Upfront License Fee. Licensee shall pay to Licensor a non-refundable "
              "upfront license fee of Three Hundred Fifty Thousand Dollars (USD $350,000) "
              "within thirty (30) days of the Effective Date."),
        ("n", ""),
        ("n", "3.2 Milestone Payments. Licensee shall pay to Licensor the following "
              "non-refundable milestone payments upon first achievement of each milestone: "
              "(a) USD $750,000 upon initiation of the first Phase I clinical trial; "
              "(b) USD $2,000,000 upon initiation of the first Phase III clinical trial; "
              "(c) USD $5,000,000 upon first FDA regulatory approval of a Licensed Product; "
              "(d) USD $3,000,000 upon first regulatory approval in the EU or Japan."),
        ("n", ""),
        ("n", "3.3 Running Royalties. Licensee shall pay running royalties on Net Sales at: "
              "(a) 3.0% on the first USD $50,000,000 of annual Net Sales; "
              "(b) 4.0% on annual Net Sales between USD $50,000,001 and USD $200,000,000; "
              "(c) 5.0% on annual Net Sales exceeding USD $200,000,000."),
        ("n", ""),
        ("h1", "SECTION 4. INTELLECTUAL PROPERTY"),
        ("n", ""),
        ("n", "4.1 Licensor Improvements. As between the Parties, Licensor shall retain all "
              "right, title, and interest in and to any improvements to the Patent Rights "
              "conceived or reduced to practice solely by Licensor's employees or agents."),
        ("n", ""),
        ("n", "4.2 Licensee Improvements. Licensee shall own all right, title, and interest "
              "in and to any improvements to the Licensed Products or processes conceived or "
              "reduced to practice solely by Licensee's employees during the Term."),
        ("n", ""),
        ("n", "4.3 Joint Improvements. Improvements conceived jointly by employees of both "
              "Parties shall be jointly owned, and each Party shall have the right to exploit "
              "such joint improvements without accounting to the other Party, subject to the "
              "royalty obligations set forth in Section 3."),
        ("n", ""),
        ("h1", "SECTION 5. TERM AND TERMINATION"),
        ("n", ""),
        ("n", "5.1 Term. This Agreement shall commence on the Effective Date and shall "
              "continue in full force and effect until the expiration of the last-to-expire "
              "Valid Claim of the Patent Rights, unless earlier terminated under this Section 5."),
        ("n", ""),
        ("TERMINATION_CLAUSE_PLACEHOLDER", "TERMINATION_CLAUSE_PLACEHOLDER"),
        ("n", ""),
        ("n", "5.3 Licensor may terminate this Agreement immediately upon written notice "
              "to Licensee if Licensee: (a) files for bankruptcy, becomes insolvent, or makes "
              "an assignment for the benefit of creditors; (b) challenges the validity or "
              "enforceability of any Patent Rights; or (c) fails to pay any amounts due "
              "hereunder within fifteen (15) days after receiving notice of such failure."),
        ("n", ""),
        ("n", "5.4 Effects of Termination. Upon termination or expiration, all licenses "
              "granted hereunder shall immediately terminate, Licensee shall cease all "
              "manufacture, use, and sale of Licensed Products, and each Party shall return "
              "or destroy the other Party's Confidential Information within thirty (30) days."),
        ("n", ""),
        ("h1", "SECTION 6. REPRESENTATIONS AND WARRANTIES"),
        ("n", ""),
        ("n", "6.1 Mutual Representations. Each Party represents and warrants that: "
              "(a) it is duly organized, validly existing, and in good standing; "
              "(b) it has full power and authority to execute this Agreement; "
              "(c) this Agreement constitutes a valid and binding obligation; and "
              "(d) its execution will not conflict with any other agreement to which it is a party."),
        ("n", ""),
        ("n", "6.2 Licensor Representations. Licensor represents and warrants that: "
              "(a) to Licensor's knowledge, it is the sole owner of the Patent Rights; "
              "(b) Licensor has not previously granted licenses that would conflict with "
              "the licenses granted herein; and (c) to Licensor's knowledge, the Patent "
              "Rights are valid and enforceable."),
        ("n", ""),
        ("h1", "SECTION 7. INDEMNIFICATION AND LIABILITY"),
        ("n", ""),
        ("n", "7.1 Indemnification by Licensee. Licensee shall indemnify, defend, and hold "
              "harmless Licensor and its officers, directors, employees, and agents from and "
              "against any and all claims, damages, losses, costs, and expenses (including "
              "reasonable attorneys' fees) arising out of: (a) Licensee's development, "
              "manufacture, use, sale, or importation of Licensed Products; (b) any breach "
              "by Licensee of this Agreement; or (c) the gross negligence of Licensee."),
        ("n", ""),
        ("n", "7.2 LIMITATION OF LIABILITY. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR "
              "INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, PUNITIVE, OR EXEMPLARY DAMAGES. "
              "EACH PARTY'S CUMULATIVE LIABILITY SHALL NOT EXCEED AMOUNTS PAID IN THE TWELVE "
              "(12) MONTHS PRECEDING THE CLAIM."),
        ("n", ""),
        ("h1", "SECTION 8. MISCELLANEOUS"),
        ("n", ""),
        ("n", "8.1 Governing Law. This Agreement shall be governed by the laws of the State "
              "of Delaware, without regard to conflict of laws provisions. Any dispute shall "
              "be subject to the exclusive jurisdiction of the federal and state courts in "
              "Wilmington, Delaware."),
        ("n", ""),
        ("n", "8.2 Entire Agreement. This Agreement constitutes the entire agreement between "
              "the Parties with respect to the subject matter hereof, and supersedes all prior "
              "negotiations, representations, warranties, and understandings of the Parties."),
        ("n", ""),
        ("n", "8.3 Amendment. This Agreement may not be amended or modified except by a "
              "written instrument signed by authorized representatives of both Parties."),
        ("n", ""),
        ("n", "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date "
              "first written above."),
        ("n", ""),
        ("n", "BIOSYNTHETICS CORPORATION                    MEDTECH INNOVATIONS LLC"),
        ("n", "By: ___________________________              By: ___________________________"),
        ("n", "Name:                                        Name:"),
        ("n", "Title:                                       Title:"),
        ("n", "Date:                                        Date:"),
    ]

    for tag, text in paragraphs:
        if tag == "h0":
            doc.add_heading(text, 0)
        elif tag == "h1":
            doc.add_heading(text, 1)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin:
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    xml_str = data.decode("utf-8")
                    xml_str = _inject_patent_track_changes(xml_str)
                    data = xml_str.encode("utf-8")
                zout.writestr(item, data)

    path = os.path.join(DATA_DIR, "patent_license_draft_tracked.docx")
    with open(path, "wb") as f:
        f.write(out_buf.getvalue())
    print(f"Created: {path}")


def create_oncology_protocol_raw():
    """Phase III oncology clinical trial protocol - all plain text, no formatting.

    Content: ICH E6(R2) GCP-compliant protocol structure for a fictional
    Phase III trial of 'HORIZIMAB' in KRAS G12C-mutant NSCLC.
    Structure based on ICH E6(R2) Guideline for Good Clinical Practice
    (public domain, international guidance document).

    Agent task: apply Heading 1 / Heading 2 styles, insert Table of Contents,
    add a formatted Schedule of Assessments table, and add a CONFIDENTIAL footer.
    Save as oncology_protocol_final.docx.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = [
        "CLINICAL STUDY PROTOCOL",
        "",
        "Study Title: A Phase III, Randomized, Double-Blind, Placebo-Controlled,"
        " Multicenter Study of Horizimab Plus Pembrolizumab Versus Pembrolizumab"
        " Plus Placebo as First-Line Treatment for Patients with KRAS G12C-Mutant,"
        " Metastatic Non-Small Cell Lung Cancer (NSCLC)",
        "",
        "Protocol Number: HRZ-NSCLC-301",
        "Version: 2.1",
        "Date: October 14, 2024",
        "Sponsor: Vericel Oncology, Inc., 1 Harbor Plaza, Boston, MA 02210",
        "IND Number: 147,382",
        "",
        "CONFIDENTIALITY NOTICE",
        "",
        "This document contains confidential and proprietary information of Vericel "
        "Oncology, Inc. This protocol is provided to clinical investigators, Institutional "
        "Review Boards, and regulatory authorities solely for the purpose of evaluating "
        "the clinical study. Disclosure to any third party is strictly prohibited without "
        "prior written consent of the Sponsor.",
        "",
        "PROTOCOL SYNOPSIS",
        "",
        "Study Phase: Phase III",
        "Study Design: Randomized, double-blind, placebo-controlled, parallel-group, "
        "multicenter study",
        "Primary Objective: To demonstrate that the addition of horizimab to pembrolizumab "
        "significantly improves progression-free survival (PFS) compared to pembrolizumab "
        "plus placebo in patients with treatment-naive KRAS G12C-mutant metastatic NSCLC.",
        "Primary Endpoint: Progression-Free Survival (PFS) per RECIST v1.1 as assessed "
        "by blinded independent central review (BICR).",
        "Secondary Endpoints: Overall Survival (OS); Objective Response Rate (ORR) per "
        "RECIST v1.1; Duration of Response (DoR); Disease Control Rate (DCR); "
        "Patient-Reported Outcomes (PROs) per EORTC QLQ-C30 and QLQ-LC13.",
        "Study Population: Adults (>=18 years) with histologically confirmed, previously "
        "untreated, stage IV NSCLC harboring KRAS G12C mutation, ECOG PS 0-1, no prior "
        "systemic treatment for metastatic disease, and PD-L1 TPS >=1%.",
        "Sample Size: 480 patients (240 per arm), 80% power, one-sided alpha of 0.025.",
        "Study Duration: Approximately 48 months (36 months enrollment + 12 months follow-up).",
        "Study Sites: Approximately 120 investigational sites in the United States, "
        "European Union, Japan, South Korea, and Australia.",
        "",
        "1. INTRODUCTION AND RATIONALE",
        "",
        "1.1 Background and Disease Overview",
        "",
        "Lung cancer is the leading cause of cancer-related mortality worldwide, accounting "
        "for approximately 1.8 million deaths annually. Non-small cell lung cancer (NSCLC) "
        "represents approximately 85% of all lung cancer cases. Despite advances in targeted "
        "therapy and immunotherapy, the prognosis for patients with metastatic NSCLC remains "
        "poor, with a median overall survival of approximately 12-18 months with standard "
        "first-line pembrolizumab-based regimens.",
        "",
        "KRAS mutations are the most common oncogenic driver mutations in NSCLC, occurring "
        "in approximately 25-30% of cases. Among KRAS mutations, the G12C variant accounts "
        "for approximately 13% of all NSCLC cases, representing a patient population of "
        "approximately 25,000 newly diagnosed patients annually in the United States alone. "
        "KRAS G12C mutation results in a constitutively active GTP-bound form of KRAS, "
        "which drives tumor cell proliferation through activation of the MAPK and PI3K "
        "signaling pathways.",
        "",
        "1.2 Horizimab: Mechanism of Action and Preclinical Data",
        "",
        "Horizimab is a first-in-class, highly selective, small molecule inhibitor of "
        "KRAS G12C that covalently binds to the mutant cysteine residue in the inactive "
        "GDP-bound state of KRAS G12C, locking the protein in an inactive conformation. "
        "This mechanism exploits a unique allosteric switch II pocket (SIIP) that is "
        "present only in the mutant KRAS G12C protein and not in wild-type KRAS.",
        "",
        "In preclinical studies, horizimab demonstrated potent anti-tumor activity in "
        "KRAS G12C-mutant cell lines and xenograft models, with IC50 values in the "
        "low nanomolar range (2-8 nM). The combination of horizimab with anti-PD-1 therapy "
        "demonstrated synergistic anti-tumor activity in syngeneic mouse models, with "
        "complete tumor regression observed in 60-75% of treated animals compared to "
        "20-30% with either agent alone.",
        "",
        "1.3 Clinical Data Supporting This Study",
        "",
        "Phase I/II clinical data from the ongoing HRZ-NSCLC-101 study demonstrated that "
        "horizimab at the recommended Phase III dose (800 mg once daily, orally) achieved "
        "an ORR of 47% (95% CI: 38-56%) in patients with previously treated KRAS G12C-mutant "
        "NSCLC (n=112), with a median PFS of 8.2 months (95% CI: 6.4-10.1 months). The "
        "safety profile was generally manageable, with Grade 3-4 treatment-related adverse "
        "events occurring in 28% of patients, most commonly diarrhea (8%), fatigue (6%), "
        "and nausea (5%).",
        "",
        "1.4 Rationale for Combination with Pembrolizumab",
        "",
        "Pembrolizumab (KEYTRUDA) is an anti-PD-1 monoclonal antibody approved by the FDA "
        "as first-line treatment for NSCLC with PD-L1 TPS >=1%. The combination of KRAS "
        "G12C inhibition with anti-PD-1 therapy is scientifically supported by: (1) evidence "
        "that KRAS pathway activation suppresses anti-tumor immunity through increased "
        "expression of PD-L1 and immunosuppressive cytokines; (2) demonstration that KRAS "
        "G12C inhibition restores anti-tumor T-cell function in preclinical models; and "
        "(3) early clinical signals of enhanced response rates in patients receiving "
        "KRAS G12C inhibitors in combination with checkpoint inhibitors.",
        "",
        "2. STUDY OBJECTIVES AND ENDPOINTS",
        "",
        "2.1 Primary Objective and Endpoint",
        "",
        "Primary Objective: To demonstrate superiority of horizimab plus pembrolizumab "
        "versus placebo plus pembrolizumab in progression-free survival.",
        "Primary Endpoint: PFS defined as the time from randomization to the first "
        "documented disease progression per RECIST v1.1 as assessed by BICR, or death "
        "from any cause, whichever occurs first.",
        "",
        "2.2 Secondary Objectives and Endpoints",
        "",
        "Overall Survival: OS defined as time from randomization to death from any cause.",
        "Objective Response Rate: ORR defined as the proportion of patients with confirmed "
        "complete response (CR) or partial response (PR) per RECIST v1.1 by BICR.",
        "Duration of Response: DoR defined as time from first documented response to "
        "first documented disease progression or death.",
        "Patient-Reported Outcomes: Change from baseline in EORTC QLQ-C30 global health "
        "status/QoL scale score and QLQ-LC13 dyspnea scale score.",
        "Safety and Tolerability: Incidence, nature, and severity of adverse events (AEs), "
        "serious adverse events (SAEs), and laboratory abnormalities per CTCAE v5.0.",
        "",
        "2.3 Exploratory Objectives",
        "",
        "Biomarker Analysis: Correlation of KRAS G12C allele frequency, co-occurring "
        "mutations (STK11, KEAP1, TP53), PD-L1 TPS, and tumor mutational burden (TMB) "
        "with clinical outcomes.",
        "Pharmacokinetics: Population PK parameters of horizimab and any metabolites.",
        "Pharmacodynamics: Changes in circulating tumor DNA (ctDNA) and immune cell "
        "subsets as early markers of response.",
        "",
        "3. STUDY DESIGN",
        "",
        "3.1 Overview",
        "",
        "This is a Phase III, randomized, double-blind, placebo-controlled, parallel-group, "
        "multicenter study. Eligible patients will be randomized 1:1 to receive either "
        "horizimab 800 mg once daily plus pembrolizumab 200 mg every 3 weeks (Arm A) or "
        "matching placebo once daily plus pembrolizumab 200 mg every 3 weeks (Arm B).",
        "",
        "3.2 Randomization and Stratification",
        "",
        "Randomization will be performed using an interactive web response system (IWRS) "
        "with stratification by: (1) PD-L1 TPS (1-49% vs >=50%); (2) Geographic region "
        "(North America/Western Europe vs Rest of World); (3) Histology (squamous vs "
        "non-squamous). Block randomization with variable block sizes will be used to "
        "maintain balance across treatment arms.",
        "",
        "3.3 Blinding",
        "",
        "This study will be conducted in a double-blind fashion. Horizimab and matching "
        "placebo will be identical in appearance. Pembrolizumab will be administered as an "
        "open-label infusion per standard institutional practice. An independent Data Safety "
        "Monitoring Board (DSMB) will review unblinded safety and efficacy data at "
        "predefined intervals.",
        "",
        "4. PATIENT SELECTION",
        "",
        "4.1 Inclusion Criteria",
        "",
        "Patients must meet ALL of the following criteria for enrollment:",
        "IC-01: Age >= 18 years at the time of signing the informed consent form.",
        "IC-02: Histologically or cytologically confirmed, stage IV (metastatic) NSCLC "
        "per the 8th edition of the AJCC Cancer Staging Manual.",
        "IC-03: KRAS G12C mutation confirmed by a validated, FDA-approved or "
        "laboratory-developed test on tumor tissue or blood.",
        "IC-04: No prior systemic treatment for metastatic NSCLC. Prior adjuvant or "
        "neoadjuvant chemotherapy is permitted if completed >=12 months before enrollment.",
        "IC-05: PD-L1 TPS >= 1% as determined by a validated PD-L1 assay.",
        "IC-06: ECOG Performance Status 0 or 1.",
        "IC-07: At least one measurable lesion per RECIST v1.1.",
        "IC-08: Adequate organ function as defined in the protocol laboratory parameters.",
        "IC-09: No known EGFR, ALK, ROS1, RET, MET exon 14 skipping, or NTRK mutations "
        "for which approved targeted therapies exist.",
        "IC-10: Life expectancy >= 12 weeks per investigator assessment.",
        "",
        "4.2 Exclusion Criteria",
        "",
        "Patients who meet ANY of the following criteria will be excluded from this study:",
        "EC-01: Prior treatment with a KRAS G12C inhibitor, anti-PD-1, anti-PD-L1, "
        "or anti-CTLA-4 antibody in the metastatic setting.",
        "EC-02: Active autoimmune disease requiring systemic treatment within the past "
        "2 years (replacement therapy, e.g., thyroxine or insulin, is not considered "
        "a form of systemic treatment).",
        "EC-03: Active or prior documented inflammatory bowel disease (Crohn's disease "
        "or ulcerative colitis).",
        "EC-04: History of (non-infectious) pneumonitis that required steroids, or current "
        "pneumonitis or interstitial lung disease.",
        "EC-05: Known brain metastases or leptomeningeal disease that are untreated, "
        "symptomatic, or require therapy to control symptoms.",
        "EC-06: Receipt of any investigational agent or participation in another "
        "interventional clinical study within 4 weeks prior to randomization.",
        "EC-07: Major surgery within 28 days of first dose of study treatment.",
        "EC-08: Pregnant or breastfeeding women, or women of childbearing potential "
        "unwilling to use effective contraception.",
        "",
        "5. STUDY TREATMENTS",
        "",
        "5.1 Horizimab (Investigational Product)",
        "",
        "Formulation: Horizimab 200 mg film-coated tablets for oral administration.",
        "Dose: 800 mg (four 200-mg tablets) once daily, administered with or without food.",
        "Route of Administration: Oral.",
        "Treatment Cycles: Continuous daily dosing in 21-day cycles, aligned with "
        "pembrolizumab administration cycles.",
        "Storage: Store at 20-25 degrees C (68-77 degrees F); excursions permitted to "
        "15-30 degrees C (59-86 degrees F).",
        "",
        "5.2 Pembrolizumab (Standard of Care)",
        "",
        "Pembrolizumab will be administered at 200 mg as a 30-minute intravenous infusion "
        "on Day 1 of each 21-day cycle. Treatment will continue until disease progression "
        "per RECIST v1.1, unacceptable toxicity, withdrawal of consent, or a maximum of "
        "35 cycles (approximately 2 years).",
        "",
        "5.3 Dose Modifications",
        "",
        "Horizimab Dose Reductions: If dose reduction is required due to toxicity, "
        "horizimab may be reduced to 600 mg once daily (first reduction) or 400 mg once "
        "daily (second reduction). No further dose reductions are permitted; if additional "
        "reduction is required, horizimab must be permanently discontinued.",
        "Pembrolizumab Dose Modifications: Pembrolizumab doses should not be reduced. "
        "Pembrolizumab may be withheld for up to 12 weeks for management of immune-related "
        "adverse events (irAEs) per current prescribing information. Pembrolizumab should "
        "be permanently discontinued for severe or life-threatening irAEs.",
        "",
        "6. STUDY ASSESSMENTS",
        "",
        "6.1 Schedule of Assessments",
        "",
        "Assessment Schedule: The following assessments will be performed at each visit "
        "per the schedule below. All assessment windows are defined relative to the first "
        "day of each treatment cycle (Day 1 of Cycle N).",
        "",
        "Screening (within 28 days prior to Cycle 1 Day 1): Complete medical history; "
        "physical examination; vital signs; ECOG PS; 12-lead ECG; clinical laboratory "
        "tests (hematology, chemistry, coagulation, urinalysis, thyroid function); "
        "tumor tissue for biomarker testing; CT/MRI of chest, abdomen, and pelvis "
        "(with contrast); brain MRI; pulmonary function tests; EORTC QLQ-C30 and QLQ-LC13; "
        "pregnancy test; HIV, HBV, HCV serology.",
        "",
        "Cycle 1 Day 1: Physical examination; vital signs; ECOG PS; clinical laboratory "
        "tests; administration of study treatment; PK sample (pre-dose and 2h post-dose); "
        "AE assessment.",
        "",
        "Cycle 1 Day 8 (+/- 2 days): Vital signs; clinical laboratory tests; AE assessment.",
        "",
        "Cycle 1 Day 15 (+/- 2 days): Vital signs; clinical laboratory tests; AE assessment; "
        "PK sample (pre-dose).",
        "",
        "Day 1 of Cycles 2-4: Physical examination; vital signs; ECOG PS; clinical "
        "laboratory tests; administration of study treatment; AE assessment; concomitant "
        "medication review.",
        "",
        "Day 1 of Cycles 5 and Beyond (every 3 cycles): Tumor imaging (CT/MRI) per RECIST "
        "v1.1; physical examination; vital signs; ECOG PS; clinical laboratory tests; "
        "administration of study treatment; PRO assessment; AE assessment.",
        "",
        "End of Treatment (within 30 days after last dose): Physical examination; vital "
        "signs; ECOG PS; clinical laboratory tests; tumor imaging; PRO assessment; AE "
        "assessment; survival status.",
        "",
        "Follow-up (every 12 weeks until death, withdrawal, or study closure): "
        "Survival status; subsequent anti-cancer therapy.",
        "",
        "6.2 Efficacy Assessments",
        "",
        "Tumor assessments will be performed by CT scan (with contrast) of the chest, "
        "abdomen, and pelvis every 6 weeks (two cycles) from randomization through Week 18, "
        "then every 9 weeks (three cycles) until disease progression. Brain MRI will be "
        "performed every 12 weeks. All imaging will be submitted to the central imaging "
        "vendor for BICR evaluation per RECIST v1.1.",
        "",
        "6.3 Safety Assessments",
        "",
        "Adverse events will be collected from the time of signed informed consent through "
        "90 days after the last dose of study treatment. All AEs will be graded per NCI "
        "CTCAE v5.0. Serious adverse events will be reported to the Sponsor within 24 hours "
        "of investigator awareness. A DSMB will conduct formal safety reviews after every "
        "50 patients have completed at least 3 cycles of treatment.",
        "",
        "7. STATISTICAL ANALYSIS PLAN",
        "",
        "7.1 Sample Size and Power",
        "",
        "With a target hazard ratio (HR) of 0.65 for PFS (horizimab + pembrolizumab vs "
        "placebo + pembrolizumab), assuming a median PFS of 9.0 months in the control arm, "
        "a two-sided log-rank test at alpha = 0.05 (one-sided alpha = 0.025) will have "
        "approximately 80% power with 256 PFS events (128 per arm). Accounting for a 10% "
        "dropout rate, 480 patients (240 per arm) will be enrolled.",
        "",
        "7.2 Analysis Populations",
        "",
        "Full Analysis Set (FAS): All randomized patients who received at least one dose "
        "of study treatment. This is the primary analysis population for efficacy endpoints.",
        "Per-Protocol Set (PPS): All FAS patients without major protocol deviations. "
        "This population will be used for sensitivity analyses.",
        "Safety Analysis Set (SAS): All patients who received at least one dose of study "
        "treatment. This is the primary analysis population for safety endpoints.",
        "Pharmacokinetic Analysis Set (PKAS): All patients who received at least one dose "
        "of horizimab and have at least one evaluable PK sample.",
        "",
        "7.3 Primary Efficacy Analysis",
        "",
        "The primary analysis of PFS will be conducted using a stratified log-rank test "
        "at a one-sided significance level of 0.025, stratified by the factors used for "
        "randomization. The hazard ratio and 95% confidence interval will be estimated "
        "using a Cox proportional hazards model. A Kaplan-Meier plot of PFS will be "
        "presented for each treatment arm.",
        "",
        "8. ETHICAL CONSIDERATIONS",
        "",
        "8.1 Good Clinical Practice",
        "",
        "This study will be conducted in compliance with the protocol, International "
        "Conference on Harmonization (ICH) E6(R2) Good Clinical Practice (GCP) guidelines, "
        "the Declaration of Helsinki (2013 revision), and applicable local regulations. "
        "The protocol and all amendments will be reviewed and approved by each site's "
        "Institutional Review Board (IRB) or Independent Ethics Committee (IEC) prior to "
        "patient enrollment.",
        "",
        "8.2 Informed Consent",
        "",
        "Each patient must provide written informed consent prior to any study-specific "
        "procedures. The informed consent form (ICF) will be written in non-technical "
        "language and translated into the local language of each participating country. "
        "Patients will be given adequate time to review the ICF and ask questions. Patients "
        "may withdraw consent at any time without penalty.",
    ]

    for line in lines:
        doc.add_paragraph(line)

    path = os.path.join(DATA_DIR, "oncology_protocol_raw.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_ap_walkthrough_raw():
    """SOX 404 Accounts Payable walkthrough narrative - all plain text.

    Content: Modeled on PCAOB AS 2201 and COSO 2013 Internal Control framework
    (PCAOB standards are public domain; COSO framework structure is widely published).
    Fictional company: Meridian Global Industries, Inc.

    Agent task: apply heading styles, convert control descriptions into 3+ formatted
    tables (Control#, Objective, Risk, Control Type, Testing Approach columns),
    create a custom paragraph style named 'Control Description', add 3+ footnotes
    citing COSO/PCAOB, and add a document header with company name and classification.
    Save as ap_walkthrough_final.docx.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = [
        "SOX 404 INTERNAL CONTROL WALKTHROUGH",
        "Accounts Payable and Disbursements Process",
        "Meridian Global Industries, Inc.",
        "For the Fiscal Year Ended December 31, 2024",
        "Prepared by: Internal Audit Department",
        "Review Date: November 8, 2024",
        "Classification: CONFIDENTIAL",
        "",
        "1. PURPOSE AND SCOPE",
        "",
        "This document presents the results of the Sarbanes-Oxley Section 404 internal "
        "control walkthrough for the Accounts Payable (AP) and Disbursements process of "
        "Meridian Global Industries, Inc. (\"Meridian\" or the \"Company\"). The walkthrough "
        "was performed in accordance with Public Company Accounting Oversight Board (PCAOB) "
        "Auditing Standard 2201, An Audit of Internal Control Over Financial Reporting That "
        "Is Integrated with An Audit of Financial Statements, and the Committee of "
        "Sponsoring Organizations of the Treadway Commission (COSO) Internal Control - "
        "Integrated Framework (2013).",
        "",
        "The scope of this walkthrough encompasses the complete AP cycle including: "
        "(1) vendor master file management and new vendor onboarding; (2) purchase order "
        "issuance and approval; (3) invoice receipt, coding, and approval; (4) three-way "
        "match processing (PO, goods receipt, invoice); (5) payment processing and "
        "disbursement; (6) month-end AP accruals and cut-off procedures; and (7) periodic "
        "management review and reconciliation of AP sub-ledger to general ledger.",
        "",
        "This walkthrough covers approximately $2.4 billion in annual AP transactions "
        "processed through SAP S/4HANA (ERP system), representing 95% of all Company "
        "disbursements. The walkthrough excludes payroll, employee expense reimbursements, "
        "and intercompany transactions, which are covered under separate walkthroughs.",
        "",
        "2. PROCESS DESCRIPTION",
        "",
        "2.1 Vendor Management",
        "",
        "The AP process begins with vendor onboarding. The Procurement department initiates "
        "new vendor requests through the Vendor Management Portal (VMP), a module integrated "
        "with SAP S/4HANA. Each new vendor must submit: (a) completed W-9 form (or W-8BEN "
        "for foreign vendors); (b) IRS EIN verification or DUNS number; (c) banking "
        "information for ACH payment setup; (d) certificate of insurance; and (e) signed "
        "Meridian Supplier Code of Conduct acknowledgment.",
        "",
        "Vendor master data changes (including banking information updates, address changes, "
        "and payment term modifications) require approval from both the Procurement Manager "
        "and the Controller. The system enforces dual approval through SAP workflow before "
        "any changes become effective in the vendor master file.",
        "",
        "2.2 Purchase Order Process",
        "",
        "All purchases over $2,500 require a pre-approved Purchase Order (PO) issued "
        "through SAP S/4HANA. Purchase requisitions are submitted by department cost center "
        "owners and routed for approval based on the Company's Delegation of Authority (DOA) "
        "matrix, which establishes approval thresholds by dollar value and cost center:",
        "",
        "Purchases $2,500 - $24,999: Cost Center Manager approval required.",
        "Purchases $25,000 - $99,999: Director-level approval required.",
        "Purchases $100,000 - $499,999: VP-level approval required.",
        "Purchases $500,000 and above: CFO and CEO dual approval required.",
        "",
        "Purchases below $2,500 are processed through corporate purchasing cards (P-cards) "
        "subject to monthly spending limits and transaction controls. P-card statements are "
        "reviewed and approved monthly by each cardholder's direct supervisor.",
        "",
        "2.3 Invoice Processing",
        "",
        "Vendor invoices are received via three channels: (1) electronic data interchange "
        "(EDI) directly into SAP; (2) email to the centralized AP inbox (ap@meridian.com), "
        "processed through optical character recognition (OCR) software; and (3) paper "
        "invoices received at the AP Processing Center in Columbus, Ohio, which are scanned "
        "and digitized within 24 hours of receipt.",
        "",
        "AP Specialists code each invoice to the appropriate cost center, general ledger "
        "account, and project code based on the associated PO. Invoices without a valid PO "
        "reference are flagged as exceptions and routed to the Procurement department for "
        "resolution before processing.",
        "",
        "3. KEY CONTROLS",
        "",
        "3.1 Invoice Approval and Three-Way Match Control",
        "",
        "Control Identifier: AP-CTRL-001",
        "Control Type: Automated preventive",
        "Control Frequency: Transaction-level (continuous)",
        "Control Owner: Director of Accounts Payable",
        "Financial Statement Risk: Existence / Occurrence; Completeness; Valuation",
        "COSO Component: Control Activities",
        "Description: SAP S/4HANA automatically matches each vendor invoice against the "
        "associated Purchase Order and Goods Receipt / Service Entry Sheet (GR/SES) before "
        "permitting payment processing. The three-way match validates: (1) the vendor on "
        "the invoice matches the vendor on the PO; (2) the quantity invoiced does not exceed "
        "the quantity received on the GR/SES by more than 2%; (3) the unit price on the "
        "invoice does not exceed the unit price on the PO by more than 0.5% (price tolerance "
        "for acceptable rounding differences). Invoices that fail three-way match are "
        "automatically blocked in SAP with a payment block indicator (R = Invoice "
        "Verification) and routed to the responsible AP Specialist and Procurement Buyer "
        "for resolution. Payment is not permitted until the block is removed by an authorized "
        "approver. System tolerance parameters are maintained by SAP Basis and are subject "
        "to quarterly review and approval by the Controller.",
        "Testing Approach: Inspect system tolerance parameter configuration in SAP (t-code "
        "OMR6). Select a sample of 25 paid invoices and verify three-way match was completed "
        "in SAP. Select a sample of 15 blocked invoices and verify appropriate resolution "
        "and approval prior to payment release.",
        "",
        "3.2 Segregation of Duties Control",
        "",
        "Control Identifier: AP-CTRL-002",
        "Control Type: Automated and manual preventive",
        "Control Frequency: Continuous (system) / Quarterly (review)",
        "Control Owner: Controller / IT Security Manager",
        "Financial Statement Risk: Existence / Occurrence; Valuation",
        "COSO Component: Control Environment; Control Activities",
        "Description: SAP S/4HANA enforces segregation of duties (SoD) through role-based "
        "access controls (RBAC) that prevent any single user from performing incompatible "
        "combinations of functions in the AP cycle. Specifically, SAP system controls "
        "prevent a user who has access to create or modify vendor master data from also "
        "approving invoices or releasing payments. Similarly, users with payment release "
        "authority cannot also create invoices or modify vendor banking information. "
        "Access to the vendor master file is restricted to three designated Vendor "
        "Management Analysts in the Procurement department, none of whom have payment "
        "release authority in SAP. The SoD ruleset is maintained in SAP Access Control "
        "(GRC) and monitored through automated weekly conflict reports reviewed by the "
        "Controller. Any SoD conflicts identified are remediated within five business days "
        "or subject to formal compensating control documentation and CFO approval.",
        "Testing Approach: Obtain SAP GRC SoD conflict report for AP cycle as of the "
        "walkthrough date. Review quarterly SoD review evidence (sign-off by Controller). "
        "Test access for 5 AP users to confirm alignment with job function and no SoD "
        "conflicts. Attempt to perform conflicting transactions in test environment to "
        "confirm system prevention.",
        "",
        "3.3 Payment Authorization Control",
        "",
        "Control Identifier: AP-CTRL-003",
        "Control Type: Manual and automated preventive",
        "Control Frequency: Transaction-level",
        "Control Owner: Vice President of Finance",
        "Financial Statement Risk: Existence / Occurrence; Valuation",
        "COSO Component: Control Activities",
        "Description: All ACH and wire payments over $50,000 require dual authorization "
        "in both SAP (payment run approval) and the Company's bank portal (Bank of America "
        "CashPro). The dual-bank-portal authorization requires approval by two of the "
        "following: Controller, VP Finance, or CFO. Wire transfers over $500,000 require "
        "phone confirmation with the bank in addition to portal authorization. ACH payments "
        "below $50,000 require single authorization in SAP by the Controller or designee. "
        "Check payments (limited to vendors without ACH capability) require two physical "
        "signatures for amounts over $25,000: Controller and VP Finance. The bank portal "
        "maintains a whitelist of approved vendor bank accounts; wire transfers to accounts "
        "not on the whitelist require additional approval from the CFO and a 48-hour "
        "mandatory hold period.",
        "Testing Approach: Obtain population of all payments over $50,000 during the "
        "quarter. Select a sample of 20 payments and verify dual authorization in SAP "
        "and bank portal. For wire transfers over $500,000, verify phone confirmation "
        "documentation. Test 5 payments below $50,000 for single authorization. Verify "
        "bank whitelist maintenance procedures.",
        "",
        "3.4 Vendor Master File Change Control",
        "",
        "Control Identifier: AP-CTRL-004",
        "Control Type: Manual and automated preventive / detective",
        "Control Frequency: Transaction-level (preventive); Monthly (detective review)",
        "Control Owner: Procurement Manager",
        "Financial Statement Risk: Existence / Occurrence; Completeness",
        "COSO Component: Control Activities; Monitoring Activities",
        "Description: Changes to the vendor master file in SAP, including creation of new "
        "vendors and modifications to existing vendor banking information, require dual "
        "approval through SAP workflow from the Procurement Manager and the Controller. "
        "The system prevents changes from becoming effective until both approvals are "
        "obtained. Additionally, a monthly detective control is performed by the Internal "
        "Audit Manager: a complete log of all vendor master changes made during the month "
        "is extracted from SAP change log table CDHDR and reviewed for: (1) appropriate "
        "dual approval documented in the workflow; (2) supporting documentation (vendor "
        "correspondence or banking verification); (3) absence of self-approval (approver "
        "cannot approve their own request); and (4) bank account verification calls for "
        "banking information changes exceeding $10,000 monthly payment volume. Results of "
        "the monthly review are documented on the Vendor Master Change Log Review "
        "workpaper and signed off by the Controller.",
        "Testing Approach: Extract vendor master change log for the quarter (CDHDR). "
        "Select a sample of 20 changes including at least 10 banking information changes. "
        "Verify dual approval workflow completion in SAP. Inspect supporting documentation. "
        "Verify monthly review workpapers are completed and signed off.",
        "",
        "3.5 Month-End AP Accrual Control",
        "",
        "Control Identifier: AP-CTRL-005",
        "Control Type: Manual preventive and detective",
        "Control Frequency: Monthly",
        "Control Owner: Assistant Controller - AP",
        "Financial Statement Risk: Completeness; Cutoff; Valuation",
        "COSO Component: Control Activities",
        "Description: At each month-end, the Assistant Controller - AP performs the "
        "following cutoff and accrual procedures to ensure that all liabilities for goods "
        "and services received but not yet invoiced are recorded in the correct period: "
        "(1) Goods Receipt/Invoice Receipt (GR/IR) account reconciliation: all open GR/IR "
        "items over $5,000 that are more than 30 days old are reviewed and accrued or "
        "reversed with supporting explanation; (2) Vendor invoice cut-off: all invoices "
        "received within 5 business days after month-end with invoice dates in the prior "
        "month are accrued in the prior period; (3) Open purchase orders: all POs with "
        "goods receipt recorded but invoice not yet received are accrued at the PO value; "
        "(4) Recurring accruals: standard journal entries are posted for recurring items "
        "(utilities, maintenance contracts, professional fees) based on contractual amounts "
        "or prior period actuals with appropriate adjustments. The completed accrual "
        "workpaper is reviewed and approved by the Controller before journal entries are "
        "posted. All journal entries include supporting workpaper references and are "
        "posted by individuals with appropriate SAP access.",
        "Testing Approach: Obtain month-end accrual workpaper for at least 3 months "
        "during the review period. Verify Controller approval and signature. Trace "
        "accrual amounts to supporting documentation (GR/IR report, vendor invoices, "
        "contracts). Test 10 journal entries for appropriate SAP posting authorization.",
        "",
        "4. CONTROL DEFICIENCIES AND REMEDIATION",
        "",
        "4.1 Findings from Current Walkthrough",
        "",
        "During the walkthrough, the following observations were noted that require "
        "follow-up assessment to determine whether they constitute control deficiencies:",
        "",
        "Observation 1: Three AP Specialists were identified with SAP access to both "
        "enter invoices (t-code MIRO) and release payment blocks (t-code MRBR). While "
        "this represents a SoD conflict, compensating controls exist in the form of "
        "the weekly GRC conflict report reviewed by the Controller. The Internal Audit "
        "team will assess whether this represents a significant deficiency under PCAOB "
        "AS 2201 and whether additional compensating controls are required.",
        "",
        "Observation 2: Month-end accrual workpapers for July and August 2024 were not "
        "signed off by the Controller until 8 and 6 business days after month-end close, "
        "respectively, exceeding the Company's 3-business-day target for review completion. "
        "The delay was attributed to the Controller's absence due to a scheduled surgery. "
        "No formal backup approval procedure was followed; the VP Finance subsequently "
        "approved the workpapers retroactively. Internal Audit recommends formalizing "
        "a documented backup approval procedure for key SOX controls during planned absences.",
        "",
        "4.2 Prior Year Findings Status",
        "",
        "Significant Deficiency 2023-AP-01 (Vendor Master Change Log Review): In the "
        "prior year, the monthly vendor master change log review was not consistently "
        "completed, with 3 of 12 monthly reviews either not performed or not documented. "
        "Management implemented corrective action effective January 2024, including "
        "assignment of backup reviewer and calendar reminders. The current walkthrough "
        "confirmed that all 10 monthly reviews completed through October 2024 are "
        "documented and signed off. This finding is considered remediated.",
        "",
        "5. CONCLUSION",
        "",
        "Based on the procedures performed during this walkthrough, the AP and "
        "Disbursements process control environment at Meridian Global Industries, Inc. "
        "is generally effective. Five key controls were identified as in-scope for SOX "
        "404 testing, and preliminary walkthrough evidence supports the design effectiveness "
        "of all five controls. Two observations have been identified for further assessment. "
        "Operating effectiveness testing will be performed on all five key controls "
        "during the fourth quarter testing window (October-December 2024). Final "
        "conclusions on control effectiveness will be communicated in the Internal "
        "Audit SOX Testing Completion Memo by January 31, 2025.",
    ]

    for line in lines:
        doc.add_paragraph(line)

    path = os.path.join(DATA_DIR, "ap_walkthrough_raw.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_solicitation_raw():
    """CISA-style RFP for Managed Security Services - all plain text.

    Content: Modeled on standard Federal Acquisition Regulation (FAR) Part 12
    commercial services solicitation format. FAR is US law / public domain.
    Fictional solicitation by the fictional 'Cybersecurity and Infrastructure
    Resilience Agency' (CIRA), a fictional civilian agency (not the real CISA).

    Agent task: apply heading styles (H1/H2), create a compliance matrix table
    referencing FAR 52.212-3, add the solicitation number in a document header,
    add section numbering, and save as mssoc_proposal_response.docx.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = [
        "SOLICITATION FOR MANAGED SECURITY OPERATIONS CENTER (MSOC) SERVICES",
        "",
        "Solicitation Number: CIRA-2025-MSOC-0047",
        "Issuing Office: Cybersecurity and Infrastructure Resilience Agency (CIRA)",
        "Office of Acquisition Management, 245 Murray Lane SW, Washington, DC 20598",
        "Contracting Officer: Jennifer T. Holloway, 202-555-0183, jholloway@cira.gov",
        "NAICS Code: 541519 - Other Computer Related Services",
        "PSC Code: D302 - IT and Telecom - Cybersecurity",
        "Solicitation Issue Date: December 2, 2024",
        "Proposal Due Date: January 17, 2025 at 4:00 PM Eastern Time",
        "Period of Performance: February 1, 2025 through January 31, 2030 "
        "(base year plus four option years)",
        "Estimated Contract Value: $45,000,000 - $75,000,000 (IDIQ, single award)",
        "",
        "SECTION A - SUPPLIES OR SERVICES AND PRICE/COST",
        "",
        "A.1 Overview",
        "",
        "The Cybersecurity and Infrastructure Resilience Agency (CIRA or Agency) is "
        "soliciting proposals for a single-award, Indefinite Delivery / Indefinite "
        "Quantity (IDIQ) contract for the provision of comprehensive Managed Security "
        "Operations Center (MSOC) Services. The Government seeks a highly qualified "
        "contractor to design, implement, operate, and continuously improve a 24x7x365 "
        "Security Operations Center capability that will monitor, detect, analyze, and "
        "respond to cybersecurity threats targeting CIRA's enterprise information "
        "technology infrastructure and mission-critical operational technology (OT) "
        "systems.",
        "",
        "A.2 Contract Line Item Numbers (CLINs)",
        "",
        "CLIN 0001 - Base Year MSOC Operations (24x7 monitoring and response): "
        "Firm Fixed Price, 12 months.",
        "CLIN 0002 - Incident Response Services (on-demand): Time and Material, "
        "200 hours estimated per year.",
        "CLIN 0003 - Threat Intelligence Subscriptions and Feeds: Firm Fixed Price, "
        "12 months.",
        "CLIN 0004 - Forensic Analysis and Expert Witness Services: Time and Material, "
        "100 hours estimated per year.",
        "CLIN 0005 through CLIN 0020 - Option Year CLINs (same structure as CLINs "
        "0001-0004 for each of four option years).",
        "",
        "SECTION B - CONTRACT CLAUSES",
        "",
        "B.1 Applicable FAR Clauses",
        "",
        "This contract incorporates the following Federal Acquisition Regulation clauses "
        "by reference. Clauses incorporated by reference have the same force and effect as "
        "if they were given in full text. The full text of each clause is available "
        "electronically at https://www.acquisition.gov/far/.",
        "",
        "FAR 52.212-1, Instructions to Offerors - Commercial Products and Commercial "
        "Services (September 2023).",
        "FAR 52.212-2, Evaluation - Commercial Products and Commercial Services "
        "(November 2021): The Government will award a contract resulting from this "
        "solicitation to the responsible offeror whose offer represents the Best Value "
        "to the Government, price and other factors considered.",
        "FAR 52.212-3, Offeror Representations and Certifications - Commercial Products "
        "and Commercial Services (May 2024): An offeror shall complete only paragraphs (b) "
        "through (v) of this provision if the offeror has completed the annual "
        "representations and certifications electronically via the System for Award "
        "Management (SAM) website. If an offeror has not completed the annual "
        "representations and certifications electronically at the SAM website, the "
        "offeror shall complete only paragraphs (b) through (v) of this provision.",
        "FAR 52.212-4, Contract Terms and Conditions - Commercial Products and Commercial "
        "Services (November 2023).",
        "FAR 52.212-5, Contract Terms and Conditions Required to Implement Statutes or "
        "Executive Orders - Commercial Products and Commercial Services (November 2024).",
        "FAR 52.204-7, System for Award Management (October 2018).",
        "FAR 52.204-10, Reporting Executive Compensation and First-Tier Subcontract Awards "
        "(June 2020).",
        "FAR 52.209-6, Protecting the Government's Interest When Subcontracting with "
        "Contractors Debarred, Suspended, or Proposed for Debarment (November 2021).",
        "FAR 52.222-26, Equal Opportunity (September 2016).",
        "FAR 52.222-50, Combating Trafficking in Persons (November 2021).",
        "FAR 52.227-14, Rights in Data-General (May 2014).",
        "FAR 52.232-33, Payment by Electronic Funds Transfer-System for Award Management "
        "(October 2018).",
        "FAR 52.239-1, Privacy or Security Safeguards (August 1996).",
        "",
        "SECTION C - STATEMENT OF WORK",
        "",
        "C.1 Background",
        "",
        "CIRA is responsible for the cybersecurity of the Nation's critical infrastructure "
        "sectors including communications, energy, financial services, healthcare, and "
        "transportation. The Agency operates a hybrid cloud and on-premises IT environment "
        "comprising approximately 8,500 endpoints, 340 servers, 25 cloud accounts across "
        "AWS and Azure, and 12 operational technology network segments connected to "
        "critical national infrastructure. The Agency currently processes and generates "
        "approximately 4.2 terabytes of security log data per day from over 1,200 distinct "
        "data sources.",
        "",
        "C.2 Scope of Work",
        "",
        "The Contractor shall provide the following MSOC services:",
        "",
        "C.2.1 Continuous Security Monitoring",
        "",
        "The Contractor shall operate a 24x7x365 Security Operations Center staffed by "
        "qualified security analysts, threat hunters, and incident responders. The MSOC "
        "shall collect, aggregate, normalize, and analyze security telemetry from all "
        "CIRA data sources using a Government-furnished or Contractor-provided Security "
        "Information and Event Management (SIEM) platform. The Contractor shall develop "
        "and maintain a minimum library of 500 detection use cases, mapped to the MITRE "
        "ATT&CK Enterprise framework, and shall review and tune use cases quarterly to "
        "reduce false positive rates below 3% for high-fidelity alerting.",
        "",
        "C.2.2 Threat Detection and Incident Response",
        "",
        "The Contractor shall detect, triage, investigate, and respond to cybersecurity "
        "incidents in accordance with CIRA's Incident Response Plan (IRP) and the National "
        "Institute of Standards and Technology (NIST) Special Publication 800-61 Revision 2 "
        "(Computer Security Incident Handling Guide). Target response time SLAs are: "
        "(i) Critical (P1) incidents - initial contact within 15 minutes, containment "
        "actions initiated within 2 hours, full incident report within 24 hours; "
        "(ii) High (P2) incidents - initial contact within 1 hour, containment within "
        "8 hours, full report within 72 hours; (iii) Medium (P3) incidents - initial "
        "contact within 4 hours, full report within 5 business days.",
        "",
        "C.2.3 Threat Intelligence",
        "",
        "The Contractor shall maintain subscriptions to and integrate intelligence feeds "
        "from at least four (4) government-approved threat intelligence providers, including "
        "CISA's AIS (Automated Indicator Sharing) platform, and shall enrich all alerts "
        "with relevant threat intelligence context. The Contractor shall produce daily "
        "tactical threat intelligence reports, weekly operational threat summaries, and "
        "monthly strategic threat assessments tailored to CIRA's mission environment.",
        "",
        "C.2.4 Vulnerability Management Support",
        "",
        "The Contractor shall support CIRA's vulnerability management program by: "
        "(i) providing weekly vulnerability scan results analysis and prioritization "
        "recommendations based on exploitability in the wild and asset criticality; "
        "(ii) integrating vulnerability data with threat intelligence to identify actively "
        "exploited vulnerabilities in CIRA's asset inventory; and (iii) tracking "
        "remediation status in a vulnerability management dashboard accessible to CIRA "
        "stakeholders.",
        "",
        "C.2.5 Security Metrics and Reporting",
        "",
        "The Contractor shall provide: (i) a real-time security dashboard accessible "
        "24x7 to designated CIRA personnel showing current threat landscape, active "
        "incidents, and SLA performance; (ii) monthly operational metrics reports "
        "including MTTD, MTTR, false positive rates, and coverage statistics; "
        "(iii) quarterly Executive Summary reports suitable for presentation to the "
        "CIRA Director and Deputy Director; (iv) an Annual Security Posture Report "
        "benchmarking CIRA's security maturity against NIST CSF and CMMC Level 2 "
        "requirements.",
        "",
        "SECTION D - TECHNICAL REQUIREMENTS",
        "",
        "D.1 Mandatory Technical Requirements",
        "",
        "D.1.1 All personnel assigned to this contract must hold at minimum an active "
        "Secret level security clearance prior to the contract start date. Key personnel "
        "in senior analytical and management roles must hold Top Secret / Sensitive "
        "Compartmented Information (TS/SCI) clearances. The Contractor shall maintain "
        "clearance status for all personnel and notify the Contracting Officer within "
        "24 hours of any clearance suspension or revocation.",
        "",
        "D.1.2 The SOC facility used to perform services under this contract must be "
        "located in the continental United States (CONUS) and must be a cleared facility "
        "with at least a Secret accreditation. No services may be performed from overseas "
        "locations or by foreign nationals.",
        "",
        "D.1.3 The Contractor's SIEM platform must be FedRAMP Authorized at the High "
        "Impact level. A current Authorization to Operate (ATO) letter must be provided "
        "with the proposal.",
        "",
        "D.1.4 The Contractor shall demonstrate experience operating a SOC in support "
        "of at least two (2) civilian federal agencies within the past five (5) years, "
        "with at least one engagement involving OT/ICS security monitoring. Relevant "
        "contract references with Contracting Officer contact information must be "
        "included in the proposal.",
        "",
        "D.1.5 The Contractor must have and maintain the following certifications "
        "across its proposed team: at least three (3) Certified Information Systems "
        "Security Professional (CISSP) holders; at least two (2) GIAC Certified "
        "Incident Handler (GCIH) holders; at least two (2) GIAC Security Essentials "
        "Certification (GSEC) holders; and at least one (1) GIAC Critical Infrastructure "
        "Protection (GICSP) holder for OT-specific monitoring.",
        "",
        "D.2 Preferred Technical Capabilities",
        "",
        "Contractors who possess the following additional capabilities will receive "
        "higher technical evaluation scores:",
        "",
        "D.2.1 Demonstrated capability to perform threat hunting using hypothesis-driven "
        "methodologies in federal civilian agency environments.",
        "D.2.2 Existing integrations with CISA's EINSTEIN 3A intrusion detection system "
        "or National Cybersecurity Protection System (NCPS).",
        "D.2.3 Artificial intelligence and machine learning-based User and Entity "
        "Behavior Analytics (UEBA) capabilities integrated with the SIEM.",
        "D.2.4 Deception technology (honeypots, honey tokens) deployment and management "
        "capability.",
        "D.2.5 Zero Trust Architecture advisory and implementation support capability.",
        "",
        "SECTION E - EVALUATION CRITERIA",
        "",
        "E.1 Basis for Award",
        "",
        "The Government will award a contract to the responsible offeror whose proposal "
        "represents the Best Value to the Government. The evaluation factors and their "
        "relative importance are as follows:",
        "",
        "Factor 1 - Technical Approach: Most Important. Offerors must demonstrate a "
        "thorough understanding of CIRA's requirements and provide a detailed, credible "
        "technical approach for delivering MSOC services. Sub-factors include: (a) SOC "
        "Operations Concept; (b) Detection and Response Methodology; (c) Technology "
        "Platform; (d) Staffing Plan and Key Personnel Qualifications.",
        "",
        "Factor 2 - Past Performance: Second Most Important. Offerors must provide at "
        "least three (3) and no more than five (5) relevant contract references from "
        "the past five (5) years demonstrating comparable MSOC services for federal "
        "agencies or large enterprise environments. The Government reserves the right "
        "to contact references and review CPARS records.",
        "",
        "Factor 3 - Price: Least Important, but must be fair and reasonable. Offerors "
        "must complete the Government-provided Price/Cost Workbook. Price realism "
        "analysis will be performed to determine whether the proposed prices reflect "
        "a clear understanding of the work and are consistent with the proposed technical "
        "approach.",
        "",
        "E.2 Evaluation Process",
        "",
        "Proposals will be evaluated by a Source Selection Evaluation Board (SSEB) "
        "consisting of technical, past performance, and price/cost evaluation teams. "
        "The SSEB will assign adjectival ratings of Outstanding, Good, Acceptable, "
        "Marginal, or Unacceptable to each technical sub-factor and an overall rating "
        "for technical approach and past performance. The Source Selection Authority "
        "(SSA) will make the final award decision based on a comparative assessment "
        "of offerors' strengths, weaknesses, risks, and prices.",
        "",
        "SECTION F - PROPOSAL SUBMISSION REQUIREMENTS",
        "",
        "F.1 Proposal Format and Page Limits",
        "",
        "Volume I - Technical Approach: Maximum 50 pages, 12-point Times New Roman font, "
        "1-inch margins, single-spaced. Diagrams and charts count toward the page limit.",
        "Volume II - Past Performance: Maximum 20 pages plus three to five reference "
        "sheets (one page per reference, not counted in page limit).",
        "Volume III - Price/Cost: No page limit. Must use Government-provided workbook.",
        "Volume IV - Representations and Certifications: Complete FAR 52.212-3 in full "
        "or provide SAM.gov annual certification date with DUNS/UEI number.",
        "",
        "F.2 Submission Instructions",
        "",
        "Proposals must be submitted electronically via the CIRA Secure Proposal Portal "
        "at proposals.cira.gov by January 17, 2025 at 4:00 PM Eastern Time. "
        "Late proposals will not be accepted. Technical questions must be submitted "
        "via the portal's Q&A function no later than December 20, 2024. "
        "The Government will post answers to all questions on SAM.gov by January 3, 2025.",
    ]

    for line in lines:
        doc.add_paragraph(line)

    path = os.path.join(DATA_DIR, "solicitation_raw.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_fms_requirements_raw():
    """Collins Aerospace FMS-4000 Software Requirements Specification - plain text.

    Content: Modeled on IEEE 29148-2018 System and Software Requirements Engineering
    standard (publicly published standard; content of requirements document is
    original/fictional). Fictional product: FMS-4000 Flight Management System.

    Agent task: apply heading styles (H1/H2/H3) per IEEE 29148 SRS structure,
    insert a Table of Contents, convert requirements into a numbered list,
    add a Requirements Traceability Matrix (RTM) table with at least 10 rows,
    and add a PROPRIETARY header and footer with 'Collins Aerospace / FMS-4000 SRS'.
    Save as fms_srs_final.docx.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    lines = [
        "SOFTWARE REQUIREMENTS SPECIFICATION",
        "FMS-4000 Flight Management System",
        "Document Number: CA-FMS4-SRS-001",
        "Revision: C",
        "Date: November 22, 2024",
        "Classification: PROPRIETARY - COLLINS AEROSPACE",
        "",
        "Prepared by: Systems Engineering, Avionics Products Division",
        "Collins Aerospace, 400 Collins Road NE, Cedar Rapids, IA 52498",
        "",
        "Reviewed by: Chief Systems Engineer, Software Safety Manager, "
        "DER (Designated Engineering Representative)",
        "",
        "WARNING: This document contains proprietary information of Collins Aerospace "
        "and is furnished for the sole purpose of evaluating the FMS-4000 Flight "
        "Management System development program. Reproduction, disclosure, or use "
        "without specific written authorization of Collins Aerospace is prohibited.",
        "",
        "REVISION HISTORY",
        "",
        "Revision A - June 10, 2023: Initial release for internal review.",
        "Revision B - February 14, 2024: Incorporated CDR action items; updated "
        "navigation requirements per FAA AC 20-138D.",
        "Revision C - November 22, 2024: Added performance computation requirements "
        "for EASA CS-25 Amendment 27; revised display requirements per ARINC 661.",
        "",
        "TABLE OF CONTENTS",
        "(Insert Table of Contents here)",
        "",
        "1. INTRODUCTION",
        "",
        "1.1 Purpose",
        "",
        "This Software Requirements Specification (SRS) defines the software requirements "
        "for the FMS-4000 Flight Management System. The FMS-4000 is an avionics system "
        "designed for installation in transport category aircraft certified under FAR Part 25 "
        "and EASA CS-25. This document establishes requirements for the FMS-4000 software "
        "at the system level and serves as the top-level software requirements baseline "
        "for the program. It shall be used as the reference document for software design, "
        "implementation, testing, and certification activities conducted in accordance "
        "with DO-178C Software Considerations in Airborne Systems and Equipment Certification "
        "and DO-254 Design Assurance Guidance for Airborne Electronic Hardware.",
        "",
        "1.2 Scope",
        "",
        "The FMS-4000 software includes the following functional subsystems: Navigation "
        "Management System (NMS), Performance Computation System (PCS), Route Management "
        "System (RMS), Display Management System (DMS), Data Communication System (DCS), "
        "and Built-In Test Equipment (BITE). This SRS covers all software components at "
        "Software Level A (catastrophic failure condition) and Software Level B (hazardous "
        "failure condition) as determined by the System Safety Assessment (SSA) per "
        "ARP4761 Guidelines and Methods for Conducting the Safety Assessment Process on "
        "Civil Airborne Systems and Equipment.",
        "",
        "1.3 Applicable Documents",
        "",
        "The following documents are applicable to the requirements in this SRS:",
        "FAA Advisory Circular AC 20-138D, Airworthiness Approval of Positioning and "
        "Navigation Systems (2016).",
        "FAA Technical Standard Order TSO-C115e, Airborne Area Navigation Equipment "
        "Using Multi-Sensor Inputs (2012).",
        "RTCA DO-178C, Software Considerations in Airborne Systems and Equipment "
        "Certification (2011).",
        "RTCA DO-254, Design Assurance Guidance for Airborne Electronic Hardware (2000).",
        "RTCA DO-236C, Minimum Aviation System Performance Standards: Required Navigation "
        "Performance for Area Navigation (2013).",
        "RTCA DO-283B, Minimum Operational Performance Standards for Required Navigation "
        "Performance for Area Navigation (2022).",
        "ARINC 702A-3, Advanced Flight Management Computer System (2006).",
        "ARINC 661-6, Cockpit Display System Interfaces to User Systems (2015).",
        "ARP4761, Guidelines and Methods for Conducting the Safety Assessment Process "
        "on Civil Airborne Systems and Equipment (1996).",
        "EASA CS-25 Amendment 27, Certification Specifications and Acceptable Means "
        "of Compliance for Large Aeroplanes (2023).",
        "",
        "2. OVERALL DESCRIPTION",
        "",
        "2.1 Product Perspective",
        "",
        "The FMS-4000 is the successor to the FMS-3600, which has accumulated over "
        "12 million flight hours on narrowbody and widebody commercial transport aircraft. "
        "The FMS-4000 provides enhanced four-dimensional (4D) trajectory management, "
        "Required Navigation Performance with Authorization Required (RNP AR) capability "
        "down to 0.1 nautical mile total system error, Advanced Continuous Descent Approach "
        "(ACDA) procedures, and real-time datalink integration with ACARS and ADS-C. "
        "The system interfaces with the following aircraft systems: Inertial Reference "
        "System (IRS), Global Navigation Satellite System (GNSS), VHF Navigation Radio "
        "(VOR/DME/ILS), Air Data Computer (ADC), Engine Control system (FADEC), "
        "Autopilot Flight Director System (AFDS), Control Display Units (CDU), "
        "Multifunction Control Display Units (MCDU), and Aircraft Communications "
        "Addressing and Reporting System (ACARS).",
        "",
        "2.2 Product Functions",
        "",
        "The FMS-4000 provides the following primary functions: four-dimensional trajectory "
        "computation and optimization; continuous position determination using sensor fusion "
        "of GNSS, IRS, and radio navigation inputs; route planning, modification, and "
        "management including SID/STAR/approach procedure loading from the navigation "
        "database; fuel and performance computation including takeoff performance, "
        "en-route fuel predictions, and landing performance; guidance command generation "
        "for lateral and vertical navigation; display of navigation, performance, and "
        "flight plan data on CDU/MCDU screens; and flight crew alerting for system "
        "anomalies and navigation integrity conditions.",
        "",
        "2.3 User Characteristics",
        "",
        "The primary users of the FMS-4000 are certificated airline transport pilots "
        "(ATP) and commercial pilots operating under FAR Part 121 and FAR Part 135 "
        "regulations. All users are assumed to have completed aircraft type-specific "
        "FMS training in accordance with the airline's FAA-approved training program. "
        "The FMS-4000 shall be operable by flight crew wearing standard aviation gloves "
        "using the CDU/MCDU line select keys and alphanumeric keyboard.",
        "",
        "3. NAVIGATION MANAGEMENT REQUIREMENTS",
        "",
        "3.1 Position Determination Requirements",
        "",
        "The Navigation Management System shall compute a best estimate position using "
        "a Kalman filter-based multi-sensor fusion algorithm that continuously integrates "
        "inputs from all available and healthy navigation sensors.",
        "",
        "The FMS shall maintain a navigation solution with a total system error (TSE) "
        "not exceeding 0.1 nautical miles (185 meters) with a probability of 10 to the "
        "power of negative 5 per flight hour for RNP AR operations, as required by "
        "RTCA DO-283B Section 2.1.2.",
        "",
        "The FMS shall compute and display Estimated Position Uncertainty (EPU) for the "
        "cross-track (XTK) axis and along-track (ATK) axis at a minimum update rate "
        "of 1 Hz.",
        "",
        "The FMS shall declare a navigation solution invalid and alert the flight crew "
        "within 6 seconds if EPU exceeds the current Required Navigation Performance "
        "(RNP) value, as specified in the active procedure or entered by the flight crew.",
        "",
        "The FMS shall provide GNSS-based position computation compliant with ICAO "
        "Annex 10 Volume I, using all available constellations (GPS, GLONASS, Galileo, "
        "BeiDou) with SBAS augmentation (WAAS, EGNOS) when available.",
        "",
        "The FMS shall implement Receiver Autonomous Integrity Monitoring (RAIM) and "
        "Fault Detection and Exclusion (FDE) algorithms compliant with RTCA DO-229F "
        "Section 2.1.1.",
        "",
        "The FMS shall maintain inertial position propagation using IRS inputs for a "
        "minimum of 15 minutes following complete GNSS signal loss, with TSE not "
        "exceeding 0.5 nautical miles per hour of IRS propagation.",
        "",
        "3.2 Navigation Database Requirements",
        "",
        "The FMS shall store and process a navigation database conforming to ARINC 424 "
        "Navigation System Data Base Standard (current cycle).",
        "",
        "The navigation database shall support storage of at least 250,000 waypoints, "
        "including latitude/longitude-defined waypoints, NAVAID-defined waypoints, "
        "airport waypoints, runway threshold waypoints, and fix-radial-distance waypoints.",
        "",
        "The FMS shall support navigation database updates on a 28-day AIRAC cycle "
        "without requiring software version changes. Database updates shall be loadable "
        "via the Portable Data Loader (PDL) interface in less than 8 minutes.",
        "",
        "The FMS shall include procedures for all airports with ICAO airport identifiers "
        "in the AIRAC cycle, including SIDs, STARs, instrument approach procedures, "
        "and airport ground routes.",
        "",
        "The FMS shall perform a cyclic redundancy check (CRC-32) on the loaded navigation "
        "database and alert the flight crew if database integrity is compromised.",
        "",
        "4. PERFORMANCE COMPUTATION REQUIREMENTS",
        "",
        "4.1 Takeoff Performance",
        "",
        "The FMS shall compute V-speeds (V1, VR, V2) and runway-limited takeoff weight "
        "for the active aircraft model, considering: pressure altitude and ambient "
        "temperature at the departure airport; selected runway length, slope, and "
        "surface condition; selected engine anti-ice and air conditioning pack selection; "
        "actual takeoff weight entered by the flight crew; and active obstacle/climb "
        "gradient requirements.",
        "",
        "The FMS shall display a takeoff performance advisory within 60 seconds of "
        "flight crew confirmation of takeoff data, and shall alert the flight crew if "
        "computed V1 speed is greater than VR speed.",
        "",
        "4.2 En-Route Fuel and Performance",
        "",
        "The FMS shall compute predicted fuel burn, time, and distance for each flight "
        "plan leg using the aircraft performance model stored in the operational flight "
        "program (OFP) database.",
        "",
        "The FMS shall continuously update the predicted fuel at destination (EFOB) "
        "during flight and shall alert the flight crew when EFOB falls below the "
        "minimum reserve fuel quantity entered by the flight crew.",
        "",
        "The FMS shall compute optimal cruise altitude using step-climb optimization "
        "over the route, considering planned fuel load, aircraft gross weight, and "
        "applicable airspace constraints.",
        "",
        "The FMS shall compute and display the Cost Index (CI)-based optimal cruise "
        "speed (ECON speed) for values of CI from 0 to 999, where CI = 0 represents "
        "maximum range cruise and CI = 999 represents maximum speed cruise.",
        "",
        "The fuel prediction error, defined as the difference between predicted and "
        "actual fuel consumption for a given flight segment, shall not exceed plus or "
        "minus 2% for flight segments between 500 and 3000 nautical miles in length "
        "under standard atmosphere conditions.",
        "",
        "4.3 Landing Performance",
        "",
        "The FMS shall compute in-flight landing performance advisory data including "
        "Vref approach speed, landing distance on the selected runway, and required "
        "runway length for the current landing configuration.",
        "",
        "The FMS shall update landing performance advisory data when the flight crew "
        "changes the selected runway, landing weight, flap setting, or brake setting.",
        "",
        "5. ROUTE MANAGEMENT REQUIREMENTS",
        "",
        "5.1 Flight Plan Entry and Modification",
        "",
        "The FMS shall provide flight plan entry capability through the CDU/MCDU keyboard "
        "interface, accepting origin and destination airports, alternate airports, "
        "company routes, ATC clearance routes in ICAO format, and individual waypoints "
        "entered by identifier, latitude/longitude, or place-bearing-distance.",
        "",
        "The FMS shall allow flight crew modification of the active flight plan in flight, "
        "including: direct-to routing; insertion, deletion, or repositioning of waypoints; "
        "altitude and speed constraint entry; and procedure modification.",
        "",
        "The FMS shall maintain a secondary (COPY) flight plan with independent editing "
        "capability, allowing preparation of alternate routing without affecting the "
        "active flight plan.",
        "",
        "Modifications to the active flight plan shall be activated by explicit flight "
        "crew selection of the EXECUTE function key. The FMS shall not automatically "
        "modify the active flight plan trajectory without flight crew confirmation.",
        "",
        "The FMS shall compute and display the effect of proposed flight plan changes "
        "on the predicted arrival time, fuel consumption, and fuel at destination before "
        "the flight crew executes the modification.",
        "",
        "5.2 Special Procedure Handling",
        "",
        "The FMS shall support all RNAV and RNP approach procedure types including "
        "RNAV (GPS) approaches to LNAV, LNAV/VNAV, LPV, and LP minima; RNP AR "
        "approaches with curved path segments (radius-to-fix legs) and vertical "
        "path angle (VPA) requirements; and RNAV (RNP) departure and arrival procedures.",
        "",
        "The FMS shall enforce RF leg (Radius to Fix) path accuracy requirements and "
        "shall alert the flight crew if the navigation solution is insufficient to "
        "maintain the required path accuracy for an RNP AR procedure.",
        "",
        "6. DISPLAY MANAGEMENT REQUIREMENTS",
        "",
        "6.1 CDU/MCDU Display Interface",
        "",
        "The FMS shall interface with Control Display Units (CDU) and Multifunction "
        "Control Display Units (MCDU) conforming to ARINC 739A and ARINC 661-6 "
        "interface specifications.",
        "",
        "The FMS shall display information using the standard ARINC 739A character set "
        "and shall support display of information in the active flight crew language "
        "(English, French, German, Spanish, Chinese Simplified, Japanese) based on "
        "the aircraft operator configuration.",
        "",
        "The FMS CDU/MCDU display pages shall be organized according to the page "
        "hierarchy defined in ARINC 702A-3 Appendix A, with context-sensitive page "
        "access based on the current phase of flight.",
        "",
        "Navigation map display data shall be formatted according to the ARINC 661-6 "
        "User System (US) protocol and shall include: active route with waypoints and "
        "constraints; RNP/EPU arc display for current navigation performance; traffic "
        "data from TCAS/ADS-B; terrain data from the EGPWS database; weather overlay "
        "from datalink weather services; and airport map display when within 30 nautical "
        "miles of a database airport.",
        "",
        "6.2 Alerting Requirements",
        "",
        "The FMS shall generate flight crew alerts conforming to the EICAS/ECAM "
        "alerting philosophy defined in the aircraft-level Human Factors Design Standard.",
        "",
        "FMS Alert Level A (Warning - Red): Conditions requiring immediate flight crew "
        "awareness and action. Alert shall be accompanied by aural alert and master "
        "warning indication. Navigation integrity failure during approach below 1000 ft "
        "AGL shall be classified as Level A.",
        "",
        "FMS Alert Level B (Caution - Amber): Conditions requiring timely flight crew "
        "awareness. Alert shall be accompanied by aural chime and master caution "
        "indication. EPU exceeding RNP value during en-route operations shall be "
        "classified as Level B.",
        "",
        "FMS Alert Level C (Advisory - Cyan/Blue): Conditions requiring flight crew "
        "awareness but not immediate action. No aural alert required.",
        "",
        "The FMS shall suppress nuisance alerts using a minimum 5-second filter for "
        "transient sensor anomalies before generating Level A or Level B alerts.",
        "",
        "7. SAFETY AND RELIABILITY REQUIREMENTS",
        "",
        "7.1 Software Safety Requirements",
        "",
        "The FMS-4000 software shall be developed to DO-178C Software Level A (DAL A) "
        "for all software components whose failure would result in a catastrophic failure "
        "condition as determined by the System Safety Assessment.",
        "",
        "The FMS shall perform Built-In Test Equipment (BITE) power-on self-test (POST) "
        "within 90 seconds of power application and shall inhibit FMS output commands "
        "until POST completes successfully.",
        "",
        "The FMS shall implement watchdog monitoring to detect software execution "
        "anomalies and shall transition to a safe state within 100 milliseconds of "
        "detecting a software hang or computational error.",
        "",
        "The FMS shall implement cross-channel monitoring with an independent monitoring "
        "channel that continuously compares navigation solution outputs and alerts the "
        "flight crew if the deviation between the active and monitoring channels exceeds "
        "the defined threshold.",
        "",
        "7.2 Reliability Requirements",
        "",
        "The FMS shall achieve a Mean Time Between Failure (MTBF) of not less than "
        "20,000 flight hours for hardware-related failures.",
        "",
        "The FMS software shall be designed to achieve a probability of Loss of Function "
        "not exceeding 10 to the power of negative 9 per flight hour for catastrophic "
        "failure conditions.",
        "",
        "The FMS shall include Bite Fault Isolation capability that isolates detected "
        "hardware failures to the applicable line replaceable unit (LRU) with a false "
        "removal rate not exceeding 5%.",
        "",
        "8. INTERFACE REQUIREMENTS",
        "",
        "8.1 External Interface Summary",
        "",
        "The FMS-4000 shall receive inputs from and transmit outputs to the following "
        "aircraft systems via the defined ARINC 429 and ARINC 664 (AFDX) interfaces:",
        "",
        "IRS-1, IRS-2, IRS-3 (ARINC 429 High Speed): Position, velocity, attitude, "
        "heading, accelerations. Label codes per ARINC 429 Chapter 10.",
        "",
        "GNSS Receiver-1, GNSS Receiver-2 (ARINC 429 High Speed): GPS position, "
        "velocity, EPE, satellite status, RAIM availability.",
        "",
        "VHF Navigation Radio-1, VHF Navigation Radio-2 (ARINC 429 High Speed): "
        "VOR bearing and DME distance, ILS localizer deviation and glide slope deviation, "
        "station identification.",
        "",
        "Air Data Computer-1, Air Data Computer-2 (ARINC 429 High Speed): "
        "Barometric altitude, indicated airspeed, true airspeed, total air temperature, "
        "Mach number, static pressure.",
        "",
        "Engine Management System / FADEC-1, FADEC-2, FADEC-3, FADEC-4 (ARINC 429): "
        "N1/N2 speed, fuel flow, EGT, thrust mode.",
        "",
        "Autopilot Flight Director System (AFDX ARINC 664): FMS guidance commands "
        "including lateral steering command, vertical speed command, airspeed target, "
        "and flight mode annunciation data.",
        "",
        "Control Display Units - CDU-1, CDU-2, CDU-3 (ARINC 429 High Speed): "
        "Crew inputs and display data exchange per ARINC 739A.",
        "",
        "ACARS Management Unit (ARINC 429): Datalink uplinks (oceanic clearances, "
        "weather, D-ATIS, PDC) and downlinks (position reports, OOOI times, ADS-C).",
        "",
        "REQUIREMENTS TRACEABILITY MATRIX (RTM)",
        "",
        "(Insert Requirements Traceability Matrix table here. The RTM shall map each "
        "requirement identifier in this SRS to the corresponding high-level system "
        "requirement in the System Requirements Specification CA-FMS4-SYRS-001, the "
        "applicable regulatory source document, and the planned verification method. "
        "Use columns: Req ID, Req Title, Parent Req (SYRS), Regulatory Source, "
        "Verification Method, Verification Procedure.)",
        "",
        "RTM-001: NMS-001 (Position TSE <= 0.1 NM) -> SYRS-NAV-001 -> DO-283B 2.1.2 "
        "-> Test -> TP-FMS4-NAV-001",
        "RTM-002: NMS-002 (EPU Display 1 Hz) -> SYRS-NAV-002 -> ARINC 702A-3 -> "
        "Demonstration -> TP-FMS4-NAV-002",
        "RTM-003: NMS-003 (Navigation Alert <6 sec) -> SYRS-NAV-003 -> DO-236C -> "
        "Test -> TP-FMS4-NAV-003",
        "RTM-004: NMS-004 (RAIM/FDE) -> SYRS-NAV-004 -> DO-229F 2.1.1 -> "
        "Analysis+Test -> TP-FMS4-NAV-004",
        "RTM-005: NMS-005 (IRS Propagation 15 min) -> SYRS-NAV-005 -> AC 20-138D -> "
        "Test -> TP-FMS4-NAV-005",
        "RTM-006: NDB-001 (Database Capacity 250k WPT) -> SYRS-DAT-001 -> ARINC 424 -> "
        "Test -> TP-FMS4-NDB-001",
        "RTM-007: NDB-002 (AIRAC Update <8 min) -> SYRS-DAT-002 -> Internal -> "
        "Demonstration -> TP-FMS4-NDB-002",
        "RTM-008: PCS-001 (EFOB Alert) -> SYRS-PERF-001 -> CS-25.1337 -> "
        "Test -> TP-FMS4-PCS-001",
        "RTM-009: PCS-002 (Fuel Prediction Error +-2%) -> SYRS-PERF-002 -> Internal -> "
        "Analysis+Test -> TP-FMS4-PCS-002",
        "RTM-010: PCS-003 (CI Range 0-999) -> SYRS-PERF-003 -> ARINC 702A-3 -> "
        "Test -> TP-FMS4-PCS-003",
        "RTM-011: RMS-001 (Flight Plan Modification) -> SYRS-RMS-001 -> ARINC 702A-3 -> "
        "Test -> TP-FMS4-RMS-001",
        "RTM-012: RMS-002 (EXECUTE Confirmation) -> SYRS-RMS-002 -> Human Factors -> "
        "Test -> TP-FMS4-RMS-002",
        "RTM-013: RMS-003 (RNP AR Procedures) -> SYRS-RMS-003 -> DO-283B -> "
        "Test -> TP-FMS4-RMS-003",
        "RTM-014: DMS-001 (ARINC 661-6 Compliance) -> SYRS-DMS-001 -> ARINC 661-6 -> "
        "Analysis+Test -> TP-FMS4-DMS-001",
        "RTM-015: DMS-002 (Alert Suppression 5 sec) -> SYRS-DMS-002 -> HF Standard -> "
        "Test -> TP-FMS4-DMS-002",
        "RTM-016: SAF-001 (POST <90 sec) -> SYRS-SAF-001 -> DO-178C -> "
        "Test -> TP-FMS4-SAF-001",
        "RTM-017: SAF-002 (Watchdog 100 ms) -> SYRS-SAF-002 -> DO-178C DAL-A -> "
        "Test -> TP-FMS4-SAF-002",
        "RTM-018: REL-001 (MTBF >20000 hr) -> SYRS-REL-001 -> Internal -> "
        "Analysis -> TP-FMS4-REL-001",
    ]

    for line in lines:
        doc.add_paragraph(line)

    path = os.path.join(DATA_DIR, "fms_requirements_raw.docx")
    doc.save(path)
    print(f"Created: {path}")


def _make_placeholder_png(w=480, h=50):
    """Create a solid light-gray PNG in memory (no PIL needed)."""
    raw = b""
    for _ in range(h):
        raw += b"\x00" + b"\xcc\xcc\xcc" * w  # filter-byte + gray RGB pixels
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)

    def chunk(t, d):
        c = t + d
        return struct.pack(">I", len(d)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def create_environmental_compliance_report_raw():
    """Annual Environmental Compliance Report - unformatted, with tables and figures.

    Content modeled on EPA NPDES/CAA compliance reporting requirements
    and Washington State Department of Ecology submission standards.
    All data is fictional but uses realistic values, units, and regulatory
    references drawn from public-domain EPA guidance documents.

    Agent task: apply heading styles, insert section breaks with per-section
    page numbering (Roman/Arabic), TOC, captions on tables/figures,
    cross-references replacing [see Table X]/[see Figure X] placeholders,
    and a Table of Figures. Save as environmental_compliance_final.docx.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    placeholder_png = _make_placeholder_png()

    # ---- COVER PAGE ----
    cover = [
        "",
        "",
        "PACIFIC NORTHWEST ENVIRONMENTAL SERVICES, INC.",
        "",
        "2024 ANNUAL ENVIRONMENTAL COMPLIANCE REPORT",
        "",
        "Cascade Industrial Park Facility",
        "7200 Marine View Drive",
        "Tacoma, Washington 98422",
        "",
        "NPDES Permit No. WA-0024571",
        "Air Quality Permit No. AQ-PSR-2019-0382",
        "RCRA Generator ID WAD 091 837 216",
        "",
        "Prepared for:",
        "Washington State Department of Ecology",
        "Southwest Regional Office",
        "",
        "Prepared by:",
        "Pacific Northwest Environmental Services, Inc.",
        "Environmental Compliance Division",
        "",
        "December 2024",
        "",
        "",
    ]
    for line in cover:
        doc.add_paragraph(line)

    # ---- EXECUTIVE SUMMARY ----
    exec_summary = [
        "Executive Summary",
        "",
        "This Annual Environmental Compliance Report summarizes the environmental "
        "monitoring activities, regulatory compliance status, and corrective actions "
        "undertaken at the Cascade Industrial Park facility during calendar year 2024. "
        "The facility, operated by Pacific Northwest Environmental Services, Inc. "
        "(PNWES), occupies 150 acres in Tacoma, Washington and employs approximately "
        "340 personnel across manufacturing, warehousing, and administrative operations.",
        "",
        "Overall compliance performance for the reporting period was 98.7 percent, "
        "calculated as the ratio of compliant monitoring events to total required "
        "monitoring events across all media (air, water, and waste). The facility "
        "received three minor Notices of Violation from the Department of Ecology "
        "during 2024, all of which were addressed through corrective actions completed "
        "within the required timeframes.",
        "",
        "Key accomplishments during 2024 include: a 12 percent reduction in total "
        "hazardous waste generation compared to 2023 baseline; successful completion "
        "of the Phase II groundwater remediation milestone for benzene contamination "
        "in the former solvent storage area; and installation of a regenerative thermal "
        "oxidizer (RTO) on Paint Line 3 that reduced volatile organic compound (VOC) "
        "emissions by 87 percent.",
        "",
        "All ambient air monitoring stations recorded pollutant concentrations below "
        "National Ambient Air Quality Standards (NAAQS) throughout the reporting period. "
        "Effluent discharge quality at all three permitted outfalls consistently met "
        "NPDES permit limits with the exception of a single pH exceedance at Outfall 002 "
        "in March 2024, which was attributed to an equipment malfunction in the "
        "neutralization system and corrected within 48 hours.",
        "",
    ]
    for line in exec_summary:
        doc.add_paragraph(line)

    # ---- SECTION 1: INTRODUCTION ----
    sec1 = [
        "1. Introduction",
        "",
        "1.1 Purpose and Scope",
        "",
        "This report has been prepared in accordance with the requirements of NPDES "
        "Permit No. WA-0024571 (Section S9.A), Air Quality Permit No. AQ-PSR-2019-0382 "
        "(Condition 12.3), and Washington Administrative Code (WAC) Chapter 173-303 "
        "Dangerous Waste Regulations. The report provides a comprehensive summary of "
        "environmental monitoring data, compliance status, and corrective actions for "
        "the twelve-month period from January 1 through December 31, 2024.",
        "",
        "1.2 Regulatory Framework",
        "",
        "Facility operations are regulated under the following federal and state "
        "environmental statutes and implementing regulations:",
        "",
        "The Clean Water Act (CWA), as implemented through the Washington Water "
        "Pollution Control Act (RCW 90.48) and the NPDES permit program administered "
        "by the Department of Ecology. The facility holds Individual NPDES Permit "
        "WA-0024571, which authorizes discharge of treated process wastewater and "
        "stormwater from three permitted outfalls to Commencement Bay.",
        "",
        "The Clean Air Act (CAA), as implemented through the Washington Clean Air Act "
        "(RCW 70A.15) and Puget Sound Clean Air Agency (PSCAA) Regulation I. The "
        "facility operates under Air Quality Permit AQ-PSR-2019-0382, which establishes "
        "emission limits for particulate matter (PM2.5 and PM10), nitrogen oxides (NOx), "
        "sulfur dioxide (SO2), and volatile organic compounds (VOCs).",
        "",
        "The Resource Conservation and Recovery Act (RCRA), as implemented through "
        "WAC 173-303 Dangerous Waste Regulations. The facility is classified as a "
        "Large Quantity Generator (LQG) under EPA Generator ID WAD 091 837 216.",
        "",
        "The Comprehensive Environmental Response, Compensation, and Liability Act "
        "(CERCLA) and the Washington Model Toxics Control Act (MTCA, RCW 70A.305). "
        "A portion of the facility is subject to ongoing remediation under MTCA "
        "Agreed Order DE 14-TCPSR-4271 for historical solvent contamination.",
        "",
        "1.3 Facility Description",
        "",
        "The Cascade Industrial Park facility is located at 7200 Marine View Drive in "
        "Tacoma, Pierce County, Washington (Latitude 47.2629 N, Longitude 122.4443 W). "
        "The 150-acre campus includes three manufacturing buildings (Buildings A, B, and "
        "C), two warehouses, an administrative office complex, a wastewater treatment "
        "plant, a hazardous waste storage facility (HWSF), and approximately 35 acres "
        "of stormwater management infrastructure including two retention ponds and "
        "an engineered bioswale treatment system.",
        "",
        "Primary manufacturing operations include metal fabrication (Building A), "
        "surface finishing and painting (Building B), and electronic assembly "
        "(Building C). The facility operates two shifts per day, five days per week, "
        "with occasional weekend operations during peak production periods.",
        "",
    ]
    for line in sec1:
        doc.add_paragraph(line)

    # ---- SECTION 2: AIR QUALITY ----
    doc.add_paragraph("2. Air Quality Monitoring")
    doc.add_paragraph("")
    doc.add_paragraph("2.1 Emission Sources and Controls")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The facility operates 14 permitted emission units, including three natural "
        "gas-fired boilers (EU-01 through EU-03, rated at 25 MMBtu/hr each), two paint "
        "spray booths with dry filter particulate controls (EU-04, EU-05), one "
        "regenerative thermal oxidizer controlling VOC emissions from Paint Line 3 "
        "(EU-06), four parts washers with vapor recovery systems (EU-07 through EU-10), "
        "and four fugitive dust sources associated with material handling operations "
        "(EU-11 through EU-14). The RTO installed on Paint Line 3 in Q2 2024 achieves "
        "a destruction efficiency of 98.2 percent for VOCs, as verified by EPA Method "
        "25A performance testing conducted in June 2024."
    )
    doc.add_paragraph("")
    doc.add_paragraph("2.2 Ambient Monitoring Results")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Three ambient air monitoring stations (AMS-North, AMS-South, and AMS-East) "
        "operated continuously during the reporting period. Monthly average "
        "concentrations for PM2.5 and NOx are summarized in the following table. "
        "As shown in [see Table 1], all monthly averages remained below applicable "
        "NAAQS limits (PM2.5 annual standard: 12.0 ug/m3; NOx annual standard not "
        "to exceed 53 ppb)."
    )
    doc.add_paragraph("")

    # TABLE 1: Air Quality Monitoring Data
    tbl1 = doc.add_table(rows=13, cols=3, style="Table Grid")
    tbl1.cell(0, 0).text = "Month"
    tbl1.cell(0, 1).text = "PM2.5 (ug/m3)"
    tbl1.cell(0, 2).text = "NOx (ppb)"
    monthly_data = [
        ("January", "8.4", "31.2"),
        ("February", "7.9", "28.7"),
        ("March", "9.1", "33.4"),
        ("April", "7.2", "26.8"),
        ("May", "6.8", "24.1"),
        ("June", "5.9", "21.3"),
        ("July", "8.3", "19.8"),
        ("August", "10.7", "22.4"),
        ("September", "9.2", "25.6"),
        ("October", "7.6", "29.3"),
        ("November", "8.1", "32.1"),
        ("December", "8.8", "34.7"),
    ]
    for i, (month, pm, nox) in enumerate(monthly_data, start=1):
        tbl1.cell(i, 0).text = month
        tbl1.cell(i, 1).text = pm
        tbl1.cell(i, 2).text = nox
    doc.add_paragraph("")

    doc.add_paragraph("2.3 Compliance Status")
    doc.add_paragraph("")
    doc.add_paragraph(
        "All emission units operated within permitted limits throughout the reporting "
        "period. Stack testing conducted in accordance with 40 CFR Part 60 confirmed "
        "compliance for all tested units. The annual emissions trend is shown in "
        "[see Figure 1] below."
    )
    doc.add_paragraph("")

    # FIGURE 1 placeholder
    doc.add_picture(io.BytesIO(placeholder_png), width=Inches(5))
    doc.add_paragraph("")

    # ---- SECTION 3: WATER QUALITY ----
    doc.add_paragraph("3. Water Quality and Stormwater Management")
    doc.add_paragraph("")
    doc.add_paragraph("3.1 Discharge Monitoring")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The facility discharges treated process wastewater and stormwater through "
        "three NPDES-permitted outfalls. Outfall 001 receives treated process wastewater "
        "from the on-site wastewater treatment plant (WWTP). Outfall 002 receives "
        "non-contact cooling water from the manufacturing buildings. Outfall 003 "
        "receives stormwater runoff from the western parking lot and loading dock area "
        "after passage through the engineered bioswale treatment system."
    )
    doc.add_paragraph("")
    doc.add_paragraph("3.2 Effluent Monitoring Data")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Quarterly composite sampling results for Outfall 001 are presented in "
        "[see Table 2]. With the exception of the March 2024 pH exceedance "
        "(pH 5.8, below the permit minimum of 6.0), all parameters met NPDES "
        "permit limits throughout the year."
    )
    doc.add_paragraph("")

    # TABLE 2: Effluent Monitoring Data
    tbl2 = doc.add_table(rows=5, cols=5, style="Table Grid")
    headers2 = ["Quarter", "BOD5 (mg/L)", "TSS (mg/L)", "pH", "Oil & Grease (mg/L)"]
    for j, h in enumerate(headers2):
        tbl2.cell(0, j).text = h
    effluent_data = [
        ("Q1 2024", "18.3", "22.1", "5.8*", "4.2"),
        ("Q2 2024", "14.7", "19.8", "7.2", "3.8"),
        ("Q3 2024", "16.2", "21.4", "7.0", "4.5"),
        ("Q4 2024", "15.1", "18.6", "7.1", "3.6"),
    ]
    for i, row_data in enumerate(effluent_data, start=1):
        for j, val in enumerate(row_data):
            tbl2.cell(i, j).text = val
    doc.add_paragraph("* Exceedance of NPDES permit minimum pH limit (6.0)")
    doc.add_paragraph("")

    doc.add_paragraph("3.3 Stormwater BMP Effectiveness")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The site drainage infrastructure and best management practices (BMPs) are "
        "illustrated in [see Figure 2]. Stormwater monitoring at Outfall 003 during "
        "the four quarterly sampling events demonstrated TSS removal efficiency "
        "averaging 82 percent through the bioswale system, exceeding the 80 percent "
        "benchmark established in the facility's Stormwater Pollution Prevention Plan "
        "(SWPPP). The bioswale was expanded in Q3 2024 to accommodate increased runoff "
        "from the newly paved employee parking area."
    )
    doc.add_paragraph("")

    # FIGURE 2 placeholder
    doc.add_picture(io.BytesIO(placeholder_png), width=Inches(5))
    doc.add_paragraph("")

    # ---- SECTION 4: WASTE MANAGEMENT ----
    doc.add_paragraph("4. Waste Management")
    doc.add_paragraph("")
    doc.add_paragraph("4.1 Hazardous Waste Generation")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The facility generated a total of 47.3 tons of hazardous waste during "
        "2024, a 12 percent reduction from the 53.8 tons generated in 2023. "
        "Waste stream details are provided in [see Table 3]. All hazardous waste "
        "was manifested and shipped to permitted Treatment, Storage, and Disposal "
        "Facilities (TSDFs) within the 90-day accumulation period required by "
        "WAC 173-303-200."
    )
    doc.add_paragraph("")

    # TABLE 3: Hazardous Waste Streams
    tbl3 = doc.add_table(rows=7, cols=4, style="Table Grid")
    headers3 = ["Waste Stream", "RCRA Code", "Quantity (tons)", "Disposal Method"]
    for j, h in enumerate(headers3):
        tbl3.cell(0, j).text = h
    waste_data = [
        ("Spent halogenated solvents", "F001", "12.4", "Fuel blending"),
        ("Chromium-bearing sludge", "D007", "8.7", "Stabilization/landfill"),
        ("Paint waste (ignitable)", "D001", "11.2", "Incineration"),
        ("Spent acid solutions", "D002", "6.8", "Neutralization/treatment"),
        ("Waste trichloroethylene", "F001/U228", "4.1", "Solvent recovery"),
        ("Mixed metal hydroxide sludge", "D006/D007", "4.1", "Stabilization/landfill"),
    ]
    for i, row_data in enumerate(waste_data, start=1):
        for j, val in enumerate(row_data):
            tbl3.cell(i, j).text = val
    doc.add_paragraph("")

    doc.add_paragraph("4.2 Waste Minimization Initiatives")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The facility implemented three waste minimization projects during 2024 as "
        "part of the five-year Pollution Prevention Plan required under WAC 173-307. "
        "First, the solvent recovery system on Paint Line 2 was upgraded from a "
        "single-stage to a double-stage distillation unit, increasing solvent recovery "
        "rate from 68 percent to 91 percent and reducing F001 waste generation by "
        "approximately 3.2 tons per year. Second, the chromium plating bath in "
        "Building A was converted from hexavalent chromium (Cr VI) to trivalent "
        "chromium (Cr III), eliminating D007 waste from that process line. Third, "
        "Building C transitioned from solvent-based conformal coatings to UV-curable "
        "alternatives, eliminating VOC emissions and associated D001 waste from "
        "electronic assembly operations."
    )
    doc.add_paragraph("")
    doc.add_paragraph("4.3 Storage and Disposal Compliance")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The Hazardous Waste Storage Facility (HWSF) was inspected weekly throughout "
        "2024 in accordance with WAC 173-303-320. All inspections documented proper "
        "container labeling, secondary containment integrity, and compliance with "
        "accumulation time limits. The facility layout showing waste storage areas and "
        "designated satellite accumulation points is shown in [see Figure 3]."
    )
    doc.add_paragraph("")

    # FIGURE 3 placeholder
    doc.add_picture(io.BytesIO(placeholder_png), width=Inches(5))
    doc.add_paragraph("")

    # ---- SECTION 5: COMPLIANCE SUMMARY ----
    doc.add_paragraph("5. Compliance Summary and Corrective Actions")
    doc.add_paragraph("")
    doc.add_paragraph("5.1 Notices of Violation")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Three Notices of Violation (NOVs) were issued by the Department of Ecology "
        "during the 2024 reporting period:"
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "NOV-2024-SW-0312 (March 15, 2024): pH exceedance at Outfall 002. "
        "Measured pH of 5.8 fell below the NPDES permit minimum of 6.0. Root cause: "
        "malfunction of the automated pH adjustment system in the neutralization tank. "
        "Corrective action: replaced pH sensor probe and recalibrated controller. "
        "Completed March 17, 2024."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "NOV-2024-AQ-0647 (June 28, 2024): Late submission of semi-annual VOC "
        "emission report. Report was submitted 4 business days past the June 15 "
        "deadline. Root cause: staffing transition in the environmental compliance "
        "department. Corrective action: implemented automated compliance calendar "
        "with 30-day and 7-day advance notifications. Completed July 2024."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "NOV-2024-HW-0891 (September 3, 2024): Container labeling deficiency in "
        "satellite accumulation area (SAA) in Building B paint mixing room. One "
        "55-gallon drum of spent thinner lacked accumulation start date. Root cause: "
        "employee oversight during container staging. Corrective action: conducted "
        "refresher training for all Building B personnel on WAC 173-303-174 SAA "
        "requirements; installed additional labeling stations. Completed September "
        "2024."
    )
    doc.add_paragraph("")

    doc.add_paragraph("5.2 Corrective Actions Summary")
    doc.add_paragraph("")
    doc.add_paragraph(
        "All three NOVs were resolved within the timeframes specified by the "
        "Department of Ecology. Penalty amounts totaling $4,750 were assessed "
        "and paid. No repeat violations occurred for any of the cited conditions "
        "during the remainder of the reporting period."
    )
    doc.add_paragraph("")

    doc.add_paragraph("5.3 Planned Improvements for 2025")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The following environmental compliance improvement projects are planned "
        "for fiscal year 2025: upgrade of the WWTP clarifier system to improve TSS "
        "removal by an estimated 15 percent (budget: $380,000, Q2 2025); installation "
        "of continuous emission monitoring systems (CEMS) on Boilers EU-01 and EU-02 "
        "as required by the 2025 permit renewal (budget: $245,000, Q3 2025); and "
        "completion of Phase III groundwater remediation activities under MTCA Agreed "
        "Order DE 14-TCPSR-4271, including installation of three additional extraction "
        "wells and expansion of the on-site air stripping system (budget: $520,000, "
        "Q1-Q4 2025)."
    )
    doc.add_paragraph("")

    # ---- APPENDICES ----
    doc.add_paragraph("Appendices")
    doc.add_paragraph("")
    doc.add_paragraph("Appendix A: Laboratory Analytical Methods")
    doc.add_paragraph("")
    doc.add_paragraph(
        "All environmental samples collected during the 2024 monitoring program were "
        "analyzed by Cascade Analytical Laboratories, Inc. (Tacoma, WA), a Washington "
        "State Department of Ecology accredited laboratory (Accreditation No. "
        "C-934). The following EPA-approved analytical methods were employed:"
    )
    doc.add_paragraph("")
    doc.add_paragraph("EPA Method 524.2 - Volatile organic compounds in water by purge-and-trap GC/MS")
    doc.add_paragraph("EPA Method 200.7 - Metals in water by ICP-OES")
    doc.add_paragraph("EPA Method 8260C - Volatile organic compounds in soil and groundwater by GC/MS")
    doc.add_paragraph("EPA Method 8270E - Semi-volatile organic compounds by GC/MS")
    doc.add_paragraph("EPA Method 9045D - Soil and waste pH")
    doc.add_paragraph("SM 5210 B - Biochemical oxygen demand (BOD5)")
    doc.add_paragraph("SM 2540 D - Total suspended solids (TSS)")
    doc.add_paragraph("")

    doc.add_paragraph("Appendix B: Monitoring Equipment Calibration Records")
    doc.add_paragraph("")
    doc.add_paragraph(
        "All field monitoring equipment used during the 2024 sampling program was "
        "calibrated in accordance with manufacturer specifications and applicable EPA "
        "guidance. Calibration records are maintained on-site and available for "
        "regulatory inspection upon request. Equipment includes: YSI ProDSS "
        "multiparameter water quality meter (calibrated quarterly), Met One BAM-1020 "
        "continuous PM2.5 monitor (calibrated monthly per 40 CFR Part 58 Appendix A), "
        "and Thermo Scientific 42i NOx analyzer (calibrated weekly with certified "
        "span gas per 40 CFR Part 58 Appendix A)."
    )
    doc.add_paragraph("")

    path = os.path.join(DATA_DIR, "environmental_compliance_report_raw.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_census_press_release()
    create_meeting_notes_raw()
    create_company_memo_draft()
    create_patent_license_draft()
    create_oncology_protocol_raw()
    create_ap_walkthrough_raw()
    create_solicitation_raw()
    create_fms_requirements_raw()
    create_environmental_compliance_report_raw()
    print("All data files created successfully.")
