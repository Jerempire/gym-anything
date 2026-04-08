#!/usr/bin/env python3
"""
Verifier for resume_ats_remediation task.

Checks:
1. Contact info moved from Header to Body.
2. Dates aligned using Tabs (not spaces) with correct Tab Stop.
3. Heading styles applied.
4. VLM visual verification of alignment.
"""

import json
import os
import sys
import tempfile
import logging
import shutil

# Add utils path
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from writer_verification_utils import (
    copy_and_parse_document,
    get_document_text,
    vlm_verify_screenshot
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_resume_remediation(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Load metadata
    metadata = task_info.get('metadata', {})
    expected_tab_pos = metadata.get('expected_tab_position_inches', 6.0)
    
    # 1. Check Export Result (Timestamps & Existence)
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            export_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load export result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not export_result.get('output_exists', False):
        return {"passed": False, "score": 0, "feedback": "Output file resume_final.docx not found."}
    
    if not export_result.get('file_fresh', False):
        return {"passed": False, "score": 0, "feedback": "Output file was not saved during the task session."}

    # 2. Parse the DOCX
    success, doc, error, temp_dir = copy_and_parse_document(
        "/home/ga/Documents/resume_final.docx", copy_from_env, file_format='docx'
    )
    
    if not success:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse document: {error}"}

    score = 0
    feedback = []
    
    try:
        # --- Criterion A: Header Empty (20 pts) ---
        # "Move contact info... Header must be left empty"
        header_text = ""
        for section in doc.sections:
            for hp in section.header.paragraphs:
                header_text += hp.text.strip()
        
        if len(header_text) == 0:
            score += 20
            feedback.append("Header is empty (Success)")
        else:
            feedback.append(f"Header still contains text: '{header_text[:20]}...' (Fail)")

        # --- Criterion B: Contact Info in Body (20 pts) ---
        # "JORDAN LEE" should be in the first few paragraphs of the body
        contact_found = False
        # Check first 5 paragraphs
        for i, p in enumerate(doc.paragraphs[:10]):
            if "JORDAN LEE" in p.text:
                contact_found = True
                break
        
        if contact_found:
            score += 20
            feedback.append("Contact info found in body (Success)")
        else:
            feedback.append("Contact Name 'JORDAN LEE' not found in top of body (Fail)")

        # --- Criterion C: Tab Usage & Space Removal (20 pts) ---
        # Check paragraphs containing dates. They should have a \t character and NO run of >3 spaces.
        date_paragraphs = []
        target_substrings = ["June 2020", "Jan 2018", "May 2017", "Present"]
        
        tabs_used = 0
        spaces_removed = 0
        relevant_paras = 0

        for p in doc.paragraphs:
            # Identify if this is a job line
            if any(sub in p.text for sub in target_substrings) and ("TechFlow" not in p.text): 
                # TechFlow is the company line, usually separate. We target the Title + Date line.
                relevant_paras += 1
                
                if '\t' in p.text:
                    tabs_used += 1
                
                # Check for space runs (more than 3 spaces is suspicious if they used tabs)
                if "    " not in p.text:
                    spaces_removed += 1
        
        # We expect at least 3 such lines (2 jobs + 1 education)
        if relevant_paras >= 3:
            if tabs_used >= relevant_paras:
                score += 10
                feedback.append("Tabs used for separation (Success)")
            else:
                feedback.append(f"Tabs missing in {relevant_paras - tabs_used} lines (Fail)")
                
            if spaces_removed >= relevant_paras:
                score += 10
                feedback.append("Space runs removed (Success)")
            else:
                feedback.append("Excessive spaces still present (Fail)")
        else:
            feedback.append("Could not locate Job Title lines to verify tabs.")

        # --- Criterion D: Right Tab Stop (20 pts) ---
        # Check tab stops in relevant paragraphs
        # Expected: ~6.0 inches (approx 5486400 EMU)
        # Alignment: RIGHT (2)
        
        tab_stops_correct = 0
        from docx.enum.text import WD_TAB_ALIGNMENT
        
        for p in doc.paragraphs:
            if any(sub in p.text for sub in target_substrings) and ("TechFlow" not in p.text):
                has_correct_stop = False
                for stop in p.paragraph_format.tab_stops:
                    # Check position (within 0.5 inch tolerance)
                    pos_in = stop.position.inches
                    if 5.5 <= pos_in <= 6.5:
                        # Check alignment (Right is usually enum 2, but we check property)
                        if stop.alignment == WD_TAB_ALIGNMENT.RIGHT:
                            has_correct_stop = True
                
                if has_correct_stop:
                    tab_stops_correct += 1

        if tab_stops_correct >= 3:
            score += 20
            feedback.append(f"Right Tab Stops set correctly in {tab_stops_correct} lines (Success)")
        elif tab_stops_correct > 0:
            score += 10
            feedback.append(f"Right Tab Stops partial: {tab_stops_correct} lines (Partial)")
        else:
            feedback.append("No Right Tab Stops found at ~6.0 inches (Fail)")

        # --- Criterion E: Heading Styles (10 pts) ---
        headings = ["EXPERIENCE", "EDUCATION", "SKILLS"]
        headings_correct = 0
        for h_text in headings:
            for p in doc.paragraphs:
                if p.text.strip() == h_text:
                    if p.style and "Heading 1" in p.style.name:
                        headings_correct += 1
                    break
        
        if headings_correct == 3:
            score += 10
            feedback.append("Heading styles applied correctly (Success)")
        else:
            feedback.append(f"Heading styles mismatch: {headings_correct}/3 correct (Fail)")

        # --- Criterion F: VLM Check (10 pts) ---
        # Verify visual result
        vlm_result = vlm_verify_screenshot(env_info, traj, 
            "Analyze this resume document. 1. Is the header empty? 2. Is the name 'JORDAN LEE' at the top of the body? 3. Are the dates on the right side aligned cleanly?"
        )
        if vlm_result.get("passed", False): # Simplified check assuming wrapper handles boolean logic
             score += 10
             feedback.append("Visual verification passed")
        
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    passed = score >= 70
    
    return {
        "passed": passed,
        "score": score,
        "feedback": "; ".join(feedback)
    }