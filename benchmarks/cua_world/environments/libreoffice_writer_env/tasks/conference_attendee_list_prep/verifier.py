#!/usr/bin/env python3
"""
Verifier for conference_attendee_list_prep task.

Verifies:
1. Output file exists and was created during the task.
2. Page orientation is Landscape.
3. Content is converted to a Table.
4. Table has correct columns (4).
5. Header row is formatted (Bold + Shading).
6. Data is sorted by Company then Last Name.
7. VLM visual verification of table layout.
"""

import json
import os
import sys
import tempfile
import logging
from typing import Dict, Any, List, Tuple

# Import VLM utilities
from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try importing python-docx (should be available in environment/verifier container)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not found. Verification will fail.")

def check_sorting(rows: List[List[str]]) -> Tuple[float, float, List[str]]:
    """
    Check if rows are sorted by Company (index 2) then Last Name (index 0).
    Returns (company_sort_score, name_sort_score, feedback).
    """
    if len(rows) < 2:
        return 0, 0, ["Table has insufficient data for sorting check"]
    
    # Extract header and data
    # Assuming row 0 is header, but we'll check sorting on the whole body
    header = rows[0]
    data = rows[1:]
    
    # 1. Check Primary Sort: Company (Column 2)
    companies = [r[2].strip().lower() for r in data if len(r) > 2]
    sorted_companies = sorted(companies)
    
    # Calculate how many are in order relative to the full sorted list
    # A simple way is to count inversions or just check if it matches sorted
    is_company_sorted = (companies == sorted_companies)
    
    # 2. Check Secondary Sort: Name (Column 0) within same company
    # Group by company
    company_groups = {}
    for r in data:
        if len(r) > 2:
            comp = r[2].strip().lower()
            name = r[0].strip().lower()
            if comp not in company_groups:
                company_groups[comp] = []
            company_groups[comp].append(name)
            
    total_groups = 0
    sorted_groups = 0
    
    for comp, names in company_groups.items():
        if len(names) > 1:
            total_groups += 1
            if names == sorted(names):
                sorted_groups += 1
                
    feedback = []
    company_score = 1.0 if is_company_sorted else 0.0
    if not is_company_sorted:
        feedback.append("Data is NOT sorted by Company (Column 3)")
    else:
        feedback.append("Data IS sorted by Company")
        
    name_score = 1.0
    if total_groups > 0:
        name_score = sorted_groups / total_groups
        if name_score < 1.0:
            feedback.append(f"Secondary sort by Last Name failed in {total_groups - sorted_groups}/{total_groups} groups")
        else:
            feedback.append("Secondary sort by Last Name correct")
            
    return company_score, name_score, feedback

def verify_conference_attendee_list_prep(traj, env_info, task_info):
    """
    Verify conference attendee list preparation.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "Verifier missing python-docx library"}

    # Load result metadata
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_meta = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Basic checks
    if not result_meta.get('output_exists'):
        return {"passed": False, "score": 0, "feedback": "Output file check_in_sheet.docx not found"}
    
    if not result_meta.get('file_created_during_task'):
        return {"passed": False, "score": 0, "feedback": "Output file was not modified during the task (timestamps match start)"}

    # Copy output DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/check_in_sheet.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output DOCX: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 0
    feedback_parts = []
    
    # 1. Check Landscape Orientation (15 pts)
    # Page width should be > Page height
    # Usually in section properties
    try:
        section = doc.sections[0]
        # In python-docx, dimensions are Emu or integers. 
        # Landscape: width > height
        if section.page_width > section.page_height:
            score += 15
            feedback_parts.append("Orientation: Landscape (Pass)")
        else:
            feedback_parts.append("Orientation: Portrait (Fail)")
    except Exception:
        feedback_parts.append("Orientation: Could not determine")

    # 2. Check Table Conversion (20 pts)
    if len(doc.tables) > 0:
        score += 20
        feedback_parts.append(f"Table found: Yes ({len(doc.tables)})")
        table = doc.tables[0]
        
        # 3. Check Column Count (10 pts)
        # Check first row
        if len(table.rows) > 0 and len(table.cells) > 0: # table.cells is flattened, use rows[0].cells
            col_count = len(table.rows[0].cells)
            if col_count == 4:
                score += 10
                feedback_parts.append("Columns: 4 (Pass)")
            else:
                feedback_parts.append(f"Columns: {col_count} (Expected 4)")
        
        # Extract data for sorting check
        table_data = []
        for row in table.rows:
            # Get text from each cell
            row_text = [cell.text.strip() for cell in row.cells]
            table_data.append(row_text)
            
        # 4. Sorting Check (25 pts total)
        # 15 for primary, 10 for secondary
        comp_score, name_score, sort_feedback = check_sorting(table_data)
        score += int(comp_score * 15)
        score += int(name_score * 10)
        feedback_parts.extend(sort_feedback)
        
        # 5. Header Formatting (10 pts)
        # Checking Bold on first row
        is_bold = False
        if len(table.rows) > 0:
            # Check first cell's first paragraph
            try:
                # Check run level
                for paragraph in table.rows[0].cells[0].paragraphs:
                    for run in paragraph.runs:
                        if run.bold:
                            is_bold = True
                # Check style level (e.g. "Strong" or "Heading")
                style_name = table.rows[0].cells[0].paragraphs[0].style.name.lower()
                if "bold" in style_name or "heading" in style_name:
                    is_bold = True
            except:
                pass
        
        if is_bold:
            score += 10
            feedback_parts.append("Header Bold: Yes")
        else:
            feedback_parts.append("Header Bold: No")
            
    else:
        feedback_parts.append("Table found: No (Text to Table failed)")

    # 6. VLM Verification (20 pts)
    # Check for shading (hard to check in python-docx) and general layout
    frames = sample_trajectory_frames(traj, n=4)
    final_screenshot = get_final_screenshot(traj)
    
    vlm_prompt = """
    Analyze these screenshots of LibreOffice Writer.
    The user was supposed to:
    1. Create a table with 4 columns.
    2. Set page orientation to Landscape.
    3. Make the header row Bold with a Light Gray background.
    
    Looking at the final state:
    - Is there a table visible?
    - Does the page look wider than it is tall (Landscape)?
    - Does the top row have a gray/shaded background?
    
    Return JSON:
    {
        "table_visible": true/false,
        "landscape_orientation": true/false,
        "header_shaded": true/false
    }
    """
    
    vlm_result = query_vlm(
        images=frames + [final_screenshot],
        prompt=vlm_prompt
    )
    
    if vlm_result.get("success"):
        parsed = vlm_result.get("parsed", {})
        if parsed.get("table_visible"):
            score += 5
        if parsed.get("landscape_orientation"):
            score += 5
        if parsed.get("header_shaded"):
            score += 10
            feedback_parts.append("Header Shading: Visible")
        else:
            feedback_parts.append("Header Shading: Not visible")
    else:
        feedback_parts.append("VLM verification failed")

    passed = (score >= 75)
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }