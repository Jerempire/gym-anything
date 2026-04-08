#!/usr/bin/env python3
"""
Verifier for clinical_table_formatting task.
Checks for:
1. File existence and creation time.
2. Table structure: Merged headers (gridSpan).
3. formatting: Decimal tabs in data columns.
4. Visual style: Absence of vertical borders (via XML or VLM).
"""

import json
import os
import sys
import logging
import tempfile
import shutil
from pathlib import Path

# Add utils directory to path to import writer_verification_utils
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
try:
    from writer_verification_utils import copy_and_parse_document, vlm_verify_screenshot
except ImportError:
    # Fallback if running standalone
    pass

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import python-docx
try:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def verify_clinical_table_formatting(traj, env_info, task_info):
    """
    Verify the clinical table formatting task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # 1. Check Output File Existence & Timestamp
    # ----------------------------------------------------------------
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_stats = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result stats: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not result_stats.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file not found."}
    
    if not result_stats.get("file_created_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file was not modified/created during the task."}

    # 2. Parse the DOCX
    # ----------------------------------------------------------------
    container_path = task_info['metadata']['output_file']
    success, doc, error, temp_dir = copy_and_parse_document(
        container_path, copy_from_env, file_format='docx'
    )

    if not success:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX: {error}"}

    score = 20 # Points for valid file creation
    feedback = ["File created successfully."]
    passed = False

    try:
        if len(doc.tables) == 0:
            return {"passed": False, "score": score, "feedback": "No tables found in document."}

        table = doc.tables[0]
        
        # 3. Check Header Merging (GridSpan)
        # ----------------------------------------------------------------
        # Look for cells in the first 2 rows that span columns.
        # Ideally, we find "Placebo" and "Sarilumab" spanning 2 columns each.
        merged_found = 0
        merge_details = []
        
        # We verify by checking the XML of the cells in the first row(s)
        # gridSpan is stored in tcPr (table cell properties)
        header_rows_to_check = table.rows[:2]
        for row in header_rows_to_check:
            for cell in row.cells:
                tc = cell._tc
                if tc.grid_span > 1:
                    text = cell.text.strip()
                    if "Placebo" in text or "Sarilumab" in text:
                        merged_found += 1
                        merge_details.append(f"Merged header found: '{text}' (span={tc.grid_span})")

        if merged_found >= 2:
            score += 25
            feedback.append("Header cells merged correctly.")
        elif merged_found == 1:
            score += 10
            feedback.append(f"Partial header merging found. {merge_details}")
        else:
            feedback.append("No merged header cells found for 'Placebo' or 'Sarilumab'.")

        # 4. Check Decimal Tabs
        # ----------------------------------------------------------------
        # Check data rows (skip headers). We look for tab stops with 'DECIMAL' alignment.
        decimal_tabs_found = 0
        rows_checked = 0
        
        # Assume data starts at row 2 (if 2 header rows) or row 1
        start_row = 2 if len(table.rows) > 5 else 1
        
        for row in table.rows[start_row:]:
            rows_checked += 1
            row_has_decimal_tab = False
            # Check cells 1-5 (numerical columns)
            for cell in row.cells[1:]: 
                for paragraph in cell.paragraphs:
                    for tab in paragraph.paragraph_format.tab_stops:
                        if tab.alignment == WD_TAB_ALIGNMENT.DECIMAL:
                            row_has_decimal_tab = True
                            break
                    if row_has_decimal_tab: break
                if row_has_decimal_tab: break
            
            if row_has_decimal_tab:
                decimal_tabs_found += 1

        if decimal_tabs_found >= (rows_checked * 0.8): # Allow some margin
            score += 25
            feedback.append("Decimal tabs applied to data rows.")
        elif decimal_tabs_found > 0:
            score += 10
            feedback.append(f"Some decimal tabs found ({decimal_tabs_found}/{rows_checked} rows).")
        else:
            feedback.append("No decimal tabs detected in data rows.")

        # 5. Check Borders (XML inspection)
        # ----------------------------------------------------------------
        # We look for <w:insideV w:val="nil"> or similar in the table properties
        tbl_xml = table._element.xml
        vertical_borders_removed = False
        if 'w:insideV w:val="nil"' in tbl_xml or 'w:insideV w:val="none"' in tbl_xml:
            vertical_borders_removed = True
        
        # Top/Bottom borders (top/bottom w:val="single" and w:sz > 12 usually)
        # This is hard to parse robustly, so we'll give partial credit for vertical removal
        # and rely on VLM for the full border style check.
        if vertical_borders_removed:
            score += 15
            feedback.append("Vertical borders removed (XML verified).")
        else:
            feedback.append("Vertical borders might still be present (XML check failed).")

        # 6. VLM Verification (Visual Check)
        # ----------------------------------------------------------------
        # We ask VLM to confirm the "academic" look
        vlm_prompt = """
        Analyze the table in this screenshot. 
        1. Does the table have vertical lines separating the columns? (Academic tables should NOT have them).
        2. Are the numbers in the columns aligned on the decimal point?
        3. Is there a clear header structure with "Placebo" and "Sarilumab" grouping columns?
        Reply in JSON: {"no_vertical_lines": bool, "decimal_aligned": bool, "grouped_headers": bool}
        """
        
        vlm_result = vlm_verify_screenshot(env_info, traj, vlm_prompt)
        vlm_data = vlm_result.get("parsed", {})
        
        vlm_score = 0
        if vlm_data.get("no_vertical_lines", False):
            vlm_score += 5
            feedback.append("VLM: Confirmed no vertical lines.")
        else:
             # If XML check failed, this confirms it. If XML passed, maybe VLM is wrong, but penalty stands.
             pass

        if vlm_data.get("decimal_aligned", False):
            vlm_score += 5
            feedback.append("VLM: Numbers appear visually aligned.")

        if vlm_data.get("grouped_headers", False):
            vlm_score += 5
            feedback.append("VLM: Header grouping visible.")
            
        score += vlm_score

    except Exception as e:
        feedback.append(f"Error during verification logic: {e}")
    finally:
        # Cleanup temp dir from docx parsing
        if temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

    # Final Evaluation
    passed = score >= 70
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " ".join(feedback)
    }