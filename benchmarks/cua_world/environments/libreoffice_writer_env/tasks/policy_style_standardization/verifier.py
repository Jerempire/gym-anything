#!/usr/bin/env python3
"""
Verifier for Policy Style Standardization task.
Checks:
1. File exists and was created during task.
2. Direct formatting (Comic Sans, etc.) is removed.
3. 'Heading 1' style is redefined correctly (Arial, 16pt, Bold, Dark Blue).
4. 'Text Body' (or Normal) style is redefined correctly (Verdana, 10pt, Justified).
5. Paragraphs use the correct styles.
"""

import json
import os
import tempfile
import logging
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Helper to check approx color match (Dark Blue #00008B is R0 G0 B139)
def is_dark_blue(color_obj):
    if not color_obj or not color_obj.rgb:
        return False
    r, g, b = color_obj.rgb
    # Accept range of dark blues
    return r < 50 and g < 50 and b > 100

def verify_policy_style_standardization(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Load result metadata
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_meta = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not result_meta.get("output_exists"):
        return {"passed": False, "score": 0, "feedback": "Output file remote_work_policy_clean.docx not found."}

    # Load the document content
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/remote_work_policy_clean.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output DOCX: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 0
    feedback = []

    # --- Criteria 1: File Existence & Creation (10 pts) ---
    if result_meta.get("created_during_task"):
        score += 10
        feedback.append("File saved correctly.")
    else:
        feedback.append("File exists but timestamp suggests it wasn't modified/created during task.")

    # --- Criteria 2: Clean Formatting (No Comic Sans/Courier) (25 pts) ---
    forbidden_fonts = ["Comic Sans MS", "Courier New", "Liberation Mono"]
    found_forbidden = []
    
    # Check styles for forbidden fonts (unlikely but possible)
    for style in doc.styles:
        if style.font.name in forbidden_fonts:
            found_forbidden.append(f"Style '{style.name}' uses {style.font.name}")

    # Check direct formatting in paragraphs
    direct_formatting_dirty = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name in forbidden_fonts:
                found_forbidden.append(f"Text '{run.text[:10]}...' uses {run.font.name}")
                direct_formatting_dirty = True
    
    if not found_forbidden:
        score += 25
        feedback.append("Direct formatting cleaned (no forbidden fonts found).")
    else:
        feedback.append(f"Found forbidden fonts (Direct formatting not cleared?): {', '.join(found_forbidden[:3])}")

    # --- Criteria 3: Heading 1 Style Definition (20 pts) ---
    # User must have MODIFIED the style, not just applied direct formatting to text
    h1_style = doc.styles['Heading 1']
    h1_score = 0
    
    # Font Face (Arial/Liberation Sans)
    if h1_style.font.name and h1_style.font.name.lower() in ['arial', 'liberation sans']:
        h1_score += 5
    else:
        feedback.append(f"Heading 1 font is '{h1_style.font.name}' (expected Arial)")

    # Size (16pt)
    if h1_style.font.size and h1_style.font.size.pt == 16.0:
        h1_score += 5
    else:
        size = h1_style.font.size.pt if h1_style.font.size else "Default"
        feedback.append(f"Heading 1 size is {size} (expected 16pt)")

    # Color (Dark Blue)
    if is_dark_blue(h1_style.font.color):
        h1_score += 5
    else:
        feedback.append("Heading 1 color is not Dark Blue")
    
    # Bold
    if h1_style.font.bold:
        h1_score += 5
    else:
        feedback.append("Heading 1 is not bold")

    score += h1_score

    # --- Criteria 4: Body Style Definition (20 pts) ---
    # Could be 'Normal' or 'Text Body' depending on what agent used
    body_style_names = ['Text Body', 'Normal']
    target_body_style = None
    
    # Heuristic: find which style is actually used by the body paragraphs
    used_styles = {}
    for para in doc.paragraphs:
        if len(para.text) > 30: # Likely body text
            sname = para.style.name
            used_styles[sname] = used_styles.get(sname, 0) + 1
    
    # Pick the most used style among candidates
    best_candidate = max(used_styles.items(), key=lambda x: x[1])[0] if used_styles else 'Normal'
    
    if best_candidate not in body_style_names and 'Body' not in best_candidate:
        feedback.append(f"Body text seems to use unexpected style: '{best_candidate}'")
    
    body_style = doc.styles[best_candidate]
    body_score = 0

    # Font (Verdana)
    if body_style.font.name == 'Verdana':
        body_score += 5
    else:
        feedback.append(f"Body style ({best_candidate}) font is '{body_style.font.name}' (expected Verdana)")

    # Size (10pt)
    if body_style.font.size and body_style.font.size.pt == 10.0:
        body_score += 5
    else:
        size = body_style.font.size.pt if body_style.font.size else "Default"
        feedback.append(f"Body style size is {size} (expected 10pt)")

    # Justified
    if body_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        body_score += 5
    else:
        feedback.append(f"Body style alignment is not Justified")

    # Spacing (1.15 lines is approx 276 line rule in docx XML, or 1.15 multiplier)
    # python-docx stores line_spacing as float for multipliers
    if body_style.paragraph_format.line_spacing and 1.1 <= body_style.paragraph_format.line_spacing <= 1.2:
        body_score += 5
    else:
        spacing = body_style.paragraph_format.line_spacing
        feedback.append(f"Body style spacing is {spacing} (expected 1.15)")
    
    score += body_score

    # --- Criteria 5: Style Application (15 pts) ---
    # Check if headers actually use Heading 1
    headers_correct = 0
    header_titles = ["Objective", "Scope", "Equipment", "Security"]
    for title in header_titles:
        found = False
        for para in doc.paragraphs:
            if title in para.text:
                if para.style.name == 'Heading 1':
                    headers_correct += 1
                    found = True
                    break
        if not found:
            feedback.append(f"Section '{title}' not found or not styled as Heading 1")
    
    if headers_correct >= 3:
        score += 15
    elif headers_correct > 0:
        score += 5
    
    # --- Criteria 6: Line Spacing (10 pts) ---
    # (Redundant with style check, but confirms final result appearance)
    if body_score >= 15: # If style was mostly right
        score += 10
    else:
        # Fallback: check direct formatting on paragraphs if style failed
        # This gives partial credit if they manually formatted instead of modifying style (bad practice but visible result)
        manual_checks_passed = 0
        for para in doc.paragraphs:
            if len(para.text) > 30:
                if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    manual_checks_passed += 1
        if manual_checks_passed > 2:
            score += 5
            feedback.append("Partial credit: Text is justified via direct formatting (should be Style).")

    return {
        "passed": score >= 75,
        "score": score,
        "feedback": " | ".join(feedback)
    }