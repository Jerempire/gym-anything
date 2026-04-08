#!/usr/bin/env python3
"""
Verifier for price_list_conversion_sort task.

Checks:
1. Output file exists and is a valid DOCX.
2. Content is organized into a Table.
3. Header row exists and is formatted (Bold/Heading).
4. Data is sorted by Price in Descending order.
"""

import json
import os
import re
import tempfile
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import docx (installed in env)
try:
    from docx import Document
except ImportError:
    Document = None

def parse_currency(value_str):
    """Parse a currency string like '$1,250.00' into a float."""
    try:
        # Remove '$', ',', and whitespace
        clean = re.sub(r'[$,\s]', '', value_str)
        return float(clean)
    except ValueError:
        return 0.0

def verify_price_list_sort(traj, env_info, task_info):
    """
    Verify the price list conversion and sorting task.
    """
    # 0. Setup
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "System error: copy_from_env missing"}
    
    if Document is None:
        return {"passed": False, "score": 0, "feedback": "System error: python-docx not installed"}

    metadata = task_info.get('metadata', {})
    output_path = metadata.get('output_path', '/home/ga/Documents/formatted_price_list.docx')
    expected_headers = metadata.get('expected_headers', ["Product Name", "Category", "SKU", "Price"])

    # 1. Get result JSON from export script
    task_result = {}
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            task_result = json.load(f)
    except Exception as e:
        logger.warning(f"Could not load task_result.json: {e}")
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # 2. Check basic file existence
    if not task_result.get('output_exists', False):
        return {
            "passed": False, 
            "score": 0, 
            "feedback": "Output file 'formatted_price_list.docx' was not created."
        }
    
    if not task_result.get('file_created_during_task', False):
        return {
            "passed": False, 
            "score": 0, 
            "feedback": "Output file exists but was not modified during the task session."
        }

    # 3. Retrieve and Parse the DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(output_path, temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {
            "passed": False,
            "score": 10,
            "feedback": f"Created file is not a valid DOCX document: {e}"
        }
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # 4. Verification Logic
    score = 10 # Base score for creating valid file
    feedback = ["File created."]
    
    # Check A: Is there a table?
    if len(doc.tables) == 0:
        return {
            "passed": False,
            "score": score,
            "feedback": "The text was not converted into a table. No tables found in document."
        }
    
    table = doc.tables[0]
    score += 20
    feedback.append("Table detected.")

    # Check B: Verify Headers (Row 0)
    # Allow some fuzzy matching, but column 3 (Price) is critical
    rows = table.rows
    if len(rows) < 2:
        return {"passed": False, "score": score, "feedback": "Table is empty or missing data rows."}

    header_cells = [cell.text.strip() for cell in rows[0].cells]
    
    # Verify header content
    headers_match = True
    for expected in expected_headers:
        if not any(expected.lower() in h.lower() for h in header_cells):
            headers_match = False
            break
            
    if headers_match:
        score += 10
        feedback.append("Header row content matches.")
    else:
        feedback.append(f"Header row mismatch. Found: {header_cells}")

    # Check C: Verify Header Styling (Bold)
    # Check the runs in the first cell of the header
    try:
        first_cell_runs = rows[0].cells[0].paragraphs[0].runs
        is_bold = any(run.bold for run in first_cell_runs)
        # Also check style name (e.g., "Table Heading")
        style_name = rows[0].cells[0].paragraphs[0].style.name
        is_heading_style = "Heading" in style_name or "Header" in style_name
        
        if is_bold or is_heading_style:
            score += 10
            feedback.append("Header formatting applied (Bold/Style).")
        else:
            feedback.append("Header row does not appear to be bold or styled.")
    except Exception:
        pass # Skip if structure is weird

    # Check D: Verify Data Sorting
    # Find the Price column index
    price_col_idx = -1
    for i, h in enumerate(header_cells):
        if "Price" in h:
            price_col_idx = i
            break
            
    if price_col_idx == -1:
        feedback.append("Could not locate 'Price' column in header.")
        # Try to guess - usually last column based on our data
        price_col_idx = 3

    # Extract prices
    prices = []
    raw_values = []
    
    # Iterate from row 1 (skip header)
    for row in rows[1:]:
        if len(row.cells) > price_col_idx:
            val_text = row.cells[price_col_idx].text.strip()
            raw_values.append(val_text)
            p_val = parse_currency(val_text)
            if p_val > 0: # filter out empty rows or parse failures
                prices.append(p_val)
    
    if len(prices) < 5:
        return {
            "passed": False, 
            "score": score, 
            "feedback": " | ".join(feedback) + " | Too few data rows found to verify sorting."
        }

    score += 10 # Data extraction successful
    
    # Verify Descending Order
    # We check if sorted(prices, reverse=True) matches prices
    # We allow slight deviations? No, sort should be exact.
    
    sorted_desc = sorted(prices, reverse=True)
    sorted_asc = sorted(prices)
    
    is_sorted_desc = (prices == sorted_desc)
    is_sorted_asc = (prices == sorted_asc)
    
    if is_sorted_desc:
        score += 40
        feedback.append("Data is correctly sorted by Price (Descending).")
    elif is_sorted_asc:
        feedback.append("Data is sorted Ascending (Lowest to Highest). Task required Descending.")
        score += 10 # Partial credit for finding the sort button
    else:
        # Check if it matches original order (roughly)
        # Original: 349, 899, 1250, 189... -> mixed
        feedback.append("Data appears unsorted.")
        # Debug info
        feedback.append(f"Top 3 prices found: {prices[:3]}")

    passed = (score >= 70)
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }