#!/usr/bin/env python3
"""
Verifier for book_chapter_page_styles task.

Checks for:
1. Output file existence and validity (DOCX).
2. Mirror margins (Inside 1.25", Outside 1.0").
3. Different First Page enabled (Empty header).
4. Even/Odd Headers enabled (Distinct content).
5. Formatting: Heading 1, 12pt Serif, Double Spacing.
"""

import json
import os
import tempfile
import logging
from typing import Dict, Any

# Import shared utils (assuming they are available in the env as per instruction)
# If running locally for testing, mocks might be needed.
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
except ImportError:
    pass # Handled in setup checks usually

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_book_chapter_page_styles(traj, env_info, task_info):
    """
    Verify Book Chapter Page Styles task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    
    # 1. Retrieve Result JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_data = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result JSON: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not result_data.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file not found at /home/ga/Documents/chapter1_formatted.docx"}

    if not result_data.get("file_created_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file was not modified during the task session (Anti-gaming)"}

    # 2. Retrieve Output DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/chapter1_formatted.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output DOCX: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # --- Scoring ---
    score = 0
    feedback = []
    
    # Check 1: Mirror Margins (10 pts)
    # 1.25 inch = 1143000 EMU, 1.0 inch = 914400 EMU
    # Tolerance: ~0.05 inch (45720 EMU)
    TOLERANCE = 50000
    section = doc.sections[0]
    
    # Note: In python-docx, if mirror margins are on, left_margin often maps to inside
    # and right_margin to outside, or check xml for <w:mirrorMargins>.
    # We check if *either* mirror margins is explicitly on OR the values match the specialized request.
    
    has_mirror_xml = "mirrorMargins" in section._sectPr.xml
    
    left_emu = section.left_margin if section.left_margin else 0
    right_emu = section.right_margin if section.right_margin else 0
    top_emu = section.top_margin if section.top_margin else 0
    bottom_emu = section.bottom_margin if section.bottom_margin else 0

    target_inside = 1143000
    target_outside = 914400
    target_vertical = 914400

    # We accept if Left=Inside and Right=Outside (standard mapping)
    margins_correct = (
        abs(left_emu - target_inside) < TOLERANCE and 
        abs(right_emu - target_outside) < TOLERANCE and
        abs(top_emu - target_vertical) < TOLERANCE and
        abs(bottom_emu - target_vertical) < TOLERANCE
    )

    if has_mirror_xml or margins_correct:
        score += 10
        feedback.append("Mirror margins set correctly")
    else:
        feedback.append(f"Margins incorrect (Left: {left_emu}, Right: {right_emu}, Expected Inside: {target_inside})")

    # Check 2: Different First Page (15 pts)
    if section.different_first_page_header_footer:
        score += 15
        feedback.append("Different First Page enabled")
        
        # Verify First Page Header Empty
        fp_header = section.first_page_header
        fp_text = "".join([p.text for p in fp_header.paragraphs]).strip()
        if not fp_text:
            score += 5
            feedback.append("First page header is empty")
        else:
            feedback.append(f"First page header should be empty, found: '{fp_text}'")
            
        # Verify First Page Footer has content (page number)
        # Difficult to check specifically for a field, but check for non-empty
        fp_footer = section.first_page_footer
        if any(p.text.strip() or p.runs for p in fp_footer.paragraphs):
            score += 5
            feedback.append("First page footer has content")
        else:
            feedback.append("First page footer appears empty (expected page number)")
    else:
        feedback.append("Different First Page NOT enabled")

    # Check 3: Even/Odd Headers (20 pts)
    # python-docx: doc.settings.odd_and_even_pages_header_footer
    if doc.settings.odd_and_even_pages_header_footer:
        score += 20
        feedback.append("Even/Odd Headers enabled")
        
        # Check Even Header (Pride and Prejudice, Left)
        even_header = section.even_page_header
        even_text = " ".join([p.text for p in even_header.paragraphs]).strip()
        if "Pride and Prejudice" in even_text:
            score += 5
            feedback.append("Even header text correct")
            # Check alignment (Left is usually default 0 or None)
            align = even_header.paragraphs[0].alignment
            if align in [WD_ALIGN_PARAGRAPH.LEFT, None]:
                score += 5
                feedback.append("Even header alignment correct")
        else:
            feedback.append(f"Even header text missing/wrong: '{even_text}'")

        # Check Odd (Default) Header (Chapter I, Right)
        odd_header = section.header
        odd_text = " ".join([p.text for p in odd_header.paragraphs]).strip()
        if "Chapter" in odd_text:
            score += 5
            feedback.append("Odd header text correct")
            # Check alignment (Right is 2)
            if odd_header.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                score += 5
                feedback.append("Odd header alignment correct")
            else:
                feedback.append("Odd header alignment incorrect (expected Right)")
        else:
            feedback.append(f"Odd header text missing/wrong: '{odd_text}'")
    else:
        feedback.append("Even/Odd Headers NOT enabled")

    # Check 4: Chapter Title Style (Heading 1 + Center) (10 pts)
    # Find the paragraph with "Chapter I"
    chapter_para = None
    for p in doc.paragraphs:
        if "Chapter I" in p.text:
            chapter_para = p
            break
    
    if chapter_para:
        style_name = chapter_para.style.name if chapter_para.style else "Normal"
        if "Heading 1" in style_name:
            score += 5
            feedback.append("Chapter title uses Heading 1")
        else:
            feedback.append(f"Chapter title uses '{style_name}' (Expected Heading 1)")
            
        if chapter_para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            score += 5
            feedback.append("Chapter title centered")
        else:
            feedback.append("Chapter title not centered")
    else:
        feedback.append("Chapter title 'Chapter I' not found")

    # Check 5: Body Formatting (12pt Serif, Double Spacing) (15 pts)
    # Check a sample paragraph (e.g., the 5th one)
    sample_para = None
    if len(doc.paragraphs) > 5:
        sample_para = doc.paragraphs[5]
    
    if sample_para:
        # Check Spacing
        # Double spacing is 2.0
        line_spacing = sample_para.paragraph_format.line_spacing
        if line_spacing == 2.0:
            score += 5
            feedback.append("Double spacing applied")
        else:
            feedback.append(f"Line spacing incorrect: {line_spacing}")

        # Check Font
        # We need to check runs or style
        # Assuming direct formatting or style update
        font_name = None
        font_size = None
        
        # Check runs first
        for run in sample_para.runs:
            if run.font.name: font_name = run.font.name
            if run.font.size: font_size = run.font.size.pt
        
        # Fallback to style if not on runs
        if not font_name and sample_para.style.font.name:
            font_name = sample_para.style.font.name
        if not font_size and sample_para.style.font.size:
            font_size = sample_para.style.font.size.pt

        # Serif check (heuristic)
        serifs = ["Times", "Liberation Serif", "Georgia", "Palatino", "Garamond"]
        if font_name and any(s.lower() in font_name.lower() for s in serifs):
            score += 5
            feedback.append(f"Serif font used ({font_name})")
        else:
            feedback.append(f"Font incorrect or not serif: {font_name}")
            
        if font_size and font_size == 12.0:
            score += 5
            feedback.append("Font size 12pt")
        else:
            feedback.append(f"Font size incorrect: {font_size}")

    return {
        "passed": score >= 60,
        "score": score,
        "feedback": "; ".join(feedback)
    }