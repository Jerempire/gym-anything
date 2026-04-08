#!/usr/bin/env python3
"""
Verifier for exam_paper_formatting task.
Checks:
1. Output file existence and valid DOCX format
2. Title block formatting (Centering, Bold, Sizes)
3. Heading styles (Heading 2)
4. Body text formatting (Font, Size, Line Spacing)
5. Page margins (1 inch)
6. Footer presence
7. End marker formatting
"""

import json
import os
import logging
import tempfile
import shutil

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import python-docx
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def verify_exam_paper_formatting(traj, env_info, task_info):
    """Verify the formatting of the biology exam paper."""
    
    # 0. Infrastructure Check
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "System error: copy_from_env missing"}
        
    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "System error: python-docx not installed in verifier"}

    # 1. Load result JSON from container
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix=".json")
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name) as f:
            result_stats = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load task result: {e}"}
    finally:
        os.unlink(temp_json.name)

    # 2. Check basic file status
    if not result_stats.get("output_exists"):
        return {"passed": False, "score": 0, "feedback": "Output file 'bio101_final_exam.docx' not found."}
    
    if not result_stats.get("file_created_during_task"):
        return {"passed": False, "score": 0, "feedback": "Output file timestamp is too old (pre-dates task start)."}

    if result_stats.get("raw_file_modified"):
        # Penalty for modifying the source file instead of saving a copy
        # We'll deduct points later or fail immediately depending on strictness.
        # Let's deduct 10 points.
        source_modified_penalty = 10
    else:
        source_modified_penalty = 0

    # 3. Retrieve and Parse the DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    try:
        copy_from_env("/home/ga/Documents/bio101_final_exam.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to open DOCX file: {e} (File may be corrupted)"}
    
    score = 0
    feedback = []
    
    # --- CRITERION 1: Title Block Formatting (20 pts) ---
    # Expect lines 0 and 1 to be Centered and Bold.
    # Line 0: Westfield... (16pt), Line 1: Biology... (14pt)
    title_score = 0
    try:
        p0 = doc.paragraphs[0]
        p1 = doc.paragraphs[1]
        
        # Check alignment
        if p0.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            title_score += 5
        
        # Check Bold & Size p0
        # Runs can be fragmented, so we check if ANY run has the property or if style has it
        p0_bold = any(r.bold for r in p0.runs) or (p0.style.font.bold if p0.style else False)
        p0_size = any((r.font.size and r.font.size.pt == 16.0) for r in p0.runs)
        
        if p0_bold: title_score += 3
        if p0_size: title_score += 2
        
        # Check Bold & Size p1
        p1_bold = any(r.bold for r in p1.runs)
        p1_size = any((r.font.size and r.font.size.pt == 14.0) for r in p1.runs)
        
        if p1_bold: title_score += 5
        if p1_size: title_score += 5
        
        if title_score >= 15:
            feedback.append("Title block formatted correctly.")
        else:
            feedback.append(f"Title block issues (Score: {title_score}/20). Check centering, bolding, and font sizes.")
            
    except IndexError:
        feedback.append("Document too short to check title block.")
    
    score += title_score

    # --- CRITERION 2: Instructions Italicized (10 pts) ---
    # Find paragraph starting with "INSTRUCTIONS:"
    instr_score = 0
    instr_found = False
    for p in doc.paragraphs:
        if p.text.strip().startswith("INSTRUCTIONS:"):
            instr_found = True
            # Check if all text runs are italic
            # Note: docx reports None for italic if it inherits. We look for explicit True or style.
            is_italic = True
            for r in p.runs:
                if not r.italic and not (p.style.font.italic):
                    # Strict check: run must be italic
                    is_italic = False
                    break
            if is_italic:
                instr_score = 10
            break
    
    if instr_found:
        score += instr_score
        if instr_score == 10:
            feedback.append("Instructions italicized.")
        else:
            feedback.append("Instructions paragraph found but not fully italicized.")
    else:
        feedback.append("Instructions paragraph not found.")

    # --- CRITERION 3: Section Headings (Heading 2) (15 pts) ---
    # Look for "Part A", "Part B", "Part C"
    heading_score = 0
    headings_found = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Part A") or text.startswith("Part B") or text.startswith("Part C"):
            if p.style and "Heading 2" in p.style.name:
                headings_found += 1
    
    if headings_found >= 3:
        heading_score = 15
        feedback.append("All section headings styles correct.")
    elif headings_found > 0:
        heading_score = 5 * headings_found
        feedback.append(f"Some section headings missing style ({headings_found}/3).")
    else:
        feedback.append("Section headings do not use 'Heading 2' style.")
    
    score += heading_score

    # --- CRITERION 4: Body Text Formatting (Font, Size, Spacing) (25 pts) ---
    # Check paragraphs that are NOT headings/title
    body_font_score = 0
    body_size_score = 0
    body_spacing_score = 0
    body_paras_checked = 0
    
    # Allowable fonts (loose matching)
    valid_fonts = ["liberation serif", "times new roman", "times"]
    
    for p in doc.paragraphs:
        # Skip empty or headings
        if not p.text.strip() or (p.style and "Heading" in p.style.name) or p in doc.paragraphs[:4]:
            continue
            
        body_paras_checked += 1
        
        # Check Font Name (in runs)
        # Note: if font is set in style 'Normal', run.font.name might be None
        # We'll check runs first.
        font_match = False
        for r in p.runs:
            if r.font.name and r.font.name.lower() in valid_fonts:
                font_match = True
                break
        if not font_match and p.style and p.style.font.name:
             if p.style.font.name.lower() in valid_fonts:
                 font_match = True
        if font_match: body_font_score += 1

        # Check Font Size (12pt)
        size_match = False
        for r in p.runs:
            if r.font.size and r.font.size.pt == 12.0:
                size_match = True
                break
        if size_match: body_size_score += 1
        
        # Check Line Spacing (1.5)
        # 1.5 spacing is usually stored as float 1.5 OR rule MULTIPLE
        pf = p.paragraph_format
        # 1.5 lines can be represented in multiple ways in python-docx
        # often line_spacing = 1.5
        if pf.line_spacing == 1.5:
            body_spacing_score += 1

    # Normalize scores based on number of paragraphs checked
    if body_paras_checked > 0:
        # Cap at max points
        s_font = min(10, int((body_font_score / body_paras_checked) * 10))
        s_size = min(5, int((body_size_score / body_paras_checked) * 5))
        s_space = min(10, int((body_spacing_score / body_paras_checked) * 10))
        
        score += s_font + s_size + s_space
        feedback.append(f"Body Formatting: Font({s_font}/10), Size({s_size}/5), Spacing({s_space}/10).")
    else:
        feedback.append("No body paragraphs found to check.")

    # --- CRITERION 5: Page Margins (10 pts) ---
    # 1 inch = 914400 EMU
    margin_score = 0
    try:
        sect = doc.sections[0]
        # Tolerance ~0.1 inch (91440 EMU)
        target = 914400
        tol = 91440
        
        m_left = abs((sect.left_margin or 0) - target) < tol
        m_right = abs((sect.right_margin or 0) - target) < tol
        m_top = abs((sect.top_margin or 0) - target) < tol
        m_bottom = abs((sect.bottom_margin or 0) - target) < tol
        
        if m_left and m_right and m_top and m_bottom:
            margin_score = 10
            feedback.append("Margins correct (1 inch).")
        else:
            feedback.append(f"Margins incorrect. Found: L={sect.left_margin}, R={sect.right_margin}...")
            # Partial credit
            if m_left: margin_score += 2
            if m_right: margin_score += 2
            if m_top: margin_score += 2
            if m_bottom: margin_score += 2
            
    except Exception as e:
        feedback.append(f"Could not check margins: {e}")
        
    score += margin_score

    # --- CRITERION 6: Footer (10 pts) ---
    # Check if footer exists and contains "Biology 101"
    footer_score = 0
    try:
        # Check all headers/footers in first section
        sect = doc.sections[0]
        footer = sect.footer
        if footer:
            footer_text = " ".join([p.text for p in footer.paragraphs])
            if "Biology 101" in footer_text:
                footer_score = 10
                feedback.append("Footer content verified.")
            else:
                feedback.append(f"Footer exists but text missing 'Biology 101'. Found: '{footer_text}'")
        else:
            feedback.append("No footer found.")
    except:
        feedback.append("Error checking footer.")
        
    score += footer_score

    # --- CRITERION 7: End Marker (10 pts) ---
    # Last paragraph aligned center and bold
    end_score = 0
    try:
        last_p = doc.paragraphs[-1]
        if "END OF EXAMINATION" in last_p.text:
            if last_p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                end_score += 5
            if any(r.bold for r in last_p.runs):
                end_score += 5
            
            if end_score == 10:
                feedback.append("End marker correct.")
            else:
                feedback.append("End marker found but formatting incorrect.")
        else:
            # Try second to last in case of trailing newline
            last_p = doc.paragraphs[-2]
            if "END OF EXAMINATION" in last_p.text:
                 if last_p.alignment == WD_ALIGN_PARAGRAPH.CENTER: end_score += 5
                 if any(r.bold for r in last_p.runs): end_score += 5
            else:
                feedback.append("End marker text not found at end of document.")
    except:
        pass
    
    score += end_score
    score -= source_modified_penalty

    # Cleanup
    try:
        os.unlink(temp_docx.name)
    except:
        pass

    # Final Result
    passed = score >= 60
    return {
        "passed": passed,
        "score": max(0, score), # Ensure no negative score
        "feedback": " | ".join(feedback)
    }