#!/usr/bin/env python3
"""
Verifier for Incident After-Action Report Table Formatting Task.

Verifies:
1. Heading Styles (Heading 1 and Heading 2)
2. Table creation (plaintext -> table conversion)
3. Page Margins (1.0 inch)
4. Font formatting (Liberation Sans/Arial 11pt)
5. Title Alignment (Center)
6. Page Numbers (Footer)
"""

import json
import os
import sys
import tempfile
import logging
import shutil

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Verification will be limited.")


def verify_incident_report(traj, env_info, task_info):
    """
    Main verification function for incident_report_table_formatting.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "Verifier environment missing python-docx"}

    # Get metadata
    metadata = task_info.get('metadata', {})
    output_path = metadata.get('output_path', '/home/ga/Documents/aar_formatted.docx')
    req_h1 = metadata.get('required_headings_1', [])
    req_h2 = metadata.get('required_headings_2', [])

    # Load result JSON
    temp_result_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result_json.name)
        with open(temp_result_json.name, 'r') as f:
            task_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load task result: {e}"}
    finally:
        if os.path.exists(temp_result_json.name):
            os.unlink(temp_result_json.name)

    # Check existence
    if not task_result.get('output_exists', False):
        return {"passed": False, "score": 0, "feedback": "Output file not found"}

    # Copy output DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_docx.close()
    
    try:
        copy_from_env(output_path, temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)
        return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX file: {e}"}

    score = 0
    feedback = []
    
    # --------------------------------------------------------------------------
    # CRITERION 1: Heading Styles (20 pts)
    # --------------------------------------------------------------------------
    h1_found = 0
    h2_found = 0
    
    for para in doc.paragraphs:
        if not para.style:
            continue
        style_name = para.style.name
        text = para.text.strip()
        
        # Check Heading 1
        if any(h in text for h in req_h1) and 'Heading 1' in style_name:
            h1_found += 1
            
        # Check Heading 2
        if any(h in text for h in req_h2) and 'Heading 2' in style_name:
            h2_found += 1

    # Scoring Headings
    if h1_found >= 5: score += 10
    elif h1_found >= 3: score += 5
    
    if h2_found >= 3: score += 10
    elif h2_found >= 2: score += 5
    
    feedback.append(f"Headings: Found {h1_found}/{len(req_h1)} H1 and {h2_found}/{len(req_h2)} H2")

    # --------------------------------------------------------------------------
    # CRITERION 2: Table Creation (20 pts)
    # --------------------------------------------------------------------------
    # Expecting at least one table with >= 3 columns and >= 8 rows
    valid_table_found = False
    table_details = "No table found"
    
    if len(doc.tables) > 0:
        table = doc.tables[0] # Assuming the first table is the timeline
        rows = len(table.rows)
        cols = len(table.columns)
        
        # Check content in cells to ensure it's not empty
        has_content = False
        try:
            if len(table.rows) > 1:
                text_sample = table.cell(1, 0).text
                if text_sample.strip(): has_content = True
        except:
            pass
            
        table_details = f"Table found: {rows} rows, {cols} cols"
        
        if cols >= 3 and rows >= 8 and has_content:
            valid_table_found = True
            score += 20
        elif cols >= 3 and rows >= 5:
            score += 10 # Partial credit
            feedback.append("Table found but rows/content may be incomplete")

    feedback.append(table_details)

    # --------------------------------------------------------------------------
    # CRITERION 3: Page Margins (15 pts)
    # --------------------------------------------------------------------------
    # 1.0 inch = 914400 EMU. Tolerance ~0.05 inch (45720 EMU)
    target_emu = 914400
    tolerance = 50000
    
    try:
        section = doc.sections[0]
        margins = [
            section.left_margin, 
            section.right_margin, 
            section.top_margin, 
            section.bottom_margin
        ]
        
        # Filter None values
        valid_margins = [m for m in margins if m is not None]
        correct_margins = sum(1 for m in valid_margins if abs(m - target_emu) < tolerance)
        
        if correct_margins == 4:
            score += 15
            feedback.append("Margins: All 4 correct (1.0 inch)")
        elif correct_margins >= 2:
            score += 7
            feedback.append(f"Margins: {correct_margins}/4 correct")
        else:
            feedback.append(f"Margins: Incorrect (Found {valid_margins[0] if valid_margins else 'None'} EMU)")
            
    except Exception as e:
        feedback.append(f"Margin check failed: {e}")

    # --------------------------------------------------------------------------
    # CRITERION 4: Title Alignment (10 pts)
    # --------------------------------------------------------------------------
    # Check if first paragraph is centered
    title_centered = False
    try:
        if len(doc.paragraphs) > 0:
            p1 = doc.paragraphs[0]
            # WD_ALIGN_PARAGRAPH.CENTER is 1
            if p1.alignment == 1 or p1.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                title_centered = True
                score += 10
                feedback.append("Title: Centered")
            else:
                feedback.append(f"Title: Not centered (alignment={p1.alignment})")
    except:
        pass

    # --------------------------------------------------------------------------
    # CRITERION 5: Font & Size (15 pts)
    # --------------------------------------------------------------------------
    # Check random sample of body paragraphs
    compliant_runs = 0
    total_runs_checked = 0
    
    approved_fonts = ["liberation sans", "arial", "helvetica"]
    
    for i, para in enumerate(doc.paragraphs):
        if i > 20: break # Check first 20 paragraphs
        if para.style.name.startswith("Heading") or "Title" in para.style.name:
            continue
            
        for run in para.runs:
            if not run.text.strip(): continue
            total_runs_checked += 1
            
            font_name = (run.font.name or "").lower()
            size = run.font.size
            
            # If font name not set on run, might be in style (hard to check fully here)
            # but we assume direct formatting or style update
            name_ok = any(f in font_name for f in approved_fonts)
            
            # 11pt is 139700 EMU (approx). 1 pt = 12700 EMU
            # python-docx size.pt returns float
            size_ok = False
            if size and size.pt >= 10.5:
                size_ok = True
                
            if name_ok and size_ok:
                compliant_runs += 1

    font_score = 0
    if total_runs_checked > 0:
        ratio = compliant_runs / total_runs_checked
        if ratio > 0.5: score += 15
        elif ratio > 0.2: score += 5
        feedback.append(f"Font compliance: {int(ratio*100)}%")
    else:
        feedback.append("Font check: No text runs found to check")

    # --------------------------------------------------------------------------
    # CRITERION 6: Page Numbers / Footer (10 pts)
    # --------------------------------------------------------------------------
    footer_found = False
    try:
        for section in doc.sections:
            footer = section.footer
            if not footer: continue
            
            # Check for PAGE field or text
            xml = footer._element.xml
            text = "".join([p.text for p in footer.paragraphs])
            
            if "w:fldSimple" in xml and "PAGE" in xml:
                footer_found = True
            elif "w:instrText" in xml and "PAGE" in xml:
                footer_found = True
            elif any(char.isdigit() for char in text):
                # Fallback: if they just typed a number manually
                footer_found = True
                
            if footer_found: break
            
        if footer_found:
            score += 10
            feedback.append("Footer: Page number detected")
        else:
            feedback.append("Footer: No page number field found")
            
    except Exception as e:
        feedback.append(f"Footer check failed: {e}")

    # --------------------------------------------------------------------------
    # CRITERION 7: Content Integrity (10 pts)
    # --------------------------------------------------------------------------
    full_text = " ".join([p.text for p in doc.paragraphs]).lower()
    integrity_phrases = [
        "1847 industrial parkway",
        "metro city fire department",
        "battalion 2"
    ]
    phrases_found = sum(1 for p in integrity_phrases if p in full_text)
    
    if phrases_found >= 2:
        score += 10
        feedback.append("Content Integrity: Passed")
    else:
        feedback.append("Content Integrity: Failed (Text appears replaced)")

    # Cleanup
    if os.path.exists(temp_docx.name):
        os.unlink(temp_docx.name)

    passed = score >= 60
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }