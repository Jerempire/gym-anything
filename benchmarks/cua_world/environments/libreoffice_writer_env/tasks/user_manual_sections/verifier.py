#!/usr/bin/env python3
"""
Verifier for user_manual_sections task.

Criteria:
1. Output file exists and is a valid DOCX.
2. Document has multiple sections (at least 4).
3. First section (Cover) has NO header.
4. Subsequent sections have header "AquaTrack User Manual".
5. Heading 1/2 styles applied to chapters/subsections.
6. Body text is formatted to Liberation Sans/Arial 11pt.
"""

import json
import os
import tempfile
import logging
from typing import Dict, Any, List

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try importing python-docx (should be available in verifier env)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.error("python-docx not installed in verifier environment")

def verify_user_manual_sections(traj: Dict[str, Any], env_info: Dict[str, Any], task_info: Dict[str, Any]) -> Dict[str, Any]:
    """Verify the restructured user manual."""
    
    # 0. Check infrastructure
    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "Verifier Error: python-docx not available"}
        
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Verifier Error: copy_from_env not available"}

    # 1. Parse result metadata
    try:
        temp_result = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
        copy_from_env("/tmp/task_result.json", temp_result.name)
        with open(temp_result.name, 'r') as f:
            result_meta = json.load(f)
        os.unlink(temp_result.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load task result JSON: {e}"}

    if not result_meta.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file not found"}

    if not result_meta.get("file_created_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file exists but was not created during this task"}

    # 2. Retrieve and parse the DOCX file
    try:
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        copy_from_env("/home/ga/Documents/aquatrack_manual_formatted.docx", temp_docx.name)
        doc = Document(temp_docx.name)
        os.unlink(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output DOCX: {e}"}

    # 3. Evaluate Criteria
    score = 0
    feedback = []
    
    # Criterion A: Section Breaks (15 pts)
    # Require at least 4 sections (Title + 3 chapters minimum split)
    section_count = len(doc.sections)
    if section_count >= 4:
        score += 15
        feedback.append(f"Pass: Document split into {section_count} sections")
    else:
        feedback.append(f"Fail: Document has only {section_count} sections (expected >= 4)")

    # Criterion B: Header Configuration (30 pts)
    # First section should have empty header
    # Subsequent sections should have "AquaTrack"
    
    # Check Section 1 (Title Page)
    header1 = doc.sections[0].header
    header1_text = " ".join([p.text for p in header1.paragraphs]).strip()
    
    if len(header1_text) == 0:
        score += 15
        feedback.append("Pass: Cover page header is empty")
    else:
        feedback.append(f"Fail: Cover page header contains text: '{header1_text}'")

    # Check subsequent sections
    body_headers_ok = False
    for i in range(1, len(doc.sections)):
        h = doc.sections[i].header
        htext = " ".join([p.text for p in h.paragraphs]).strip()
        if "AquaTrack" in htext:
            body_headers_ok = True
            break
    
    if body_headers_ok:
        score += 15
        feedback.append("Pass: Body sections contain 'AquaTrack' in header")
    else:
        feedback.append("Fail: No body section headers contain 'AquaTrack'")

    # Criterion C: Heading Styles (25 pts)
    # Check for "Heading 1" and "Heading 2" usage
    h1_count = 0
    h2_count = 0
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            h1_count += 1
        elif p.style.name == "Heading 2":
            h2_count += 1
            
    if h1_count >= 4:
        score += 15
        feedback.append(f"Pass: Found {h1_count} 'Heading 1' paragraphs")
    else:
        feedback.append(f"Fail: Only found {h1_count} 'Heading 1' paragraphs (expected >= 4)")

    if h2_count >= 6: # There are about 8 subsections in the text
        score += 10
        feedback.append(f"Pass: Found {h2_count} 'Heading 2' paragraphs")
    else:
        feedback.append(f"Fail: Only found {h2_count} 'Heading 2' paragraphs (expected >= 6)")

    # Criterion D: Font Formatting (30 pts)
    # Check body text (not headings) for font family and size
    valid_fonts = ["liberation sans", "arial"]
    body_runs_checked = 0
    body_runs_passed = 0
    
    for p in doc.paragraphs:
        if "Heading" in p.style.name or "Title" in p.style.name:
            continue
        if not p.text.strip():
            continue
            
        for run in p.runs:
            if not run.text.strip():
                continue
                
            body_runs_checked += 1
            
            # Check font name
            font_name = run.font.name
            if not font_name and p.style.font:
                font_name = p.style.font.name
            
            # Check font size
            font_size = run.font.size
            if not font_size and p.style.font:
                font_size = p.style.font.size
            
            # Normalize checks
            name_ok = False
            if font_name and font_name.lower() in valid_fonts:
                name_ok = True
            
            size_ok = False
            if font_size:
                # 11pt is 140970 EMU. Allow small tolerance.
                pt_val = font_size.pt
                if 10.5 <= pt_val <= 11.5:
                    size_ok = True
            
            if name_ok and size_ok:
                body_runs_passed += 1

    percent_compliant = 0
    if body_runs_checked > 0:
        percent_compliant = (body_runs_passed / body_runs_checked) * 100
        
    if percent_compliant >= 50: # Allow some misses/partial runs
        score += 30
        feedback.append(f"Pass: Body text formatting is {percent_compliant:.1f}% compliant")
    else:
        feedback.append(f"Fail: Body text formatting is only {percent_compliant:.1f}% compliant (expected > 50%)")

    # Final Result
    passed = score >= 60 # Threshold
    
    return {
        "passed": passed,
        "score": score,
        "feedback": "; ".join(feedback)
    }