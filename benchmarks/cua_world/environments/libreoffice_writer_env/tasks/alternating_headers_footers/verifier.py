#!/usr/bin/env python3
"""
Verifier for alternating_headers_footers task.
Checks that headers/footers are correctly configured with
different first page and odd/even content.
"""

import os
import sys
import json
import logging
import tempfile
import shutil

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try imports
try:
    from docx import Document
    from lxml import etree
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx or lxml not available")


def verify_alternating_headers(traj, env_info, task_info):
    """
    Main verification function for alternating headers/footers task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "Verification dependencies missing (python-docx/lxml)"}

    # Load task result metadata
    temp_result_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result_json.name)
        with open(temp_result_json.name, 'r') as f:
            task_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read task result: {e}"}
    finally:
        if os.path.exists(temp_result_json.name):
            os.unlink(temp_result_json.name)

    # Check basic file existence
    if not task_result.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file not found"}
    
    if not task_result.get("file_created_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file was not created/modified during the task session"}

    if task_result.get("original_modified", False):
        return {"passed": False, "score": 0, "feedback": "Failed: You overwrote the original file instead of saving as a new one"}

    # Copy the output document for analysis
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/ind_annual_report_formatted.docx", temp_doc.name)
        doc = Document(temp_doc.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output document: {e}"}
    
    # Initialize scoring
    score = 0
    feedback_parts = []
    
    # Get the first section (assuming single section document)
    section = doc.sections[0]
    
    # === Criterion 1: Different First Page Enabled (15 pts) ===
    if section.different_first_page_header_footer:
        score += 15
        feedback_parts.append("Different first page enabled")
    else:
        feedback_parts.append("Different first page NOT enabled")
        
    # === Criterion 2: Odd/Even Pages Enabled (15 pts) ===
    # Check w:evenAndOddHeaders in settings.xml or section properties
    even_odd_enabled = False
    try:
        # Check settings.xml
        settings = doc.settings.element
        if settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}evenAndOddHeaders') is not None:
            even_odd_enabled = True
        
        # If not in settings, check if even header has distinct content (fallback)
        if not even_odd_enabled:
            # If even header exists and has content different from odd header
            if section.even_page_header and len(section.even_page_header.paragraphs) > 0:
                even_odd_enabled = True
    except Exception:
        pass

    if even_odd_enabled:
        score += 15
        feedback_parts.append("Odd/Even headers enabled")
    else:
        feedback_parts.append("Odd/Even headers NOT enabled")

    # === Criterion 3: First Page Header Empty (10 pts) ===
    # Note: If different first page is enabled, we check first_page_header
    fp_header_text = ""
    try:
        fp_header = section.first_page_header
        fp_header_text = "".join([p.text for p in fp_header.paragraphs]).strip()
    except:
        pass
    
    if section.different_first_page_header_footer and fp_header_text == "":
        score += 10
        feedback_parts.append("First page header empty")
    elif not section.different_first_page_header_footer:
        feedback_parts.append("First page header not checked (setting disabled)")
    else:
        feedback_parts.append(f"First page header not empty ('{fp_header_text}')")

    # === Criterion 4: First Page Footer Content (10 pts) ===
    fp_footer_text = ""
    try:
        fp_footer = section.first_page_footer
        fp_footer_text = "".join([p.text for p in fp_footer.paragraphs]).strip()
    except:
        pass

    if "CONFIDENTIAL" in fp_footer_text and "Meridian" in fp_footer_text:
        score += 10
        feedback_parts.append("First page footer correct")
    else:
        feedback_parts.append("First page footer missing required text")

    # === Criterion 5: Odd (Default) Header (10 pts) ===
    odd_header_text = ""
    try:
        odd_header = section.header
        odd_header_text = "".join([p.text for p in odd_header.paragraphs]).strip()
    except:
        pass
        
    if "Meridian" in odd_header_text and "MRD-4721" in odd_header_text:
        score += 10
        feedback_parts.append("Odd header content correct")
    else:
        feedback_parts.append("Odd header missing required text")

    # === Criterion 6: Even Header (10 pts) ===
    even_header_text = ""
    try:
        even_header = section.even_page_header
        even_header_text = "".join([p.text for p in even_header.paragraphs]).strip()
    except:
        pass
        
    if "IND 123456" in even_header_text:
        score += 10
        feedback_parts.append("Even header content correct")
    else:
        feedback_parts.append("Even header missing required text")

    # === Criterion 7 & 8: Page Numbers in Footers (30 pts total) ===
    # We check for XML fields indicating page numbers
    def has_page_field(footer_obj):
        try:
            xml = etree.tostring(footer_obj._element, encoding='unicode')
            return 'w:fldSimple' in xml or 'PAGE' in xml or 'w:instrText' in xml
        except:
            return False

    odd_footer_has_page = has_page_field(section.footer)
    even_footer_has_page = has_page_field(section.even_page_footer)

    if odd_footer_has_page:
        score += 15
        feedback_parts.append("Odd footer page number found")
    else:
        feedback_parts.append("Odd footer page number missing")

    if even_footer_has_page:
        score += 15
        feedback_parts.append("Even footer page number found")
    else:
        feedback_parts.append("Even footer page number missing")

    # Clean up
    if os.path.exists(temp_doc.name):
        os.unlink(temp_doc.name)

    passed = score >= 65 and section.different_first_page_header_footer and even_odd_enabled
    
    return {
        "passed": passed,
        "score": score,
        "feedback": "; ".join(feedback_parts)
    }