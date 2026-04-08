#!/usr/bin/env python3
"""
Verifier for screenplay_formatting task.
Checks if the document was reformatted to industry standards.
"""

import json
import os
import sys
import tempfile
import logging
from typing import Dict, Any

# Add utils path
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from writer_verification_utils import (
    copy_and_parse_document,
    cleanup_verification_temp,
    get_paragraph_styles,
    vlm_verify_screenshot
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants for verification
# 1 inch = 914400 EMU
INCH_EMU = 914400
TOLERANCE_EMU = 60000 # ~1/16th inch tolerance

def verify_screenplay_formatting(traj: Dict[str, Any], env_info: Dict[str, Any], task_info: Dict[str, Any]) -> Dict[str, Any]:
    """
    Verify the screenplay formatting task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "System Error: Copy function not available"}

    metadata = task_info.get('metadata', {})
    
    # 1. Check Output File Existence & Timestamp
    # -----------------------------------------
    # Get the export result json
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            export_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Could not retrieve task result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not export_result.get('output_exists', False):
        return {"passed": False, "score": 0, "feedback": "Output file 'screenplay_formatted.docx' not found."}

    if not export_result.get('file_created_during_task', False):
        return {"passed": False, "score": 0, "feedback": "Output file was not modified during the task session."}

    # 2. Parse the Document Content
    # -----------------------------
    container_path = "/home/ga/Documents/screenplay_formatted.docx"
    success, doc, error_msg, temp_dir = copy_and_parse_document(container_path, copy_from_env, file_format='docx')

    if not success:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output document: {error_msg}"}

    score = 0
    max_score = 100
    feedback = []
    
    try:
        # A. Verify Font (Courier New/Courier/Liberation Mono) - 15 pts
        # -----------------------------------------------------------
        accepted_fonts = [f.lower() for f in metadata.get('accepted_fonts', [])]
        target_size = metadata.get('target_font_size', 12)
        
        font_matches = 0
        size_matches = 0
        total_runs = 0
        
        for para in doc.paragraphs:
            if not para.text.strip(): continue
            for run in para.runs:
                if not run.text.strip(): continue
                total_runs += 1
                
                # Check font name
                run_font = (run.font.name or "").lower()
                # If run has no font, check style font (not implemented deep here, assuming direct format or style)
                # Fallback to checking if style implies it, but task asked for direct change or style change
                if any(af in run_font for af in accepted_fonts):
                    font_matches += 1
                
                # Check font size
                # python-docx returns None if inherited, or Pt object
                if run.font.size is not None and run.font.size.pt == target_size:
                    size_matches += 1
                # If None, it inherits. We can be strict or lenient. Let's be lenient if style was updated.
                elif run.font.size is None:
                    # Assuming agent updated "Default Paragraph Style" or "Normal"
                    pass 

        # We'll use a threshold of 50% of text content having correct properties
        # (Since some might be missed or handled via style inheritance we can't easily parse)
        if total_runs > 0:
            if (font_matches / total_runs) > 0.5:
                score += 15
                feedback.append("Font family Correct (Courier/Mono).")
            else:
                feedback.append(f"Font family incorrect or inconsistent ({font_matches}/{total_runs} runs matched).")
                
            if (size_matches / total_runs) > 0.5:
                score += 15
                feedback.append("Font size Correct (12pt).")
            else:
                feedback.append(f"Font size incorrect ({size_matches}/{total_runs} runs matched).")
        else:
            feedback.append("Document appears empty.")

        # B. Verify Margins - 25 pts total
        # --------------------------------
        # Left: 1.5 inch, Others: 1.0 inch
        section = doc.sections[0]
        
        # Left Margin (1.5 inch)
        expected_left = 1.5 * INCH_EMU
        if abs(section.left_margin - expected_left) <= TOLERANCE_EMU:
            score += 15
            feedback.append("Left margin Correct (1.5\").")
        else:
            actual = section.left_margin / INCH_EMU if section.left_margin else 0
            feedback.append(f"Left margin incorrect (Found {actual:.2f}\").")

        # Other Margins (1.0 inch)
        expected_std = 1.0 * INCH_EMU
        margins_ok = True
        for m in [section.right_margin, section.top_margin, section.bottom_margin]:
            if abs(m - expected_std) > TOLERANCE_EMU:
                margins_ok = False
        
        if margins_ok:
            score += 10
            feedback.append("Top/Bottom/Right margins Correct (1.0\").")
        else:
            feedback.append("Other margins incorrect.")

        # C. Verify Styles & Alignment - 35 pts total
        # -------------------------------------------
        # Scene Headings -> Heading 1
        # Character Names -> Centered
        # Transitions -> Right Aligned
        
        scene_headings = metadata.get('scene_headings', [])
        character_names = metadata.get('character_names', [])
        transitions = metadata.get('transitions', [])
        
        # Heading 1 check
        h1_count = 0
        for para in doc.paragraphs:
            txt = para.text.strip()
            if any(sh in txt for sh in scene_headings):
                if para.style and "Heading 1" in para.style.name:
                    h1_count += 1
        
        if h1_count >= 3: # Allow 1 miss
            score += 15
            feedback.append("Scene headings have 'Heading 1' style.")
        else:
            feedback.append(f"Scene headings missing 'Heading 1' style (Found {h1_count}).")

        # Centered Character Names
        center_count = 0
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt in character_names:
                # WD_ALIGN_PARAGRAPH.CENTER is 1
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    center_count += 1
        
        if center_count >= 3:
            score += 10
            feedback.append("Character names are centered.")
        else:
            feedback.append(f"Character names not centered (Found {center_count}).")

        # Right Aligned Transitions
        right_count = 0
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt in transitions:
                # WD_ALIGN_PARAGRAPH.RIGHT is 2
                if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    right_count += 1
                    
        if right_count >= 2: # Allow 1 miss
            score += 10
            feedback.append("Transitions are right-aligned.")
        else:
            feedback.append(f"Transitions not right-aligned (Found {right_count}).")

        # D. VLM Trajectory Verification - 10 pts
        # ---------------------------------------
        # Check if the agent actually worked in the UI (formatting, menus)
        # Use VLM to look for "Page Style" dialog or "Paragraph" formatting interaction
        
        vlm_result = vlm_verify_screenshot(
            env_info, traj,
            prompt="Does this screenshot show LibreOffice Writer with a screenplay document? "
                   "Look for formatting actions like: centered text, courier font, or the Page Style dialog open. "
                   "Answer JSON with keys: 'is_writer', 'formatting_visible'."
        )
        
        if vlm_result.get("parsed", {}).get("is_writer", False):
            score += 10
            feedback.append("Visual verification passed.")
        else:
            feedback.append("Visual verification inconclusive.")

    except Exception as e:
        logger.error(f"Verification logic error: {e}")
        feedback.append(f"Verification error: {e}")
    finally:
        cleanup_verification_temp(temp_dir)

    return {
        "passed": score >= 60,
        "score": score,
        "feedback": " ".join(feedback)
    }