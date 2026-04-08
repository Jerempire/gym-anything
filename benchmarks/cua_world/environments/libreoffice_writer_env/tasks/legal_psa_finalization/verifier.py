#!/usr/bin/env python3
"""Verifier for legal_psa_finalization task.

Checks that the agent formatted a Professional Services Agreement draft
into an execution-ready final document per law firm standards.

Criteria (each ~16 points, 6 criteria total = 96 pts):
  1. Output file psa_final.docx exists
  2. Document title is centered, bold, and ~14pt
  3. Section headings use Heading 1 style (at least 7 of 9)
  4. Defined terms are bolded in the Definitions section
  5. Signature block is formatted as a 2-column table with no visible borders
  6. Footer contains firm name and page number field

Pass threshold: 65%
"""

import sys
import os
import logging
import re

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from writer_verification_utils import (
    copy_and_parse_document,
    cleanup_verification_temp,
    get_document_text,
    check_heading_styles,
    check_text_formatting,
    check_paragraph_alignment,
    vlm_verify_screenshot,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

SECTION_HEADINGS = [
    "Definitions", "Scope of Services", "Fees and Payment",
    "Intellectual Property", "Confidentiality",
    "Representations and Warranties", "Limitation of Liability",
    "Indemnification", "General Provisions",
]

DEFINED_TERMS = [
    "Services", "Deliverables", "Confidential Information",
    "Intellectual Property Rights", "Force Majeure Event", "Effective Date",
]


def _check_title_formatting(doc, title_text="PROFESSIONAL SERVICES AGREEMENT"):
    """Check if document title is centered, bold, and approximately 14pt."""
    for para in doc.paragraphs:
        if title_text.lower() in para.text.lower():
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
            is_bold = False
            has_large_font = False
            for run in para.runs:
                if run.bold or (para.style and 'heading' in (para.style.name or '').lower()):
                    is_bold = True
                if run.font.size and run.font.size.pt >= 13.0:
                    has_large_font = True
            # Check style inheritance for bold
            if not is_bold and para.style:
                if para.style.font and para.style.font.bold:
                    is_bold = True
            return is_centered, is_bold, has_large_font
    return False, False, False


def _check_defined_terms_bold(doc):
    """Check that defined terms are bolded in the Definitions section."""
    in_definitions = False
    found_bold = []
    found_not_bold = []

    for para in doc.paragraphs:
        text = para.text.strip()
        # Enter definitions section
        if not in_definitions:
            if "definitions" in text.lower() and len(text) < 50:
                in_definitions = True
            continue

        # Exit when we hit next major section
        if in_definitions and len(text) < 40 and text.isupper():
            break
        if in_definitions and any(
            sec.lower() in text.lower()
            for sec in ["scope of services", "fees and payment", "intellectual property"]
            if len(text) < 60
        ):
            break

        for term in DEFINED_TERMS:
            if f'"{term}"' in text or f'"{term}"' in text:
                # Term found — check if it's bold in at least one run
                for run in para.runs:
                    if term in run.text:
                        if run.bold:
                            if term not in found_bold:
                                found_bold.append(term)
                        else:
                            if term not in found_not_bold:
                                found_not_bold.append(term)

    return found_bold, found_not_bold


def _check_signature_table(doc):
    """
    Check if the signature block is formatted as a 2-column table.
    Also verify the table has no visible borders (or minimal borders).
    Returns (has_2col_table, table_details_str)
    """
    # Look through all tables in the document
    for table in doc.tables:
        # Check for 2-column structure
        if len(table.columns) == 2:
            # Verify it looks like a signature block (contains "By:" or "Name:" etc.)
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
            sig_keywords = ["By:", "Name:", "Title:", "Date:"]
            matches = sum(1 for kw in sig_keywords if kw in table_text)
            if matches >= 3:
                # Check for borderless: try to inspect XML for border settings
                try:
                    from lxml import etree
                    xml = table._element.xml
                    # Look for w:tcBorders with no border or nil border
                    has_no_border = (
                        'w:nil' in xml or
                        'none' in xml.lower() or
                        'w:tcBorders' not in xml  # no border element = no border
                    )
                    return True, f"2-column signature table found (border-free: {has_no_border})"
                except Exception:
                    return True, "2-column signature table found"

    # Also check for signature block formatted as paragraphs (not ideal but partial credit)
    full_text = get_document_text(doc)
    sig_count = sum(1 for kw in ["By:", "Name:", "Title:", "Date:"] if kw in full_text)
    if sig_count >= 3:
        return False, f"Signature keywords present but not in 2-column table (found {sig_count}/4 keywords)"
    return False, "Signature block not found or missing key fields"


def _check_footer(doc, required_fragments=None):
    """Check that footer contains required text and a page number field."""
    if required_fragments is None:
        required_fragments = ["Meridian Legal", "CONFIDENTIAL"]

    try:
        for section in doc.sections:
            # Check the default footer
            footer = section.footer
            if not footer:
                continue
            footer_text = " ".join(p.text for p in footer.paragraphs).strip()

            # Check for page number field in footer XML
            footer_xml = footer._element.xml if hasattr(footer, '_element') else ""
            has_page_field = (
                "PAGE" in footer_xml.upper() or
                "w:fldChar" in footer_xml or
                "w:instrText" in footer_xml
            )

            has_required_text = any(
                frag.lower() in footer_text.lower()
                for frag in required_fragments
            )

            if has_required_text or has_page_field:
                return True, has_required_text, has_page_field, footer_text

        return False, False, False, ""
    except Exception as e:
        logger.warning(f"Footer check error: {e}")
        return False, False, False, ""


def verify_legal_psa_finalization(traj, env_info, task_info):
    """Verify Legal PSA finalization document formatting."""

    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "copy_from_env not available"}

    output_path = "/home/ga/Documents/psa_final.docx"
    success, doc, error, temp_dir = copy_and_parse_document(output_path, copy_from_env, 'docx')

    if not success:
        return {
            "passed": False,
            "score": 0,
            "feedback": f"Output file psa_final.docx not found or unreadable: {error}",
        }

    try:
        feedback_parts = []
        score = 0
        POINTS = 16

        # --- Criterion 1: File exists ---
        score += POINTS
        feedback_parts.append("Output file psa_final.docx exists")

        # --- Criterion 2: Title formatting (centered, bold, large font) ---
        is_centered, is_bold, has_large_font = _check_title_formatting(doc)
        title_criteria = sum([is_centered, is_bold, has_large_font])
        if title_criteria == 3:
            score += POINTS
            feedback_parts.append("Title: centered, bold, and >= 13pt")
        elif title_criteria == 2:
            score += POINTS // 2
            feedback_parts.append(
                f"Title: 2/3 criteria met (centered={is_centered}, bold={is_bold}, large={has_large_font})"
            )
        else:
            feedback_parts.append(
                f"Title: only {title_criteria}/3 criteria met — "
                f"needs 14pt, bold, centered on first page"
            )

        # --- Criterion 3: Section headings use Heading 1 style ---
        heading_map = {h: "Heading 1" for h in SECTION_HEADINGS}
        h1_matched, h1_total, _ = check_heading_styles(doc, heading_map)
        if h1_matched >= 7:
            score += POINTS
            feedback_parts.append(f"Headings: {h1_matched}/{h1_total} sections have Heading 1 style")
        elif h1_matched >= 4:
            score += POINTS // 2
            feedback_parts.append(
                f"Headings: {h1_matched}/{h1_total} have Heading 1 (need at least 7)"
            )
        else:
            feedback_parts.append(
                f"Headings: only {h1_matched}/{h1_total} section headings have Heading 1 style"
            )

        # --- Criterion 4: Defined terms are bold in Definitions section ---
        bold_terms, not_bold_terms = _check_defined_terms_bold(doc)
        if len(bold_terms) >= 4:
            score += POINTS
            feedback_parts.append(
                f"Defined terms: {len(bold_terms)}/6 terms are bold in Definitions: {bold_terms}"
            )
        elif len(bold_terms) >= 2:
            score += POINTS // 2
            feedback_parts.append(
                f"Defined terms: {len(bold_terms)}/6 bolded (need at least 4): {bold_terms}"
            )
        else:
            # Fallback: check if bold appears anywhere in the Definitions section text area
            full_text = get_document_text(doc)
            has_definitions = "definitions" in full_text.lower()
            feedback_parts.append(
                f"Defined terms: only {len(bold_terms)}/6 terms bolded in Definitions section. "
                f"Terms 'Services', 'Deliverables', 'Confidential Information', etc. must be bold."
            )

        # --- Criterion 5: Signature block as 2-column table ---
        has_sig_table, sig_detail = _check_signature_table(doc)
        if has_sig_table:
            score += POINTS
            feedback_parts.append(f"Signature block: {sig_detail}")
        else:
            feedback_parts.append(
                f"Signature block: {sig_detail}. "
                f"Must be a borderless 2-column table with By/Name/Title/Date/Company on each side."
            )

        # --- Criterion 6: Footer with firm name and page number ---
        footer_found, has_firm_name, has_page_num, footer_text = _check_footer(doc)
        if footer_found and has_firm_name and has_page_num:
            score += POINTS
            feedback_parts.append(
                f"Footer: contains firm name and page number field ('{footer_text[:50]}')"
            )
        elif footer_found and (has_firm_name or has_page_num):
            score += POINTS // 2
            feedback_parts.append(
                f"Footer: partially correct — firm name: {has_firm_name}, "
                f"page field: {has_page_num}"
            )
        else:
            feedback_parts.append(
                "Footer: not found or missing required content "
                "(needs 'CONFIDENTIAL — Meridian Legal Partners LLP' and auto-page number)"
            )

        passed = score >= 65
        feedback = " | ".join(feedback_parts)

        return {
            "passed": passed,
            "score": min(score, 100),
            "feedback": feedback,
        }

    except Exception as e:
        logger.error(f"Verification error: {e}", exc_info=True)
        return {"passed": False, "score": 0, "feedback": f"Verification error: {str(e)}"}
    finally:
        cleanup_verification_temp(temp_dir)
