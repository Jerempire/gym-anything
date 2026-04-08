#!/usr/bin/env python3
"""
Verifier for bilingual_safety_alignment task.

Criteria:
1. File exists and was created/modified during task.
2. Document contains a table with 2 columns.
3. Header row exists ("English" / "Español").
4. Semantic Alignment: Specific keywords from English/Spanish rules appear in the SAME row.
5. Cleanup: Original sequential text is removed from the document body.
"""

import json
import os
import tempfile
import logging
from collections import Counter

# Import shared Writer verification utilities
# Assumes the environment provides these in utils/
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_bilingual_alignment(traj, env_info, task_info):
    """
    Verify the transformation of sequential text to a side-by-side table.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "python-docx library not available in verifier"}

    # 1. Load result JSON from export script
    task_result = {}
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            task_result = json.load(f)
    except Exception as e:
        logger.warning(f"Failed to load task result JSON: {e}")
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # 2. Check file existence
    if not task_result.get('output_exists', False):
        return {
            "passed": False, 
            "score": 0, 
            "feedback": "Output file 'forklift_safety_aligned.docx' not found."
        }
    
    if not task_result.get('file_fresh', False):
         return {
            "passed": False, 
            "score": 0, 
            "feedback": "Output file exists but was not modified during the task."
        }

    # 3. Download and parse the DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/forklift_safety_aligned.docx", temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {
            "passed": False, 
            "score": 0, 
            "feedback": f"Failed to parse output document: {e}"
        }
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 0
    feedback_parts = []
    
    # --- Check 1: Table Structure (20 pts) ---
    tables = doc.tables
    if not tables:
        return {"passed": False, "score": 10, "feedback": "File exists, but no table found."}
    
    # Assume the main table is the first one found that has content
    target_table = None
    for t in tables:
        if len(t.columns) == 2 and len(t.rows) >= 5:
            target_table = t
            break
    
    if target_table:
        score += 20
        feedback_parts.append("Valid 2-column table found")
    else:
        # Penalize but continue checking (maybe they used 3 cols?)
        feedback_parts.append("No suitable 2-column table found (cols!=2 or rows<5)")
        target_table = tables[0] # Fallback to check content anyway

    # --- Check 2: Header Row (10 pts) ---
    # Check first few rows for headers
    header_found = False
    header_row_idx = 0
    
    for i in range(min(3, len(target_table.rows))):
        row_text = [cell.text.lower().strip() for cell in target_table.rows[i].cells]
        # Check for English/Spanish keywords in separate cells
        if any("english" in t for t in row_text) and any("spa" in t for t in row_text):
            header_found = True
            header_row_idx = i
            score += 10
            feedback_parts.append("Header row found")
            break
            
    if not header_found:
        feedback_parts.append("Header row missing or incorrect")

    # --- Check 3: Semantic Alignment (50 pts) ---
    # We check if corresponding EN/ES keywords appear in the SAME row
    alignment_keywords = task_info.get('metadata', {}).get('alignment_keywords', [])
    aligned_count = 0
    
    # Iterate rows starting after header
    for row in target_table.rows[header_row_idx+1:]:
        if len(row.cells) < 2: 
            continue
            
        # Get text from first 2 cells
        c1_text = row.cells[0].text.lower()
        c2_text = row.cells[1].text.lower()
        row_content = c1_text + " " + c2_text
        
        # Check against our expected pairs
        for pair in alignment_keywords:
            en_kw = pair['en'].lower()
            es_kw = pair['es'].lower()
            
            # Check strict separation (best case): En in Col 1, Es in Col 2 (or vice versa)
            strict_match = (en_kw in c1_text and es_kw in c2_text) or \
                           (en_kw in c2_text and es_kw in c1_text)
            
            if strict_match:
                aligned_count += 1
                # Mark this pair as found so we don't double count if rows duplicate? 
                # For simplicity, we just count hits.
    
    # Scoring alignment: proportional to expected rules
    expected_count = len(alignment_keywords)
    if aligned_count >= expected_count:
        score += 50
        feedback_parts.append(f"Perfect alignment ({aligned_count}/{expected_count} rules)")
    elif aligned_count > 0:
        partial_pts = int((aligned_count / expected_count) * 50)
        score += partial_pts
        feedback_parts.append(f"Partial alignment ({aligned_count}/{expected_count} rules)")
    else:
        feedback_parts.append("No content alignment detected")

    # --- Check 4: Cleanup (10 pts) ---
    # Ensure the original text list is NOT in the document body (paragraphs outside table)
    # We assume if the user did it right, the body should be mostly empty or just the title
    body_text = " ".join([p.text.lower() for p in doc.paragraphs])
    
    # Check for a few keywords that should have moved into the table
    leftover_keywords = 0
    check_words = ["buckle", "horquillas", "intersections", "bocina"]
    for w in check_words:
        if w in body_text:
            leftover_keywords += 1
            
    if leftover_keywords == 0:
        score += 10
        feedback_parts.append("Cleanup complete")
    else:
        feedback_parts.append("Original text still present in document body")

    # --- Check 5: File Creation (10 pts) ---
    # We already checked fresh file exists, give points
    score += 10

    passed = score >= 70
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }