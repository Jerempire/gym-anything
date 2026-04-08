#!/usr/bin/env python3
"""
Verifier for chem_lab_character_formatting task.
Checks if chemical formulas and scientific notation have proper subscript/superscript formatting.
"""

import json
import os
import sys
import re
import hashlib
import tempfile
import logging
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_chem_lab_formatting(traj, env_info, task_info):
    """
    Verify the formatting of the chemistry lab report.
    
    Criteria:
    1. Output file exists and was created during task (10 pts)
    2. Text content matches original (no text deletion/corruption) (10 pts)
    3. Original draft file is unmodified (5 pts)
    4. Subscript usage count >= 10 (15 pts)
    5. Superscript usage count >= 8 (15 pts)
    6. Specific formula checks:
       - H2SO4 has subscripts (15 pts)
       - Ca2+ has superscript (10 pts)
       - Scientific notation (10-3) has superscript (10 pts)
       - NO3- has mixed formatting (subscript 3, superscript -) (10 pts)
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Load task result metadata
    temp_result = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result.name)
        with open(temp_result.name, 'r') as f:
            task_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load task result: {e}"}
    finally:
        os.unlink(temp_result.name)

    score = 0
    feedback = []

    # 1. File Existence & Creation (10 pts)
    if not task_result.get('output_exists', False):
        return {"passed": False, "score": 0, "feedback": "Output file not found."}
    
    if not task_result.get('file_created_during_task', False):
        feedback.append("WARNING: Output file timestamp indicates it wasn't created during task.")
    else:
        score += 10
        feedback.append("Output file created successfully.")

    # 2. Original Draft Integrity (5 pts)
    if not task_result.get('original_draft_modified', True):
        score += 5
        feedback.append("Original draft preserved.")
    else:
        feedback.append("Original draft was modified (should have saved as new file).")

    # Load the document for content analysis
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/lab_report_formatted.docx", temp_doc.name)
        doc = Document(temp_doc.name)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse output DOCX: {e}"}
    finally:
        pass # Keep temp file for now or unlink if strictly needed, system cleans up usually

    # 3. Content Preservation (10 pts)
    # Compare text content with expected hash from setup
    current_text = '\n'.join([p.text for p in doc.paragraphs])
    current_hash = hashlib.sha256(current_text.encode('utf-8')).hexdigest()
    
    # Get original hash
    temp_hash = tempfile.NamedTemporaryFile(delete=False)
    try:
        copy_from_env("/tmp/original_text_hash.txt", temp_hash.name)
        with open(temp_hash.name, 'r') as f:
            original_hash = f.read().strip()
    except:
        original_hash = ""
    finally:
        os.unlink(temp_hash.name)

    if current_hash == original_hash:
        score += 10
        feedback.append("Text content preserved exactly.")
    else:
        # Allow slight deviations (whitespace)
        # We assume if the hash doesn't match, maybe they added a space, but if length is close it's OK
        feedback.append("Text content modified (hash mismatch).")

    # Analyze formatting
    subscript_count = 0
    superscript_count = 0
    
    # Specific Checks
    h2so4_correct = False
    ca2_correct = False
    sci_not_correct = False
    no3_mixed_correct = False

    for para in doc.paragraphs:
        runs = para.runs
        
        # General counters
        for run in runs:
            if run.font.subscript:
                subscript_count += 1
            if run.font.superscript:
                superscript_count += 1

        # Specific Logic: Reconstruct formatting map for the paragraph
        # This is complex because "H2SO4" might be split into 4 runs: H, 2(sub), SO, 4(sub)
        
        # Check H2SO4
        if "H2SO4" in para.text or "H₂SO₄" in para.text: # text might contain unicode if they used symbols?
            # Look for '2' and '4' with subscript in this paragraph
            has_sub_2 = any('2' in r.text and r.font.subscript for r in runs)
            has_sub_4 = any('4' in r.text and r.font.subscript for r in runs)
            if has_sub_2 and has_sub_4:
                h2so4_correct = True

        # Check Ca2+
        if "Ca2+" in para.text:
            # Look for '2+' or '2' and '+' with superscript
            has_sup_2plus = any('2+' in r.text and r.font.superscript for r in runs)
            has_sup_parts = (any('2' in r.text and r.font.superscript for r in runs) and 
                             any('+' in r.text and r.font.superscript for r in runs))
            if has_sup_2plus or has_sup_parts:
                ca2_correct = True

        # Check Scientific Notation (look for negative exponents like -3, -1)
        # Pattern: ×10-3
        if "10-" in para.text or "10" in para.text:
            # Look for superscript run containing '-' and a digit
            for r in runs:
                if r.font.superscript and re.search(r'-[1-9]', r.text):
                    sci_not_correct = True

        # Check NO3- (mixed)
        if "NO3-" in para.text:
            has_sub_3 = any('3' in r.text and r.font.subscript for r in runs)
            has_sup_min = any('-' in r.text and r.font.superscript for r in runs)
            if has_sub_3 and has_sup_min:
                no3_mixed_correct = True

    # 4. Subscript Count (15 pts)
    if subscript_count >= 10:
        score += 15
        feedback.append(f"Subscript usage good ({subscript_count} runs).")
    elif subscript_count >= 5:
        score += 7
        feedback.append(f"Subscript usage partial ({subscript_count} runs).")
    else:
        feedback.append(f"Insufficient subscript usage ({subscript_count} runs).")

    # 5. Superscript Count (15 pts)
    if superscript_count >= 8:
        score += 15
        feedback.append(f"Superscript usage good ({superscript_count} runs).")
    elif superscript_count >= 4:
        score += 7
        feedback.append(f"Superscript usage partial ({superscript_count} runs).")
    else:
        feedback.append(f"Insufficient superscript usage ({superscript_count} runs).")

    # 6. Specific Formulas (45 pts total)
    if h2so4_correct:
        score += 15
        feedback.append("H2SO4 correctly formatted.")
    if ca2_correct:
        score += 10
        feedback.append("Ca2+ correctly formatted.")
    if sci_not_correct:
        score += 10
        feedback.append("Scientific notation exponent formatted.")
    if no3_mixed_correct:
        score += 10
        feedback.append("NO3- correctly formatted (mixed sub/super).")

    os.unlink(temp_doc.name)

    passed = score >= 60
    return {
        "passed": passed,
        "score": score,
        "feedback": " ".join(feedback)
    }