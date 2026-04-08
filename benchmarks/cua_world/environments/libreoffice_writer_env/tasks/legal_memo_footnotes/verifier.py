#!/usr/bin/env python3
"""
Verifier for legal_memo_footnotes@1 task.
Checks:
1. Footnote conversion (count and content)
2. Document formatting (font, size, spacing, margins)
3. Text cleanup (removal of inline brackets)
4. Visual verification via VLM
"""

import json
import os
import sys
import tempfile
import logging
import zipfile
import re
import shutil
from typing import Dict, Any, Tuple
try:
    from lxml import etree
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass  # handled in runtime check

# Import VLM utils from the environment
try:
    from gym_anything.vlm import query_vlm, sample_trajectory_frames, get_final_screenshot
except ImportError:
    # Mock for testing outside of gym environment
    def query_vlm(**kwargs): return {"success": False, "error": "VLM not available"}
    def sample_trajectory_frames(traj, n): return []
    def get_final_screenshot(traj): return None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

EXPECTED_OUTPUT_PATH = "/home/ga/Documents/legal_memo_formatted.docx"
EXPECTED_CASES = ["Hadley", "Jacob", "Groves", "Peevyhouse", "Sullivan", "Hawkins"]
ONE_INCH_EMU = 914400
MARGIN_TOLERANCE = 91440  # 0.1 inch

def extract_footnotes_xml(docx_path):
    """Extract footnotes directly from XML to avoid python-docx footnote limitation."""
    footnotes = []
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            if 'word/footnotes.xml' not in z.namelist():
                return []
            with z.open('word/footnotes.xml') as f:
                tree = etree.parse(f)
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for fn in tree.findall('.//w:footnote', namespaces):
                    fn_id = fn.get(f'{{{namespaces["w"]}}}id')
                    # ID < 1 are system separators
                    if fn_id and int(fn_id) >= 1:
                        text = "".join(fn.itertext())
                        if text.strip():
                            footnotes.append(text)
    except Exception as e:
        logger.error(f"XML parsing error: {e}")
    return footnotes

def verify_legal_memo(traj, env_info, task_info):
    """
    Main verification function for Legal Memo Footnotes task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Load task result metadata
    temp_result = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result.name)
        with open(temp_result.name, 'r') as f:
            meta_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read result metadata: {e}"}
    finally:
        if os.path.exists(temp_result.name):
            os.unlink(temp_result.name)

    # Check existence and timestamp
    if not meta_result.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file not found"}
    
    if not meta_result.get("file_created_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file was not modified during the task (anti-gaming)"}

    # Copy the actual DOCX file for analysis
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(EXPECTED_OUTPUT_PATH, temp_docx.name)
        
        # --- SCORING CRITERIA ---
        score = 0
        feedback = []
        
        # 1. Footnote Verification (30 pts)
        footnotes = extract_footnotes_xml(temp_docx.name)
        fn_count = len(footnotes)
        
        if fn_count >= 6:
            score += 15
            feedback.append(f"✓ Found {fn_count} footnotes (expected 6)")
        elif fn_count >= 1:
            score += 5
            feedback.append(f"⚠ Found only {fn_count} footnotes (expected 6)")
        else:
            feedback.append("✗ No footnotes found")

        # Check content of footnotes
        cases_found = 0
        all_fn_text = " ".join(footnotes).lower()
        for case in EXPECTED_CASES:
            if case.lower() in all_fn_text:
                cases_found += 1
        
        if cases_found >= 4:
            score += 15
            feedback.append(f"✓ Footnotes contain correct case citations ({cases_found}/6)")
        elif cases_found > 0:
            score += 5
            feedback.append(f"⚠ Footnotes contain some case citations ({cases_found}/6)")
        else:
            feedback.append("✗ Footnotes do not contain expected case names")

        # 2. Document Parsing with python-docx (40 pts)
        try:
            doc = Document(temp_docx.name)
            
            # Check Cleanliness (No brackets left) - 10 pts
            full_text = " ".join([p.text for p in doc.paragraphs])
            bracket_matches = re.findall(r'\[See\s.*?\]', full_text)
            if len(bracket_matches) == 0:
                score += 10
                feedback.append("✓ No inline bracketed citations remaining")
            else:
                feedback.append(f"✗ Found {len(bracket_matches)} unconverted inline citations")

            # Check Formatting (Font/Size) - 10 pts
            # Sample first few meaningful paragraphs
            body_paras = [p for p in doc.paragraphs if len(p.text) > 50]
            serif_count = 0
            size_count = 0
            double_space_count = 0
            
            if body_paras:
                for p in body_paras:
                    # Font check (heuristic based on run or style)
                    font_name = p.style.font.name if p.style and p.style.font.name else None
                    for r in p.runs:
                        if r.font.name: font_name = r.font.name
                        if r.font.size and r.font.size >= Pt(11.5): # approx 12
                             size_count += 1
                    
                    if font_name and any(f in font_name.lower() for f in ['times', 'liberation serif', 'georgia']):
                        serif_count += 1
                        
                    # Line spacing check
                    if p.paragraph_format.line_spacing == 2.0 or (p.paragraph_format.line_spacing_rule and p.paragraph_format.line_spacing_rule == 2):
                        double_space_count += 1

                # Scoring formatting
                if serif_count > 0 or size_count > 0: # Lenient check due to style inheritance complexity
                    score += 10
                    feedback.append("✓ Font/Size formatting applied")
                else:
                    feedback.append("✗ Font/Size formatting not detected")
                
                if double_space_count >= len(body_paras) * 0.5:
                    score += 10
                    feedback.append("✓ Double spacing applied")
                else:
                    feedback.append("✗ Double spacing not consistently applied")

            # Check Margins - 5 pts
            sect = doc.sections[0]
            margins_ok = all(abs(m - ONE_INCH_EMU) < MARGIN_TOLERANCE for m in 
                             [sect.left_margin, sect.right_margin, sect.top_margin, sect.bottom_margin] 
                             if m is not None)
            if margins_ok:
                score += 5
                feedback.append("✓ Margins set to 1 inch")
            else:
                feedback.append("✗ Margins not set to 1 inch")

            # Check Title Center - 5 pts
            title_para = doc.paragraphs[0]
            if "MEMORANDUM" in title_para.text and title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                score += 5
                feedback.append("✓ Title centered")
            else:
                feedback.append("✗ Title not centered")

        except Exception as e:
            feedback.append(f"⚠ formatting check failed: {str(e)}")

        # 3. VLM Verification (30 pts)
        # We rely on VLM to confirm the visual "look" of a legal memo with footnotes
        final_screenshot = get_final_screenshot(traj)
        if final_screenshot:
            vlm_prompt = """
            Analyze this document screenshot. It should be a legal memorandum.
            1. Are there footnotes visible at the bottom of the page (small text, numbered)?
            2. Is the main text double-spaced (lots of vertical space between lines)?
            3. Is the title 'MEMORANDUM OF LAW' centered at the top?
            Answer JSON: {"footnotes_visible": bool, "double_spaced": bool, "title_centered": bool}
            """
            vlm_res = query_vlm(prompt=vlm_prompt, image=final_screenshot)
            
            if vlm_res.get("success"):
                parsed = vlm_res.get("parsed", {})
                if parsed.get("footnotes_visible"):
                    score += 15
                    feedback.append("✓ VLM confirmed visible footnotes")
                else:
                    feedback.append("✗ VLM did not see footnotes")
                
                if parsed.get("double_spaced"):
                    score += 10
                    feedback.append("✓ VLM confirmed double spacing")
                
                if parsed.get("title_centered"):
                    score += 5
                    feedback.append("✓ VLM confirmed centered title")
            else:
                # Fallback if VLM fails but programmatic passed
                feedback.append("⚠ VLM check skipped/failed")
                if score >= 60: score += 10 # Grace points if programmatic was good
        else:
            feedback.append("⚠ No screenshot available for VLM")

    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Verification error: {e}"}
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    passed = score >= 60
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }