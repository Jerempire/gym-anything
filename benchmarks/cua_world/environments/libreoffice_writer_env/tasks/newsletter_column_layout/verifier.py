#!/usr/bin/env python3
"""
Verifier for newsletter_column_layout task.
Verifies document structure (sections, columns), styling, and content.
"""

import json
import os
import tempfile
import logging
from typing import Dict, Any

# Import gym_anything utils if available, or define mocks for standalone testing
try:
    from gym_anything.vlm import query_vlm, get_final_screenshot
except ImportError:
    # Fallback for testing environment
    def get_final_screenshot(traj): return None
    def query_vlm(**kwargs): return {"success": False, "error": "VLM not available"}

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try importing python-docx (should be installed in environment)
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

def verify_newsletter_layout(traj: Dict[str, Any], env_info: Dict[str, Any], task_info: Dict[str, Any]) -> Dict[str, Any]:
    """
    Verify the newsletter formatting task.
    
    Scoring Criteria:
    1. Output file exists and was created during task (10 pts)
    2. Document has at least 2 sections (Required for mixed column layout) (15 pts)
    3. At least one section has 2 columns (15 pts)
    4. Title is Heading 1 and Centered (10 pts)
    5. Subtitle is Italic and Centered (10 pts)
    6. Article headings use Heading 2 (10 pts)
    7. Footer contains correct contact info (15 pts)
    8. VLM Visual Confirmation of columns (15 pts)
    """
    
    # 1. Setup and Load Data
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    feedback = []
    score = 0
    max_score = 100
    
    # Retrieve metadata
    metadata = task_info.get('metadata', {})
    output_path = metadata.get('output_path', '/home/ga/Documents/diabetes_newsletter.docx')
    
    # Get the export result JSON
    task_result = {}
    with tempfile.NamedTemporaryFile(suffix=".json") as tf:
        try:
            copy_from_env("/tmp/task_result.json", tf.name)
            tf.seek(0)
            task_result = json.load(tf)
        except Exception as e:
            logger.error(f"Failed to load task result: {e}")
            return {"passed": False, "score": 0, "feedback": "Failed to retrieve task execution data"}

    # 2. Verify File Existence (10 pts)
    if not task_result.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file diabetes_newsletter.docx not found."}
    
    if not task_result.get("file_created_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file exists but was not created/modified during this session (Anti-Gaming Trigger)."}
        
    score += 10
    feedback.append("Output file created successfully.")

    # 3. Parse Document content
    if Document is None:
        return {"passed": False, "score": score, "feedback": "Verifier Error: python-docx not installed."}

    doc_temp_path = tempfile.mktemp(suffix=".docx")
    try:
        copy_from_env(output_path, doc_temp_path)
        doc = Document(doc_temp_path)
    except Exception as e:
        return {"passed": False, "score": score, "feedback": f"Failed to parse DOCX file: {e}"}

    # 4. Check Section Breaks (15 pts)
    # To have mixed columns (1-col title, 2-col body), there MUST be a section break.
    # Therefore, len(doc.sections) must be >= 2.
    num_sections = len(doc.sections)
    if num_sections >= 2:
        score += 15
        feedback.append(f"Document structure correct: {num_sections} sections found.")
    else:
        feedback.append(f"Structure Incorrect: Found {num_sections} section. Mixed column layouts require at least 2 sections (Title vs Body).")

    # 5. Check Column Layout (15 pts)
    # python-docx doesn't easily expose column count in the high-level API, 
    # we need to inspect the XML of the section properties (sectPr).
    has_two_columns = False
    try:
        for section in doc.sections:
            sectPr = section._sectPr
            # Check for <w:cols w:num="2">
            cols = sectPr.xpath('./w:cols')
            if cols:
                num_cols = cols[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num')
                if num_cols == '2':
                    has_two_columns = True
                    break
    except Exception as e:
        logger.warning(f"XML parsing for columns failed: {e}")

    if has_two_columns:
        score += 15
        feedback.append("Two-column layout detected in section properties.")
    else:
        feedback.append("Two-column layout NOT detected in XML properties.")

    # 6. Check Title Styling (10 pts)
    # Find paragraph "Harmony Health Quarterly"
    title_para = None
    for p in doc.paragraphs:
        if "Harmony Health Quarterly" in p.text:
            title_para = p
            break
            
    if title_para:
        is_heading1 = title_para.style.name == 'Heading 1'
        is_centered = title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        
        if is_heading1 and is_centered:
            score += 10
            feedback.append("Title styled correctly (Heading 1 + Center).")
        else:
            feedback.append(f"Title found but styling incorrect. Style: {title_para.style.name}, Alignment: {title_para.alignment}")
    else:
        feedback.append("Title text 'Harmony Health Quarterly' not found.")

    # 7. Check Subtitle Styling (10 pts)
    subtitle_para = None
    for p in doc.paragraphs:
        if "Your Guide to Living Well" in p.text:
            subtitle_para = p
            break
            
    if subtitle_para:
        # Check italic (either in style or runs)
        is_italic = False
        # Check explicit runs
        if any(r.italic for r in subtitle_para.runs):
            is_italic = True
        # Check style
        if subtitle_para.style and subtitle_para.style.font.italic:
            is_italic = True
            
        is_centered = subtitle_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        
        if is_italic and is_centered:
            score += 10
            feedback.append("Subtitle styled correctly (Italic + Center).")
        elif is_italic or is_centered:
            score += 5
            feedback.append("Subtitle partially correct.")
        else:
            feedback.append("Subtitle styling incorrect.")

    # 8. Check Headings (10 pts)
    expected_headings = metadata.get('expected_headings', [])
    found_headings = 0
    for p in doc.paragraphs:
        if any(h in p.text for h in expected_headings):
            if p.style.name == 'Heading 2':
                found_headings += 1
    
    if found_headings >= 3:
        score += 10
        feedback.append(f"Article headings styled correctly ({found_headings}/{len(expected_headings)}).")
    else:
        feedback.append(f"Article headings missing Heading 2 style (Found {found_headings}).")

    # 9. Check Footer (15 pts)
    footer_text_found = False
    for section in doc.sections:
        try:
            footer = section.footer
            full_footer_text = " ".join([p.text for p in footer.paragraphs])
            if "(555) 234-5678" in full_footer_text:
                footer_text_found = True
                break
        except:
            continue
            
    if footer_text_found:
        score += 15
        feedback.append("Footer with contact info found.")
    else:
        feedback.append("Footer missing or incorrect content.")

    # 10. VLM Verification for Columns (15 pts)
    # Columns are hard to verify via XML sometimes (defaults), so visual check is crucial
    vlm_prompt = """
    Analyze this document screenshot. 
    Does the document show a newsletter layout with:
    1. A centered title at the top spanning the full width?
    2. The main body text arranged in TWO columns?
    3. A clear visual distinction between the single-column header and two-column body?
    
    Return JSON: {"has_two_columns": bool, "has_centered_title": bool}
    """
    
    final_screenshot = get_final_screenshot(traj)
    if final_screenshot:
        vlm_res = query_vlm(prompt=vlm_prompt, image=final_screenshot)
        if vlm_res.get("success"):
            parsed = vlm_res.get("parsed", {})
            if parsed.get("has_two_columns"):
                score += 15
                feedback.append("VLM confirmed two-column layout.")
            else:
                feedback.append("VLM did NOT see two-column layout.")
        else:
            # Fallback if VLM fails: check programmatic signal again
            if has_two_columns: 
                score += 15 # Trust XML if VLM fails
                feedback.append("VLM failed, trusting XML check.")
    else:
        if has_two_columns:
            score += 15
            feedback.append("No screenshot, trusting XML check.")

    # Cleanup
    if os.path.exists(doc_temp_path):
        os.remove(doc_temp_path)

    passed = score >= 80
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }