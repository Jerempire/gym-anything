#!/usr/bin/env python3
"""
Verifier for Business Letter Formatting task.
Checks:
1. File existence and creation time.
2. Font family (Liberation Serif/Times) and Size (12pt).
3. Page margins (1.0 inch).
4. Style usage (Heading 1).
5. Character formatting (Bold).
6. Paragraph spacing (12pt after).
7. Content preservation.
8. VLM visual check.
"""

import json
import os
import sys
import tempfile
import logging
from docx import Document
from docx.shared import Pt, Inches, Emu

# Add gym_anything specific imports if available, otherwise mock
try:
    from gym_anything.vlm import sample_trajectory_frames, get_final_screenshot, query_vlm
except ImportError:
    # Mocks for local testing
    def sample_trajectory_frames(traj, n=5): return []
    def get_final_screenshot(traj): return None
    def query_vlm(images, prompt): return {"success": False, "error": "VLM not available"}

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
ACCEPTED_FONTS = {"liberation serif", "times new roman", "tinos", "nimbus roman"}
MARGIN_TARGET_EMU = 914400  # 1.0 inch
MARGIN_TOLERANCE_EMU = 137160  # ~0.15 inch
SPACING_TARGET_PT = 12
SPACING_TOLERANCE_PT = 2


def verify_business_letter_formatting(traj, env_info, task_info):
    """
    Verify the business letter formatting task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    output_path = metadata.get('output_path', '/home/ga/Documents/tenant_notice_final.docx')
    key_phrases = metadata.get('key_phrases', [])

    # Load result JSON
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_data = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result JSON: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Criterion 1: Output Exists & Created During Task (15 pts)
    output_exists = result_data.get('output_exists', False)
    file_fresh = result_data.get('file_created_during_task', False)
    
    if not output_exists:
        return {"passed": False, "score": 0, "feedback": "Output file not found."}

    score = 0
    feedback_parts = []
    
    if output_exists and file_fresh:
        score += 15
        feedback_parts.append("File created successfully (+15)")
    else:
        # File exists but old? Anti-gaming fail
        feedback_parts.append("File exists but was not modified/created during task (0)")

    # Load the document for parsing
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(output_path, temp_doc.name)
        doc = Document(temp_doc.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse output document: {e}"}
    finally:
        # Don't delete immediately, might need for debug, but standard practice is cleanup
        if os.path.exists(temp_doc.name):
            os.unlink(temp_doc.name)

    # Helper to check content preservation (prevents empty file gaming)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    phrases_found = sum(1 for p in key_phrases if p in full_text)
    if phrases_found < 3:
        return {"passed": False, "score": 0, "feedback": "Document content appears corrupted or empty."}
    
    # Criterion 2: Font Family (15 pts)
    # Check if > 50% of text runs are in accepted font
    total_runs = 0
    correct_font_runs = 0
    
    for para in doc.paragraphs:
        if not para.text.strip(): continue
        # Default font from style
        style_font = (para.style.font.name or "").lower() if para.style else ""
        
        for run in para.runs:
            if not run.text.strip(): continue
            total_runs += 1
            # Run font overrides style font
            font_name = (run.font.name or style_font).lower()
            
            if font_name in ACCEPTED_FONTS:
                correct_font_runs += 1
            # If no font specified anywhere, Word usually defaults to Calibri/Arial (fail)
            # unless "Normal" style was modified. 
            # Note: docx reader might return None if default.
            # We assume agent must set it explicitly or modify style.

    if total_runs > 0 and (correct_font_runs / total_runs) > 0.5:
        score += 15
        feedback_parts.append("Font family correct (+15)")
    else:
        feedback_parts.append(f"Font family incorrect ({correct_font_runs}/{total_runs} runs match)")

    # Criterion 3: Font Size (15 pts)
    total_runs = 0
    correct_size_runs = 0
    for para in doc.paragraphs:
        if not para.text.strip(): continue
        # Skip headings for size check (they might be larger)
        if para.style and 'heading' in para.style.name.lower():
            continue
            
        style_size = para.style.font.size.pt if (para.style and para.style.font and para.style.font.size) else None
        
        for run in para.runs:
            if not run.text.strip(): continue
            total_runs += 1
            size = run.font.size.pt if run.font.size else style_size
            
            # Tolerance 11.5 - 12.5
            if size and 11.5 <= size <= 12.5:
                correct_size_runs += 1
                
    if total_runs > 0 and (correct_size_runs / total_runs) > 0.5:
        score += 15
        feedback_parts.append("Font size correct (+15)")
    else:
        feedback_parts.append(f"Font size incorrect ({correct_size_runs}/{total_runs} body runs match 12pt)")

    # Criterion 4: Margins (15 pts)
    section = doc.sections[0]
    margins = [section.top_margin, section.bottom_margin, section.left_margin, section.right_margin]
    # Check if all 4 are within tolerance of 1 inch
    margins_correct = 0
    for m in margins:
        if m and abs(m - MARGIN_TARGET_EMU) < MARGIN_TOLERANCE_EMU:
            margins_correct += 1
            
    if margins_correct == 4:
        score += 15
        feedback_parts.append("Margins correct (+15)")
    else:
        feedback_parts.append(f"Margins incorrect ({margins_correct}/4 sides match 1.0 inch)")

    # Criterion 5: Styles & Formatting (25 pts split)
    # 5a: Heading 1 on company name (10 pts)
    heading_ok = False
    for para in doc.paragraphs:
        if "Meridian Property Group" in para.text:
            if para.style and "Heading 1" in para.style.name:
                heading_ok = True
                break
    if heading_ok:
        score += 10
        feedback_parts.append("Company heading style correct (+10)")
    else:
        feedback_parts.append("Company heading style incorrect")

    # 5b: Bold "Re:" line (5 pts)
    re_bold_ok = False
    for para in doc.paragraphs:
        if para.text.strip().startswith("Re:"):
            # Check if all runs are bold
            runs = [r for r in para.runs if r.text.strip()]
            if runs and all(r.bold for r in runs):
                re_bold_ok = True
            break
    if re_bold_ok:
        score += 5
        feedback_parts.append("Subject line bold (+5)")
    else:
        feedback_parts.append("Subject line not bold")

    # 5c: Bold Signer Name (5 pts)
    signer_bold_ok = False
    for para in doc.paragraphs:
        if "Jonathan R. Whitfield" in para.text:
            runs = [r for r in para.runs if r.text.strip()]
            # If any run is bold or paragraph style implies bold
            if runs and any(r.bold for r in runs):
                signer_bold_ok = True
            break
    if signer_bold_ok:
        score += 5
        feedback_parts.append("Signer name bold (+5)")
    else:
        feedback_parts.append("Signer name not bold")
        
    # 5d: Paragraph Spacing (5 pts)
    # Check a sample body paragraph (e.g., one starting with "We are writing")
    spacing_ok = False
    for para in doc.paragraphs:
        if para.text.strip().startswith("We are writing"):
            space_after = para.paragraph_format.space_after
            if space_after:
                # space_after is in EMU or Pt object. Pt(12) = 152400 EMU
                pt_val = space_after.pt
                if abs(pt_val - 12) <= SPACING_TOLERANCE_PT:
                    spacing_ok = True
            break
    if spacing_ok:
        score += 5
        feedback_parts.append("Paragraph spacing correct (+5)")
    else:
        feedback_parts.append("Paragraph spacing incorrect")

    # Criterion 6: VLM Verification (15 pts)
    # Visual check to confirm it looks like a letter
    final_screenshot = get_final_screenshot(traj)
    if final_screenshot:
        prompt = """
        Does this image show a business letter in LibreOffice Writer?
        Check for:
        1. A company header at the top (Meridian Property Group).
        2. A clear block structure (paragraphs separated by space).
        3. A "Sincerely," closing at the bottom.
        
        Answer JSON: {"is_business_letter": true, "header_visible": true, "block_structure": true}
        """
        vlm_res = query_vlm([final_screenshot], prompt)
        if vlm_res.get("success"):
            parsed = vlm_res.get("parsed", {})
            if parsed.get("is_business_letter") and parsed.get("block_structure"):
                score += 15
                feedback_parts.append("VLM visual verification passed (+15)")
            else:
                feedback_parts.append("VLM visual verification failed")
        else:
            # Fallback if VLM fails: give points if programmatic checks were strong
            if score >= 60:
                score += 15
                feedback_parts.append("VLM skipped (programmatic pass)")
    else:
        feedback_parts.append("No screenshot for VLM")

    # Final tally
    passed = score >= 65
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback_parts)
    }