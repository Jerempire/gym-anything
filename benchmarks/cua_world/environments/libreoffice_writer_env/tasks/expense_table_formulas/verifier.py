#!/usr/bin/env python3
"""
Verifier for expense_table_formulas task.

Checks:
1. Output file exists and is a valid DOCX.
2. Table structure matches 8 rows x 5 columns.
3. Total cells contain correct SUM values (tolerance ±2).
4. Total cells contain actual formula fields (XML check), not just typed text.
5. "Monthly Total" row is bold.
"""

import json
import os
import logging
import tempfile
import re
from zipfile import ZipFile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import python-docx if available (it should be in the env)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def check_cell_formula_xml(cell):
    """
    Check if a table cell contains a field/formula in its XML.
    LibreOffice formulas often appear as <w:fldSimple> or <w:instrText> containing ' = ' or 'SUM'.
    """
    xml = cell._element.xml
    # Common patterns for formulas in OOXML produced by LO
    # 1. w:instrText contains "=SUM" or similar
    # 2. w:fldSimple w:instr="=..."
    # 3. LO might export a simplified version, but usually fields persist.
    
    # We look for indications of a field + formula syntax
    has_field = 'w:fldChar' in xml or 'w:fldSimple' in xml
    has_formula_syntax = ('SUM' in xml or 'sum' in xml or 
                          'LEFT' in xml or 'left' in xml or 
                          'ABOVE' in xml or 'above' in xml or
                          'sum(' in xml.lower() or
                          ' = ' in xml)
                          
    return has_field and has_formula_syntax


def verify_expense_table_formulas(traj, env_info, task_info):
    """Verify the expense report table formulas and formatting."""
    
    # 1. Setup and Load
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "System error: copy_from_env missing"}

    metadata = task_info.get('metadata', {})
    expected_row_totals = metadata.get('expected_values', {}).get('row_totals', [])
    expected_col_totals = metadata.get('expected_values', {}).get('col_totals', [])
    expected_grand_total = metadata.get('expected_values', {}).get('grand_total', 0)
    tolerance = metadata.get('tolerance', 2)

    # Load result JSON
    try:
        temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_data = json.load(f)
        os.unlink(temp_json.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result metadata: {e}"}

    if not result_data.get('output_exists'):
        return {"passed": False, "score": 0, "feedback": "Output file not found (expense_report_q1_complete.docx)"}

    if not result_data.get('file_created_during_task'):
        return {"passed": False, "score": 0, "feedback": "Output file was not created/modified during the task session."}

    # Load DOCX
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(result_data['output_path'], temp_docx.name)
        if not DOCX_AVAILABLE:
            return {"passed": False, "score": 0, "feedback": "Verifier error: python-docx not installed"}
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to open DOCX file: {e}"}
    finally:
        # Keep file for a moment if needed, else cleanup
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # 2. Verify Table Structure
    if len(doc.tables) < 1:
        return {"passed": False, "score": 10, "feedback": "No table found in document."}
    
    table = doc.tables[0]
    # Expected: 8 rows (Header + 6 data + 1 total), 5 cols
    if len(table.rows) != 8 or len(table.columns) != 5:
        return {
            "passed": False, 
            "score": 10, 
            "feedback": f"Table dimensions incorrect. Expected 8x5, found {len(table.rows)}x{len(table.columns)}"
        }

    score = 15  # Base score for file + table structure

    feedback_log = []
    
    # 3. Verify Row Totals (Col index 4, Rows 1-6)
    # Expected: [3334, 10550, 16800, 5600, 9950, 1201]
    row_score = 0
    formula_count = 0
    
    for i, expected_val in enumerate(expected_row_totals):
        row_idx = i + 1  # Skip header
        cell = table.cell(row_idx, 4)
        
        # Check Value
        try:
            val_text = cell.text.strip().replace(',', '').replace('$', '')
            val = float(val_text) if val_text else 0
            if abs(val - expected_val) <= tolerance:
                row_score += 5
            else:
                feedback_log.append(f"Row {row_idx} total: Expected ~{expected_val}, got {val}")
        except ValueError:
            feedback_log.append(f"Row {row_idx} total: Not a number ('{cell.text}')")

        # Check Formula XML
        if check_cell_formula_xml(cell):
            formula_count += 2
    
    # 4. Verify Column Totals (Row index 7, Cols 1-3)
    # Expected: [21490, 11244, 14701]
    col_score = 0
    last_row = table.rows[7]
    
    for j, expected_val in enumerate(expected_col_totals):
        col_idx = j + 1 # Skip Category col
        cell = last_row.cells[col_idx]
        
        # Check Value
        try:
            val_text = cell.text.strip().replace(',', '').replace('$', '')
            val = float(val_text) if val_text else 0
            if abs(val - expected_val) <= tolerance:
                col_score += 5
            else:
                feedback_log.append(f"Col {col_idx} total: Expected ~{expected_val}, got {val}")
        except ValueError:
            feedback_log.append(f"Col {col_idx} total: Not a number ('{cell.text}')")

        # Check Formula XML
        if check_cell_formula_xml(cell):
            formula_count += 2

    # 5. Verify Grand Total (Row 7, Col 4)
    grand_total_score = 0
    gt_cell = last_row.cells[4]
    try:
        val_text = gt_cell.text.strip().replace(',', '').replace('$', '')
        val = float(val_text) if val_text else 0
        if abs(val - expected_grand_total) <= tolerance:
            grand_total_score += 10
        else:
            feedback_log.append(f"Grand Total: Expected ~{expected_grand_total}, got {val}")
    except ValueError:
        feedback_log.append(f"Grand Total: Not a number ('{gt_cell.text}')")
    
    if check_cell_formula_xml(gt_cell):
        formula_count += 2

    # 6. Verify Bold Formatting on Last Row
    # Check if all data cells in last row are bold
    bold_score = 0
    bold_cells = 0
    cells_checked = 0
    
    # Check cells 1-4 (skip label "Monthly Total" at 0 if you want, or check all)
    for cell in last_row.cells:
        # Check direct formatting or style
        is_bold = False
        # Check paragraphs
        for p in cell.paragraphs:
            if p.style and p.style.font.bold:
                is_bold = True
            for run in p.runs:
                if run.bold:
                    is_bold = True
                    break
            if is_bold: break
        
        if is_bold:
            bold_cells += 1
        cells_checked += 1

    if bold_cells >= 4: # Allow some leniency, e.g. empty cells might not show bold prop easily
        bold_score = 10
    else:
        feedback_log.append(f"Bold formatting missing on Monthly Total row (found {bold_cells}/{cells_checked} bold cells)")

    # Consolidate Score
    # File/Table: 15
    # Row Vals: 30 (5 * 6)
    # Col Vals: 15 (5 * 3)
    # Grand Val: 10
    # Formulas: 20 (2 * 10 cells)
    # Bold: 10
    # Total: 100
    
    total_score = score + row_score + col_score + grand_total_score + formula_count + bold_score
    
    # Cap score at 100
    total_score = min(100, total_score)
    
    passed = total_score >= 60 and (formula_count > 0) # Must use at least some formulas to pass
    
    final_feedback = "Task completed."
    if feedback_log:
        final_feedback = "Issues found: " + "; ".join(feedback_log)
    if formula_count == 0:
        final_feedback += " [FAIL: No formulas detected - manual entry suspected]"

    return {
        "passed": passed,
        "score": total_score,
        "feedback": final_feedback,
        "details": {
            "row_val_score": row_score,
            "col_val_score": col_score,
            "formula_score": formula_count,
            "bold_score": bold_score
        }
    }