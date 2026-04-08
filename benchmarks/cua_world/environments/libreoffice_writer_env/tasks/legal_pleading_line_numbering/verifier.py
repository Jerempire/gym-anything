#!/usr/bin/env python3
"""
Verifier for legal_pleading_line_numbering task.

Checks:
1. Output file exists and was created during task.
2. Line Numbering is enabled (CRITICAL).
3. Font is Times New Roman (or metric equivalent) and 12pt.
4. Line Spacing is Double.
5. Left Margin is ~1.5 inches.
6. Page Numbers are present in footer.
7. Section Headings use 'Heading 1' style.
"""

import sys
import os
import logging
import json
import tempfile
import shutil

# Import shared verification utils if available, or define minimal needed imports
# Assuming the environment has python-docx and lxml installed via setup_writer.sh
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError:
    pass

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_legal_pleading(traj, env_info, task_info):
    """
    Verify legal pleading formatting compliance.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    output_path = metadata.get('output_path', '/home/ga/Documents/motion_formatted.docx')
    
    # ------------------------------------------------------------------
    # 1. Fetch and Load Document
    # ------------------------------------------------------------------
    temp_dir = tempfile.mkdtemp()
    local_path = os.path.join(temp_dir, "motion_formatted.docx")
    
    try:
        # Check task_result.json first for basic file existence signals
        result_json_path = os.path.join(temp_dir, "task_result.json")
        try:
            copy_from_env("/tmp/task_result.json", result_json_path)
            with open(result_json_path, 'r') as f:
                task_result = json.load(f)
        except Exception:
            task_result = {}

        if not task_result.get('output_exists', False):
            return {
                "passed": False, 
                "score": 0, 
                "feedback": "FAILED: Output file motion_formatted.docx not found."
            }

        # Copy the actual docx file
        copy_from_env(output_path, local_path)
        
        if not os.path.exists(local_path) or os.path.getsize(local_path) == 0:
            return {
                "passed": False, 
                "score": 0, 
                "feedback": "FAILED: Output file is empty or could not be downloaded."
            }
            
        try:
            doc = Document(local_path)
        except Exception as e:
            return {
                "passed": False, 
                "score": 0, 
                "feedback": f"FAILED: Could not parse DOCX file. Corrupt? Error: {str(e)}"
            }

        # ------------------------------------------------------------------
        # Verification Logic
        # ------------------------------------------------------------------
        score = 0
        feedback = []
        
        # Criterion 1: Line Numbering (20 pts)
        # Check for <w:lnNumType> in section properties
        has_line_numbering = False
        try:
            for section in doc.sections:
                sectPr = section._sectPr
                lnNumType = sectPr.find(qn('w:lnNumType'))
                if lnNumType is not None:
                    has_line_numbering = True
                    break
        except Exception as e:
            logger.warning(f"Error checking line numbering: {e}")
            
        if has_line_numbering:
            score += 20
            feedback.append("Line numbering enabled (20/20)")
        else:
            feedback.append("MISSING: Line numbering not detected (0/20)")

        # Criterion 2: Margins (Left 1.5", others 1.0") (17 pts)
        # 1 inch = 914400 EMU
        # 1.5 inch = 1371600 EMU
        # Tolerance ~ 0.1 inch = 91440 EMU
        TOLERANCE = 100000
        EXPECTED_LEFT = 1371600
        EXPECTED_OTHER = 914400
        
        try:
            section = doc.sections[0]
            left = section.left_margin
            right = section.right_margin
            
            # Left margin check
            if left is not None and abs(left - EXPECTED_LEFT) < TOLERANCE:
                score += 10
                feedback.append("Left margin correct (1.5\") (10/10)")
            else:
                feedback.append(f"Left margin incorrect (expected 1.5\", got {left/914400:.2f}\")")
                
            # Right/Top/Bottom check (checking Right as proxy for others to keep it simple)
            if right is not None and abs(right - EXPECTED_OTHER) < TOLERANCE:
                score += 7
                feedback.append("Other margins correct (1.0\") (7/7)")
            else:
                feedback.append("Other margins incorrect")
        except Exception:
            feedback.append("Could not verify margins")

        # Criterion 3: Font Face and Size (25 pts)
        # Check body paragraphs (approximate check of >50% of content)
        # Expected: Times New Roman or Liberation Serif, 12pt
        compliant_runs = 0
        total_runs = 0
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            # Skip likely headings
            if para.style.name.startswith("Heading"):
                continue
                
            for run in para.runs:
                if not run.text.strip():
                    continue
                total_runs += 1
                
                # Check Font Name
                font_name = run.font.name
                if not font_name and para.style.font:
                    font_name = para.style.font.name
                if not font_name and doc.styles['Normal'].font:
                    font_name = doc.styles['Normal'].font.name
                
                font_name = (font_name or "").lower()
                font_ok = any(n in font_name for n in ["times", "liberation serif"])
                
                # Check Font Size
                font_size = run.font.size
                if not font_size and para.style.font:
                    font_size = para.style.font.size
                
                # 12pt = 152400 EMU or 12.0 Pt
                size_ok = False
                if font_size:
                    if hasattr(font_size, 'pt'):
                        size_ok = abs(font_size.pt - 12.0) < 0.5
                    else:
                        size_ok = abs(font_size - 152400) < 6350 # 0.5pt tolerance
                
                if font_ok: 
                    compliant_runs += 0.5
                if size_ok: 
                    compliant_runs += 0.5

        font_score_ratio = compliant_runs / total_runs if total_runs > 0 else 0
        font_points = int(font_score_ratio * 25)
        score += font_points
        feedback.append(f"Font compliance (Times/12pt): {int(font_score_ratio*100)}% ({font_points}/25)")

        # Criterion 4: Line Spacing (Double) (15 pts)
        double_spaced_paras = 0
        body_paras = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            if para.style.name.startswith("Heading"):
                continue
            
            body_paras += 1
            pf = para.paragraph_format
            # Double spacing can be:
            # 1. line_spacing_rule = WD_LINE_SPACING.DOUBLE (2)
            # 2. line_spacing = 2.0
            is_double = False
            if pf.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
                is_double = True
            elif pf.line_spacing == 2.0:
                is_double = True
            
            if is_double:
                double_spaced_paras += 1
        
        spacing_ratio = double_spaced_paras / body_paras if body_paras > 0 else 0
        if spacing_ratio > 0.6:
            score += 15
            feedback.append("Double spacing applied correctly (15/15)")
        else:
            feedback.append(f"Double spacing missing or inconsistent ({int(spacing_ratio*100)}% paragraphs)")

        # Criterion 5: Page Numbers in Footer (8 pts)
        has_page_numbers = False
        try:
            for section in doc.sections:
                if section.footer:
                    # Check XML for 'w:fldChar' with 'PAGE' or simple text check
                    for p in section.footer.paragraphs:
                        if 'PAGE' in p._element.xml or any(run.text.strip().isdigit() for run in p.runs):
                            has_page_numbers = True
                            break
                if has_page_numbers: break
        except Exception:
            pass
            
        if has_page_numbers:
            score += 8
            feedback.append("Page numbers detected in footer (8/8)")
        else:
            feedback.append("Page numbers missing from footer (0/8)")

        # Criterion 6: Section Headings (Heading 1) (7 pts)
        expected_headings = [
            "MEMORANDUM OF POINTS AND AUTHORITIES",
            "STATEMENT OF FACTS",
            "ARGUMENT",
            "CONCLUSION"
        ]
        headings_matched = 0
        for expected in expected_headings:
            found = False
            for para in doc.paragraphs:
                if expected in para.text:
                    if para.style and "Heading 1" in para.style.name:
                        found = True
                    break
            if found:
                headings_matched += 1
        
        if headings_matched >= 3:
            score += 7
            feedback.append(f"Section headings formatted correctly ({headings_matched}/{len(expected_headings)}) (7/7)")
        else:
            feedback.append(f"Section headings missing 'Heading 1' style ({headings_matched}/{len(expected_headings)})")

        # Criterion 7: Original Draft Preserved (8 pts)
        if task_result.get('source_preserved', False):
            score += 8
            feedback.append("Original draft preserved (8/8)")
        else:
            feedback.append("Original draft was modified or deleted (0/8)")

        # Final Score
        passed = (score >= 60) and has_line_numbering
        
        return {
            "passed": passed,
            "score": score,
            "feedback": " | ".join(feedback)
        }

    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Verification error: {str(e)}"}
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)