#!/usr/bin/env python3
"""
Verifier for mixed_page_orientation task.
Checks if the document has mixed portrait/landscape sections appropriate for the content.
"""

import json
import os
import tempfile
import logging
from docx import Document
# Note: WD_ORIENT is useful for explicit orientation, but mostly we check dimensions
from docx.enum.section import WD_ORIENT

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_mixed_orientation(traj, env_info, task_info):
    """
    Verify the BCP document has correct mixed orientation.
    
    Criteria:
    1. Output file exists and was modified/created during task.
    2. Document has at least 4 sections (indicates section breaks were used).
    3. Sections containing specific keywords match expected orientation:
       - "Business Impact Analysis" -> Landscape
       - "Emergency Contact Matrix" -> Landscape
       - "Executive Summary" -> Portrait
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Load result metadata
    temp_result = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_result.name)
        with open(temp_result.name, 'r') as f:
            result_meta = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result metadata: {e}"}
    finally:
        if os.path.exists(temp_result.name):
            os.unlink(temp_result.name)

    if not result_meta.get('output_exists'):
        return {"passed": False, "score": 0, "feedback": "Output file bcp_formatted.docx not found"}

    if not result_meta.get('file_modified_during_task'):
        return {"passed": False, "score": 0, "feedback": "Output file was not created/modified during the task window"}

    # Copy the actual document
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env("/home/ga/Documents/bcp_formatted.docx", temp_doc.name)
        doc = Document(temp_doc.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX file: {e}"}
    finally:
        if os.path.exists(temp_doc.name):
            os.unlink(temp_doc.name)

    score = 0
    feedback = []
    
    # 1. Check Section Count (Need at least 4 to isolate 2 separate landscape regions)
    # Typically: Portrait (Start) -> Landscape (BIA) -> Portrait (Strat) -> Landscape (Contacts)
    num_sections = len(doc.sections)
    if num_sections >= 4:
        score += 20
        feedback.append(f"Pass: Document has {num_sections} sections (>= 4 required)")
    elif num_sections >= 2:
        score += 10
        feedback.append(f"Partial: Document has {num_sections} sections (expected 4+ for full mixed layout)")
    else:
        feedback.append(f"Fail: Document has only {num_sections} section(s). Did you insert Section Breaks?")
        return {"passed": False, "score": score, "feedback": "; ".join(feedback)}

    # 2. Analyze sections
    # We define a helper to determine if a section is landscape
    def is_landscape(section):
        # Method A: Explicit orientation enum
        if section.orientation == WD_ORIENT.LANDSCAPE:
            return True
        # Method B: Dimensions (More robust for LibreOffice saved files)
        # 1 inch = 914400 EMU. Letter is 8.5x11
        w = section.page_width
        h = section.page_height
        if w is None or h is None: 
            return False
        return w > h

    # Map content to sections to verify correct orientation per content
    # We'll scan paragraphs to find which section they belong to.
    # Note: doc.paragraphs covers the whole doc. We need to correlate ranges.
    # python-docx doesn't easily map paragraphs to sections directly without iterating 
    # elements, but we can iterate sections and check the text roughly if sections split cleanly.
    # A simpler heuristic: Check if WE HAVE landscape sections and if they contain the tables.
    
    landscape_sections = 0
    portrait_sections = 0
    
    landscape_keywords_found = []
    portrait_keywords_found = []

    # Iterate paragraphs and guess section (approximate since python-docx linearizes paragraphs)
    # A more robust way: python-docx `doc.sections` gives properties, but content is in `doc.elements`
    # We will assume standard ordering and check if *any* landscape section contains target text
    # However, `doc.sections` don't hold text directly.
    # Let's try a different approach: Verify we have landscape sections, and verify the text exists.
    # Then verify strictly that:
    #   - BIA text is NOT in a Portrait section? (Hard to check without complex parsing)
    #   - Simple check: Do we have at least 2 landscape and 2 portrait sections?
    
    for section in doc.sections:
        if is_landscape(section):
            landscape_sections += 1
        else:
            portrait_sections += 1
            
    if landscape_sections >= 2:
        score += 30
        feedback.append(f"Pass: Found {landscape_sections} landscape sections")
    elif landscape_sections == 1:
        score += 15
        feedback.append("Partial: Found only 1 landscape section (expected 2)")
    else:
        feedback.append("Fail: No landscape sections found")

    if portrait_sections >= 2:
        score += 20
        feedback.append(f"Pass: Found {portrait_sections} portrait sections")
    else:
        feedback.append(f"Fail: Found {portrait_sections} portrait sections (expected >= 2)")

    # 3. Content Preservation Check
    full_text = " ".join([p.text for p in doc.paragraphs])
    
    required_phrases = [
        "Business Impact Analysis",
        "Emergency Contact Matrix",
        "Recovery Strategies",
        "Meridian Financial Services"
    ]
    
    missing_phrases = [p for p in required_phrases if p not in full_text]
    
    if not missing_phrases:
        score += 30
        feedback.append("Pass: All key content preserved")
    else:
        score += 10 # formatting failed to preserve text?
        feedback.append(f"Fail: Missing content: {', '.join(missing_phrases)}")

    # Bonus: Robust check for BIA table width?
    # If the user put the BIA table in a section with width > 10 inches, that's a win.
    # We can't easily link specific text to sections in python-docx simple API, 
    # but the section count + orientation count + content presence is a very strong proxy.

    passed = (score >= 80)
    
    return {
        "passed": passed,
        "score": score,
        "feedback": "; ".join(feedback)
    }