#!/usr/bin/env python3
"""
Verifier for HVAC Index Creation task.

Verifies:
1. Output file exists and is a valid DOCX.
2. Content integrity (original sections preserved).
3. Index Entry fields (XE) exist for specific terms.
4. Index field (INDEX) exists (the compiled table).
5. File was modified during the task.
"""

import json
import os
import sys
import tempfile
import logging
import re
from typing import List, Dict, Set

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    logger.error("python-docx not installed")
    DOCX_AVAILABLE = False


def extract_field_codes(doc) -> List[str]:
    """
    Extract field codes from a python-docx Document object.
    
    Word stores fields in <w:instrText> elements within <w:fldChar> blocks.
    We need to scan the XML for these instructions.
    """
    field_codes = []
    
    # Namespace map for XML finding
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    try:
        # Iterate over all paragraphs in the document body
        for para in doc.paragraphs:
            # Access the underlying XML element
            # instrText elements contain the field code string (e.g., ' XE "Compressor" ')
            instr_texts = para._element.findall('.//w:instrText', nsmap)
            for instr in instr_texts:
                if instr.text:
                    field_codes.append(instr.text)
    except Exception as e:
        logger.error(f"Error parsing XML for fields: {e}")
        
    return field_codes


def check_content_preservation(doc, required_headings: List[str]) -> int:
    """Check if original section headings are still present."""
    found_headings = 0
    doc_text = "\n".join([p.text for p in doc.paragraphs])
    
    for heading in required_headings:
        # Simple text check usually suffices for heading preservation
        if heading in doc_text:
            found_headings += 1
            
    return found_headings


def verify_hvac_index_creation(traj, env_info, task_info):
    """
    Main verification function.
    """
    # 1. Setup and Environment Check
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "System error: copy_from_env not available"}
    
    if not DOCX_AVAILABLE:
        return {"passed": False, "score": 0, "feedback": "System error: python-docx not available"}

    # Load metadata
    metadata = task_info.get('metadata', {})
    required_terms = [t.lower() for t in metadata.get('required_terms', [])]
    required_headings = metadata.get('required_section_headings', [])
    output_path = metadata.get('output_path', "/home/ga/Documents/hvac_manual_indexed.docx")

    # Load result JSON
    task_result = {}
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            task_result = json.load(f)
    except Exception:
        pass  # Result might not exist if script failed, we'll check file existence next
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # 2. Retrieve Document
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(output_path, temp_docx.name)
        file_exists = True
        file_size = os.path.getsize(temp_docx.name)
    except Exception as e:
        logger.warning(f"Could not copy output file: {e}")
        file_exists = False
        file_size = 0
    
    # 3. Scoring Logic
    score = 0
    feedback = []
    
    # Criterion 1: File Existence & Validity (10 pts)
    if not file_exists:
        return {"passed": False, "score": 0, "feedback": "Output file not found."}
    
    if file_size < 10000: # ~10KB minimum for this doc
        return {"passed": False, "score": 0, "feedback": "Output file is empty or corrupted."}
    
    try:
        doc = Document(temp_docx.name)
        score += 10
        feedback.append("Valid DOCX file submitted.")
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to parse DOCX file: {e}"}

    # Criterion 2: Content Preservation (10 pts)
    # Prevent creating a new empty doc with just an index
    headings_found = check_content_preservation(doc, required_headings)
    if headings_found >= len(required_headings) - 2: # Allow missing 1-2
        score += 10
        feedback.append("Original content preserved.")
    else:
        feedback.append(f"Content missing (found {headings_found}/{len(required_headings)} headings).")

    # Parse Fields
    field_codes = extract_field_codes(doc)
    logger.info(f"Found {len(field_codes)} field codes")

    # Criterion 3: Check for XE (Index Entry) fields (30 pts)
    # XE fields look like: ' XE "Compressor" '
    found_terms = set()
    xe_count = 0
    
    for code in field_codes:
        if 'XE' in code:
            # Extract content inside quotes
            match = re.search(r'XE\s+"([^"]+)"', code)
            if match:
                term = match.group(1).lower()
                # Check if it matches one of our required terms (fuzzy match)
                for req in required_terms:
                    if req in term:
                        found_terms.add(req)
                        xe_count += 1
                        break
    
    # Scoring for terms
    term_score = 0
    if len(found_terms) >= len(required_terms) * 0.8: # 8+ terms
        term_score = 30
    elif len(found_terms) >= len(required_terms) * 0.5: # 5+ terms
        term_score = 15
    elif len(found_terms) >= 1:
        term_score = 5
    
    score += term_score
    feedback.append(f"Marked {len(found_terms)}/{len(required_terms)} required terms.")

    # Criterion 4: Check for INDEX field (30 pts)
    # INDEX fields look like: ' INDEX \e "	" \h "A" \c "2" '
    has_index = any('INDEX' in code for code in field_codes)
    
    if has_index:
        score += 30
        feedback.append("Alphabetical Index inserted.")
    else:
        feedback.append("Alphabetical Index MISSING (INDEX field not found).")

    # Criterion 5: Index Position (10 pts)
    # Check if INDEX field is in the last 20% of paragraphs or after the last heading
    if has_index:
        # Simple heuristic: is it near the end?
        total_paras = len(doc.paragraphs)
        index_pos = -1
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        for i, para in enumerate(doc.paragraphs):
            if para._element.findall('.//w:instrText', nsmap):
                 if any('INDEX' in t.text for t in para._element.findall('.//w:instrText', nsmap) if t.text):
                     index_pos = i
                     break
        
        if index_pos > -1 and index_pos > total_paras * 0.7:
             score += 10
             feedback.append("Index positioned at end of document.")
        else:
             feedback.append("Index found but not at the end.")

    # Criterion 6: Anti-Gaming / Timestamp (10 pts)
    is_new = task_result.get("file_created_during_task", False)
    if is_new:
        score += 10
        feedback.append("File modified during task.")
    else:
        feedback.append("File modification time suspiciously old.")

    # Cleanup
    if os.path.exists(temp_docx.name):
        os.unlink(temp_docx.name)

    # Final Result
    passed = (score >= 60) and has_index and (len(found_terms) >= 5)
    
    return {
        "passed": passed,
        "score": score,
        "feedback": " ".join(feedback)
    }