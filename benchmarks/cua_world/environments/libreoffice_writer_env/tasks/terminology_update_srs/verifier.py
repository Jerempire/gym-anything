#!/usr/bin/env python3
"""
Verifier for terminology_update_srs task.

Checks:
1. Output file exists and was created during task.
2. Original file remains untouched (or at least exists).
3. "DataSync Pro", "DS Pro", "DataSync", "Nextera Solutions" are GONE.
4. "CloudBridge Enterprise", "CloudBridge Inc." are PRESENT.
5. Version is 3.0 (Title page only ideally, but we check text presence).
6. Date is January 15, 2025.
"""

import json
import os
import tempfile
import logging
import re

# Import utility for docx parsing
# This assumes the environment has python-docx installed (as per setup_writer.sh)
try:
    from docx import Document
except ImportError:
    # Fallback/Mock for local testing if needed, but in env it will exist
    Document = None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_terminology_update(traj, env_info, task_info):
    """
    Verify the SRS rebranding task.
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    expected_output_path = metadata.get('expected_output_path', '/home/ga/Documents/CloudBridge_SRS_v3.0.docx')

    # Load basic result info from export_result.sh
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_data = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load task result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Scoring Variables
    score = 0
    feedback = []
    
    # 1. Check File Existence (20 pts)
    if not result_data.get('output_exists', False):
        return {"passed": False, "score": 0, "feedback": "Output file CloudBridge_SRS_v3.0.docx not found."}
    
    if not result_data.get('output_created_during_task', False):
        feedback.append("Warning: Output file timestamp is older than task start.")
        # We might penalize or fail, but let's allow it if content is perfect (maybe clock skew), 
        # though usually this is anti-gaming.
        # Let's deduct points.
        score += 10 # Half points for existence
    else:
        score += 20
        feedback.append("Output file created successfully.")

    # 2. Check Original File Preservation (10 pts)
    if result_data.get('original_exists', False):
        score += 10
        feedback.append("Original file preserved.")
    else:
        feedback.append("Original file missing.")

    # 3. Content Analysis (70 pts)
    # Copy the docx to host
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(expected_output_path, temp_docx.name)
        
        if Document is None:
             return {"passed": False, "score": score, "feedback": "python-docx not installed in verifier environment."}

        doc = Document(temp_docx.name)
        full_text = []
        for p in doc.paragraphs:
            full_text.append(p.text)
        # Also check tables if any (none in setup script, but good practice)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        full_text_str = "\n".join(full_text)
        
        # Check Replacements
        
        # A. Forbidden Terms (Old Brand) - 30 pts
        # Terms: DataSync Pro, DS Pro, DataSync, Nextera Solutions
        forbidden = ["DataSync Pro", "DS Pro", "Nextera Solutions"]
        # Note: "DataSync" handles both "DataSync" and "DataSync Pro", but checking specifics helps feedback
        
        # We check "DataSync" case-insensitively or sensitively? Description implied sensitive.
        # Real world rebranding is usually sensitive but "datasync" likely doesn't exist.
        
        errors_found = 0
        if "DataSync" in full_text_str:
            count = full_text_str.count("DataSync")
            feedback.append(f"Found {count} instances of 'DataSync' remaining.")
            errors_found += 1
        
        if "DS Pro" in full_text_str:
            count = full_text_str.count("DS Pro")
            feedback.append(f"Found {count} instances of 'DS Pro' remaining.")
            errors_found += 1
            
        if "Nextera Solutions" in full_text_str:
            count = full_text_str.count("Nextera Solutions")
            feedback.append(f"Found {count} instances of 'Nextera Solutions' remaining.")
            errors_found += 1

        if errors_found == 0:
            score += 30
            feedback.append("Old branding completely removed.")
        else:
            # Partial credit?
            score += max(0, 30 - (errors_found * 10))

        # B. Required Terms (New Brand) - 20 pts
        required = ["CloudBridge Enterprise", "CloudBridge Inc."]
        found_count = 0
        for term in required:
            if term in full_text_str:
                found_count += 1
            else:
                feedback.append(f"Missing new term: '{term}'")
        
        if found_count == len(required):
            score += 20
            feedback.append("New branding terms present.")
        else:
            score += (found_count * 10)

        # C. Version and Date (20 pts)
        # Version 3.0
        if "Version: 3.0" in full_text_str:
            score += 10
            feedback.append("Version updated to 3.0.")
        else:
            feedback.append("Version 3.0 not found on title page format.")

        # Date January 15, 2025
        if "January 15, 2025" in full_text_str:
            score += 10
            feedback.append("Date updated to January 15, 2025.")
        else:
            feedback.append("Target date not found.")

        # D. Safety Check (Bonus/Negative)
        # Check if Section 2.3 became "Section 3.0" (bad regex replace)
        # We look for "2.3 Assumptions" - if it's gone, they might have broke it.
        if "2.3 Assumptions" not in full_text_str:
             # Check if it became "3.0 Assumptions"
             if "3.0 Assumptions" in full_text_str:
                 score -= 10
                 feedback.append("PENALTY: Replaced section number 2.3 with 3.0 incorrectly.")
             else:
                 # Maybe they just deleted it?
                 pass 

    except Exception as e:
        feedback.append(f"Error parsing document: {str(e)}")
    finally:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    # Normalize score
    score = max(0, min(100, score))
    passed = score >= 70

    return {
        "passed": passed,
        "score": score,
        "feedback": " ".join(feedback)
    }