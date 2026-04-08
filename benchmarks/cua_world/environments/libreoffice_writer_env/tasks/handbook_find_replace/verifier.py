#!/usr/bin/env python3
"""
Verifier for handbook_find_replace task.
Checks if specific text patterns were replaced correctly in the output document.
"""

import json
import os
import re
import logging
import tempfile
import shutil

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_handbook_corrections(traj, env_info, task_info):
    """
    Verify that the employee handbook was corrected according to instructions.
    
    Criteria:
    1. Output file exists and is a valid DOCX.
    2. File was created/modified during the task window.
    3. Document content is preserved (anti-gaming: didn't just delete text).
    4. "Greenfield industries" -> "Greenfield Industries"
    5. Double spaces -> Single spaces
    6. "@greenfield-ind.com" -> "@greenfieldind.com"
    7. "dept." -> "Department"
    8. "2019" -> "2024"
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "System error: copy_from_env not available"}

    metadata = task_info.get('metadata', {})
    checks = metadata.get('checks', {})
    
    # 1. Check metadata from export script
    try:
        temp_meta = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
        copy_from_env("/tmp/task_result.json", temp_meta.name)
        with open(temp_meta.name, 'r') as f:
            task_result = json.load(f)
        os.unlink(temp_meta.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load task result metadata: {str(e)}"}

    if not task_result.get("output_exists", False):
        return {"passed": False, "score": 0, "feedback": "Output file 'employee_handbook_corrected.docx' not found."}
        
    if not task_result.get("file_modified_during_task", False):
        return {"passed": False, "score": 0, "feedback": "Output file was not created or modified during the task session."}

    # 2. Copy and parse the document
    temp_dir = tempfile.mkdtemp()
    local_docx = os.path.join(temp_dir, "corrected.docx")
    
    try:
        copy_from_env(metadata.get("output_path"), local_docx)
        
        # Import here to avoid dependency issues if environment is different (though usually verifier runs in env with dependencies)
        try:
            from docx import Document
            doc = Document(local_docx)
        except ImportError:
            return {"passed": False, "score": 0, "feedback": "System error: python-docx not installed in verifier environment"}
        except Exception as e:
            return {"passed": False, "score": 0, "feedback": f"Failed to parse output DOCX file: {str(e)}"}

        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also check tables if any (though this task uses paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text_content = "\n".join(full_text)
        
        # ANTI-GAMING: Check if content looks preserved
        if len(doc.paragraphs) < 15: # Original has ~20-25 paragraphs
            return {"passed": False, "score": 0, "feedback": "Document content appears significantly truncated. Do not delete content."}

        score = 0
        feedback_parts = []
        
        # --- Check 1: Output file validity (already passed if we are here) ---
        score += 10
        feedback_parts.append("File valid (10/10)")
        
        # --- Check 2: Content preservation ---
        score += 8
        feedback_parts.append("Structure preserved (8/8)")

        # --- Check 3: Double spaces ---
        # Count "  "
        double_space_count = text_content.count("  ")
        if double_space_count == 0:
            score += 16
            feedback_parts.append("No double spaces (16/16)")
        else:
            feedback_parts.append(f"Found {double_space_count} double spaces left (0/16)")

        # --- Check 4: Company Name Casing ---
        # Bad: "Greenfield industries"
        # Good: "Greenfield Industries"
        bad_case = text_content.count("Greenfield industries")
        good_case = text_content.count("Greenfield Industries")
        
        if bad_case == 0 and good_case >= 8:
            score += 16
            feedback_parts.append("Company name casing fixed (16/16)")
        elif bad_case > 0:
            feedback_parts.append(f"Found {bad_case} lowercase 'industries' (0/16)")
        else:
            # If bad_case is 0 but good_case is low, they might have deleted the text
            feedback_parts.append("Company name missing from document (0/16)")

        # --- Check 5: Email Domain ---
        # Bad: "@greenfield-ind.com"
        # Good: "@greenfieldind.com"
        bad_email = text_content.count("@greenfield-ind.com")
        good_email = text_content.count("@greenfieldind.com")
        
        if bad_email == 0 and good_email >= 5:
            score += 16
            feedback_parts.append("Email domain updated (16/16)")
        elif bad_email > 0:
            feedback_parts.append(f"Found {bad_email} old email domains (0/16)")
        else:
            feedback_parts.append("Emails missing from document (0/16)")

        # --- Check 6: Abbreviations (Case Insensitive) ---
        # Bad: "dept." (regex needed for case insensitive check or just lower())
        # Good: "Department"
        bad_abbrev_count = len(re.findall(r"dept\.", text_content, re.IGNORECASE))
        good_full_count = text_content.count("Department")
        
        if bad_abbrev_count == 0 and good_full_count >= 7:
            score += 16
            feedback_parts.append("Abbreviations expanded (16/16)")
        elif bad_abbrev_count > 0:
            feedback_parts.append(f"Found {bad_abbrev_count} abbreviations (0/16)")
        else:
            feedback_parts.append("Departments missing from document (0/16)")

        # --- Check 7: Year Update ---
        # Bad: "2019"
        # Good: "2024"
        bad_year = text_content.count("2019")
        good_year = text_content.count("2024")
        
        if bad_year == 0 and good_year >= 4:
            score += 18
            feedback_parts.append("Year updated (18/18)")
        elif bad_year > 0:
            feedback_parts.append(f"Found {bad_year} instances of 2019 (0/18)")
        else:
            feedback_parts.append("Years missing from document (0/18)")

        passed = score >= 66
        
        return {
            "passed": passed,
            "score": score,
            "feedback": "; ".join(feedback_parts)
        }
        
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Verification failed with error: {str(e)}"}
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)