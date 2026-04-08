#!/usr/bin/env python3
"""
Verifier for regulatory_page_numbering task.
Checks section breaks, page numbering formats, and headers/footers using XML parsing.
"""

import os
import sys
import logging
import json
import tempfile
import shutil

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import verifier utils
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '../../', 'utils'))
from writer_verification_utils import copy_and_parse_document, cleanup_verification_temp

# OOXML namespace
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def get_page_num_type(section):
    """Extract pgNumType attributes from a section's sectPr element."""
    # section._sectPr is the XML element
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(f'{{{W_NS}}}pgNumType')
    if pg_num_type is None:
        return None, None
    fmt = pg_num_type.get(f'{{{W_NS}}}fmt')
    start = pg_num_type.get(f'{{{W_NS}}}start')
    return fmt, start


def footer_has_page_field(section):
    """Check if any footer in the section contains a PAGE field."""
    try:
        footer = section.footer
        if footer is None:
            return False
        for para in footer.paragraphs:
            xml_str = para._element.xml
            # Check for simple field or complex field char
            if 'PAGE' in xml_str and ('w:fldChar' in xml_str or 'w:instrText' in xml_str):
                return True
            # Fallback: check for page number in text (less reliable but catches manual entry)
            if para.text.strip().isdigit():
                return True
    except Exception:
        pass
    return False


def footer_is_empty(section):
    """Check if footer has no visible content."""
    try:
        footer = section.footer
        if footer is None:
            return True
        all_text = ''.join(p.text for p in footer.paragraphs).strip()
        # Also check XML for hidden fields
        for para in footer.paragraphs:
            xml_str = para._element.xml
            if 'PAGE' in xml_str or 'w:fldChar' in xml_str:
                return False
        return all_text == ''
    except Exception:
        return True


def header_text(section):
    """Get concatenated header text from a section."""
    try:
        header = section.header
        if header is None:
            return ''
        return ' '.join(p.text for p in header.paragraphs).strip()
    except Exception:
        return ''


def verify_regulatory_page_numbering(traj, env_info, task_info):
    """
    Main verification function for regulatory_page_numbering.
    """
    # 1. Setup access to file in container
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    metadata = task_info.get('metadata', {})
    output_path = metadata.get('output_path', '/home/ga/Documents/clinical_overview_formatted.docx')
    
    score = 0
    max_score = 100
    feedback = []

    # 2. Check metadata from export_result.sh (anti-gaming)
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            task_result = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to read task result: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    # Criterion 1: Output file exists (8 pts)
    if not task_result.get('output_exists', False):
        return {
            "passed": False,
            "score": 0,
            "feedback": "Output file not found at expected location."
        }
    score += 8
    feedback.append("Output file exists")

    # Anti-gaming: File created/modified during task
    if not task_result.get('file_created_during_task', False):
        feedback.append("WARNING: File was not modified during task timeframe.")
        # We don't fail immediately but this is suspicious
    
    # 3. Parse Document
    success, doc, err, temp_dir = copy_and_parse_document(output_path, copy_from_env, 'docx')
    if not success:
        return {
            "passed": False,
            "score": score,
            "feedback": f"Failed to parse document: {err}"
        }

    try:
        # Check file size/validity
        host_file = os.path.join(temp_dir, "result.docx")
        if os.path.getsize(host_file) < 10000:
            return {
                "passed": False,
                "score": score, 
                "feedback": "File too small, likely empty or corrupted."
            }

        # Verify content preservation (simple check)
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        if "Product Development Rationale" not in full_text:
            return {
                "passed": False,
                "score": score,
                "feedback": "Document content appears to be overwritten or corrupted."
            }

        # ---- Criterion 2: Exactly 3 sections (15 pts) ----
        num_sections = len(doc.sections)
        if num_sections == 3:
            score += 15
            feedback.append("Document has exactly 3 sections")
        elif num_sections >= 2:
            score += 5
            feedback.append(f"Document has {num_sections} sections (expected 3)")
        else:
            feedback.append(f"Document has only {num_sections} section (expected 3)")

        if num_sections < 2:
            return {"passed": False, "score": score, "feedback": " | ".join(feedback)}

        # ---- Criterion 3: Section breaks are New Page type (10 pts) ----
        from docx.enum.section import WD_SECTION
        new_page_count = 0
        # Check start_type for sections 1 and 2 (indices 1 and 2 in doc.sections)
        # Note: Section 0 start type is usually irrelevant, check transition to 1 and 2
        for i in range(1, num_sections):
            start_type = doc.sections[i].start_type
            if start_type == WD_SECTION.NEW_PAGE or start_type is None:
                new_page_count += 1
            elif start_type in (WD_SECTION.ODD_PAGE, WD_SECTION.EVEN_PAGE):
                new_page_count += 1
        
        expected_breaks = min(num_sections - 1, 2)
        if new_page_count >= expected_breaks:
            score += 10
            feedback.append("Section breaks are correct (New Page)")
        elif new_page_count > 0:
            score += 5
            feedback.append("Some section breaks correct")
        else:
            feedback.append("Sections do not start on new pages")

        # ---- Criterion 4: Section 1 footer has no page number (14 pts) ----
        section1 = doc.sections[0]
        if footer_is_empty(section1):
            score += 14
            feedback.append("Section 1 footer is empty (Correct)")
        elif not footer_has_page_field(section1):
            score += 10
            feedback.append("Section 1 footer text present but no page number (Acceptable)")
        else:
            feedback.append("Section 1 footer contains page number (Incorrect)")

        # ---- Criterion 5: Section 2 uses lowercase Roman numerals (16 pts) ----
        if num_sections >= 2:
            section2 = doc.sections[1]
            fmt2, start2 = get_page_num_type(section2)
            
            # 'lowerRoman' is standard, check variants just in case
            if fmt2 == 'lowerRoman':
                score += 16
                feedback.append("Section 2 uses lowercase Roman numerals")
            elif fmt2 == 'upperRoman':
                score += 8
                feedback.append("Section 2 uses uppercase Roman numerals (Partial)")
            elif footer_has_page_field(section2):
                score += 4
                feedback.append(f"Section 2 has page numbers but format is {fmt2}")
            else:
                feedback.append("Section 2 has no page numbering configured")

        # ---- Criterion 6: Section 3 uses Arabic numerals starting at 1 (16 pts) ----
        if num_sections >= 3:
            section3 = doc.sections[2]
            fmt3, start3 = get_page_num_type(section3)
            
            is_arabic = fmt3 in ('decimal', None) # None often defaults to decimal
            is_start_one = (start3 == '1')
            
            if is_arabic and is_start_one:
                score += 16
                feedback.append("Section 3 uses Arabic numerals starting at 1")
            elif is_arabic and not is_start_one:
                score += 8
                feedback.append(f"Section 3 uses Arabic numerals but starts at {start3}")
            elif not is_arabic and is_start_one:
                score += 8
                feedback.append(f"Section 3 starts at 1 but format is {fmt3}")
            else:
                if footer_has_page_field(section3):
                    score += 4
                    feedback.append("Section 3 has page numbers (incorrect format/restart)")
                else:
                    feedback.append("Section 3 has no page numbering")

        # ---- Criterion 7 & 8: Headers (20 pts) ----
        required_text_parts = ["sorvimab", "215847"]
        
        # Section 2 Header
        if num_sections >= 2:
            h2 = header_text(doc.sections[1]).lower()
            if all(part in h2 for part in required_text_parts):
                score += 10
                feedback.append("Section 2 header correct")
            elif any(part in h2 for part in required_text_parts):
                score += 5
                feedback.append("Section 2 header partial match")
            else:
                feedback.append("Section 2 header missing or incorrect")
        
        # Section 3 Header
        if num_sections >= 3:
            h3 = header_text(doc.sections[2]).lower()
            if all(part in h3 for part in required_text_parts):
                score += 10
                feedback.append("Section 3 header correct")
            elif any(part in h3 for part in required_text_parts):
                score += 5
                feedback.append("Section 3 header partial match")
            else:
                feedback.append("Section 3 header missing or incorrect")

        # Section 1 Header Check (Bonus/Sanity: Should be empty)
        h1 = header_text(doc.sections[0]).strip()
        if len(h1) > 10: # Allow small noise
            score = max(0, score - 5) # Penalty for header on title page
            feedback.append("Penalty: Header found on title page")

    except Exception as e:
        logger.error(f"Verification error: {e}")
        return {"passed": False, "score": score, "feedback": f"Error during verification: {str(e)}"}
    finally:
        cleanup_verification_temp(temp_dir)

    passed = score >= 60
    return {
        "passed": passed,
        "score": score,
        "feedback": " | ".join(feedback)
    }