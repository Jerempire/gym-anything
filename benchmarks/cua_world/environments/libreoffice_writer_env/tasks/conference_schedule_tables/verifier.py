#!/usr/bin/env python3
"""
Verifier for conference_schedule_tables task.
Checks:
1. File existence and valid DOCX format.
2. Page orientation (Landscape) and margins.
3. Table structure (4 columns).
4. Merged cells for 'ALL ATTENDEES' events.
5. Header row repetition setting.
6. Visual styling (background colors).
"""

import json
import os
import tempfile
import logging
import zipfile

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_conference_schedule(traj, env_info, task_info):
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # Import docx here to handle potential import errors gracefully in the environment
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return {"passed": False, "score": 0, "feedback": "python-docx library not available in verifier"}

    # Load result metadata
    temp_json = tempfile.NamedTemporaryFile(delete=False, suffix='.json')
    try:
        copy_from_env("/tmp/task_result.json", temp_json.name)
        with open(temp_json.name, 'r') as f:
            result_meta = json.load(f)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Failed to load result metadata: {e}"}
    finally:
        if os.path.exists(temp_json.name):
            os.unlink(temp_json.name)

    if not result_meta.get("output_exists"):
        return {"passed": False, "score": 0, "feedback": "Output file conference_schedule.docx not found."}

    if not result_meta.get("created_during_task"):
        return {"passed": False, "score": 0, "feedback": "File was not modified during the task session."}

    # Retrieve the DOCX file
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    try:
        copy_from_env(result_meta["output_path"], temp_docx.name)
        doc = Document(temp_docx.name)
    except Exception as e:
        return {"passed": False, "score": 0, "feedback": f"Invalid or corrupted DOCX file: {e}"}
    finally:
        # Keep file for a moment if needed, but standard practice is to clean up
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)

    score = 0
    max_score = 100
    feedback = []
    
    # --- Criterion 1: Page Layout (20 pts) ---
    # Landscape: Width > Height
    # Margins: ~0.5 inches (457200 EMU)
    try:
        section = doc.sections[0]
        page_width = section.page_width
        page_height = section.page_height
        
        if page_width > page_height:
            score += 10
            feedback.append("✅ Orientation is Landscape")
        else:
            feedback.append("❌ Orientation is Portrait (expected Landscape)")

        # Margin check (tolerance +/- 0.1 inch = 91440 EMU)
        target_margin = 457200
        tolerance = 91440
        margins_ok = True
        for m in [section.left_margin, section.right_margin, section.top_margin, section.bottom_margin]:
            if m is None or abs(m - target_margin) > tolerance:
                margins_ok = False
                break
        
        if margins_ok:
            score += 10
            feedback.append("✅ Margins are approx 0.5 inches")
        else:
            feedback.append("❌ Margins are not 0.5 inches")
    except Exception as e:
        feedback.append(f"❌ Could not verify page layout: {e}")

    # --- Criterion 2: Table Structure (15 pts) ---
    if len(doc.tables) == 0:
        return {"passed": False, "score": score, "feedback": "No table found in document."}
    
    table = doc.tables[0]
    # Check column count
    # Note: With merged cells, len(table.columns) might raise error or be inaccurate in some versions,
    # but usually works for the underlying grid.
    try:
        col_count = len(table.columns)
        if col_count == 4:
            score += 15
            feedback.append("✅ Table has 4 columns")
        else:
            feedback.append(f"❌ Table has {col_count} columns (expected 4)")
    except:
        # Fallback: check first row cells
        if len(table.rows[0].cells) == 4:
            score += 15
            feedback.append("✅ Table has 4 columns (inferred)")
        else:
            feedback.append("❌ Table column count check failed")

    # --- Criterion 3: Header Repeat (15 pts) ---
    # Check for <w:tblHeader> in the first row's trPr
    try:
        row_xml = table.rows[0]._element.xml
        if 'w:tblHeader' in row_xml:
            score += 15
            feedback.append("✅ Header row set to repeat")
        else:
            feedback.append("❌ Header row NOT set to repeat")
    except:
        feedback.append("⚠️ Could not verify header repeat setting")

    # --- Criterion 4: Cell Merging (30 pts) ---
    # We expect rows containing "Keynote" or "Lunch" to span columns 2, 3, 4.
    # In a 4-col table, this means the row visually has 2 cells (Time, Event).
    # Programmatically, python-docx might show this as the cells sharing identity or having gridSpan.
    
    merge_points = 0
    keynote_merged = False
    lunch_merged = False
    
    for row in table.rows:
        text = " ".join(cell.text for cell in row.cells).lower()
        
        # Check specific rows
        if "opening keynote" in text:
            # Check for gridSpan in XML for the second cell (index 1)
            # OR check if cells[1] == cells[2] == cells[3]
            try:
                # Method A: Check if cell objects are identical (python-docx behavior for merged cells)
                if row.cells[1] == row.cells[2] and row.cells[2] == row.cells[3]:
                    keynote_merged = True
                # Method B: XML check if Method A fails (sometimes they are distinct objects but empty/merged)
                elif 'gridSpan' in row.cells[1]._element.xml:
                     keynote_merged = True
            except:
                pass
                
        if "lunch" in text:
            try:
                if row.cells[1] == row.cells[2] and row.cells[2] == row.cells[3]:
                    lunch_merged = True
                elif 'gridSpan' in row.cells[1]._element.xml:
                     lunch_merged = True
            except:
                pass

    if keynote_merged:
        merge_points += 15
        feedback.append("✅ Keynote cells merged correctly")
    else:
        feedback.append("❌ Keynote cells NOT merged")

    if lunch_merged:
        merge_points += 15
        feedback.append("✅ Lunch cells merged correctly")
    else:
        feedback.append("❌ Lunch cells NOT merged")
    
    score += merge_points

    # --- Criterion 5: Styling (20 pts) ---
    # Check for shading (w:shd) in the XML of rows that should be colored
    shading_found = False
    alignment_found = False
    
    table_xml = table._element.xml
    if "w:shd" in table_xml and "w:fill" in table_xml:
        shading_found = True
        score += 10
        feedback.append("✅ Background shading detected")
    else:
        feedback.append("❌ No background shading found")

    # Check vertical alignment (vAlign)
    if "w:vAlign" in table_xml:
        score += 10
        alignment_found = True
        feedback.append("✅ Vertical alignment detected")
    else:
        feedback.append("❌ Vertical alignment not explicitly set")

    # --- Final Score Calculation ---
    passed = score >= 75
    
    # Safety check: if merge failed completely, fail the task regardless of points
    if merge_points == 0:
        passed = False
        feedback.append("⛔ CRITICAL: No required cell merging detected.")

    return {
        "passed": passed,
        "score": score,
        "feedback": "\n".join(feedback)
    }