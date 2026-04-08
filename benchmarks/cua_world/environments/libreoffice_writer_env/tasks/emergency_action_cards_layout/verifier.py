#!/usr/bin/env python3
"""Verifier for emergency_action_cards_layout task."""

import json
import tempfile
import os
import logging
from typing import Dict, Any, List

# Import gym-anything utilities
from gym_anything.vlm import query_vlm, get_final_screenshot, sample_trajectory_frames

# Import docx via fallback in case not installed in system python (it is installed in env)
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
except ImportError:
    pass # Will fail gracefully inside verifier if copy fails

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def verify_emergency_cards(traj: Dict[str, Any], env_info: Dict[str, Any], task_info: Dict[str, Any]) -> Dict[str, Any]:
    """
    Verify the emergency cards document layout and content.
    
    Criteria:
    1. File creation/validity (10 pts)
    2. Page Setup: Landscape, Margins <= 0.5 (15 pts)
    3. Table Structure: 2x2 grid (15 pts)
    4. Text Content: Correct text in correct cells (10 pts)
    5. Formatting: Vertical centering, horizontal centering (20 pts)
    6. Cell Backgrounds: Detected via XML or VLM (20 pts)
    7. VLM Visual Confirmation: Colors and layout look correct (10 pts)
    """
    copy_from_env = env_info.get('copy_from_env')
    if not copy_from_env:
        return {"passed": False, "score": 0, "feedback": "Copy function not available"}

    # --- Step 1: File Retrieval & Basic Checks ---
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.close()
    
    try:
        copy_from_env("/home/ga/Documents/emergency_cards.docx", temp_file.name)
        if os.path.getsize(temp_file.name) == 0:
            raise ValueError("Empty file")
        doc = Document(temp_file.name)
        file_valid = True
    except Exception as e:
        if os.path.exists(temp_file.name):
            os.unlink(temp_file.name)
        return {
            "passed": False, 
            "score": 0, 
            "feedback": f"Failed to open output file: {str(e)}. Did you save it to /home/ga/Documents/emergency_cards.docx?"
        }

    score = 10
    feedback = ["File exists and is valid DOCX"]
    
    try:
        # --- Step 2: Page Setup (15 pts) ---
        section = doc.sections[0]
        # Check Orientation
        is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
        # Check Margins (0.5 inch = 457200 EMU) - Allow tolerance
        margin_limit = 460000 
        margins_ok = (
            section.left_margin <= margin_limit and 
            section.right_margin <= margin_limit and 
            section.top_margin <= margin_limit and 
            section.bottom_margin <= margin_limit
        )
        
        if is_landscape:
            score += 10
            feedback.append("Orientation: Landscape (Correct)")
        else:
            feedback.append("Orientation: Portrait (Incorrect, expected Landscape)")
            
        if margins_ok:
            score += 5
            feedback.append("Margins: <= 0.5 inch (Correct)")
        else:
            feedback.append("Margins: Too wide (Expected <= 0.5 inch)")

        # --- Step 3: Table Structure (15 pts) ---
        has_table = len(doc.tables) > 0
        grid_ok = False
        if has_table:
            table = doc.tables[0]
            rows = len(table.rows)
            cols = len(table.columns) if rows > 0 else 0
            if rows >= 2 and cols >= 2:
                grid_ok = True
                score += 15
                feedback.append(f"Table Structure: {rows}x{cols} (Correct)")
            else:
                feedback.append(f"Table Structure: {rows}x{cols} (Incorrect, expected 2x2)")
        else:
            feedback.append("No table found")

        # --- Step 4: Content & Formatting (30 pts) ---
        # We expect 4 cells: TL(Red), TR(Yellow), BL(Blue), BR(Green)
        # Content checks
        content_score = 0
        formatting_score = 0
        
        expected_content = [
            ("CODE RED", 0, 0),
            ("CODE YELLOW", 0, 1),
            ("CODE BLUE", 1, 0),
            ("ALL CLEAR", 1, 1) # or Green info
        ]
        
        if grid_ok:
            table = doc.tables[0]
            
            # Check Content matches quadrants
            matches = 0
            for text, r, c in expected_content:
                try:
                    cell_text = table.cell(r, c).text.upper()
                    if text in cell_text:
                        matches += 1
                except IndexError:
                    pass
            
            if matches >= 3:
                content_score = 10
                feedback.append(f"Content Placement: Correct ({matches}/4 quadrants)")
            else:
                feedback.append(f"Content Placement: Incorrect ({matches}/4 quadrants matched)")
                
            # Check Alignment (Horizontal & Vertical)
            # Sample a few cells
            center_h_count = 0
            center_v_count = 0
            checked_cells = 0
            
            for r in range(2):
                for c in range(2):
                    try:
                        cell = table.cell(r, c)
                        checked_cells += 1
                        
                        # Vertical Align
                        if cell.vertical_alignment == WD_ALIGN_VERTICAL.CENTER:
                            center_v_count += 1
                        
                        # Horizontal Align (check first paragraph)
                        if cell.paragraphs and cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            center_h_count += 1
                    except IndexError:
                        pass
            
            if center_v_count >= 2: # Allow some leniency
                formatting_score += 10
                feedback.append("Vertical Alignment: Centered")
            else:
                feedback.append("Vertical Alignment: Not Centered")
                
            if center_h_count >= 2:
                formatting_score += 10
                feedback.append("Horizontal Alignment: Centered")
            else:
                feedback.append("Horizontal Alignment: Not Centered")

        score += content_score + formatting_score

        # --- Step 5: Background Colors (XML Check) (20 pts) ---
        # Parsing raw XML for shading because python-docx doesn't expose read-access easily
        # We look for <w:shd w:fill="..."> in the cell properties
        xml_score = 0
        if grid_ok:
            cells_with_shading = 0
            table_xml = doc.tables[0]._element.xml
            if 'w:shd' in table_xml and 'w:fill' in table_xml:
                # Rudimentary check: does the XML contain shading definitions?
                # A more robust check would parse specific cells, but finding 'w:fill' 
                # usually implies the user set a background color.
                # To distinguish from "Auto", we check if fill is not "auto"
                import re
                fills = re.findall(r'w:fill="([0-9A-Fa-f]+|red|yellow|blue|green|cyan|magenta)"', table_xml)
                valid_fills = [f for f in fills if f.lower() != 'auto']
                
                if len(valid_fills) >= 3:
                    xml_score = 20
                    feedback.append(f"Cell Backgrounds: Detected {len(valid_fills)} colored cells")
                elif len(valid_fills) > 0:
                    xml_score = 10
                    feedback.append(f"Cell Backgrounds: Detected some colors ({len(valid_fills)})")
                else:
                    feedback.append("Cell Backgrounds: No specific colors detected in XML")
            else:
                feedback.append("Cell Backgrounds: No shading tags found")
        
        score += xml_score

    except Exception as e:
        feedback.append(f"Error during DOCX analysis: {e}")
        import traceback
        traceback.print_exc()

    finally:
        if os.path.exists(temp_file.name):
            os.unlink(temp_file.name)

    # --- Step 6: VLM Verification (10 pts + backup) ---
    # Use VLM to confirm the "Gestalt" of the page: 4 colored boxes, 2x2 grid
    vlm_score = 0
    final_screenshot = get_final_screenshot(traj)
    
    if final_screenshot:
        prompt = """
        Review this image of a document in LibreOffice Writer.
        Does it contain a 2x2 table with 4 distinct colored rectangular cells (Red, Yellow, Blue, Green)?
        Is the page in Landscape orientation (wider than it is tall)?
        Is the text centered in the cells?
        
        Respond in JSON:
        {
            "has_4_colored_cells": true/false,
            "colors_visible": ["list", "of", "colors"],
            "is_landscape": true/false,
            "text_centered": true/false,
            "confidence": "high/medium/low"
        }
        """
        vlm_result = query_vlm(prompt, final_screenshot)
        
        if vlm_result.get("success"):
            parsed = vlm_result.get("parsed", {})
            
            if parsed.get("has_4_colored_cells") or len(parsed.get("colors_visible", [])) >= 3:
                vlm_score += 5
                feedback.append("VLM: Verified colored cells")
            
            if parsed.get("is_landscape"):
                vlm_score += 2 # Bonus/Confirmation
            
            if parsed.get("text_centered"):
                vlm_score += 3
                feedback.append("VLM: Verified centered text")
                
            # Fallback: if XML check failed but VLM sees colors, grant partial points
            if xml_score == 0 and parsed.get("has_4_colored_cells"):
                score += 15
                feedback.append("VLM: Overrode XML check for colors")
                
    score += vlm_score
    score = min(score, 100) # Cap at 100

    return {
        "passed": score >= 75,
        "score": score,
        "feedback": " | ".join(feedback)
    }