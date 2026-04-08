#!/usr/bin/env python3
"""
LibreOffice Writer verification utilities for gym-anything tasks.
Provides helper functions to verify Writer document tasks using DOCX/ODT parsing.
"""

import logging
import os
import re
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional, Callable

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import python-docx for DOCX parsing
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - DOCX parsing will be limited")
    DOCX_AVAILABLE = False

# Import odfpy for ODT parsing
try:
    from odf import opendocument, text as odf_text, style as odf_style
    ODF_AVAILABLE = True
except ImportError:
    logger.warning("odfpy not available - ODT parsing will be limited")
    ODF_AVAILABLE = False


# ============================================================================
# Core copy/parse functions
# ============================================================================

def copy_and_parse_document(container_path: str, copy_from_env_fn: Callable,
                            file_format: str = 'auto') -> Tuple[bool, Any, str, str]:
    """
    Copy a document from the container and parse it.

    Args:
        container_path: Path to the file inside the container.
        copy_from_env_fn: Framework-provided copy function.
        file_format: 'docx', 'odt', or 'auto' (detect from extension).

    Returns:
        (success, document_object, error_message, temp_dir)
    """
    temp_dir = tempfile.mkdtemp(prefix='writer_verify_')
    ext = Path(container_path).suffix.lower()
    host_file = os.path.join(temp_dir, f"result{ext}")

    try:
        copy_from_env_fn(container_path, host_file)
    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return False, None, f"Failed to copy file: {e}", ''

    if not os.path.exists(host_file) or os.path.getsize(host_file) == 0:
        shutil.rmtree(temp_dir, ignore_errors=True)
        return False, None, f"File not found or empty: {container_path}", ''

    try:
        if file_format == 'auto':
            file_format = 'docx' if ext in ('.docx', '.doc') else 'odt'

        if file_format == 'docx' and DOCX_AVAILABLE:
            doc = Document(host_file)
            return True, doc, "", temp_dir
        elif file_format == 'odt' and ODF_AVAILABLE:
            doc = opendocument.load(host_file)
            return True, doc, "", temp_dir
        else:
            return False, None, f"Unsupported format or missing library: {ext}", temp_dir
    except Exception as e:
        return False, None, f"Parse error: {e}", temp_dir


def cleanup_verification_temp(temp_dir: Optional[str]):
    """Clean up temporary verification files."""
    if temp_dir and os.path.exists(temp_dir):
        shutil.rmtree(temp_dir, ignore_errors=True)


# ============================================================================
# Document text extraction
# ============================================================================

def get_document_text(doc) -> str:
    """Extract all text from a DOCX document."""
    return '\n'.join(para.text for para in doc.paragraphs)


def get_paragraph_styles(doc) -> List[Dict[str, Any]]:
    """Get list of paragraphs with their style and run info."""
    results = []
    for para in doc.paragraphs:
        results.append({
            'text': para.text,
            'style_name': para.style.name if para.style else None,
            'alignment': para.alignment,
            'runs': [{
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font else None,
                'font_size': run.font.size.pt if run.font and run.font.size else None,
            } for run in para.runs]
        })
    return results


# ============================================================================
# Heading and style checks
# ============================================================================

def check_heading_styles(doc, expected_headings: Dict[str, str]) -> Tuple[int, int, List[str]]:
    """
    Check that text content has correct heading styles applied.

    Args:
        doc: python-docx Document object
        expected_headings: {'text_fragment': 'Heading 1', ...}

    Returns:
        (matched_count, total_expected, feedback_list)
    """
    matched = 0
    total = len(expected_headings)
    feedback = []

    for text_fragment, expected_style in expected_headings.items():
        found = False
        for para in doc.paragraphs:
            if text_fragment.lower() in para.text.lower():
                actual_style = para.style.name if para.style else 'None'
                if expected_style.lower() in actual_style.lower():
                    matched += 1
                    found = True
                    feedback.append(f"OK: '{text_fragment[:40]}' has style '{actual_style}'")
                else:
                    feedback.append(
                        f"WRONG: '{text_fragment[:40]}' has '{actual_style}' "
                        f"(expected '{expected_style}')"
                    )
                    found = True
                break
        if not found:
            feedback.append(f"MISSING: '{text_fragment[:40]}' not found in document")

    return matched, total, feedback


def count_headings_by_level(doc) -> Dict[str, int]:
    """Count headings at each level in the document."""
    counts = {}
    for para in doc.paragraphs:
        if para.style and 'Heading' in para.style.name:
            level = para.style.name
            counts[level] = counts.get(level, 0) + 1
    return counts


# ============================================================================
# Table of Contents detection
# ============================================================================

def detect_toc_present(doc) -> bool:
    """
    Detect if a Table of Contents is present in the DOCX document.
    LibreOffice Writer stores TOC as structured fields in the XML.
    """
    try:
        from lxml import etree
    except ImportError:
        logger.warning("lxml not available for TOC detection")
        return False

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Method 1: Check for instrText with TOC
    body = doc.element.body
    instr_texts = body.findall('.//w:instrText', nsmap)
    for instr in instr_texts:
        if instr.text and 'TOC' in instr.text.upper():
            return True

    # Method 2: Check for SDT (structured document tag) with TOC
    sdts = body.findall('.//w:sdt', nsmap)
    for sdt in sdts:
        alias = sdt.find('.//w:alias', nsmap)
        if alias is not None:
            val = alias.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', ''
            )
            if 'TOC' in val.upper() or 'Table of Contents' in val:
                return True

    # Method 3: Check for fldChar with TOC
    for para in doc.paragraphs:
        xml_str = para._element.xml
        if 'TOC' in xml_str and ('w:fldChar' in xml_str or 'w:instrText' in xml_str):
            return True

    # Method 4: Heuristic - "Table of Contents" or "Contents" heading
    # followed by entries that look like TOC entries (with page numbers or tab leaders)
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        if text_lower in ('table of contents', 'contents'):
            # Require at least 2 subsequent paragraphs that look like TOC entries
            # (contain trailing page numbers, tab characters, or dot leaders)
            toc_entry_count = 0
            for j in range(i + 1, min(i + 20, len(doc.paragraphs))):
                entry_text = doc.paragraphs[j].text.strip()
                if not entry_text:
                    continue
                # TOC entries typically end with a page number
                if re.search(r'\d+\s*$', entry_text):
                    toc_entry_count += 1
                # Or contain tab characters (used for leader dots)
                elif '\t' in entry_text:
                    toc_entry_count += 1
                # Or have a TOC-style style applied
                elif doc.paragraphs[j].style and 'toc' in doc.paragraphs[j].style.name.lower():
                    toc_entry_count += 1
            if toc_entry_count >= 2:
                return True

    return False


# ============================================================================
# Text formatting checks
# ============================================================================

def _resolve_run_bold(run, para) -> Optional[bool]:
    """Resolve bold formatting considering style inheritance."""
    if run.bold is not None:
        return run.bold
    # Check paragraph style's font bold setting
    if para.style and para.style.font and para.style.font.bold is not None:
        return para.style.font.bold
    # Check if paragraph style name implies bold (Heading, Title styles)
    if para.style and para.style.name:
        style_name = para.style.name.lower()
        if any(kw in style_name for kw in ('heading', 'title')):
            return True
    return None


def _resolve_run_italic(run, para) -> Optional[bool]:
    """Resolve italic formatting considering style inheritance."""
    if run.italic is not None:
        return run.italic
    # Check paragraph style's font italic setting
    if para.style and para.style.font and para.style.font.italic is not None:
        return para.style.font.italic
    # Check if paragraph style name implies italic (Subtitle styles)
    if para.style and para.style.name:
        style_name = para.style.name.lower()
        if 'subtitle' in style_name:
            return True
    return None


def check_text_formatting(doc, text_fragment: str, bold=None, italic=None,
                          underline=None, font_size=None) -> bool:
    """
    Check if text in the document has specific formatting.
    Handles style-inherited formatting (e.g., bold from Heading style).
    Also handles run fragmentation (text split across multiple runs).

    Args:
        doc: python-docx Document object
        text_fragment: Text to search for (case-insensitive)
        bold: Expected bold state (True/False/None to skip)
        italic: Expected italic state (True/False/None to skip)
        underline: Expected underline state (True/False/None to skip)
        font_size: Expected font size in Pt (None to skip)

    Returns:
        True if matching text with all specified formatting is found.
    """
    for para in doc.paragraphs:
        # First try: match within individual runs
        for run in para.runs:
            if text_fragment.lower() in run.text.lower():
                if bold is not None:
                    resolved_bold = _resolve_run_bold(run, para)
                    if resolved_bold != bold:
                        continue
                if italic is not None:
                    resolved_italic = _resolve_run_italic(run, para)
                    if resolved_italic != italic:
                        continue
                if underline is not None and run.underline != underline:
                    continue
                if font_size is not None:
                    if run.font.size is None:
                        continue
                    if abs(run.font.size.pt - font_size) > 1.0:
                        continue
                return True

        # Second try: text may be split across multiple runs (run fragmentation)
        # Check if the fragment exists in the concatenated paragraph text
        if text_fragment.lower() in para.text.lower() and para.runs:
            # Verify formatting on all runs that overlap with the text region
            all_match = True
            if bold is not None:
                # Check that at least one run has the formatting, or it's inherited
                has_bold = any(
                    _resolve_run_bold(r, para) == bold
                    for r in para.runs if r.text.strip()
                )
                if not has_bold:
                    all_match = False
            if italic is not None:
                has_italic = any(
                    _resolve_run_italic(r, para) == italic
                    for r in para.runs if r.text.strip()
                )
                if not has_italic:
                    all_match = False
            if font_size is not None:
                has_size = any(
                    r.font.size is not None and abs(r.font.size.pt - font_size) <= 1.0
                    for r in para.runs if r.text.strip()
                )
                if not has_size:
                    all_match = False
            if all_match:
                return True
    return False


def check_paragraph_alignment(doc, text_fragment: str, alignment: str) -> bool:
    """
    Check if paragraph containing text has specified alignment.

    Args:
        doc: python-docx Document object
        text_fragment: Text to search for in paragraph
        alignment: 'left', 'center', 'right', or 'justify'

    Returns:
        True if paragraph with matching text has the specified alignment.
    """
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT if DOCX_AVAILABLE else 0,
        'center': WD_ALIGN_PARAGRAPH.CENTER if DOCX_AVAILABLE else 1,
        'right': WD_ALIGN_PARAGRAPH.RIGHT if DOCX_AVAILABLE else 2,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY if DOCX_AVAILABLE else 3,
    }
    target = alignment_map.get(alignment.lower())

    for para in doc.paragraphs:
        if text_fragment.lower() in para.text.lower():
            if para.alignment == target:
                return True
    return False


# ============================================================================
# Indentation checks
# ============================================================================

def check_hanging_indent(para) -> bool:
    """
    Check if a paragraph has hanging indent formatting.
    Hanging indent means: left_indent > 0 and first_line_indent is negative
    (so the first line is less indented than subsequent lines).

    Args:
        para: python-docx Paragraph object

    Returns:
        True if paragraph has hanging indent.
    """
    try:
        pf = para.paragraph_format
        left_indent = pf.left_indent
        first_line_indent = pf.first_line_indent

        if left_indent is None:
            # Also check XML directly for w:hanging attribute
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            ind = para._element.find('.//w:pPr/w:ind', nsmap)
            if ind is not None:
                hanging = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging')
                if hanging and int(hanging) > 0:
                    return True
            return False

        # Convert EMU to inches (914400 EMU = 1 inch)
        left_inches = left_indent / 914400

        if first_line_indent is None:
            # When first_line_indent is None, check XML for w:hanging attribute
            # A plain block indent (left only) is NOT a hanging indent
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            ind = para._element.find('.//w:pPr/w:ind', nsmap)
            if ind is not None:
                hanging = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging')
                if hanging and int(hanging) > 0:
                    return True
            return False

        first_inches = first_line_indent / 914400

        # Hanging indent: left > 0 and first line is negative (pulled back)
        return (left_inches >= 0.3) and (first_inches < 0)
    except Exception:
        return False


def check_hanging_indent_count(doc, text_fragments: Optional[List[str]] = None) -> int:
    """
    Count paragraphs with hanging indent.

    Args:
        doc: python-docx Document object
        text_fragments: Optional list of text fragments to check.
                       If None, checks all non-empty paragraphs.

    Returns:
        Number of paragraphs with hanging indent.
    """
    count = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if text_fragments:
            if not any(frag.lower() in para.text.lower() for frag in text_fragments):
                continue
        if check_hanging_indent(para):
            count += 1
    return count


# ============================================================================
# Italic text detection
# ============================================================================

def has_italicized_text(para) -> bool:
    """Check if paragraph contains any substantive italicized text."""
    for run in para.runs:
        if run.italic and len(run.text.strip()) > 3:
            return True
    return False


def count_paragraphs_with_italics(doc, min_text_length: int = 30) -> int:
    """Count paragraphs that contain italic runs (for substantial paragraphs only)."""
    count = 0
    for para in doc.paragraphs:
        if len(para.text.strip()) < min_text_length:
            continue
        if has_italicized_text(para):
            count += 1
    return count


# ============================================================================
# Mail merge verification
# ============================================================================

def check_mail_merge_output(doc, expected_names: List[str]) -> Tuple[int, int, List[str]]:
    """
    Check mail merge output contains personalized letters for all expected names.

    Args:
        doc: python-docx Document object
        expected_names: List of names that should appear in merged output.

    Returns:
        (found_count, total_expected, feedback_list)
    """
    full_text = get_document_text(doc).lower()
    found = 0
    feedback = []

    for name in expected_names:
        if name.lower() in full_text:
            found += 1
            feedback.append(f"OK: Letter for '{name}' found")
        else:
            feedback.append(f"MISSING: Letter for '{name}' not found")

    return found, len(expected_names), feedback


def check_no_raw_placeholders(doc, placeholders: List[str]) -> Tuple[int, List[str]]:
    """
    Check that no raw placeholders remain in the document.

    Args:
        doc: python-docx Document object
        placeholders: List of placeholder strings like ['{Name}', '{Address}'].

    Returns:
        (violation_count, feedback_list)
    """
    full_text = get_document_text(doc)
    violations = 0
    feedback = []

    for placeholder in placeholders:
        if placeholder in full_text:
            violations += 1
            feedback.append(f"RAW PLACEHOLDER: '{placeholder}' still in document")

    return violations, feedback


def verify_page_breaks(doc) -> int:
    """
    Count page breaks in document (useful for mail merge output).

    Returns:
        Number of page break elements found.
    """
    count = 0
    for para in doc.paragraphs:
        xml_str = para._element.xml
        if 'w:br' in xml_str and 'type="page"' in xml_str:
            count += 1
    return count


# ============================================================================
# Citation / bibliography verification
# ============================================================================

def check_apa_citation_format(text: str) -> Tuple[bool, List[str]]:
    """
    Check if a single citation text follows APA 7th edition format.

    Returns:
        (is_valid, list_of_issues)
    """
    issues = []

    # Check for year in parentheses
    if not re.search(r'\(\d{4}[a-z]?\)', text):
        issues.append("Missing year in parentheses")

    # Check for period after parenthetical year
    if re.search(r'\(\d{4}\)', text) and not re.search(r'\(\d{4}[a-z]?\)\.', text):
        issues.append("Missing period after year")

    # Check author format (Last, F. M. or Last, F.)
    if not re.search(r'^[A-Z][a-z\'-]+,\s+[A-Z]\.', text):
        issues.append("Author format may not follow APA (Last, F. M.)")

    return len(issues) == 0, issues


def check_alphabetical_order(doc, start_after: str = "References") -> Tuple[bool, List[str]]:
    """
    Check if citation entries are sorted alphabetically by first author last name.

    Args:
        doc: python-docx Document object
        start_after: Heading text after which citations begin.

    Returns:
        (is_sorted, first_words_in_order)
    """
    citations = []
    in_references = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if not in_references:
            if start_after.lower() in text.lower():
                in_references = True
            continue

        # Skip empty lines and short text
        if not text or len(text) < 20:
            continue

        # Skip instruction lines
        if text.startswith('-') or text.startswith('TASK') or text.startswith('MISSING'):
            continue

        citations.append(text)

    if len(citations) < 2:
        return False, ["Not enough citations to check order"]

    first_words = []
    for cite in citations:
        words = cite.split()
        if words:
            first_words.append(words[0].strip('.,;:()[]{}'))

    sorted_words = sorted(first_words, key=str.lower)
    is_sorted = [w.lower() for w in first_words] == [w.lower() for w in sorted_words]

    return is_sorted, first_words


def extract_citation_paragraphs(doc, start_after: str = "References") -> List:
    """
    Extract citation paragraphs from the document (after the References heading).

    Returns:
        List of paragraph objects that are citations.
    """
    citations = []
    in_references = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if not in_references:
            if start_after.lower() in text.lower():
                in_references = True
            continue

        if not text or len(text) < 20:
            continue

        if text.startswith('-') or text.startswith('TASK') or text.startswith('MISSING'):
            continue

        citations.append(para)

    return citations


# ============================================================================
# VLM cross-validation helper
# ============================================================================

def vlm_verify_screenshot(env_info, traj, prompt: str) -> Optional[Dict]:
    """
    Query the VLM with the final screenshot and a structured prompt.
    Returns parsed JSON from VLM response, or None if VLM is unavailable.

    Args:
        env_info: Environment info dict (contains query_vlm, get_final_screenshot).
        traj: Trajectory dict.
        prompt: Structured prompt asking VLM to evaluate the screenshot.

    Returns:
        Parsed dict from VLM response, or None.
    """
    query_vlm = env_info.get('query_vlm')
    get_final_screenshot = env_info.get('get_final_screenshot')

    if not query_vlm or not get_final_screenshot:
        return None

    final_frame = get_final_screenshot(traj)
    if not final_frame:
        return None

    try:
        result = query_vlm(prompt=prompt, image=final_frame)
        if result and result.get("success"):
            return result.get("parsed", {})
        logger.warning(f"VLM query failed: {result.get('error', 'unknown') if result else 'no result'}")
    except Exception as e:
        logger.warning(f"VLM query exception: {e}")

    return None
